# -*- coding: utf-8 -*-
"""组装完整专著 docx。按章节顺序登记引用，末尾按顺序编码制生成参考文献。"""
import importlib
import os
import sys

from docxbuilder import BookBuilder, count_chars, SIZE, SONG, HEI
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import references as refs

HERE = os.path.dirname(os.path.abspath(__file__))

CHAPTERS = ['content_ch01', 'content_ch02', 'content_ch03', 'content_ch04', 'content_ch05',
            'content_ch06', 'content_ch07', 'content_ch08', 'content_ch09', 'content_ch10',
            'content_appendix']


def title_page(bk):
    doc = bk.doc
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bk._append_rich(p, '全国统一大市场建设与要素配置效率', HEI, SIZE['小二'], bold=True)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bk._append_rich(p2, '——理论建构、测度体系与改革路径', HEI, SIZE['三号'])
    for _ in range(6):
        doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bk._append_rich(p3, '蔡鑫宇  著', SONG, SIZE['四号'])
    bk.add_pagebreak()


def load_blocks():
    all_blocks = []
    # 前置：前言 + 内容摘要
    import content_front
    all_blocks += content_front.preface_blocks()
    all_blocks.append({'pagebreak': True})
    all_blocks += content_front.abstract_blocks()
    all_blocks.append({'pagebreak': True})
    # 正文各章
    counts = {}
    refs.reset()
    for mod in CHAPTERS:
        m = importlib.import_module(mod)
        before = refs.n_used()
        blks = m.blocks()
        counts[mod] = count_chars(blks)
        all_blocks += blks
        all_blocks.append({'pagebreak': True})
    # 后记
    all_blocks += content_front.postscript_blocks()
    all_blocks.append({'pagebreak': True})
    return all_blocks, counts


def main():
    bk = BookBuilder(os.path.join(HERE, 'build'))
    all_blocks, counts = load_blocks()
    # 正文字数
    body_chars = sum(counts.values())
    front_chars = count_chars([b for b in all_blocks if 'h1' in b and b.get('h1') == '内容摘要']) if False else 0
    total_chars = count_chars(all_blocks)
    title_page(bk)
    bk.build(all_blocks)
    # 参考文献（按首次引用顺序）
    bk.add_references(refs.gb_ordered())
    out = os.path.join(HERE, 'build', '全国统一大市场建设与要素配置效率.docx')
    bk.save(out)
    print('=== 字数统计 ===')
    for m in CHAPTERS:
        print(f'  {m}: {counts[m]:>7} 字')
    print(f'  正文合计: {total_chars} 字（含摘要，不含参考文献英文串）')
    print(f'  参考文献: {refs.n_used()} 条（已引用） / {len(refs.REFS)} 条（库）')
    print('  saved:', out)
    return total_chars


if __name__ == '__main__':
    main()
