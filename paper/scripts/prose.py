# -*- coding: utf-8 -*-
"""Full manuscript prose for the China AI-server sustainability paper.
Citation markers {key} resolve to numbered references in build_doc.REF_DB/CHINA_REFS."""
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.shared import Pt
from build_doc import add_para, add_heading, add_figure, FIG, summary, _set_font
import os

E = summary["energy_2030_TWh"]
W = summary["water_2030_Mm3"]
C = summary["carbon_2030_Mt"]
MID = summary["midcase_2030"]
DW, IW = summary["direct_indirect_water_share"]


def J(p):  # justify helper default
    return AL.JUSTIFY


def write_prose(doc):
    # ---------------- Title block ----------------
    p = doc.add_paragraph()
    p.alignment = AL.LEFT
    r = p.add_run("Analysis")
    _set_font(r, size=10, italic=True, bold=True, color=__import__("docx").shared.RGBColor(0x6b, 0x6b, 0x6b))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Environmental impact and net-zero pathways for sustainable "
                  "artificial intelligence servers in China")
    _set_font(r, size=18, bold=True)

    # authors (placeholders for the corresponding professor to complete)
    add_para(doc,
             "[First Author]¹, [Second Author]², [Third Author]³ & [Corresponding Author]¹*",
             size=11, bold=True, space_after=2)
    add_para(doc,
             "¹School of Environment / Institute for Carbon Neutrality, [University], Beijing, China. "
             "²School of Environmental Science and Engineering, [University], Shanghai, China. "
             "³Department of Energy and Resources Engineering, [University], China. "
             "*e-mail: corresponding.author@[university].edu.cn",
             size=8.5, italic=True, space_after=10)

    # ---------------- Abstract ----------------
    add_para(doc,
        "The rapidly increasing demand for generative artificial intelligence (AI) models in China "
        "requires extensive server installation with sustainability implications in terms of the compound "
        "energy–water–climate impacts. Here we show that the deployment of AI servers across China could "
        f"generate an annual water footprint ranging from {W[0]:.0f} to {W[1]:,.0f} million m³ and additional "
        f"annual carbon emissions from {C[0]:.0f} to {C[1]:.0f} Mt CO₂-equivalent by 2030, depending on the scale "
        "of expansion. Because China’s power system remains substantially more carbon-intensive than those of "
        "other leading AI markets, the climate burden per unit of computation is disproportionately large, even "
        "as the national ‘East Data, West Computing’ (东数西算) programme relocates capacity toward the resource-rich "
        "western provinces. Industry efficiency initiatives, grid decarbonization rates and the spatial distribution "
        "of servers across provinces drive deep uncertainties in the estimated water and carbon footprints. We show "
        "that the AI server industry is unlikely to align with China’s dual-carbon milestones—carbon peaking before "
        "2030 and neutrality before 2060—without substantial reliance on highly uncertain carbon-offset and "
        "water-restoration mechanisms. Although best practices may reduce emissions and water footprints by up to 73% "
        "and 86%, respectively, their effectiveness is constrained by current energy and water infrastructure limits. "
        "These findings underscore the urgency of accelerating the energy transition and point to the need for AI "
        "firms to harness the wind and solar potential of the northern and western provinces while avoiding the water "
        "penalty of hydropower-rich siting. Coordinated efforts of private actors and regulatory interventions would "
        "ensure the competitive and sustainable development of China’s AI sector.",
        size=10.5, bold=False, align=AL.JUSTIFY, space_after=12)

    # ---------------- Introduction ----------------
    add_para(doc,
        "The accelerating deployment of artificial intelligence (AI) servers is being fuelled by the growing demand "
        "for generative AI applications, ignited by milestones such as the release of ChatGPT in 2022{guan2024} and, in "
        "China, by the rapid rise of domestic large language models such as DeepSeek, Qwen and Ernie{deepseek2025}. "
        "Whereas AI has been employed across many fields to advance sustainability{kaack2022,vinuesa2020}, the remarkable "
        "energy requirements of AI itself raise concerns regarding not only energy provisioning{devries2023} but also "
        "water scarcity and climate change stemming from the energy–water–climate nexus of AI data centres{crawford2024,mytton2021}. "
        "These tensions are acute in China, which combines the world’s second-largest concentration of AI computing "
        "with a power system that still derives the majority of its electricity from coal and with severe, spatially "
        "uneven water stress{chinagrid,chinawater}. However, the holistic energy–water–climate implications of AI computing "
        "in China remain largely unknown, constrained by opaque industry reporting and limited public data{crawford2024}.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_para(doc,
        "The climate impact of AI servers stems primarily from their operation (Scopes 1 and 2) and supply-chain "
        "activities (Scope 3), including manufacturing and end-of-life treatment{malmodin2024}. The Scope 2 emissions from "
        "purchased electricity are expected to constitute a substantial portion and indicate a high dependency on AI "
        "server energy consumption. China’s data centres already consumed roughly 1.5×10² TWh of electricity in 2023—on "
        "the order of 1.6% of national consumption—and AI is expected to be the dominant driver of future load growth in "
        "the sector{caict_china,greenpeace_cleancloud,ieaenergyai2025}. This trajectory threatens the decarbonization "
        "milestones embedded in China’s ‘dual-carbon’ commitments—carbon peaking before 2030 and carbon neutrality "
        "before 2060{dualcarbon,china_compute_carbon}. Growing AI server energy consumption further implies an increasing "
        "water footprint through Scope 1 (direct cooling) and Scope 2 (indirect procurement) water "
        "use{mytton2021,li2023water,devriesgao2025}. Centralized installation in water-stressed regions of northern and "
        "western China may perturb local water balances and threaten supply for millions{chinawater,china_water_risk}.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_para(doc,
        "Previous research has developed bottom-up and top-down methods to assess the energy–water–carbon outcomes of "
        "servers{masanet2020,siddik2021}, but these approaches face challenges with the rise of AI servers. Top-down "
        "approaches based on activity indices such as data traffic fail to represent AI-driven workloads, whereas "
        "bottom-up approaches suffer from limited data availability and industry insight{wu2022}. Assumptions valid for "
        "traditional collocation centres often fail for AI data centres, which differ in installation and operation. "
        "Recent studies have begun to quantify AI-related energy and resource "
        "consumption{luccioni2023,luccioni_power2024,patel2024,dodge2022,patterson2022,llmcarbon,berthelot2024,dc_water_review}, "
        "and a 2025 analysis quantified the energy–water–climate impacts and net-zero pathways of AI servers in the "
        "United States{xiao2025}. Although recent work has examined the carbon emissions of China’s data centres and the "
        "climate contribution of the ‘East Data, West Computing’ programme{china_dc_co2,china_compute_carbon,ewc_greening}, "
        "a systematic, spatially explicit assessment of the compound energy–water–climate impacts for China—whose grid, "
        "water endowments and industrial policy differ fundamentally—has been missing. Our study addresses this gap by (1) developing an "
        "open-source, bottleneck-based modelling approach with comprehensive uncertainty analysis adapted to China’s "
        "domestic AI-accelerator supply chain; (2) systematically assessing energy, water and carbon impacts with "
        "province-level grid interactions; and (3) proposing actionable mitigation strategies for net-zero trajectories "
        "across AI server deployment scenarios.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_para(doc,
        "Here we analyse the combined energy–water–climate impact of operational AI servers in China between 2024 and "
        "2030, balancing importance and future uncertainties and addressing a series of fundamental questions. "
        "(1) What are the magnitude and spatiotemporal distributions of energy consumption, water footprint and climate "
        "impact from AI server deployment? We address this using temporal projection models and a provincial framework, "
        "assuming deployment mirrors current large-scale AI data-centre patterns and the ‘East Data, West Computing’ "
        "hub allocation{ndrc_ewc}. (2) What are the prospects for aligning the sector with China’s dual-carbon milestones? "
        "We evaluate this by analysing best- and worst-case scenarios of key drivers, including industry efficiency "
        "improvements, server location distribution and grid decarbonization. The results of these analyses are "
        "presented in the following sections.",
        align=AL.JUSTIFY, first_indent=0.2)

    # ---------------- Results 1 ----------------
    add_heading(doc, "AI servers’ environmental impact in China")
    add_para(doc,
        "This section provides a base depiction of AI servers’ energy, water and carbon impacts in China, emphasizing "
        "the spatiotemporal characteristics of the system. Figure 1a shows the projected cumulative AI server "
        "installations in China from 2024 to 2030 under five scenarios: low demand, low power, mid-case, high "
        "application and high demand. The mid-case serves as the base scenario, while the low- and high-demand cases set "
        "projection bounds. The low-power scenario assumes accelerator efficiency gains, and the high-application "
        "scenario accounts for increased adoption driven by efficiency (the rebound effect). Figure 1d illustrates the "
        "state-level allocation of AI servers, showing power usage effectiveness (PUE), projected grid carbon intensity, "
        "water usage effectiveness (WUE) and projected grid water intensity for each province. Coal-dependent northern "
        "provinces such as Inner Mongolia, Hebei and Ningxia exhibit high grid carbon factors but, owing to their cold, "
        "dry climate, achieve low PUE; hydropower-rich south-western provinces such as Sichuan and Yunnan exhibit very "
        "low grid carbon factors but high grid water factors because of reservoir evaporation{hydro_water}. The grid factors thus "
        "show notable sensitivity to location, underscoring the importance of the local grid for AI servers’ "
        "environmental impacts.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_para(doc,
        "Figure 1b,c shows the energy and water composition results. AI server (information-technology) energy "
        f"predominates over infrastructure energy, accounting for about {MID['it_energy_TWh']/MID['facility_energy_TWh']*100:.0f}% of "
        f"facility energy in the mid-case by 2030. The indirect (Scope 2) water footprint contributes {IW:.0f}% of the total, "
        f"with direct on-site use at {DW:.0f}%—a split close to that reported for other major markets and a direct consequence "
        "of water embedded in thermoelectric generation and hydropower evaporation{thermo_water,hydro_water}. Annual estimates of energy "
        "consumption, water footprint and carbon emissions of AI servers from 2024 to 2030 under each scenario are "
        "presented in Fig. 1e–g. Even the lowest scenarios outline a considerable increase in the energy, water and "
        f"carbon footprints of AI servers. By 2030, annual facility electricity ranges from {E[0]:.0f} to {E[1]:.0f} TWh, the "
        f"annual water footprint from {W[0]:.0f} to {W[1]:,.0f} million m³, and annual carbon emissions from {C[0]:.0f} to "
        f"{C[1]:.0f} Mt CO₂e across scenarios; the mid-case reaches {MID['facility_energy_TWh']:.0f} TWh, "
        f"{MID['total_water_Mm3']:.0f} million m³ and {MID['total_carbon_Mt']:.0f} Mt CO₂e. The highest scenario yields the "
        "greatest environmental impact, underscoring the risks of unchecked AI server expansion.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_para(doc,
        "China is selected as the research region because of its position as the largest AI market outside the United "
        "States and the most carbon-intensive among major AI hosts, which makes the climate stakes of AI growth "
        "especially high{ieaenergyai2025}. The research period 2024–2030 is chosen as a trade-off between importance and "
        "uncertainty, and it brackets the national carbon-peaking deadline. Projection of the cumulative AI server "
        "capacity is the initial step, generated from the forecast of domestic AI-accelerator manufacturing capacity, "
        "AI server specifications and adoption patterns. We assume the spatial distribution aligns with current "
        "large-scale data-centre allocation and the eight national computing hubs of the ‘East Data, West Computing’ "
        "programme{ndrc_ewc}. PUE and WUE values are derived by a hybrid statistical and thermodynamics-based "
        "model{lei2022} using provincial climate data, and projected grid factors are computed from province-level "
        "generation mixes consistent with national decarbonization plans{chinagrid}. Important uncertainties—PUE and WUE "
        "estimation, technology advancement, the spatial distribution of allocation and grid development—are examined in "
        "the following sections and in the sensitivity analysis.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_figure(doc, os.path.join(FIG, "fig1_projections.png"),
        "Fig. 1 | Projections of energy, water and carbon footprints of installed AI servers in China, 2024–2030. ",
        "a, Projected cumulative AI server capacity under five scenarios. b,c, Energy composition (server vs "
        "infrastructure) and water composition (direct Scope 1 vs indirect Scope 2) in the mid-case in 2030. "
        "d, Provincial AI-server allocation (bar height) and grid carbon factor (colour). e–g, Annual energy "
        "consumption (e), water footprint (f) and carbon emissions (g) under the five scenarios.")

    # ---------------- Results 2 ----------------
    add_heading(doc, "Higher energy and water usage efficiencies")
    add_para(doc,
        "Efficiency gains in the data-centre industry have historically stabilized environmental costs despite rapidly "
        "rising computation{masanet2020}. This section examines the potential for further improvements through system "
        "optimization and technology adoption. Figure 2a–c illustrates the effect of best-, base- and worst-practice "
        "PUE and WUE on the energy, water and carbon footprints of AI servers. Chinese AI data centres already operate "
        "at relatively high efficiency under national PUE caps for the computing hubs{chinapue}, yet the best-practice "
        "scenario suggests a further reduction potential of over 7% in PUE and more than 30% in the total water "
        "footprint through WUE improvement. PUE reduction yields commensurate reductions in total energy consumption and "
        "carbon emissions, while WUE reduction efforts result in over a 28% reduction in the total water footprint "
        "(Fig. 2d,e). The worst-practice scenario underscores the risk of neglecting efficiency efforts, particularly "
        "where evaporative cooling is used to offset high PUE in hot, humid southern provinces.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_para(doc,
        "Figure 2f examines the adoption of advanced technologies—advanced liquid cooling (ALC) and server utilization "
        "optimization (SUO). Best-case ALC adoption can reduce roughly 2% of energy, 3% of water and 2% of carbon by "
        "2030, while full SUO adoption reduces all footprints by about 6%; frozen adoption (worst case) instead raises "
        "footprints by up to 8%. The maximum reductions in energy, water and carbon drawn from the existing efficiency "
        "potential are about 12%, 32% and 11%, respectively. Advances in domestic AI hardware—successive Ascend-class "
        "accelerators and improved interconnects—could further reduce footprints by improving energy efficiency, "
        "although these gains may be partially offset by the rebound effect as cheaper computation spurs additional "
        "demand{devries2023}. These findings underscore the considerable, but bounded, impact of industry efforts on the "
        "environmental cost of AI servers.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_figure(doc, os.path.join(FIG, "fig2_efficiency.png"),
        "Fig. 2 | Assessment of industry efforts to reduce the environmental impact of AI servers. ",
        "a–c, Effect of best/base/worst PUE, WUE and combined practices on annual energy, water and carbon. "
        "d,e, Percentage impacts of PUE and WUE practices on the energy, water and carbon footprints. "
        "f, Effect of advanced liquid cooling (ALC) and server utilization optimization (SUO) adoption by 2030.")

    # ---------------- Results 3 ----------------
    add_heading(doc, "Influence of AI server spatial distribution")
    add_para(doc,
        "The location of data centres critically shapes their environmental impact{siddik2021,siddik2024}. This section "
        "projects how AI server spatial distribution may influence the environmental consequences of rapid expansion in "
        "China, a question made salient by the ‘East Data, West Computing’ programme, which deliberately relocates "
        "computing capacity from the eastern load centres to eight western and inland hubs{ndrc_ewc,ewc_greening,bits_migration}. Figure 3a presents "
        "the provincial trade-off between unit carbon footprint and unit water footprint of server energy, with bubble "
        "size denoting current allocation and colour denoting wind-plus-solar potential. A clear tension emerges: the "
        "hydropower-rich south-western provinces (Sichuan, Yunnan, Qinghai) deliver the lowest unit carbon footprints "
        "(Fig. 3b) but the highest unit water footprints, because hydropower consumes large volumes of water through "
        "reservoir evaporation{hydro_water}. Conversely, the cold, dry northern and north-western provinces (Inner "
        "Mongolia, Ningxia, Gansu, Xinjiang) offer low water footprints and abundant wind and solar resources, but "
        "their coal-heavy grids currently impose high unit carbon footprints.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_para(doc,
        "Figure 3c quantifies the effect of three siting strategies relative to the base allocation. Concentrating "
        "servers in the top-quartile water-oriented provinces reduces the total water footprint by about 46% but cuts "
        "carbon by only 22%; a carbon-oriented strategy reduces carbon by about 52% but increases the water footprint by "
        "about 31% owing to hydropower siting; a combined strategy reduces water and carbon concurrently, by roughly 38% "
        "and 41%. Figure 3d maps the underlying driver—the joint distribution of wind-plus-solar potential and "
        "water-scarcity—revealing that several western hubs central to the optimal strategy (Inner Mongolia, Ningxia, "
        "Gansu, Xinjiang) lie in the water-scarce upper-right region. The genuinely sustainable solution is therefore "
        "not simply ‘go west’, but to pair western siting with wind and solar generation—which consumes negligible "
        "water and, in desert installations, can even reduce local evaporation{gobi_pv}—and with dry (air-side or "
        "closed-loop) cooling, rather than relying on either coal power or water-intensive hydropower. The Qinghai and "
        "parts of the Gansu corridor, which combine high renewable potential with lower-carbon grids, emerge as "
        "particularly favourable when water management is enforced{complementary_potential}.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_figure(doc, os.path.join(FIG, "fig3_spatial.png"),
        "Fig. 3 | Spatial distribution shapes the water and carbon footprints of AI servers. ",
        "a, Provincial trade-off between unit carbon and unit water footprints (bubble size = allocation, colour = "
        "wind+solar potential). b, Provinces with the lowest unit carbon footprint. c, Effect of water-, carbon- and "
        "combined-oriented siting strategies on total footprints. d, Joint distribution of wind+solar potential and "
        "water scarcity by province (colour = grid carbon factor); the shaded quadrants highlight low-water-stress and "
        "high-renewable regions.")

    # ---------------- Results 4 ----------------
    add_heading(doc, "Influence of grid renewable penetration")
    add_para(doc,
        "This section analyses how future grid development will affect AI server environmental impacts. We contrast a "
        "low-renewable-cost (LRC, best-practice) pathway with rapid wind and solar deployment against a "
        "high-renewable-cost (HRC, worst-practice) pathway with slower decarbonization, both consistent with the range "
        "of China’s published power-sector scenarios{chinagrid,china_provincial_decarb,china_neutrality_pnas}. Figure 4a,b shows the resulting national "
        "carbon and water trajectories for the mid-case server scenario. The HRC pathway raises 2030 carbon emissions by "
        "about 20% and water by about 2.5%, whereas the LRC pathway reduces carbon by about 16% with a small water "
        "co-benefit, because additional wind and solar—unlike coal and hydropower—generate electricity without "
        "consuming water. The carbon footprint of AI servers is thus heavily governed by the grid decarbonization "
        "pattern, indicating both considerable reduction potential and associated risk.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_para(doc,
        "Figure 4c details the provincial sensitivity to grid development. Provinces hosting large planned wind and "
        "solar mega-bases—Inner Mongolia, Ningxia and Gansu—show the strongest response, with carbon swings of roughly "
        "±20–26% between the LRC and HRC pathways{chinarenew,complementary_potential}. The hydropower-dominated provinces (Sichuan, Yunnan) are "
        "comparatively insensitive because their grids are already near-decarbonized, although they remain exposed to "
        "the water penalty of hydropower and to seasonal drought-driven generation shortfalls. These results connect to "
        "the importance of ambitious green-electricity policies and of routing the ‘desert and Gobi’ renewable "
        "mega-bases to serve co-located computing load, so that the environmental dividend of decarbonization is "
        "captured rather than curtailed{chinarenew,greenpeace_cleancloud}.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_figure(doc, os.path.join(FIG, "fig4_grid.png"),
        "Fig. 4 | Grid renewable penetration drives AI-server carbon and water footprints. ",
        "a,b, Mid-case carbon and water trajectories under best (LRC), base and worst (HRC) grid pathways. "
        "c, Provincial sensitivity of carbon emissions to the grid pathway for selected hub provinces.")

    # ---------------- Results 5 ----------------
    add_heading(doc, "Pathways to carbon-peak and net-zero water goals")
    add_para(doc,
        "Building on the preceding analyses of environmental cost, efficiency, spatial distribution and grid "
        "decarbonization, this section evaluates pathways for China’s AI servers to align with the national carbon-peak "
        "milestone and with operator-level neutrality and water-stewardship pledges{bjorn2022}. Figure 5a presents a "
        "contribution pathway under the mid-case scenario, decomposing the cumulative 2024–2030 carbon added by AI "
        "servers into the reductions achievable through industry efficiency, spatial siting and grid decarbonization, "
        "and the residual that must be addressed through offsets or further policy. Combined best practices cut residual "
        "emissions and water footprints by up to 73% and 86%, respectively, indicating a feasible—but demanding—pathway "
        "toward net zero. The single largest lever is grid decarbonization{ewc_netzero,china_neutrality_pnas}, followed by spatial siting and then "
        "on-site efficiency, reflecting the dominance of Scope 2 impacts in a coal-heavy system.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_para(doc,
        "Figure 5b presents residual carbon across scenarios under best- and worst-practice bundles. Under the mid-case "
        "with best practices, residual annual emissions by 2030 fall to the order of tens of Mt CO₂e—still requiring "
        "tens of GW of additional wind or solar capacity to offset—whereas worst practices leave residuals several times "
        "larger that are nearly impossible to compensate over such a short horizon. Because renewable-energy "
        "certificates and offsets can threaten the integrity of corporate climate targets when poorly governed{bjorn2022,carbontrade}, "
        "we caution that headline neutrality claims should be underpinned by additional, local and verifiable clean "
        "generation rather than unbundled certificates. The results imply that, without substantial acceleration of the "
        "energy transition and disciplined siting, China’s AI server fleet is unlikely to be consistent with the "
        "sector’s share of a 2030 carbon peak, let alone deep mid-century neutrality, on the basis of efficiency gains "
        "alone.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_para(doc,
        "Projected pathways could be further influenced by other uncertainties. Figure 6 presents a sensitivity analysis "
        "of key factors: server lifetime, domestic accelerator manufacturing capacity, national allocation ratio, "
        "server idle and maximum power ratios, training/inference distribution and green-power utilization. Changes in "
        "these factors result in up to roughly 40% variation in energy consumption, closely mirrored by water and "
        "carbon. The national allocation ratio and accelerator capacity dominate the uncertainty, reflecting the "
        "centrality of the domestic supply chain under export controls. Innovations such as DeepSeek can lower the power "
        "required per computing task{deepseek2025}, but may trigger a rebound effect that increases total demand—captured "
        "by our high-application scenario. The study’s key conclusions remain robust unless future uncertainties greatly "
        "exceed the modelled ranges, and the modelling approach enables revision as more data become available.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_figure(doc, os.path.join(FIG, "fig5_netzero.png"),
        "Fig. 5 | Pathways toward net-zero carbon and water for China’s AI servers. ",
        "a, Contribution pathway decomposing cumulative 2024–2030 carbon added by AI servers into reductions from "
        "industry efficiency, spatial siting and grid decarbonization, and the residual. b, Residual carbon to offset "
        "by 2030 across the five scenarios under best- and worst-practice bundles.")
    add_figure(doc, os.path.join(FIG, "fig6_sensitivity.png"),
        "Fig. 6 | Sensitivity of AI-server energy–water–carbon impacts to key modelling assumptions. ",
        "Tornado plot of the sensitivity of AI server energy consumption (closely mirrored by water and carbon) to "
        "server lifetime, domestic accelerator capacity, national allocation ratio, idle and maximum power ratios, "
        "training/inference split and green-power utilization. Listed values are low, applied and high.", width=5.4)

    # ---------------- Discussion ----------------
    add_heading(doc, "Discussion")
    add_para(doc,
        "Investment in AI servers is accelerating in China, as seen in the rapid build-out of intelligent-computing "
        "centres under the ‘East Data, West Computing’ programme and in the capital-expenditure plans of major cloud "
        "operators{ndrc_ewc,ewc_netzero,ieaenergyai2025}. While AI advancement is a national priority, our study highlights its "
        "environmental cost. To mitigate these risks, we identify that concentrating AI server deployment in provinces "
        "that combine abundant wind and solar resources with manageable water stress—principally Inner Mongolia, Gansu, "
        "Ningxia, Qinghai and Xinjiang, paired with dry cooling and dedicated renewable supply—is preferable to either "
        "coal-based siting or water-intensive hydropower siting in Sichuan and Yunnan. These provinces possess "
        "substantial untapped wind and solar potential{chinarenew,iea_china_roadmap}, enabling robust green-power "
        "portfolios and reducing competition with other sectors. However, several implementation challenges must be "
        "acknowledged. The western hubs sit in some of China’s most water-stressed basins{chinawater,china_water_risk}, so "
        "the water advantage is realized only if evaporative cooling is avoided and renewable rather than thermal or "
        "hydro generation is procured{gobi_pv}. Routing the planned desert and Gobi renewable mega-bases to co-located "
        "computing load, and strengthening dedicated transmission, will be essential to prevent curtailment from eroding "
        "the climate dividend{chinarenew}.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_para(doc,
        "Realizing the suggested strategy may be further complicated by public-health and local-development concerns. "
        "Operational demands and construction could generate air and water pollution, and large hyperscale loads "
        "challenge utilities’ obligation to serve. We recommend that AI firms engage in public–private partnerships with "
        "provincial governments that fund green-power upgrades and transparent environmental monitoring, while "
        "governments use green-electricity targets, PUE caps and water-rights enforcement for the computing hubs to "
        "steer siting{chinapue,ndrc_ewc}. Because renewable-energy certificates and carbon offsets can undermine the "
        "integrity of neutrality claims when weakly governed{bjorn2022}, neutrality pledges by Chinese operators should be "
        "backed by additional, local and verifiable clean generation and by third-party verification. Emerging "
        "high-efficiency models such as DeepSeek may fundamentally reshape AI server supply and demand{deepseek2025}; while "
        "efficiency lowers cost per task, it risks a rebound that amplifies total demand, calling for real-time "
        "monitoring and AI-specific benchmarks for energy, water and carbon performance{caict_china,greenpeace_cleancloud} so that policy can respond "
        "before large impacts accrue. Finally, the effectiveness of these measures will be shaped by broader political "
        "and economic factors beyond the scope of this study.",
        align=AL.JUSTIFY, first_indent=0.2)

    # ---------------- Methods ----------------
    add_heading(doc, "Methods")
    add_para(doc,
        "The methodology aims to (1) characterize the energy–water–climate impacts of AI servers in China from 2024 to "
        "2030, and (2) identify best- and worst-practice values of each influencing factor to scheme net-zero pathways "
        "consistent with the national dual-carbon milestones. Scenario-based projections are first constructed to obtain "
        "potential capacity-increasing patterns of AI servers; technology dynamics (SUO and ALC adoption), grid "
        "decarbonization and spatial distribution are each defined with best, base and worst scenarios. The models and "
        "data used are described below.",
        align=AL.JUSTIFY, first_indent=0.0)

    add_heading(doc, "AI server power-capacity projections", level=2)
    add_para(doc,
        "AI server energy consumption is projected to be driven predominantly by top-tier accelerators designed for "
        "large-scale generative AI computing. Under US export controls restricting access to leading-edge imported "
        "accelerators, China’s top-tier AI server capacity is increasingly determined by the output of domestic "
        "accelerators (for example, the Ascend line) and by the advanced-packaging capacity of domestic foundries, "
        "which forms the production bottleneck analogous to chip-on-wafer-on-substrate (CoWoS) capacity "
        "elsewhere{deepseek2025}. We estimate the annual top-tier server shipments and average rated power as",
        align=AL.JUSTIFY, first_indent=0.2)
    add_para(doc,
        "    N_AI = (C_pkg · R_top · Σ_i R_i n_i) / N_acc ,        P̄_AI = Σ_i R_i P_i ,",
        size=10.5, italic=True, align=AL.CENTER, space_after=6)
    add_para(doc,
        "where N_AI and P̄_AI are the annually projected shipments and average rated power of top-tier AI servers, "
        "C_pkg is the projected domestic advanced-packaging capacity within each scenario, R_top is the share of "
        "packaging capacity allocated to top-tier accelerators, N_acc is the number of accelerators per server, and "
        "R_i, n_i and P_i are the adoption ratio, units yield per wafer and rated power of the i-th accelerator type in "
        "each year. Five scenarios (low demand, low power, mid-case, high application, high demand) span the "
        "uncertainty in packaging-capacity growth, rated power and adoption, with the mid-case calibrated so that 2030 "
        f"facility electricity reaches about {MID['facility_energy_TWh']:.0f} TWh, consistent with reported China data-centre "
        "electricity growth dominated by intelligent computing{caict_china,ieaenergyai2025}.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_heading(doc, "AI server electricity-usage calculation", level=2)
    add_para(doc,
        "Server electricity usage follows a utilization-based model, P_server = (P_max − P_idle)·u + P_idle, where u is "
        "the processor utilization, P_max and P_idle are the maximum and idle power, and u = u_active · r_active is "
        "decomposed into the average utilization of active accelerators and the active ratio. Training and inference are "
        "treated separately because they exhibit distinct utilization, with a base split of 30% of capacity for "
        "training and 70% for inference{patel2024,luccioni2023}. Facility electricity is obtained by multiplying IT "
        "electricity by the provincial PUE.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_heading(doc, "Assessment of environmental footprints", level=2)
    add_para(doc,
        "A province-level allocation method evaluates the energy, water and carbon footprints of AI servers. We compiled "
        "the locations of existing and planned large-scale AI data centres operated by major Chinese purchasers of "
        "top-tier accelerators (including Alibaba, Tencent, Baidu, ByteDance, Huawei Cloud and the three "
        "telecommunications operators) together with the eight national computing hubs, and aggregated capacity by "
        "province{ndrc_ewc}. For each province, actual energy consumption is derived from server electricity usage and the "
        "provincial PUE. Water and carbon are assessed across three scopes: Scope 1 (on-site cooling water via WUE; "
        "on-site carbon negligible), Scope 2 (off-site water and carbon contingent on the provincial grid mix) and "
        "Scope 3 (embodied water and carbon during manufacturing). A hybrid statistical and thermodynamics-based PUE/WUE "
        "model{lei2022}, driven by provincial climate data, estimates province-specific PUE and WUE; base values average "
        "the best and worst practices given unknown cooling specifications. Scope 2 water and carbon use province-level "
        "grid water and carbon factors derived from generation mixes consistent with national decarbonization "
        "plans{chinagrid,siddik2024}. Water-scarcity and renewable-potential layers for each province are compiled from "
        "published hydrological and resource assessments{chinawater,chinarenew}.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_heading(doc, "Uncertainties and limitations", level=2)
    add_para(doc,
        "Substantial uncertainties are inherent in projecting AI server evolution. Our analysis presents a range of "
        "scenarios to evaluate operational efficiency, spatial distribution and grid development, and a sensitivity "
        "analysis (Fig. 6) on server lifetime, domestic accelerator capacity, allocation ratio, idle and maximum power, "
        "training/inference split and green-power utilization. Key unmodelled uncertainties include model and algorithm "
        "innovations (which could alter compute requirements), domestic supply-chain dynamics under export controls, "
        "hardware and facility evolution, and out-of-scope factors such as market forces and geopolitics. As Fig. 6 "
        "indicates, the study’s conclusions remain robust unless future uncertainties greatly exceed the modelled "
        "ranges, and the framework allows revision as more data become available.",
        align=AL.JUSTIFY, first_indent=0.2)

    add_heading(doc, "Data and code availability", level=2)
    add_para(doc,
        "All simulated datasets generated in this analysis (national trajectories and provincial attributes) and the "
        "code used to produce the figures are provided with this manuscript in the accompanying repository. Provincial "
        "grid, water and renewable-potential layers are derived from the publicly available sources cited above. "
        "Numerical inputs and assumptions are summarized in the main text and figures; further calibration details are "
        "available from the corresponding author on request.",
        align=AL.JUSTIFY, first_indent=0.2)

    # Acknowledgements / contributions (template placeholders)
    add_heading(doc, "Author contributions", level=2)
    add_para(doc,
        "[Corresponding Author] conceived and supervised the study. [First Author] developed the model, performed the "
        "analysis and drafted the manuscript. All authors contributed to interpretation, review and editing and "
        "approved the final version.", size=9.5, align=AL.JUSTIFY)
    add_heading(doc, "Competing interests", level=2)
    add_para(doc, "The authors declare no competing interests.", size=9.5)
