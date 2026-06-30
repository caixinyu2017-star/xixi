#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Assemble the Word manuscript:
  "Environmental impact and net-zero pathways for sustainable artificial
   intelligence servers in China"

Nature Sustainability "Analysis" style. Citations are numbered by order of first
appearance (Nature convention). Reference entries in REF_DB are REAL works:
those carried over from the two source Nature Sustainability papers (Xiao 2025;
Lee 2026 reference lists) plus China-specific references verified via web search.
"""
import os, re, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, "..")
DATA = os.path.join(ROOT, "data")
FIG = os.path.join(ROOT, "figures")
OUT = os.path.join(ROOT, "China_AI_Server_Sustainability.docx")

summary = json.load(open(os.path.join(DATA, "summary.json")))

# ===========================================================================
# REFERENCE DATABASE  (key -> full reference string).  All entries are real.
# Numbering is assigned automatically by order of first citation in the text.
# ===========================================================================
REF_DB = {
    # --- carried over from Xiao et al. 2025 / Lee et al. 2026 (verified real) ---
    "guan2024": "Guan, L. Reaching carbon neutrality requires energy-efficient training of AI. Nature 626, 33 (2024).",
    "devries2023": "de Vries, A. The growing energy footprint of artificial intelligence. Joule 7, 2191–2194 (2023).",
    "crawford2024": "Crawford, K. Generative AI’s environmental costs are soaring—and mostly secret. Nature 626, 693 (2024).",
    "kaack2022": "Kaack, L. H. et al. Aligning artificial intelligence with climate change mitigation. Nat. Clim. Change 12, 518–527 (2022).",
    "vinuesa2020": "Vinuesa, R. et al. The role of artificial intelligence in achieving the Sustainable Development Goals. Nat. Commun. 11, 233 (2020).",
    "mytton2021": "Mytton, D. Data centre water consumption. npj Clean Water 4, 11 (2021).",
    "siddik2021": "Siddik, M. A. B., Shehabi, A. & Marston, L. The environmental footprint of data centers in the United States. Environ. Res. Lett. 16, 064017 (2021).",
    "lei2022": "Lei, N. & Masanet, E. Climate- and technology-specific PUE and WUE estimations for US data centers using a hybrid statistical and thermodynamics-based approach. Resour. Conserv. Recycl. 182, 106323 (2022).",
    "masanet2020": "Masanet, E., Shehabi, A., Lei, N., Smith, S. & Koomey, J. Recalibrating global data center energy-use estimates. Science 367, 984–986 (2020).",
    "wu2022": "Wu, C.-J. et al. Sustainable AI: environmental implications, challenges and opportunities. Proc. Mach. Learn. Syst. 4, 795–813 (2022).",
    "luccioni2023": "Luccioni, A. S., Viguier, S. & Ligozat, A.-L. Estimating the carbon footprint of BLOOM, a 176B parameter language model. J. Mach. Learn. Res. 24, 11990–12004 (2023).",
    "patel2024": "Patel, P. et al. Characterizing power management opportunities for LLMs in the cloud. In Proc. 29th ACM International Conference on Architectural Support for Programming Languages and Operating Systems Vol. 3, 207–222 (ACM, 2024).",
    "bjorn2022": "Bjørn, A., Lloyd, S. M., Brander, M. & Matthews, H. D. Renewable energy certificates threaten the integrity of corporate science-based targets. Nat. Clim. Change 12, 539–546 (2022).",
    "shehabi2024": "Shehabi, A. et al. 2024 United States Data Center Energy Usage Report (Lawrence Berkeley National Laboratory, 2024).",
    "siddik2024": "Siddik, M. A. B., Shehabi, A., Rao, P. & Marston, L. T. Spatially and temporally detailed water and carbon footprints of US electricity generation and use. Water Resour. Res. 60, e2024WR038350 (2024).",
    "berthelot2024": "Berthelot, A., Caron, E., Jay, M. & Lefèvre, L. Estimating the environmental impact of generative-AI services using an LCA-based methodology. Procedia CIRP 122, 707–712 (2024).",
    "dodge2022": "Dodge, J. et al. Measuring the carbon intensity of AI in cloud instances. In Proc. 2022 ACM Conference on Fairness, Accountability, and Transparency 1877–1894 (ACM, 2022).",
    "malmodin2024": "Malmodin, J., Lövehagen, N., Bergmark, P. & Lundén, D. ICT sector electricity consumption and greenhouse gas emissions—2020 outcome. Telecommun. Policy 48, 102701 (2024).",
    "ieaenergyai2025": "International Energy Agency. Energy and AI (IEA, 2025); https://www.iea.org/reports/energy-and-ai.",
    "iea_elec2024": "International Energy Agency. Electricity 2024: Analysis and Forecast to 2026 (IEA, 2024); https://www.iea.org/reports/electricity-2024.",
    "xiao2025": "Xiao, T., Fuso Nerini, F., Matthews, H. D., Tavoni, M. & You, F. Environmental impact and net-zero pathways for sustainable artificial intelligence servers in the USA. Nat. Sustain. 8, 1541–1553 (2025).",
    # --- China-specific references (populated/verified by the research workflow) ---
    # Filled in by merge_china_refs(); see __main__.
}

# China references confirmed real and merged after web verification.
CHINA_REFS = {}

# ===========================================================================
# Citation engine: assign numbers by first appearance across the document body.
# ===========================================================================
class Cites:
    def __init__(self):
        self.order = []
        self.index = {}
    def num(self, key):
        if key not in self.index:
            self.order.append(key)
            self.index[key] = len(self.order)
        return self.index[key]
CITE = Cites()

def _set_font(run, name="Times New Roman", size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(a), name)

CITE_RE = re.compile(r"\{([a-zA-Z0-9_,]+)\}")

def add_para(doc, text, size=10.5, bold=False, italic=False, align=None,
             space_after=6, space_before=0, first_indent=None, line=1.15, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if first_indent is not None:
        pf.first_line_indent = Inches(first_indent)
    pos = 0
    for m in CITE_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            _set_font(r, size=size, bold=bold, italic=italic)
        keys = m.group(1).split(",")
        nums = sorted({CITE.num(k) for k in keys})
        # compress consecutive runs e.g. 4,5,6 -> 4–6
        groups = []
        for n in nums:
            if groups and n == groups[-1][-1] + 1:
                groups[-1].append(n)
            else:
                groups.append([n])
        sup = ",".join(f"{g[0]}–{g[-1]}" if len(g) > 1 else f"{g[0]}" for g in groups)
        r = p.add_run(sup)
        _set_font(r, size=size - 2.5)
        r.font.superscript = True
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        _set_font(r, size=size, bold=bold, italic=italic)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_font(r, size=13 if level == 1 else 11, bold=True,
              color=RGBColor(0x1a, 0x2a, 0x44))
    return p

def add_figure(doc, path, caption_bold, caption_rest, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(10)
    r1 = cap.add_run(caption_bold)
    _set_font(r1, size=8.5, bold=True)
    r2 = cap.add_run(caption_rest)
    _set_font(r2, size=8.5)

# ===========================================================================
# Build
# ===========================================================================
def build(prose):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"; style.font.size = Pt(10.5)

    prose(doc)

    # References
    add_heading(doc, "References", 1)
    allrefs = dict(REF_DB); allrefs.update(CHINA_REFS)
    for i, key in enumerate(CITE.order, 1):
        ref = allrefs.get(key)
        if ref is None:
            ref = f"[MISSING REFERENCE for key '{key}']"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        rn = p.add_run(f"{i}. ")
        _set_font(rn, size=9, bold=True)
        rr = p.add_run(ref)
        _set_font(rr, size=9)
    doc.save(OUT)
    # report unresolved
    missing = [k for k in CITE.order if k not in allrefs]
    print(f"Saved {OUT}")
    print(f"Total citations: {len(CITE.order)}")
    if missing:
        print("MISSING REFS:", missing)
    return missing


if __name__ == "__main__":
    raise SystemExit(
        "Run 'python3 make.py' instead — it imports build_doc as a single shared module "
        "so the citation registry is not split by the __main__ double-import.")
