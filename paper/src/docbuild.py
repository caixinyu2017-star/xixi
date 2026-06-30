"""Infrastructure for building the Biomimetics-style Word paper with python-docx.
Provides: page/style setup, citation manager (auto-number in order), numbered
equation tables (native OMML), figure insertion with captions, data-table builder,
and MDPI reference list output."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import mathdocx as M
import refs_db

FIG="/tmp/claude-0/-home-user-xixi/013086be-932a-56cd-b187-9ec3125f6bc6/scratchpad/figs"

BODY_FONT="Palatino Linotype"
BODY_SIZE=10

class Paper:
    def __init__(self):
        self.doc=Document()
        self.refs=[]            # ordered list of tags as cited
        self.ref_index={}       # tag -> number
        self.eqno=0
        self.figno=0
        self.tblno=0
        self._setup()

    # ---------- setup ----------
    def _setup(self):
        d=self.doc
        sec=d.sections[0]
        sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)   # A4
        sec.top_margin=Cm(1.9); sec.bottom_margin=Cm(2.4)
        sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)
        st=d.styles["Normal"]
        st.font.name=BODY_FONT; st.font.size=Pt(BODY_SIZE)
        st.paragraph_format.space_after=Pt(0); st.paragraph_format.space_before=Pt(0)
        st.paragraph_format.line_spacing=1.15
        # east-asian font binding
        rpr=st.element.get_or_add_rPr(); rfonts=rpr.get_or_add_rFonts()
        rfonts.set(qn('w:eastAsia'),BODY_FONT)
        self._add_header_footer(sec)

    def _page_field(self, paragraph):
        for t,val in (("begin",None),("instrText","PAGE"),("end",None)):
            r=paragraph.add_run(); r.font.size=Pt(8)
            fld=OxmlElement('w:fldChar')
            if t=="instrText":
                it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=" PAGE "
                r._r.append(it); continue
            fld.set(qn('w:fldCharType'),t); r._r.append(fld)

    def _add_header_footer(self, sec):
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
        hdr=sec.header; hdr.is_linked_to_previous=False
        hp=hdr.paragraphs[0]; hp.alignment=_AL.RIGHT; hp.paragraph_format.space_after=Pt(0)
        r=hp.add_run("Biomimetics 2026, 10, x FOR PEER REVIEW"); r.italic=True; r.font.size=Pt(8)
        ftr=sec.footer; ftr.is_linked_to_previous=False
        fp=ftr.paragraphs[0]; fp.alignment=_AL.CENTER
        self._page_field(fp)

    # ---------- low level ----------
    def _p(self, text="", align=AL.JUSTIFY, size=BODY_SIZE, bold=False, italic=False,
           before=0, after=6, first_indent=None, line=1.15, font=None):
        p=self.doc.add_paragraph()
        pf=p.paragraph_format; pf.alignment=align
        pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line
        if first_indent is not None: pf.first_line_indent=Cm(first_indent)
        if text:
            r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
            if font: r.font.name=font
        return p

    def heading(self, text, level=1):
        sizes={1:11,2:10.5,3:10}
        p=self._p(text, align=AL.LEFT, size=sizes.get(level,10), bold=True,
                  before=10 if level==1 else 8, after=4)
        return p

    def body(self, segments, first_indent=0.0, after=6, align=AL.JUSTIFY):
        """segments: str OR list of ('text',s)/('math',latex)/('cite',[tags]).
        Returns paragraph. Citations auto-numbered."""
        p=self.doc.add_paragraph()
        pf=p.paragraph_format; pf.alignment=align; pf.space_after=Pt(after); pf.line_spacing=1.15
        if first_indent: pf.first_line_indent=Cm(first_indent)
        if isinstance(segments,str): segments=[('text',segments)]
        for kind,val in segments:
            if kind=='text':
                r=p.add_run(val); r.font.size=Pt(BODY_SIZE)
            elif kind=='math':
                p._p.append(M.omath_element(val))
            elif kind=='cite':
                r=p.add_run(self.cite(val)); r.font.size=Pt(BODY_SIZE)
            elif kind=='it':
                r=p.add_run(val); r.italic=True; r.font.size=Pt(BODY_SIZE)
            elif kind=='b':
                r=p.add_run(val); r.bold=True; r.font.size=Pt(BODY_SIZE)
        return p

    # ---------- citations ----------
    def cite(self, tags):
        if isinstance(tags,str): tags=[tags]
        nums=[]
        for t in tags:
            if t not in self.ref_index:
                refs_db.citation(t)   # validate
                self.refs.append(t); self.ref_index[t]=len(self.refs)
            nums.append(self.ref_index[t])
        # compress consecutive runs like [1,2,3]->[1-3]
        nums_sorted=nums
        return "["+",".join(str(n) for n in nums)+"]"

    def cnum(self, tag):
        if tag not in self.ref_index:
            refs_db.citation(tag); self.refs.append(tag); self.ref_index[tag]=len(self.refs)
        return self.ref_index[tag]

    # ---------- equations (numbered, in a borderless 2-col table) ----------
    def equation(self, latex):
        self.eqno+=1
        tbl=self.doc.add_table(rows=1,cols=2)
        tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit=False
        tbl.columns[0].width=Cm(15.0); tbl.columns[1].width=Cm(2.0)
        cell=tbl.cell(0,0); cell.width=Cm(15.0)
        pe=cell.paragraphs[0]; pe.alignment=AL.CENTER
        pe._p.append(M.omath_element(latex))
        cn=tbl.cell(0,1); cn.width=Cm(2.0)
        pn=cn.paragraphs[0]; pn.alignment=AL.RIGHT
        r=pn.add_run(f"({self.eqno})"); r.font.size=Pt(BODY_SIZE)
        # spacing around equation
        self._set_table_spacing(tbl)
        return self.eqno

    def _set_table_spacing(self,tbl):
        pass

    # ---------- figures ----------
    def figure(self, fname, caption, width_cm=14.0):
        self.figno+=1
        p=self.doc.add_paragraph(); p.alignment=AL.CENTER
        p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
        path=os.path.join(FIG,fname)
        run=p.add_run(); run.add_picture(path, width=Cm(width_cm))
        cap=self.doc.add_paragraph(); cap.alignment=AL.CENTER
        cap.paragraph_format.space_after=Pt(8)
        r=cap.add_run(f"Figure {self.figno}. "); r.bold=True; r.font.size=Pt(9)
        r2=cap.add_run(caption); r2.font.size=Pt(9)
        return self.figno

    def figure_caption_only(self, caption):
        cap=self.doc.add_paragraph(); cap.alignment=AL.CENTER
        cap.paragraph_format.space_after=Pt(8)
        r=cap.add_run(f"Figure {self.figno}. "); r.bold=True; r.font.size=Pt(9)
        r2=cap.add_run(caption); r2.font.size=Pt(9)

    # ---------- tables ----------
    def table_caption(self, caption, label=None):
        """label=None -> auto-number main tables; label='A1' -> appendix label (no counter)."""
        if label is None:
            self.tblno+=1; lab=str(self.tblno)
        else:
            lab=label
        cap=self.doc.add_paragraph(); cap.alignment=AL.CENTER
        cap.paragraph_format.space_before=Pt(8); cap.paragraph_format.space_after=Pt(2)
        r=cap.add_run(f"Table {lab}. "); r.bold=True; r.font.size=Pt(9)
        r2=cap.add_run(caption); r2.font.size=Pt(9)
        return lab

    def make_table(self, headers, rows, font_size=8, header_bold=True,
                   col_align=None, first_col_left=True, style_grid=True):
        ncol=len(headers)
        t=self.doc.add_table(rows=1,cols=ncol)
        t.alignment=WD_TABLE_ALIGNMENT.CENTER
        if style_grid:
            self._three_line(t)
        hdr=t.rows[0].cells
        for j,h in enumerate(headers):
            pp=hdr[j].paragraphs[0]; pp.alignment=AL.CENTER
            r=pp.add_run(str(h)); r.bold=header_bold; r.font.size=Pt(font_size)
        for row in rows:
            cells=t.add_row().cells
            for j,v in enumerate(row):
                pp=cells[j].paragraphs[0]
                al=AL.LEFT if (j==0 and first_col_left) else AL.CENTER
                if col_align and col_align[j] is not None: al=col_align[j]
                pp.alignment=al
                r=pp.add_run(str(v)); r.font.size=Pt(font_size)
        # tighten cell margins
        return t

    def _three_line(self, table):
        """MDPI 3-line table: top & bottom thick, header bottom thin, no verticals."""
        tbl=table._tbl
        borders=OxmlElement('w:tblBorders')
        def bd(tag,sz,val="single"):
            e=OxmlElement('w:'+tag); e.set(qn('w:val'),val); e.set(qn('w:sz'),str(sz))
            e.set(qn('w:space'),'0'); e.set(qn('w:color'),'000000'); return e
        borders.append(bd('top',12)); borders.append(bd('bottom',12))
        borders.append(bd('insideH',0,'none')); borders.append(bd('insideV',0,'none'))
        borders.append(bd('left',0,'none')); borders.append(bd('right',0,'none'))
        tblPr=tbl.tblPr; tblPr.append(borders)
        # header row bottom border
        try:
            row0=table.rows[0]
            for c in row0.cells:
                tcPr=c._tc.get_or_add_tcPr(); tcb=OxmlElement('w:tcBorders')
                b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'8')
                b.set(qn('w:space'),'0'); b.set(qn('w:color'),'000000'); tcb.append(b)
                tcPr.append(tcb)
        except Exception: pass

    # ---------- references ----------
    def references(self):
        import re
        self.heading("References",1)
        for i,tag in enumerate(self.refs,1):
            c=refs_db.citation(tag)
            p=self.doc.add_paragraph()
            pf=p.paragraph_format; pf.left_indent=Cm(0.6); pf.first_line_indent=Cm(-0.6)
            pf.space_after=Pt(2); pf.line_spacing=1.0
            r=p.add_run(f"{i}. "); r.font.size=Pt(9)
            self._emit_reference(p, c)
        return len(self.refs)

    def _emit_reference(self, p, c):
        """MDPI typography: bold the publication year, italicize the volume that follows it."""
        import re
        # locate "YEAR, VOLUME," or "YEAR," (year = 19xx/20xx) — the publication-year token
        m=re.search(r'\b(19|20)\d{2}\b', c)
        def run(txt, bold=False, italic=False):
            if not txt: return
            rr=p.add_run(txt); rr.font.size=Pt(9); rr.bold=bold; rr.italic=italic
        if not m:
            run(c); return
        ys, ye = m.start(), m.end()
        run(c[:ys])                       # everything before the year (authors, title, journal)
        run(c[ys:ye], bold=True)          # bold year
        rest=c[ye:]
        # volume: ", <number/number-range>" immediately after year
        vm=re.match(r'(,\s*)(\d+[A-Za-z]?)(\s*,)', rest)
        if vm:
            run(vm.group(1))
            run(vm.group(2), italic=True)
            run(vm.group(3))
            run(rest[vm.end():])
        else:
            run(rest)

    def page_break(self):
        from docx.enum.text import WD_BREAK
        p=self.doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

    def start_section(self, landscape=False):
        from docx.enum.section import WD_ORIENT, WD_SECTION
        sec=self.doc.add_section(WD_SECTION.NEW_PAGE)
        if landscape:
            sec.orientation=WD_ORIENT.LANDSCAPE
            sec.page_width=Cm(29.7); sec.page_height=Cm(21.0)
            sec.top_margin=Cm(1.6); sec.bottom_margin=Cm(1.6)
            sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
        else:
            sec.orientation=WD_ORIENT.PORTRAIT
            sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
            sec.top_margin=Cm(1.9); sec.bottom_margin=Cm(2.4)
            sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)
        self._add_header_footer(sec)
        return sec

    def fit_table(self, table, total_cm, first_widths):
        """Force fixed layout, set total width and per-column widths so wide tables fit the page."""
        tbl=table._tbl; tblPr=tbl.tblPr
        lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
        w=OxmlElement('w:tblW'); w.set(qn('w:w'),str(int(total_cm*567))); w.set(qn('w:type'),'dxa'); tblPr.append(w)
        ncol=len(table.columns); nfix=len(first_widths)
        rest=(total_cm-sum(first_widths))/max(1,(ncol-nfix))
        widths=[first_widths[j] if j<nfix else rest for j in range(ncol)]
        table.autofit=False
        for row in table.rows:
            for j,cell in enumerate(row.cells):
                cell.width=Cm(widths[j])

    def save(self, path):
        self.doc.save(path)
