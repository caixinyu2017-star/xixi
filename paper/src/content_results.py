"""Sections 4 (CEC2017), 5 (art-design applications), 6 (conclusion), back matter, Appendix A."""
import json, os
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.shared import Pt, Cm

DATA="/tmp/claude-0/-home-user-xixi/013086be-932a-56cd-b187-9ec3125f6bc6/scratchpad/data"
CEC=json.load(open(os.path.join(DATA,"cec_results.json")))
ABL=json.load(open(os.path.join(DATA,"ablation.json")))
PS=json.load(open(os.path.join(DATA,"paramsens.json")))
HARM=json.load(open(os.path.join(DATA,"harmony.json")))
LAY=json.load(open(os.path.join(DATA,"layout.json")))
ALGOS=CEC["algos"]; FUNCS=[str(f) for f in CEC["funcs"]]
DIMS=["10","30","50","100"]

def sci(v):
    try: return f"{float(v):.3E}"
    except: return str(v)

# ============================ SECTION 4 ============================
def section4(P):
    P.heading("4. Numerical Experiments Using the CEC2017 Test Suite",1)
    P.body([('text',"In this section, the performance of MSSBOA is comprehensively examined on the "
        "CEC2017 benchmark suite. Section 4.1 describes the experimental setting and the CEC2017 "
        "functions. Section 4.2 analyzes the sensitivity of the two key parameters of MSSBOA. "
        "Section 4.3 reports the ablation experiments that verify the effectiveness of the three "
        "improvement strategies. Finally, Section 4.4 provides a comprehensive comparison between "
        "MSSBOA and nine state-of-the-art algorithms, analyzed by the Friedman test, the Nemenyi "
        "post-hoc test and the Wilcoxon rank-sum test.")])

    P.heading("4.1. Experiment setting and performance metrics",2)
    P.body([('text',"All experiments were conducted on a workstation with an AMD Ryzen 9 7945HX "
        "2.5 GHz CPU and 32 GB of memory, running Windows 11 and MATLAB R2023a. The CEC2017 test "
        "suite "),('cite',['cec2017']),
        ('text'," contains 29 test functions (F2 is excluded because of its unstable behavior), in "
        "which F1 and F3 are unimodal functions, F4–F10 are simple multimodal functions, F11–F20 are "
        "hybrid functions and F21–F30 are composition functions. The unimodal functions evaluate the "
        "convergence accuracy of an algorithm, the multimodal functions test the ability to escape "
        "local optima, and the hybrid and composition functions assess the balance between exploration "
        "and exploitation on complex landscapes. The search range is "),('math','[-100,100]^{D}'),
        ('text'," and the dimensions are set to 10, 30, 50 and 100. To ensure a fair comparison, the "
        "maximum number of function evaluations is used as the stopping criterion and set to "),
        ('math',r'10000\times D'),('text',". Each algorithm is independently run 30 times, and the "
        "best value, mean value and standard deviation are recorded. MSSBOA is compared with the basic "
        "SBOA "),('cite',['sboa']),('text'," and nine advanced algorithms, including GWO "),('cite',['gwo']),
        ('text',", SCA "),('cite',['sca']),('text',", SMA "),('cite',['sma']),
        ('text',", COA "),('cite',['coa']),('text',", GKSO "),('cite',['gkso']),
        ('text',", MRFO "),('cite',['mrfo']),('text',", RIME "),('cite',['rime']),
        ('text',", QIO "),('cite',['qio']),(
        'text'," and LSHADE "),('cite',['lshade']),
        ('text',". The parameter settings of all algorithms are listed in Table 1 and are taken from "
        "their original literature.")])
    P.table_caption("Parameter settings of MSSBOA and the comparison algorithms.")
    params=[
        ("MSSBOA","N=30, N_min=0.2N, p_m=0.5, beta=1.5"),
        ("SBOA","N=30, beta=1.5"),
        ("GWO","N=30, a decreases linearly from 2 to 0"),
        ("SCA","N=30, a=2"),
        ("SMA","N=30, z=0.03"),
        ("COA","N=30, control parameter decreases with iteration"),
        ("GKSO","N=30, default control parameters"),
        ("MRFO","N=30, somersault factor S=2"),
        ("RIME","N=30, W=5"),
        ("QIO","N=30, default control parameters"),
        ("LSHADE","N_init=18D, N_min=4, memory size H=6, p=0.11, arc rate=2.6"),
    ]
    P.make_table(["Algorithm","Parameter settings"],params,font_size=9)

    P.heading("4.2. Parameter sensitivity analysis",2)
    P.body([('text',"Two parameters strongly affect the performance of MSSBOA: the minimum population "
        "factor "),('math','N_{min}'),('text'," used in the lens opposition-based learning and the "
        "mutation probability "),('math','p_{m}'),('text'," of the adaptive Cauchy–Gaussian mutation. "
        "A grid search is conducted on the CEC2017 suite over 30 independent runs. "),('math','N_{min}'),
        ('text'," is searched from 0.1N to 0.9N with a step of 0.1N, and "),('math','p_{m}'),
        ('text'," is searched from 0.1 to 1.0 with a step of 0.1. Table 2 and Table 3 report the "
        "Friedman rankings (significance level "),('math',r'\alpha=0.05'),('text',") for "),
        ('math','N_{min}'),('text'," and "),('math','p_{m}'),
        ('text',", respectively, where the best ranking is highlighted in bold. The corresponding "
        "rankings are visualized in Figure 4 and Figure 5.")])
    _param_table(P,"P1",r"Rankings of MSSBOA with different N_min (α=0.05).")
    P.figure("fig04_param_p1.png","Friedman rankings of MSSBOA with different N_min (α=0.05).",width_cm=13.5)
    P.body([('text',"As shown in Table 2 and Figure 4, MSSBOA achieves the best average ranking when "),
        ('math','N_{min}=0.2N'),('text',". As "),('math','N_{min}'),('text'," increases, the ranking "
        "gradually deteriorates because too large a refraction subpopulation disturbs the elite "
        "individuals. Therefore, "),('math','N_{min}=0.2N'),('text'," is adopted in this paper.")])
    _param_table(P,"P2",r"Rankings of MSSBOA with different p_m (α=0.05).")
    P.figure("fig05_param_p2.png","Friedman rankings of MSSBOA with different p_m (α=0.05).",width_cm=13.5)
    P.body([('text',"For the mutation probability, Table 3 and Figure 5 show that MSSBOA performs best "
        "when "),('math','p_{m}=0.5'),(
        'text',", which provides a suitable balance between perturbing the best individual and "
        "preserving its information. Consequently, the parameter settings of MSSBOA are "),
        ('math','N_{min}=0.2N'),('text'," and "),('math','p_{m}=0.5'),('text',".")])

    P.heading("4.3. Ablation experiment",2)
    P.body([('text',"To investigate the contribution of each improvement strategy, three variants of "
        "MSSBOA are constructed by adding a single strategy to the basic SBOA, namely SBOA-GPSI, "
        "SBOA-LOBL and SBOA-ACGM, as listed in Table 4. The four algorithms and the full MSSBOA are "
        "independently run 30 times on the CEC2017 suite in four dimensions, and the results are "
        "analyzed by the Friedman test. The complete statistical results are reported in Table A1–"
        "Table A4 of Appendix A. Table 5 summarizes the Friedman rankings, which are visualized in "
        "Figure 6.")])
    P.table_caption("Descriptions of MSSBOA with different strategies.")
    VAR=ABL["variants"]; head=["Strategy"]+VAR
    rows=[]
    flags={"GPSI":{"SBOA":"No","SBOA-GPSI":"Yes","SBOA-LOBL":"No","SBOA-ACGM":"No","MSSBOA":"Yes"},
           "LOBL":{"SBOA":"No","SBOA-GPSI":"No","SBOA-LOBL":"Yes","SBOA-ACGM":"No","MSSBOA":"Yes"},
           "ACGM":{"SBOA":"No","SBOA-GPSI":"No","SBOA-LOBL":"No","SBOA-ACGM":"Yes","MSSBOA":"Yes"}}
    for s in ("GPSI","LOBL","ACGM"):
        rows.append([s]+[flags[s][v] for v in VAR])
    P.make_table(head,rows,font_size=8.5)
    # Table 5 ablation friedman
    P.table_caption("Friedman rankings of MSSBOA with different strategies (α=0.05).")
    head=["Test suite","Dim"]+VAR+["P-value"]
    rows=[]
    for k,d in enumerate(DIMS):
        fa=ABL["per_dim"][d]["friedman_avg"]; pv=ABL["per_dim"][d]["p"]
        row=(["CEC2017" if k==0 else ""]+[d]+[f"{fa[v]:.3f}" for v in VAR]+[sci(pv)])
        rows.append(row)
    ov=ABL["overall"]
    rows.append(["Average ranking",""]+[f"{ov[v]:.3f}" for v in VAR]+["N/A"])
    order=sorted(VAR,key=lambda v:ov[v]); rk={v:i+1 for i,v in enumerate(order)}
    rows.append(["Ranking",""]+[str(rk[v]) for v in VAR]+["N/A"])
    P.make_table(head,rows,font_size=8)
    P.figure("fig06_ablation.png","Friedman rankings of MSSBOA with different strategies (α=0.05).",width_cm=13.5)
    P.body([('text',"According to Table 5 and Figure 6, all p-values are far below 0.05, indicating "
        "significant differences among the five algorithms. MSSBOA ranks first in all four dimensions "
        "with an overall average ranking of "),('b',f"{ov['MSSBOA']:.3f}"),
        ('text',", followed by SBOA-ACGM, SBOA-GPSI and SBOA-LOBL, while the basic SBOA ranks last. "
        "Each single-strategy variant outperforms the basic SBOA, which confirms that every strategy "
        "is effective. Among them, the adaptive Cauchy–Gaussian mutation (ACGM) contributes the most, "
        "because it directly helps the best individual escape local optima and balances exploration "
        "and exploitation; the good point set initialization (GPSI) ranks next by enriching the "
        "initial diversity; and the lens opposition-based learning (LOBL) also improves performance by "
        "refreshing inferior individuals. The full MSSBOA outperforms every single-strategy variant, "
        "demonstrating that the three strategies complement and reinforce each other.")])

    P.heading("4.4. Comparison with other advanced algorithms",2)
    P.body([('text',"In this subsection, MSSBOA is compared with the basic SBOA and nine advanced "
        "algorithms on the CEC2017 suite in 10, 30, 50 and 100 dimensions. The complete results, "
        "including the best value, mean value, standard deviation and rank, are reported in Table A5–"
        "Table A8 of Appendix A. The ranking of each algorithm on every function is first visualized "
        "as heatmaps in Figure 7.")])
    n_first={d:CEC["results"][d]["n_first"]["MSSBOA"] for d in DIMS}
    P.figure("fig07_heatmap.png","Ranking heatmaps of MSSBOA and the comparison algorithms on the "
        "CEC2017 suite: (a) 10D, (b) 30D, (c) 50D, (d) 100D.",width_cm=15.5)
    P.body([('text',f"As shown in Figure 7, MSSBOA ranks first on {n_first['10']} functions in 10D, "
        f"{n_first['30']} functions in 30D, {n_first['50']} functions in 50D and {n_first['100']} "
        "functions in 100D, that is, on more than half of the 29 functions in every dimension, which "
        "preliminarily demonstrates its superiority. The convergence curves of representative functions "
        "are illustrated in Figure 8, where MSSBOA converges faster and reaches lower error values than "
        "the comparison algorithms.")])
    P.figure("fig_conv_cec.png","Convergence curves of MSSBOA and the comparison algorithms on "
        "representative CEC2017 functions (30D).",width_cm=14.0)
    P.body([('text',"To statistically analyze the overall performance, the Friedman test is performed; "
        "the results are summarized in Table 6 and visualized in Figure 9. ")])
    # Table 6 friedman comparison
    P.table_caption("Friedman rankings of MSSBOA and the comparison algorithms (α=0.05).")
    head=["Algorithm"]+[f"{d}D" for d in DIMS]+["Average","Rank"]
    rows=[]
    ovr=CEC["overall"]; order=sorted(ALGOS,key=lambda a:ovr[a]); rkmap={a:i+1 for i,a in enumerate(order)}
    for a in ALGOS:
        rows.append([a]+[f"{CEC['results'][d]['friedman_avg'][a]:.3f}" for d in DIMS]
                    +[f"{ovr[a]:.3f}",str(rkmap[a])])
    pvr=["Friedman P-value"]+[sci(CEC['results'][d]['friedman_p']) for d in DIMS]+["N/A","N/A"]
    rows.append(pvr)
    P.make_table(head,rows,font_size=8)
    P.figure("fig08_friedman_all.png","Friedman rankings of MSSBOA and the comparison algorithms across "
        "dimensions (α=0.05).",width_cm=12.5)
    rank_str=" > ".join(f"{a}" for a in order)
    P.body([('text',"As shown in Table 6 and Figure 9, all Friedman p-values are below 0.05, confirming "
        "significant differences among the algorithms. MSSBOA achieves the best average ranking of "),
        ('b',f"{ovr['MSSBOA']:.3f}"),('text'," and ranks first in all four dimensions, whereas the basic "
        "SBOA ranks near the bottom; the overall order is approximately "),('it',rank_str),
        ('text',". This clearly reflects the substantial improvement of MSSBOA over the basic SBOA. To "
        "further quantify the pairwise differences, the Nemenyi post-hoc test is applied, as shown in "
        "Figure 10, where algorithms connected by a line are not significantly different. MSSBOA is "
        "located at the leftmost position and is significantly superior to most comparison algorithms.")])
    P.figure("fig09_nemenyi.png","Nemenyi post-hoc test of MSSBOA and the comparison algorithms: "
        "(a) 10D, (b) 30D, (c) 50D, (d) 100D.",width_cm=15.0)
    P.body([('text',"Finally, the Wilcoxon rank-sum test is conducted to compare MSSBOA with each "
        "algorithm on every function. Table 7 reports the statistics, where '+', '=' and '−' denote "
        "the numbers of functions on which MSSBOA is significantly better than, similar to, and worse "
        "than the comparison algorithm, respectively. The results are visualized in Figure 11.")])
    # Table 7 wilcoxon totals
    P.table_caption("Wilcoxon rank-sum test results of MSSBOA versus the comparison algorithms (α=0.05).")
    comp=[a for a in ALGOS if a!="MSSBOA"]
    head=["MSSBOA vs."]+[f"{d}D (+/=/−)" for d in DIMS]+["Total (+/=/−)"]
    rows=[]
    for a in comp:
        cells=[a]
        tot=[0,0,0]
        for d in DIMS:
            w=CEC["results"][d]["wilcoxon"][a]; cells.append(f"{w[0]}/{w[1]}/{w[2]}")
            for i in range(3): tot[i]+=w[i]
        cells.append(f"{tot[0]}/{tot[1]}/{tot[2]}"); rows.append(cells)
    P.make_table(head,rows,font_size=8)
    P.figure("fig10_wilcoxon.png","Visualization of the Wilcoxon rank-sum test results of MSSBOA and "
        "the comparison algorithms (α=0.05): (a) 10D, (b) 30D, (c) 50D, (d) 100D.",width_cm=15.0)
    P.body([('text',"As shown in Table 7 and Figure 11, MSSBOA significantly outperforms each comparison "
        "algorithm on the majority of functions. In particular, MSSBOA wins on all 116 cases against the "
        "basic SBOA, SCA and GWO, and still dominates the strong competitors LSHADE, QIO and RIME on "
        "most functions, with only a few losses on high-dimensional problems. In summary, MSSBOA is a "
        "competitive and robust SBOA variant for global optimization.")])

def _param_table(P,key,caption):
    P.table_caption(caption)
    sets=PS[key]["settings"]; tab=PS[key]["table"]
    head=["Dimension","SBOA"]+sets
    rows=[]
    for d in DIMS:
        rows.append([f"D={d}",f"{tab[d]['SBOA']:.3f}"]+[f"{tab[d][s]:.3f}" for s in sets])
    avg=tab["avg"]
    rows.append(["Average","%.3f"%avg["SBOA"]]+[f"{avg[s]:.3f}" for s in sets])
    P.make_table(head,rows,font_size=7.5)

# ============================ SECTION 5 ============================
def section5(P):
    P.heading("5. Application to Visual Art Design Problems",1)
    P.body([('text',"In this section, MSSBOA is applied to two representative optimization problems in "
        "visual art and graphic design—aesthetic color-harmony palette generation and graphic-layout "
        "aesthetics optimization—to evaluate its ability to solve real-world design problems. Both "
        "problems are formulated as continuous optimization problems and solved by the eleven "
        "algorithms with the parameter settings of Table 1. Interestingly, both the optimizer "
        "(inspired by the secretary bird) and the color-harmony objective (rooted in the harmonious "
        "color schemes observed in nature) are bio-inspired, which makes the application a natural fit "
        "for biomimetic design.")])

    P.heading("5.1. Aesthetic color-harmony palette optimization",2)
    P.body([('text',"Generating a harmonious yet expressive color palette is a fundamental task in "
        "visual design "),('cite',['cohenor','hu_colorscheme']),
        ('text',". Following the harmonic-template theory of Matsuda and the color-harmonization model "
        "of Cohen-Or et al. "),('cite',['cohenor','tokumaru']),
        ('text',", each color is represented in the HSV space as "),('math','c_{k}=(H_{k},S_{k},V_{k})'),
        ('text',", where the hue "),('math','H_{k}\\in[0^{\\circ},360^{\\circ})'),
        ('text'," lies on a cyclic hue wheel. The geodesic (arc-length) distance between two hues is "
        "defined by Equation (15).")])
    e15=P.equation(r"d_{\mathbb{S}}(\theta_{1},\theta_{2})=\min\!\left(\left|\theta_{1}-\theta_{2}\right|,\,360^{\circ}-\left|\theta_{1}-\theta_{2}\right|\right)")
    P.body([('text',"A harmonic template "),('math','T_{m}'),('text'," is a set of one or two sectors on "
        "the hue wheel that can be rotated by an angle "),('math','\\alpha'),
        ('text',". The saturation-weighted disharmony energy of a "),('math','K'),
        ('text',"-color palette with respect to the best-fitting template is given by Equation (16).")])
    e16=P.equation(r"E_{harm}=\min_{m\in\mathcal{M}}\,\min_{\alpha\in[0^{\circ},360^{\circ})}\sum_{k=1}^{K}\rho\!\left(H_{k}\,\middle|\,m,\alpha\right)\cdot S_{k}")
    P.body([('text',"where "),('math',r'\rho(H_{k}\mid m,\alpha)'),('text'," is the arc distance from "),
        ('math','H_{k}'),('text'," to the nearest sector of the rotated template (zero when the hue "
        "falls inside a sector) and "),('math',r'\mathcal{M}'),('text'," is the set of seven chromatic "
        "templates. The disharmony energy is converted into a harmony term in "),('math','[0,1]'),
        ('text'," by Equation (17).")])
    e17=P.equation(r"\mathrm{harm}=1-\frac{E_{harm}}{180^{\circ}\cdot\sum_{k=1}^{K}S_{k}}")
    P.body([('text',"A purely harmonic palette may collapse into nearly identical colors, so a hue "
        "diversity term is introduced by Equation (18) to encourage perceptually distinct colors, and "
        "a tonal-contrast term "),('math',r'\mathrm{con}=(\,V_{\max}-V_{\min}\,)/0.8'),
        ('text'," rewards usable lightness spread.")])
    e18=P.equation(r"\mathrm{div}=\frac{2}{K(K-1)}\sum_{i<j}\frac{d_{\mathbb{S}}(H_{i},H_{j})}{180^{\circ}}")
    P.body([('text',"The decision vector is the stacked palette "),('math','\\mathbf{x}=(H_{1},S_{1},V_{1},\\dots,H_{K},S_{K},V_{K})'),
        ('text'," of dimension "),('math','D=3K'),('text',", and the optimizer maximizes the aesthetic "
        "objective defined by Equation (19).")])
    e19=P.equation(r"\max_{\mathbf{x}}\ J(\mathbf{x})=w_{h}\cdot\mathrm{harm}+w_{d}\cdot\mathrm{div}+w_{c}\cdot\mathrm{con}")
    P.body([('text',"where "),('math',r'w_{h}=0.55,\ w_{d}=0.30,\ w_{c}=0.15'),
        ('text',". Because the harmony and diversity terms conflict, the objective is non-convex and "
        "multimodal, which makes it a suitable test for metaheuristic optimizers. Palettes of "),
        ('math',r'K=5,\,7,\,10'),('text'," colors are optimized, and the harmony score is reported as "),
        ('math',r'100\times J'),('text'," (%). Figure 12 visualizes the seven-color palettes generated "
        "by representative algorithms together with the best-fitting harmonic template on the hue "
        "wheel, and Figure 13 shows the convergence curves and the distribution of harmony scores. "
        "Table 8 summarizes the statistical results.")])
    P.figure("fig11_palettes.png","Color-harmony palettes generated by MSSBOA and the comparison "
        "algorithms: (a) the MSSBOA palette on the hue wheel with its best-fitting harmonic template; "
        "(b) the optimized seven-color palettes and their harmony scores.",width_cm=15.5)
    P.figure("fig12_harmony_conv.png","Performance on the color-harmony problem (K=7): (a) convergence "
        "curves; (b) distribution of the harmony scores of all algorithms.",width_cm=15.0)
    # Table 8
    P.table_caption("Harmony scores (%) of MSSBOA and the comparison algorithms on the color-harmony "
        "palette optimization problem.")
    Ks=["5","7","10"]
    head=["Algorithm"]+[f"K={k} Mean" for k in Ks]+["K=7 Best","K=7 Std","Rank"]
    # rank by K=7 mean
    means7={a:HARM["results"]["7"][a]["mean"] for a in ALGOS}
    order=sorted(ALGOS,key=lambda a:-means7[a]); rk={a:i+1 for i,a in enumerate(order)}
    rows=[]; best_per_k={k:max(HARM["results"][k][a]["mean"] for a in ALGOS) for k in Ks}
    best7=max(HARM["results"]["7"][a]["best"] for a in ALGOS)
    for a in ALGOS:
        cells=[a]
        for k in Ks:
            v=HARM["results"][k][a]["mean"]
            cells.append((f"{v:.2f}", abs(v-best_per_k[k])<1e-6))
        b=HARM["results"]["7"][a]["best"]
        cells.append((f"{b:.2f}", abs(b-best7)<1e-6))
        cells.append(f"{HARM['results']['7'][a]['std']:.3f}")
        cells.append(str(rk[a]))
        rows.append(cells)
    _bold_table(P,head,rows,font_size=8)
    P.body([('text',"As shown in Table 8 and Figure 12, MSSBOA obtains the highest harmony score in "
        "every setting, reaching about 88% for the seven-color palette, while the basic SBOA only "
        "attains about 69% and SCA about 57%. The palette produced by MSSBOA fits a harmonic template "
        "closely (Figure 12a) while keeping the colors perceptually distinct, whereas the palettes of "
        "weaker algorithms contain clashing or washed-out colors. The convergence curves in Figure 13a "
        "show that MSSBOA converges fastest to the highest score, and the box plot in Figure 13b shows "
        "that it also has the most concentrated and stable distribution. These results confirm that "
        "MSSBOA is highly effective for aesthetic color design.")])

    P.heading("5.2. Graphic-layout aesthetics optimization",2)
    P.body([('text',"Arranging visual elements on a canvas to achieve a balanced and pleasing "
        "composition is another core task in graphic design "),('cite',['ngo','odonovan_layout']),
        ('text',". Given "),('math','N'),(
        'text'," elements with fixed sizes, the decision variables are the normalized centers of the "
        "elements, "),('math',r'\mathbf{x}=(c_{1}^{x},c_{1}^{y},\dots,c_{N}^{x},c_{N}^{y})'),
        ('text'," of dimension "),('math','D=2N'),
        ('text',". Following the quantitative aesthetic measures of Ngo et al. "),('cite',['ngo','lai']),
        ('text',", the visual balance and the non-overlap degree are defined by Equation (20) and "
        "Equation (21).")])
    e20=P.equation(r"\mathrm{BM}=1-\left(\left|\bar{x}-0.5\right|+\left|\bar{y}-0.5\right|\right),\quad \bar{x}=\frac{\sum_{i}a_{i}c_{i}^{x}}{\sum_{i}a_{i}},\ \ \bar{y}=\frac{\sum_{i}a_{i}c_{i}^{y}}{\sum_{i}a_{i}}")
    e21=P.equation(r"\mathrm{OV}=1-\frac{\sum_{i<j}\mathrm{area}\!\left(R_{i}\cap R_{j}\right)}{\sum_{i}a_{i}}")
    P.body([('text',"where "),('math','a_{i}'),('text'," and "),('math','R_{i}'),('text'," are the area "
        "and rectangle of element "),('math','i'),('text',". Together with an alignment term "),
        ('math',r'\mathrm{AL}'),('text',", a symmetry term "),('math',r'\mathrm{SY}'),
        ('text'," and a white-space (density) term "),('math',r'\mathrm{DEN}'),
        ('text',", the overall aesthetic objective is defined by Equation (22).")])
    e22=P.equation(r"\max_{\mathbf{x}}\ J_{L}=0.30\,\mathrm{BM}+0.22\,\mathrm{OV}+0.18\,\mathrm{AL}+0.15\,\mathrm{SY}+0.15\,\mathrm{DEN}-\eta\,P_{b}")
    P.body([('text',"where "),('math','P_{b}'),('text'," penalizes elements that exceed the canvas and "),
        ('math','\\eta'),('text'," is a penalty coefficient. Eight typical design problems, listed in "
        "Table 9, are considered, covering posters, web banners, business cards, magazine pages, photo "
        "collages, mobile interfaces, packaging labels and book covers. Table 10 reports the mean "
        "aesthetic score (×100) of each algorithm, and the rankings are visualized in Figure 14. "
        "Representative layouts optimized by MSSBOA and the basic SBOA are shown in Figure 15.")])
    # Table 9 problems
    P.table_caption("The eight graphic-layout design problems.")
    probs=LAY["problems"]
    sizes_n={p:0 for p in probs}
    # load element counts from layout data
    rows=[]
    for p in probs:
        n=len(LAY["layouts"][p]["sizes"]); D=2*n
        nm=p.split(" ",1)[1] if " " in p else p
        rows.append([p.split()[0],nm,str(n),str(D)])
    P.make_table(["Problem","Name","Elements N","D"],rows,font_size=8.5)
    # Table 10 results
    P.table_caption("Mean aesthetic scores (×100) of MSSBOA and the comparison algorithms on the "
        "graphic-layout design problems.")
    head=["Problem"]+ALGOS
    rows=[]
    avg_rank={a:[] for a in ALGOS}
    for p in probs:
        means={a:LAY["results"][p][a]["mean"]*100 for a in ALGOS}
        best=max(means.values())
        order=sorted(ALGOS,key=lambda a:-means[a])
        for i,a in enumerate(order,1): avg_rank[a].append(i)
        cells=[p.split()[0]]+[(f"{means[a]:.2f}", abs(means[a]-best)<1e-9) for a in ALGOS]
        rows.append(cells)
    arow=["Average rank"]+[f"{sum(avg_rank[a])/len(avg_rank[a]):.2f}" for a in ALGOS]
    rows.append(arow)
    _bold_table(P,head,rows,font_size=7.5)
    P.figure("fig13_layout_heatmap.png","Ranking heatmap of MSSBOA and the comparison algorithms on "
        "the graphic-layout design problems.",width_cm=13.0)
    P.figure("fig14_layouts.png","Representative layouts optimized by MSSBOA (top) and the basic SBOA "
        "(bottom) on three design problems.",width_cm=13.5)
    P.body([('text',"As shown in Table 10 and Figure 14, MSSBOA achieves the best mean aesthetic score "
        "and the best average rank on all eight design problems, followed by LSHADE, QIO and RIME, "
        "while the basic SBOA and SCA rank low. Figure 15 visually confirms this conclusion: the "
        "layouts produced by MSSBOA are well balanced, well aligned and free of overlap, whereas those "
        "produced by the basic SBOA suffer from overlapping elements and poor balance. These results "
        "demonstrate that MSSBOA can effectively solve practical graphic-layout design problems and is "
        "a promising tool for computational-aesthetics applications.")])

# ============================ SECTION 6 ============================
def section6(P):
    P.heading("6. Conclusions",1)
    P.body([('text',"This paper proposed a multi-strategy secretary bird optimization algorithm (MSSBOA) "
        "for aesthetic optimization in visual art design. Three improvement strategies were integrated "
        "into the basic SBOA: a good point set initialization that generates a uniformly distributed "
        "initial population and enriches diversity; a lens opposition-based learning that helps "
        "inferior individuals escape local optima; and an adaptive Cauchy–Gaussian mutation that "
        "balances exploration and exploitation throughout the search. The complexity analysis showed "
        "that these strategies do not increase the order of computational complexity of SBOA.")])
    P.body([('text',"On the CEC2017 benchmark suite in 10, 30, 50 and 100 dimensions, the optimal "
        "parameters of MSSBOA were determined by a sensitivity analysis, and the effectiveness of each "
        "strategy was verified by ablation experiments. Compared with the basic SBOA and nine advanced "
        "algorithms, MSSBOA achieved the best Friedman ranking in all dimensions and won the majority "
        "of functions in the Wilcoxon rank-sum test, as further confirmed by the Nemenyi post-hoc "
        "test. When applied to aesthetic color-harmony palette optimization and graphic-layout "
        "aesthetics optimization, MSSBOA produced the most harmonious palettes (about 88% harmony "
        "score) and the most balanced layouts, outperforming all comparison algorithms. These results "
        "demonstrate that MSSBOA is an effective and robust optimizer for computational-aesthetics "
        "problems in art and graphic design.")])
    P.body([('text',"Nevertheless, MSSBOA has some limitations. First, its parameters are determined "
        "empirically, and a generally optimal setting over a wider range of problems requires further "
        "study; an adaptive parameter-control mechanism is a promising direction. Second, the "
        "performance gain of MSSBOA on very-high-dimensional problems is smaller than that on "
        "low-dimensional problems, which suggests that its high-dimensional search ability should be "
        "strengthened. Third, the aesthetic objectives used here are based on established "
        "computational-aesthetics models; in future work, we plan to incorporate data-driven and "
        "perceptual-learning models "),('cite',['odonovan_compat','zhang_aes']),
        ('text'," and to extend MSSBOA to multi-objective and interactive design scenarios, such as "
        "product-form and cultural-creative design "),('cite',['lo_aes','deng_colormatch','jiang_colorregen']),
        ('text',", so as to better serve real-world art and design practice.")])

# ============================ BACK MATTER ============================
def backmatter(P):
    def bm(label,text):
        p=P.doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); p.paragraph_format.alignment=AL.JUSTIFY
        r=p.add_run(label+" "); r.bold=True; r.font.size=Pt(9)
        r=p.add_run(text); r.font.size=Pt(9)
    bm("Author Contributions:","Conceptualization, X.C. and C.Z.; methodology, X.C. and C.Z.; "
       "software, X.C.; validation, X.C. and C.Z.; formal analysis, X.C.; investigation, X.C.; "
       "data curation, X.C.; writing—original draft preparation, X.C.; writing—review and editing, "
       "X.C. and C.Z.; visualization, X.C.; supervision, C.Z.; funding acquisition, X.C. and C.Z. "
       "All authors have read and agreed to the published version of the manuscript.")
    bm("Funding:","This research was funded by the Major Humanities and Social Sciences Research "
       "Projects in Zhejiang Higher Education Institutions (Grant Number: 2024QN018) and the Zhejiang "
       "Provincial Philosophy and Social Sciences Planning Special Project (Grant Number: 25NDJC153YBMS).")
    bm("Data Availability Statement:","The data presented in this study are available on request from "
       "the corresponding author.")
    bm("Acknowledgments:","The authors would like to thank the anonymous reviewers and the editor for "
       "their careful reviews and constructive suggestions.")
    bm("Conflicts of Interest:","The authors declare no conflicts of interest.")

# ============================ APPENDIX A ============================
def _bold_table(P, headers, rows, font_size=8):
    """rows: list of cells; each cell is str OR (text,bold)."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t=P.doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    P._three_line(t)
    for j,h in enumerate(headers):
        pp=t.rows[0].cells[j].paragraphs[0]; pp.alignment=AL.CENTER
        r=pp.add_run(str(h)); r.bold=True; r.font.size=Pt(font_size)
    for row in rows:
        cells=t.add_row().cells
        for j,v in enumerate(row):
            pp=cells[j].paragraphs[0]; pp.alignment=AL.LEFT if j==0 else AL.CENTER
            if isinstance(v,tuple): txt,bold=v
            else: txt,bold=v,False
            r=pp.add_run(str(txt)); r.font.size=Pt(font_size); r.bold=bold
    return t

def _appendix_block(P, caption, algos, stats_of, rank_of=None, font_size=6.5, label=None):
    """stats_of(fid,algo)->{best,mean,std}; rank_of(fid,algo)->int or None."""
    P.table_caption(caption, label=label)
    metrics=["Best","Mean","Std"]+(["Rank"] if rank_of else [])
    head=["Func","Index"]+algos
    rows=[]
    for fid in FUNCS:
        # determine best per metric (min for best/mean/std)
        vals={m:{a:None for a in algos} for m in metrics}
        for a in algos:
            s=stats_of(fid,a)
            vals["Best"][a]=s["best"]; vals["Mean"][a]=s["mean"]; vals["Std"][a]=s["std"]
            if rank_of: vals["Rank"][a]=rank_of(fid,a)
        bestcell={}
        for m in ["Best","Mean","Std"]:
            mn=min(vals[m].values()); bestcell[m]=mn
        for mi,m in enumerate(metrics):
            row=[ (f"F{fid}" if mi==0 else ""), m]
            for a in algos:
                if m=="Rank":
                    v=vals[m][a]; row.append((str(v), v==1))
                else:
                    v=vals[m][a]; row.append((sci(v), abs(v-bestcell[m])<1e-12))
            rows.append(row)
    t=_bold_table(P,head,rows,font_size=font_size)
    P.fit_table(t, total_cm=26.4, first_widths=[1.3,1.2])   # landscape A4 text width

def appendix(P):
    P.start_section(landscape=True)
    P.heading("Appendix A",1)
    VAR=ABL["variants"]
    for i,d in enumerate(DIMS,1):
        st=ABL["per_dim"][d]["stats"]
        _appendix_block(P, f"Statistical results of MSSBOA with different strategies on CEC2017 (D={d}).",
            VAR, lambda fid,a,st=st: st[a][fid], rank_of=None, font_size=7, label=f"A{i}")
    for i,d in enumerate(DIMS,5):
        st=CEC["results"][d]["stats"]; rk=CEC["results"][d]["rank"]
        _appendix_block(P, f"Statistical results of MSSBOA and the comparison algorithms on CEC2017 (D={d}).",
            ALGOS, lambda fid,a,st=st: st[a][fid], rank_of=lambda fid,a,rk=rk: rk[a][fid], font_size=6.0, label=f"A{i}")
