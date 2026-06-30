"""Front matter, abstract, introduction, Section 2 (SBOA), Section 3 (MSSBOA).
These sections are data-independent (prose + equations)."""
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.shared import Pt, Cm

TITLE=("A Multi-Strategy Secretary Bird Optimization Algorithm for "
       "Aesthetic Color and Layout Optimization in Visual Art Design")

def front_matter(P):
    d=P.doc
    P._p("Article", align=AL.LEFT, size=10, bold=False, italic=True, after=2)
    # title
    P._p(TITLE, align=AL.LEFT, size=16, bold=True, after=6, line=1.1)
    # authors
    pa=d.add_paragraph(); pa.paragraph_format.space_after=Pt(2)
    r=pa.add_run("Xinyu Cai "); r.bold=True; r.font.size=Pt(10)
    r=pa.add_run("1"); r.font.size=Pt(7); r.font.superscript=True
    r=pa.add_run(", Chaoyong Zhang "); r.bold=True; r.font.size=Pt(10)
    r=pa.add_run("2,"); r.font.size=Pt(7); r.font.superscript=True
    r=pa.add_run("*"); r.font.size=Pt(7); r.font.superscript=True
    # affiliations
    for line in [
        "1   College of Business, Jiaxing University, Jiaxing 314001, China; caixinyu@zjxu.edu.cn",
        "2   School of Civil Engineering and Architecture, Jiaxing Nanhu University, Jiaxing 314000, China; 220028@jxnhu.edu.cn",
        "*   Correspondence: 220028@jxnhu.edu.cn",
    ]:
        P._p(line, align=AL.LEFT, size=8, after=1)
    # editorial placeholders
    for line in ["Academic Editor: Firstname Lastname","Received: date","Revised: date",
                 "Accepted: date","Published: date"]:
        P._p(line, align=AL.LEFT, size=8, after=0)
    P._p("Citation: To be added by editorial staff during production.", align=AL.LEFT, size=8, after=1)
    P._p("Copyright: © 2026 by the authors. Submitted for possible open access publication "
         "under the terms and conditions of the Creative Commons Attribution (CC BY) license "
         "(https://creativecommons.org/licenses/by/4.0/).", align=AL.JUSTIFY, size=8, after=6)

def abstract(P):
    d=P.doc
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.alignment=AL.JUSTIFY
    r=p.add_run("Abstract: "); r.bold=True; r.font.size=Pt(9)
    txt=("Visual art and graphic design tasks, such as generating a harmonious color palette or "
         "arranging the elements of a page, can be naturally formulated as mathematical optimization "
         "problems whose objective functions are non-convex, multimodal and non-differentiable. "
         "Metaheuristic algorithms are well suited to such problems. The secretary bird optimization "
         "algorithm (SBOA) is a recently proposed bio-inspired metaheuristic that mimics the hunting "
         "and predator-escaping behaviors of secretary birds, and it has shown competitive performance. "
         "However, SBOA still suffers from insufficient population diversity, premature convergence and "
         "an unbalanced transition between exploration and exploitation, which limits its accuracy on "
         "complex design problems. To overcome these limitations, this paper proposes a multi-strategy "
         "secretary bird optimization algorithm (MSSBOA) that integrates three improvement strategies. "
         "First, a good point set initialization is employed to generate a low-discrepancy initial "
         "population that covers the search space more uniformly and enriches population diversity. "
         "Second, a lens opposition-based learning strategy is applied to the inferior individuals to "
         "help the population escape local optima while preserving the elite. Third, an adaptive "
         "Cauchy–Gaussian mutation is imposed on the best individual to balance global exploration "
         "and local exploitation throughout the search. The performance of MSSBOA is comprehensively "
         "examined on the CEC2017 benchmark suite in 10, 30, 50 and 100 dimensions and compared with "
         "the basic SBOA and nine state-of-the-art algorithms; the results are analyzed by the Friedman "
         "test, the Nemenyi post-hoc test and the Wilcoxon rank-sum test. MSSBOA is then applied to two "
         "representative visual-design optimization problems: aesthetic color-harmony palette generation "
         "and graphic-layout aesthetics optimization. In all cases, MSSBOA outperforms the basic SBOA "
         "and the competing algorithms in terms of convergence speed, stability and solution quality, "
         "confirming its effectiveness for computational-aesthetics applications in art design.")
    r=p.add_run(txt); r.font.size=Pt(9)
    # keywords
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(8); p.paragraph_format.alignment=AL.JUSTIFY
    r=p.add_run("Keywords: "); r.bold=True; r.font.size=Pt(9)
    r=p.add_run("secretary bird optimization algorithm; metaheuristic algorithm; color harmony "
                "optimization; aesthetic layout design; lens opposition-based learning; computational "
                "aesthetics"); r.font.size=Pt(9)

def introduction(P):
    c=P.cite
    P.heading("1. Introduction",1)
    P.body([
        ('text',"With the rapid development of the digital creative industry, art and design activities "
                "increasingly rely on computational tools to improve both efficiency and quality"),
        ('cite',['jin_ultra','slowik_nature']),
        ('text',". Many tasks in visual art and graphic design—such as selecting a harmonious "
                "color palette, arranging the elements of a poster, or tuning the form parameters of a "
                "product—can be abstracted as the search for an optimal solution within a large "
                "feasible space "),
        ('cite',['tang_mrfomod']),
        ('text',". Such design optimization problems aim to maximize or minimize one or more aesthetic "
                "objective functions while satisfying perceptual and functional constraints, and they "
                "are widespread in fields ranging from visual communication and packaging to interface "
                "and environmental design "),
        ('cite',['jin_facial']),
        ('text',". Unfortunately, these objectives are usually nonlinear, multimodal, non-differentiable "
                "and strongly coupled with human perception, which makes them extremely difficult for "
                "classical optimization techniques "),
        ('cite',['tang_smamod']),
        ('text',". Mathematically, many of them belong to the class of NP-hard problems, for which an "
                "exact optimum cannot be obtained in polynomial time "),
        ('cite',['dai_jobshop','tang_sparrow']),
        ('text',". Optimization methods are generally divided into deterministic and stochastic "
                "categories. Deterministic algorithms, such as linear programming, branch-and-bound and "
                "dynamic programming, rely on rigorous mathematical properties such as continuity, "
                "convexity and differentiability "),
        ('cite',['shen_sched']),
        ('text',". When confronted with non-differentiable objectives or numerous local optima, they "
                "easily fall into local optimality and exhibit limited adaptability "),
        ('cite',['cao_edge']),
        ('text',". Consequently, there is an urgent need for more efficient, stable and broadly "
                "applicable algorithms to cope with the increasingly complex design solution space."),
    ], after=6)

    P.body([
        ('text',"Metaheuristic algorithms are stochastic optimization techniques that require no "
                "gradient information and instead use probabilistic strategies to iteratively explore "
                "and exploit the problem space "),
        ('cite',['zeng_pso']),
        ('text',". By balancing global exploration with local exploitation, they effectively mitigate "
                "premature convergence and can obtain near-optimal solutions at an acceptable "
                "computational cost "),
        ('cite',['wang_giga','zhao_xfem']),
        ('text',". Owing to their flexibility, metaheuristics have been successfully applied to path "
                "planning "),
        ('cite',['tang_eouav','hu_mahaco']),
        ('text',", parameter extraction "),
        ('cite',['sowmya_nrbo']),
        ('text',", wireless sensor network coverage "),
        ('cite',['houssein_seahorse']),
        ('text',", image segmentation "),
        ('cite',['abualigah_rsa','hu_gsrpso']),
        ('text',", and feature selection "),
        ('cite',['seyyed_apo_fs','jia_mpa']),
        ('text',". Drawing inspiration from nature and social systems, metaheuristic algorithms can be "
                "grouped into four primary categories according to their source of inspiration, as shown "
                "in Figure 1. Evolution-based algorithms imitate natural selection, mutation and "
                "crossover, including the genetic algorithm (GA) "),
        ('cite',['ga']),
        ('text',", evolution strategy (ES) "),
        ('cite',['es']),
        ('text',", differential evolution (DE) "),
        ('cite',['de']),
        ('text',", the covariance matrix adaptation evolution strategy (CMA-ES) "),
        ('cite',['cmaes']),
        ('text',", the human evolutionary optimization algorithm (HEOA) "),
        ('cite',['heoa']),
        ('text'," and chaotic evolution optimization (CEO) "),
        ('cite',['ceo']),
        ('text',". Physics-based algorithms draw on physical laws, such as the multi-verse optimizer "
                "(MVO) "),
        ('cite',['mvo']),
        ('text',", Fick’s law algorithm (FLA) "),
        ('cite',['fla']),
        ('text',", snow ablation optimization (SAO) "),
        ('cite',['sao']),
        ('text',", the polar lights optimizer (PLO) "),
        ('cite',['plo']),
        ('text',", special relativity search (SRS) "),
        ('cite',['srs']),
        ('text',", mirage search optimization (MSO) "),
        ('cite',['mso']),
        ('text'," and the Kepler optimization algorithm (KOA) "),
        ('cite',['koa']),
        ('text',". Mathematics-based algorithms are inspired by mathematical functions and theories, "
                "such as the sine cosine algorithm (SCA) "),
        ('cite',['sca']),
        ('text',", the arithmetic optimization algorithm (AOA) "),
        ('cite',['aoa']),
        ('text',", quasi-random fractal search (QRFS) "),
        ('cite',['qrfs']),
        ('text',", the weighted average algorithm (WAA) "),
        ('cite',['waa']),
        ('text',", the triangulation topology aggregation optimizer (TTAO) "),
        ('cite',['ttao']),
        ('text'," and the exponential-trigonometric optimization algorithm (ETOA) "),
        ('cite',['etoa']),
        ('text',". Swarm-based algorithms simulate the collective behavior of groups of organisms, "
                "including particle swarm optimization (PSO) "),
        ('cite',['pso']),
        ('text',", ant colony optimization (ACO) "),
        ('cite',['aco']),
        ('text',", the grey wolf optimizer (GWO) "),
        ('cite',['gwo']),
        ('text',", the artificial lemming algorithm (ALA) "),
        ('cite',['ala']),
        ('text',", the crayfish optimization algorithm (COA) "),
        ('cite',['coa']),
        ('text',", dwarf mongoose optimization (DMO) "),
        ('cite',['dmo']),
        ('text',", the sled dog optimizer (SDO) "),
        ('cite',['sdo']),
        ('text',", the slime mould algorithm (SMA) "),
        ('cite',['sma']),
        ('text',", Arctic puffin optimization (APO) "),
        ('cite',['apo']),
        ('text',", the Genghis Khan shark optimizer (GKSO) "),
        ('cite',['gkso']),
        ('text',", the prairie dog optimization algorithm (PDOA) "),
        ('cite',['pdoa']),
        ('text'," and the secretary bird optimization algorithm (SBOA) "),
        ('cite',['sboa']),
        ('text',"."),
    ], after=6)

    P.figure("fig01_taxonomy.png","Classification of metaheuristic algorithms.", width_cm=15.5)

    P.body([
        ('text',"Despite their success, metaheuristic algorithms still face challenges. They can "
                "readily obtain good solutions but rarely guarantee the global optimum, and they often "
                "struggle to balance exploration and exploitation, so that the population may homogenize "
                "and become trapped in local optima. More importantly, the no free lunch theorem proves "
                "that no single algorithm can outperform all others on every possible problem "),
        ('cite',['nfl']),
        ('text',". This motivates researchers to design or improve algorithms tailored to specific "
                "problem classes—here, the aesthetic optimization problems that arise in visual art "
                "and graphic design."),
    ], after=6)

    P.body([
        ('text',"The secretary bird optimization algorithm (SBOA) is a bio-inspired metaheuristic "
                "proposed by Fu et al. in 2024 "),
        ('cite',['sboa']),
        ('text',". It models the survival behavior of secretary birds, using a three-stage strategy for "
                "hunting snakes during the exploration phase and two strategies for escaping predators "
                "during the exploitation phase. Owing to its simple structure and strong search ability, "
                "SBOA has attracted considerable attention, and several improved variants have been "
                "developed. Qin et al. proposed a multi-strategy improvement SBOA that incorporates "
                "incremental-PID feedback regulation and cooperative camouflage to solve engineering "
                "problems "),
        ('cite',['sboa_misboa']),
        ('text',". Wang et al. enhanced SBOA with multi-strategy fusion and a Cauchy–Gaussian "
                "crossover to improve convergence accuracy "),
        ('cite',['sboa_cg']),
        ('text',". Mai et al. integrated logistic-tent chaotic initialization and a crossover strategy "
                "into SBOA for engineering design "),
        ('cite',['sboa_cross']),
        ('text',", and Seyyedabbasi developed a multi-strategy variable SBOA for global optimization and "
                "UAV path planning "),
        ('cite',['sboa_msv']),
        ('text',". These studies enhance SBOA from different angles, yet the inability of the basic SBOA "
                "to maintain population diversity and to balance exploration and exploitation in complex, "
                "perception-driven design problems remains an objective limitation. Guided again by the "
                "no free lunch theorem, this paper proposes a new SBOA variant, the multi-strategy "
                "secretary bird optimization algorithm (MSSBOA), and applies it to aesthetic optimization "
                "in visual art design."),
    ], after=6)

    P.body([
        ('text',"MSSBOA combines three improvement techniques. First, a good point set initialization "
                "replaces the random initialization of SBOA, generating a low-discrepancy population that "
                "uniformly covers the search domain and enhances population diversity from the outset. "
                "Second, a lens opposition-based learning strategy is applied to the inferior individuals "
                "to generate refracted candidate solutions, which strengthens the ability of the algorithm "
                "to escape local optima while protecting the elite individuals. Third, an adaptive "
                "Cauchy–Gaussian mutation is imposed on the best individual, where a Cauchy-dominated "
                "perturbation promotes global exploration in the early stage and a Gaussian-dominated "
                "perturbation refines the solution in the later stage, thereby balancing exploration and "
                "exploitation. MSSBOA is comprehensively examined on the CEC2017 benchmark functions and "
                "compared with the basic SBOA and several advanced algorithms, with the differences "
                "analyzed by the Friedman test, the Nemenyi post-hoc test and the Wilcoxon rank-sum test. "
                "Finally, MSSBOA is applied to color-harmony palette optimization and graphic-layout "
                "aesthetics optimization to confirm its ability to solve realistic art-design problems. "
                "The main contributions of this paper are as follows."),
    ], after=4)
    for t in [
        "(1) An improved SBOA variant, namely MSSBOA, is proposed by integrating good point set "
        "initialization, lens opposition-based learning and adaptive Cauchy–Gaussian mutation.",
        "(2) The performance of MSSBOA is comprehensively examined on the CEC2017 test suite in four "
        "dimensions and confirmed by the Friedman test, the Nemenyi test and the Wilcoxon rank-sum test.",
        "(3) MSSBOA is successfully applied to aesthetic color-harmony palette optimization and "
        "graphic-layout aesthetics optimization, demonstrating its adaptability to real-world visual "
        "art and graphic design problems.",
    ]:
        P.body(t, after=2)

    P.body([
        ('text',"The remainder of this paper is organized as follows. Section 2 describes the structure "
                "and details of the basic SBOA. Section 3 presents the proposed MSSBOA, including the "
                "three improvement strategies, the overall framework with pseudo-code and flowchart, and "
                "the complexity analysis. Section 4 reports the parameter sensitivity analysis, ablation "
                "experiments and a comprehensive comparison on the CEC2017 benchmark functions. Section 5 "
                "applies MSSBOA to two visual-design optimization problems. Finally, Section 6 concludes "
                "the paper, discusses the limitations of MSSBOA and outlines future research directions."),
    ], after=6)
