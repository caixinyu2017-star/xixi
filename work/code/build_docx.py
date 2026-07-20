"""MDPI Biomimetics manuscript builder.

Rebuilds the body of the converted MDPI template docx while keeping its styles,
headers/footers, margins and first-page furniture. Display and inline equations
are native Word OMML objects produced by pandoc from LaTeX.

Content spec: list of block dicts:
  {'h1': text} {'h2': text} {'h3': text}
  {'p': text, 'style': optional}            # $latex$ inline math, **bold**, *italic*, {sup:x}
  {'eq': latex, 'num': int}                 # numbered display equation (2-col MDPI table)
  {'fig': path, 'caption': text, 'width_cm': float}
  {'table': {'caption': str, 'header': [..], 'rows': [[..]], 'widths': [..] or None,
             'fontsize': float, 'note': str or None}}
  {'algorithm': {'title': str, 'lines': [str]}}
  {'pagebreak': True}
  {'refs': [str, ...]}
"""
import copy
import os
import re
import subprocess
import tempfile
import zipfile

from lxml import etree
import docx
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# --------------------------------------------------------------- OMML via pandoc
def latex_batch_to_omml(latex_list, workdir):
    """Convert list of LaTeX strings to list of lxml <m:oMath> elements via pandoc."""
    if not latex_list:
        return []
    md_parts = []
    for i, l in enumerate(latex_list):
        md_parts.append(f'EQMARK{i:04d} ${l}$')
    md = '\n\n'.join(md_parts)
    mdp = os.path.join(workdir, '_eqs.md')
    outp = os.path.join(workdir, '_eqs.docx')
    with open(mdp, 'w') as fh:
        fh.write(md)
    subprocess.run(['pandoc', mdp, '-o', outp], check=True)
    with zipfile.ZipFile(outp) as z:
        xml = z.read('word/document.xml')
    root = etree.fromstring(xml)
    paras = root.findall('.//{%s}p' % W_NS)
    results = {}
    for p in paras:
        text = ''.join(t.text or '' for t in p.findall('.//{%s}t' % W_NS))
        m = re.search(r'EQMARK(\d{4})', text)
        if m is None:
            continue
        idx = int(m.group(1))
        om = p.find('.//{%s}oMath' % M_NS)
        if om is not None:
            results[idx] = om
    out = []
    for i in range(len(latex_list)):
        if i not in results:
            raise RuntimeError(f'equation {i} failed pandoc OMML conversion: {latex_list[i][:80]}')
        out.append(results[i])
    return out


INLINE_RE = re.compile(r'(\$[^$]+\$|\*\*[^*]+\*\*|\*[^*]+\*|\{sup:[^}]*\}|\{sub:[^}]*\})')


def collect_math(blocks):
    """Pre-scan blocks for every LaTeX snippet (inline + display), return ordered list."""
    math = []
    for b in blocks:
        if 'eq' in b:
            math.append(b['eq'])
        for key in ('p', 'caption'):
            if key in b and isinstance(b[key], str):
                for seg in INLINE_RE.findall(b[key]):
                    if seg.startswith('$'):
                        math.append(seg[1:-1])
        if 'table' in b:
            tb = b['table']
            cells = list(tb.get('header', [])) + [c for row in tb.get('rows', []) for c in row]
            if tb.get('caption'):
                cells.append(tb['caption'])
            if tb.get('note'):
                cells.append(tb['note'])
            for c in cells:
                if isinstance(c, str):
                    for seg in INLINE_RE.findall(c):
                        if seg.startswith('$'):
                            math.append(seg[1:-1])
        if 'algorithm' in b:
            for line in [b['algorithm']['title']] + b['algorithm']['lines']:
                for seg in INLINE_RE.findall(line):
                    if seg.startswith('$'):
                        math.append(seg[1:-1])
    return math


class PaperBuilder:
    def __init__(self, template_path, workdir):
        self.doc = docx.Document(template_path)
        self.workdir = workdir
        os.makedirs(workdir, exist_ok=True)
        # prototype equation table (1 row x 2 cols) from the template
        self.eq_proto = copy.deepcopy(self.doc.tables[1]._tbl)
        self._clear_after_line()
        self.body = self.doc.element.body
        self.sectPr = self.body.find(qn('w:sectPr'))
        self.omml_cache = {}

    # ---------------------------------------------------------------- setup
    def _clear_after_line(self):
        body = self.doc.element.body
        kill = False
        to_remove = []
        for el in list(body):
            if el.tag == qn('w:sectPr'):
                continue
            if kill:
                to_remove.append(el)
                continue
            if el.tag == qn('w:p'):
                pPr = el.find(qn('w:pPr'))
                if pPr is not None:
                    ps = pPr.find(qn('w:pStyle'))
                    if ps is not None and ps.get(qn('w:val')) == 'MDPI19line':
                        kill = True
        for el in to_remove:
            body.remove(el)

    def prepare_math(self, blocks):
        latexes = collect_math(blocks)
        uniq = list(dict.fromkeys(latexes))
        ommls = latex_batch_to_omml(uniq, self.workdir)
        self.omml_cache = dict(zip(uniq, ommls))

    # ---------------------------------------------------------------- front matter
    def set_front_matter(self, article_type, title, authors_runs, affiliations,
                         abstract, keywords, year='2026', volume='11'):
        """authors_runs: list of (text, is_sup) tuples. affiliations: list of texts (may contain {sup:})."""
        paras = self.doc.paragraphs
        by_style = {}
        for p in paras:
            by_style.setdefault(p.style.name, []).append(p)
        self._settext(by_style['MDPI_1.1_article_type'][0], article_type)
        self._settext(by_style['MDPI_1.2_title'][0], title)
        ap = by_style['MDPI_1.3_authornames'][0]
        for r in list(ap.runs):
            r._element.getparent().remove(r._element)
        for text, sup in authors_runs:
            r = ap.add_run(text)
            if sup:
                r.font.superscript = True
        affs = by_style['MDPI_1.6_affiliation']
        # make sure we have as many affiliation paragraphs as needed
        while len(affs) < len(affiliations):
            newp = copy.deepcopy(affs[-1]._p)
            affs[-1]._p.addnext(newp)
            affs.append(docx.text.paragraph.Paragraph(newp, affs[-1]._parent))
        for p, text in zip(affs, affiliations):
            self._set_rich(p, text)
        for extra in affs[len(affiliations):]:
            extra._p.getparent().remove(extra._p)
        self._set_rich(by_style['MDPI_1.7_abstract'][0], '**Abstract:** ' + abstract)
        self._set_rich(by_style['MDPI_1.8_keywords'][0], '**Keywords:** ' + keywords)
        # update header/footer year and volume
        for sec in self.doc.sections:
            for part in (sec.header, sec.first_page_header, sec.footer, sec.first_page_footer):
                for p in part.paragraphs:
                    if 'Biomimetics' in p.text:
                        for r in p.runs:
                            if r.text and '2025' in r.text:
                                r.text = r.text.replace('2025', year)
                            if r.text and '10' in r.text:
                                r.text = r.text.replace('10', volume, 1)
        # copyright year in margin box
        t = self.doc.tables[0]
        for p in t.cell(0, 0).paragraphs:
            for r in p.runs:
                if '2025' in (r.text or ''):
                    r.text = r.text.replace('2025', year)

    def _settext(self, para, text):
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        para.add_run(text)

    # ---------------------------------------------------------------- rich text
    def _set_rich(self, para, text, fontsize=None):
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        for child in list(para._p):
            if child.tag == qn('m:oMath') or child.tag == qn('m:oMathPara'):
                para._p.remove(child)
        self._append_rich(para, text, fontsize)

    def _append_rich(self, para, text, fontsize=None):
        pos = 0
        for m in INLINE_RE.finditer(text):
            if m.start() > pos:
                self._add_run(para, text[pos:m.start()], fontsize=fontsize)
            seg = m.group(0)
            if seg.startswith('$'):
                latex = seg[1:-1]
                om = self.omml_cache.get(latex)
                if om is None:
                    om = latex_batch_to_omml([latex], self.workdir)[0]
                    self.omml_cache[latex] = om
                para._p.append(copy.deepcopy(om))
            elif seg.startswith('**'):
                self._add_run(para, seg[2:-2], bold=True, fontsize=fontsize)
            elif seg.startswith('{sup:'):
                self._add_run(para, seg[5:-1], sup=True, fontsize=fontsize)
            elif seg.startswith('{sub:'):
                self._add_run(para, seg[5:-1], sub=True, fontsize=fontsize)
            else:
                self._add_run(para, seg[1:-1], italic=True, fontsize=fontsize)
            pos = m.end()
        if pos < len(text):
            self._add_run(para, text[pos:], fontsize=fontsize)

    def _add_run(self, para, text, bold=False, italic=False, sup=False, sub=False, fontsize=None):
        r = para.add_run(text)
        r.font.bold = bold or None
        r.font.italic = italic or None
        if sup:
            r.font.superscript = True
        if sub:
            r.font.subscript = True
        if fontsize:
            r.font.size = Pt(fontsize)
        return r

    # ---------------------------------------------------------------- blocks
    def _new_para(self, style):
        p = self.doc.add_paragraph(style=style)
        # move before sectPr (add_paragraph appends at body end, before sectPr already)
        return p

    def add_heading(self, text, level):
        style = {1: 'MDPI_2.1_heading1', 2: 'MDPI_2.2_heading2', 3: 'MDPI_2.3_heading3'}[level]
        p = self._new_para(style)
        self._append_rich(p, text)

    def add_paragraph(self, text, style='MDPI_3.1_text'):
        p = self._new_para(style)
        self._append_rich(p, text)

    def add_equation(self, latex, num):
        tbl = copy.deepcopy(self.eq_proto)
        self.body.insert(list(self.body).index(self.sectPr), tbl)
        t = docx.table.Table(tbl, self.doc)
        c0 = t.cell(0, 0)
        p = c0.paragraphs[0]
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        for child in list(p._p):
            if child.tag in (qn('m:oMath'), qn('m:oMathPara'), qn('w:r')):
                p._p.remove(child)
        om = self.omml_cache.get(latex)
        if om is None:
            om = latex_batch_to_omml([latex], self.workdir)[0]
        p._p.append(copy.deepcopy(om))
        c1 = t.cell(0, 1)
        p1 = c1.paragraphs[0]
        for r in list(p1.runs):
            r._element.getparent().remove(r._element)
        p1.add_run(f'({num})')

    def add_figure(self, img_path, caption, width_cm=13.5):
        p = self._new_para('MDPI_5.2_figure')
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        cp = self._new_para('MDPI_5.1_figure_caption')
        self._append_rich(cp, caption)

    def add_table(self, caption, header, rows, widths=None, fontsize=8, note=None,
                  header_bold=True):
        cp = self._new_para('MDPI_4.1_table_caption')
        self._append_rich(cp, caption)
        ncol = len(header)
        t = self.doc.add_table(rows=1 + len(rows), cols=ncol)
        t.autofit = True
        self._mdpi_borders(t)
        for j, txt in enumerate(header):
            cell = t.cell(0, j)
            p = cell.paragraphs[0]
            p.style = self.doc.styles['MDPI_4.2_table_body']
            self._append_rich(p, f'**{txt}**' if header_bold and txt else str(txt), fontsize=fontsize)
        for i, row in enumerate(rows):
            for j, txt in enumerate(row):
                cell = t.cell(i + 1, j)
                p = cell.paragraphs[0]
                p.style = self.doc.styles['MDPI_4.2_table_body']
                self._append_rich(p, str(txt), fontsize=fontsize)
        if widths:
            total = sum(widths)
            for j, w in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Cm(w)
        if note:
            np_ = self._new_para('MDPI_4.3_table_footer')
            self._append_rich(np_, note)
        else:
            sp = self._new_para('MDPI_4.3_table_footer')
            sp.paragraph_format.space_after = Pt(6)

    def _mdpi_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = tblPr.find(qn('w:tblBorders'))
        if borders is None:
            borders = tblPr.makeelement(qn('w:tblBorders'), {})
            tblPr.append(borders)
        for edge in list(borders):
            borders.remove(edge)
        for name, sz in (('top', 12), ('bottom', 12)):
            el = borders.makeelement(qn(f'w:{name}'), {})
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:color'), 'auto')
            borders.append(el)
        # thin rule under header row
        first_row = table.rows[0]._tr
        for tc in first_row.findall(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = tc.makeelement(qn('w:tcPr'), {})
                tc.insert(0, tcPr)
            tcB = tcPr.makeelement(qn('w:tcBorders'), {})
            bot = tcB.makeelement(qn('w:bottom'), {})
            bot.set(qn('w:val'), 'single')
            bot.set(qn('w:sz'), '6')
            bot.set(qn('w:color'), 'auto')
            tcB.append(bot)
            tcPr.append(tcB)

    def add_algorithm(self, title, lines, fontsize=9):
        t = self.doc.add_table(rows=1 + len(lines), cols=1)
        self._mdpi_borders(t)
        p = t.cell(0, 0).paragraphs[0]
        p.style = self.doc.styles['MDPI_4.2_table_body']
        self._append_rich(p, f'**{title}**', fontsize=fontsize)
        for i, line in enumerate(lines):
            p = t.cell(i + 1, 0).paragraphs[0]
            p.style = self.doc.styles['MDPI_4.2_table_body']
            p.alignment = 0
            self._append_rich(p, line, fontsize=fontsize)
        sp = self._new_para('MDPI_4.3_table_footer')
        sp.paragraph_format.space_after = Pt(6)

    def add_refs(self, refs):
        self.add_heading('References', 1)
        for ref in refs:
            p = self._new_para('MDPI_8.1_references')
            self._append_rich(p, ref)

    def add_backmatter(self, items):
        for text in items:
            p = self._new_para('MDPI_6.2_back_matter')
            self._append_rich(p, text)

    def add_pagebreak(self):
        from docx.enum.text import WD_BREAK
        p = self._new_para('MDPI_3.1_text')
        p.add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------------- driver
    def build(self, blocks):
        self.prepare_math(blocks)
        for b in blocks:
            if 'h1' in b:
                self.add_heading(b['h1'], 1)
            elif 'h2' in b:
                self.add_heading(b['h2'], 2)
            elif 'h3' in b:
                self.add_heading(b['h3'], 3)
            elif 'p' in b:
                style = b.get('style')
                if style is None:
                    style = ('MDPI_3.2_text_no_indent' if b['p'].startswith('where ')
                             else 'MDPI_3.1_text')
                self.add_paragraph(b['p'], style)
            elif 'eq' in b:
                self.add_equation(b['eq'], b['num'])
            elif 'fig' in b:
                self.add_figure(b['fig'], b['caption'], b.get('width_cm', 13.5))
            elif 'table' in b:
                tb = b['table']
                self.add_table(tb['caption'], tb['header'], tb['rows'],
                               tb.get('widths'), tb.get('fontsize', 8), tb.get('note'))
            elif 'algorithm' in b:
                self.add_algorithm(b['algorithm']['title'], b['algorithm']['lines'])
            elif 'refs' in b:
                self.add_refs(b['refs'])
            elif 'backmatter' in b:
                self.add_backmatter(b['backmatter'])
            elif 'pagebreak' in b:
                self.add_pagebreak()

    def save(self, path):
        self.doc.save(path)
