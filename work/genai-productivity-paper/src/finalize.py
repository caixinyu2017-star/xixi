# -*- coding: utf-8 -*-
"""对 pandoc 生成的 docx 做期刊化排版后处理：中文字体、字号、
版面、表格与题注样式等。"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC = sys.argv[1] if len(sys.argv) > 1 else "paper.docx"

doc = Document(SRC)

# ---------- 页面：A4，常规页边距 ----------
for sec in doc.sections:
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(2.6)


def set_fonts(style, east="宋体", ascii_font="Times New Roman", size=None,
              bold=None, color_black=True):
    f = style.font
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    f.italic = False
    f.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east)
    if color_black:
        f.color.rgb = RGBColor(0, 0, 0)


def para_fmt(style, align=None, first_indent=None, before=None, after=None,
             line=None):
    pf = style.paragraph_format
    if align is not None:
        pf.alignment = align
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line


S = {s.name: s for s in doc.styles}

# ---------- 基础正文 ----------
for name in ("Normal", "Body Text", "First Paragraph"):
    if name in S:
        set_fonts(S[name], east="宋体", size=10.5)
        para_fmt(S[name], align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0, line=1.5)
for name in ("Body Text", "First Paragraph"):
    if name in S:
        para_fmt(S[name], first_indent=Cm(0.74))

# ---------- 标题层级 ----------
if "Heading 1" in S:
    set_fonts(S["Heading 1"], east="黑体", ascii_font="Times New Roman",
              size=12, bold=True)
    para_fmt(S["Heading 1"], align=WD_ALIGN_PARAGRAPH.LEFT, before=10,
             after=6, line=1.5, first_indent=Cm(0))
for name in ("Heading 2", "Heading 3"):
    if name in S:
        set_fonts(S[name], east="黑体", ascii_font="Times New Roman",
                  size=10.5, bold=True)
        para_fmt(S[name], align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3,
                 line=1.5, first_indent=Cm(0))

# ---------- 自定义样式 ----------
custom = {
    "TitleCN": dict(east="黑体", size=15, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=12),
    "AuthorCN": dict(east="仿宋", size=10.5, bold=False,
                     align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=3),
    "AffilCN": dict(east="宋体", size=9, bold=False,
                    align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10),
    "AbstractCN": dict(east="楷体", size=9, bold=False,
                       align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4),
    "TableNote": dict(east="宋体", size=8, bold=False,
                      align=WD_ALIGN_PARAGRAPH.LEFT, before=2, after=8),
    "RefItem": dict(east="宋体", size=9, bold=False,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=2),
    "Table Caption": dict(east="黑体", size=9, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=3),
    "Image Caption": dict(east="黑体", size=9, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=8),
    "Captioned Figure": dict(east="宋体", size=10.5, bold=False,
                             align=WD_ALIGN_PARAGRAPH.CENTER, before=8,
                             after=0),
    "Compact": dict(east="宋体", size=9, bold=False,
                    align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1),
}
for name, cfg in custom.items():
    if name not in S:
        continue
    set_fonts(S[name], east=cfg["east"], size=cfg["size"], bold=cfg["bold"])
    para_fmt(S[name], align=cfg["align"], before=cfg["before"],
             after=cfg["after"], line=1.2, first_indent=Cm(0))

# ---------- 表格：居中、单元格对齐 ----------
for tbl in doc.tables:
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

doc.save(SRC)
print("finalize done:", SRC)
