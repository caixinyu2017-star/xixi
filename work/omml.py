# -*- coding: utf-8 -*-
"""极简 LaTeX → OMML（Office Math Markup Language）转换器。

生成 Word 原生公式对象（m:oMath），而非文本模拟，
因此在 Word 中可用公式编辑器直接打开与编辑。

支持：^{} _{} \frac{}{} \sqrt{} \sum_{}^{} \prod_{}^{} \int_{}^{}
      \left( \right) \text{} \mathrm{} 希腊字母、常用运算符与关系符。
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
import re

# python-docx 默认未注册 math 命名空间
nsmap.setdefault("m", "http://schemas.openxmlformats.org/officeDocument/2006/math")

GREEK = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ", "epsilon": "ε",
    "zeta": "ζ", "eta": "η", "theta": "θ", "iota": "ι", "kappa": "κ",
    "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ", "pi": "π", "rho": "ρ",
    "sigma": "σ", "tau": "τ", "upsilon": "υ", "phi": "φ", "chi": "χ",
    "psi": "ψ", "omega": "ω",
    "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ", "Xi": "Ξ",
    "Pi": "Π", "Sigma": "Σ", "Phi": "Φ", "Psi": "Ψ", "Omega": "Ω",
}
SYM = {
    "cdot": "·", "times": "×", "div": "÷", "pm": "±", "mp": "∓",
    "leq": "≤", "le": "≤", "geq": "≥", "ge": "≥", "neq": "≠", "ne": "≠",
    "approx": "≈", "equiv": "≡", "propto": "∝", "sim": "∼",
    "in": "∈", "notin": "∉", "subset": "⊂", "subseteq": "⊆",
    "infty": "∞", "partial": "∂", "nabla": "∇", "forall": "∀", "exists": "∃",
    "rightarrow": "→", "to": "→", "leftarrow": "←", "Rightarrow": "⇒",
    "ldots": "…", "cdots": "⋯", "quad": " ", "qquad": "  ",
    "%": "%", "&": "&", "#": "#", "_": "_", "{": "{", "}": "}", ",": " ",
}
FUNCS = {"exp", "ln", "log", "max", "min", "sin", "cos", "tan", "arg",
         "lim", "det", "dim", "gcd", "sup", "inf", "Pr"}


def _el(tag, *children, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k.replace("__", ":")), v)
    for c in children:
        if c is not None:
            e.append(c)
    return e


def _r(text, italic=None):
    """一个数学 run。italic=False 用于函数名与直立文本。"""
    rpr = None
    if italic is False:
        rpr = _el("m:rPr", _el("m:sty", **{"m__val": "p"}))
    t = _el("m:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    return _el("m:r", rpr, t)


class _Tok:
    def __init__(self, s):
        self.s, self.i = s, 0

    def eof(self):
        return self.i >= len(self.s)

    def peek(self):
        return self.s[self.i] if not self.eof() else ""

    def next(self):
        c = self.s[self.i]; self.i += 1
        return c

    def group(self):
        """读取一个 {...} 组或单个字符/命令。"""
        while not self.eof() and self.s[self.i] == " ":
            self.i += 1
        if self.eof():
            return ""
        if self.peek() == "{":
            depth, self.i, start = 0, self.i, self.i
            while not self.eof():
                c = self.next()
                if c == "{":
                    depth += 1
                elif c == "}":
                    depth -= 1
                    if depth == 0:
                        return self.s[start + 1:self.i - 1]
            return self.s[start + 1:]
        if self.peek() == "\\":
            j = self.i + 1
            while j < len(self.s) and self.s[j].isalpha():
                j += 1
            if j == self.i + 1:
                j += 1
            out = self.s[self.i:j]; self.i = j
            return out
        return self.next()


def _parse(src):
    """把 LaTeX 片段解析为 OMML 元素列表。"""
    out, tk = [], _Tok(src)
    buf = ""

    def flush(it=None):
        nonlocal buf
        if buf:
            out.append(_r(buf, it)); buf = ""

    while not tk.eof():
        c = tk.peek()
        if c == "\\":
            cmd = tk.group()[1:]
            if cmd == "frac":
                flush()
                num, den = tk.group(), tk.group()
                out.append(_el("m:f",
                               _el("m:fPr", _el("m:type", **{"m__val": "bar"})),
                               _el("m:num", *_parse(num)),
                               _el("m:den", *_parse(den))))
            elif cmd == "sqrt":
                flush()
                out.append(_el("m:rad",
                               _el("m:radPr", _el("m:degHide", **{"m__val": "1"})),
                               _el("m:deg"), _el("m:e", *_parse(tk.group()))))
            elif cmd in ("sum", "prod", "int"):
                flush()
                ch = {"sum": "∑", "prod": "∏", "int": "∫"}[cmd]
                sub = sup = None
                while tk.peek() in ("_", "^"):
                    k = tk.next()
                    g = tk.group()
                    if k == "_":
                        sub = g
                    else:
                        sup = g
                if sub is not None and sup is not None:
                    out.append(_el("m:nary",
                                   _el("m:naryPr", _el("m:chr", **{"m__val": ch}),
                                       _el("m:limLoc", **{"m__val": "undOvr"}),
                                       _el("m:grow", **{"m__val": "1"})),
                                   _el("m:sub", *_parse(sub)),
                                   _el("m:sup", *_parse(sup)),
                                   _el("m:e")))
                elif sub is not None:
                    out.append(_el("m:nary",
                                   _el("m:naryPr", _el("m:chr", **{"m__val": ch}),
                                       _el("m:limLoc", **{"m__val": "undOvr"}),
                                       _el("m:supHide", **{"m__val": "1"}),
                                       _el("m:grow", **{"m__val": "1"})),
                                   _el("m:sub", *_parse(sub)),
                                   _el("m:sup"), _el("m:e")))
                else:
                    buf += ch
            elif cmd in ("hat", "bar", "tilde", "vec", "dot", "overline"):
                flush()
                ch = {"hat": "\u0302", "bar": "\u0304", "overline": "\u0304",
                      "tilde": "\u0303", "vec": "\u20d7", "dot": "\u0307"}[cmd]
                out.append(_el("m:acc",
                               _el("m:accPr", _el("m:chr", **{"m__val": ch})),
                               _el("m:e", *_parse(tk.group()))))
            elif cmd in ("text", "mathrm", "operatorname", "mathbf"):
                flush()
                out.append(_r(tk.group(), italic=False))
            elif cmd in ("left", "right"):
                buf += tk.group()
            elif cmd in GREEK:
                buf += GREEK[cmd]
            elif cmd in SYM:
                buf += SYM[cmd]
            elif cmd in FUNCS:
                flush(); out.append(_r(cmd, italic=False))
            elif cmd == "\\":
                pass
            else:
                buf += cmd
        elif c in ("^", "_"):
            k = tk.next()
            if buf:
                if buf[-1] in ")]}":
                    close = buf[-1]
                    op = {")": "(", "]": "[", "}": "{"}[close]
                    depth, k0 = 0, None
                    for idx in range(len(buf) - 1, -1, -1):
                        if buf[idx] == close:
                            depth += 1
                        elif buf[idx] == op:
                            depth -= 1
                            if depth == 0:
                                k0 = idx; break
                    if k0 is None:
                        k0 = len(buf) - 1
                    base_txt, buf = buf[k0:], buf[:k0]
                else:
                    base_txt, buf = buf[-1], buf[:-1]
                flush()
                base = _r(base_txt)
            else:
                base = out.pop() if out else _r("")
            g = tk.group()
            tag = "m:sSup" if k == "^" else "m:sSub"
            part = "m:sup" if k == "^" else "m:sub"
            # 支持连续的 _{}^{}
            second = None
            if tk.peek() in ("^", "_") and tk.peek() != k:
                k2 = tk.next(); second = (k2, tk.group())
            if second is None:
                out.append(_el(tag, _el("m:e", base), _el(part, *_parse(g))))
            else:
                sub_s = g if k == "_" else second[1]
                sup_s = g if k == "^" else second[1]
                out.append(_el("m:sSubSup", _el("m:e", base),
                               _el("m:sub", *_parse(sub_s)),
                               _el("m:sup", *_parse(sup_s))))
        else:
            buf += tk.next()
    flush()
    return out


def omath(latex):
    """返回一个 m:oMath 元素。"""
    return _el("m:oMath", *_parse(latex))


def add_equation(doc, latex, number=None, indent_cm=0.0):
    """插入居中公式段落；number 不为空时在行末以圆括号给出编号。

    通过一个三格无框表格实现：左格空、中格公式（居中）、右格编号（右对齐），
    这样编号严格落在公式行末，且公式本身保持水平居中。
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import docbuild as D

    if number is None:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6); pf.space_after = Pt(6); pf.line_spacing = 1.5
        p._p.append(omath(latex))
        return p

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        D._set_border(t, edge, 0, "none")
    widths = [2.0, 11.0, 2.0]
    for j, w in enumerate(widths):
        t.rows[0].cells[j].width = Cm(w)
    cells = t.rows[0].cells
    for j, al in enumerate((WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.RIGHT)):
        p = cells[j].paragraphs[0]
        pf = p.paragraph_format
        pf.alignment = al
        pf.space_before = Pt(6); pf.space_after = Pt(6); pf.line_spacing = 1.5
    cells[1].paragraphs[0]._p.append(omath(latex))
    D._font(cells[2].paragraphs[0].add_run(f"（{number}）"), D.SONG, Pt(12))
    return t
