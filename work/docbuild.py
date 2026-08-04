# -*- coding: utf-8 -*-
"""Word 生成工具：严格执行指定的字体字号规范。

一级标题（章）：黑体 小三(15pt) 加粗
二级标题（节）：黑体 四号(14pt)
三级及以下（条）：黑体 小四(12pt)
正文：宋体 小四(12pt)，段前段后间距为 0，首行缩进 2 字符，1.5 倍行距
图名：置于图下方，五号(10.5pt) 黑体，居中
表名：置于表上方，五号(10.5pt) 黑体，居中；表格为三线表
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

HEI = "黑体"
SONG = "宋体"
LATIN = "Times New Roman"

SZ_H1 = Pt(15)      # 小三
SZ_H2 = Pt(14)      # 四号
SZ_H3 = Pt(12)      # 小四
SZ_BODY = Pt(12)    # 小四
SZ_CAP = Pt(10.5)   # 五号


def _font(run, cn, size, bold=False, latin=LATIN, color=None):
    run.font.name = latin
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), cn)


def _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0,
          line=1.5, indent=None, keep_next=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = indent
    if keep_next:
        p.paragraph_format.keep_with_next = True
    return p


# ---------- 富文本：**加粗**；^[n] 渲染为上标角标引用 ----------
def _add_rich(p, text, cn, size, bold=False):
    """把 **粗体** 与 ^[1] / ^[1-3] / ^[1,3] 形式的角标拆开渲染。"""
    for seg in re.split(r"(\*\*.+?\*\*|\^\[[0-9,\-]+\])", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            _font(p.add_run(seg[2:-2]), cn, size, bold=True)
        elif seg.startswith("^[") and seg.endswith("]"):
            r = p.add_run(seg[1:])
            _font(r, cn, size, bold=False)
            r.font.superscript = True
        else:
            _font(p.add_run(seg), cn, size, bold=bold)


def h1(doc, text, before=12, after=6):
    p = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, before, after, 1.5, keep_next=True)
    _font(p.add_run(text), HEI, SZ_H1, bold=True)
    return p


def h2(doc, text, before=8, after=4):
    p = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, before, after, 1.5, keep_next=True)
    _font(p.add_run(text), HEI, SZ_H2, bold=False)
    return p


def h3(doc, text, before=6, after=3):
    p = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, before, after, 1.5, keep_next=True)
    _font(p.add_run(text), HEI, SZ_H3, bold=False)
    return p


def body(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = _para(doc, align, 0, 0, 1.5, indent=Pt(24) if indent else Pt(0))
    _add_rich(p, text, SONG, SZ_BODY)
    return p


def title(doc, text):
    p = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, 0, 10, 1.5)
    _font(p.add_run(text), HEI, Pt(18), bold=True)
    return p


def subtitle(doc, text, size=Pt(12), cn=SONG, bold=False, after=0):
    p = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, 0, after, 1.5)
    _font(p.add_run(text), cn, size, bold=bold)
    return p


def figure(doc, path, caption, width_cm=14.0):
    """插入图片 + 图名（图下方，五号黑体，居中）。"""
    p = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, 6, 0, 1.0, keep_next=True)
    p.add_run().add_picture(path, width=Cm(width_cm))
    cp = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, 2, 8, 1.0)
    _font(cp.add_run(caption), HEI, SZ_CAP)
    return cp


def _set_border(tc_or_tbl, edge, sz, val="single"):
    el = tc_or_tbl._element
    pr = el.find(qn("w:tcPr")) if el.tag.endswith("tc") else el.find(qn("w:tblPr"))
    if pr is None:
        pr = OxmlElement("w:tcPr" if el.tag.endswith("tc") else "w:tblPr")
        el.insert(0, pr)
    tag = "w:tcBorders" if el.tag.endswith("tc") else "w:tblBorders"
    b = pr.find(qn(tag))
    if b is None:
        b = OxmlElement(tag); pr.append(b)
    e = b.find(qn(f"w:{edge}"))
    if e is None:
        e = OxmlElement(f"w:{edge}"); b.append(e)
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), "000000")


def table(doc, caption, header, rows, widths=None, fs=Pt(10.5),
          align_first_left=True, note=None):
    """三线表：表名在表上方（五号黑体居中）；顶线/表头下横线/底线，无竖线。"""
    cp = _para(doc, WD_ALIGN_PARAGRAPH.CENTER, 8, 2, 1.0, keep_next=True)
    _font(cp.add_run(caption), HEI, SZ_CAP)

    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    # 关闭默认边框
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set_border(t, edge, 0, "none")

    def fill(cells, vals, bold=False):
        for j, (c, v) in enumerate(zip(cells, vals)):
            c.text = ""
            p = c.paragraphs[0]
            pf = p.paragraph_format
            pf.alignment = (WD_ALIGN_PARAGRAPH.LEFT if (j == 0 and align_first_left)
                            else WD_ALIGN_PARAGRAPH.CENTER)
            pf.space_before = Pt(2); pf.space_after = Pt(2); pf.line_spacing = 1.0
            _font(p.add_run(str(v)), SONG, fs, bold=bold)

    fill(t.rows[0].cells, header, bold=True)
    for i, r in enumerate(rows):
        fill(t.rows[i + 1].cells, r)

    # 三线
    for c in t.rows[0].cells:
        _set_border(c, "top", 12)
        _set_border(c, "bottom", 6)
    for c in t.rows[-1].cells:
        _set_border(c, "bottom", 12)

    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    if note:
        np_ = _para(doc, WD_ALIGN_PARAGRAPH.LEFT, 2, 8, 1.0, indent=Pt(0))
        _font(np_.add_run(note), SONG, Pt(9))
    else:
        _para(doc, WD_ALIGN_PARAGRAPH.LEFT, 0, 8, 1.0)
    return t


def new_doc(margins=(2.5, 2.5, 3.0, 3.0)):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = LATIN
    st.font.size = SZ_BODY
    st.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)
    for s in doc.sections:
        s.top_margin = Cm(margins[0]); s.bottom_margin = Cm(margins[1])
        s.left_margin = Cm(margins[2]); s.right_margin = Cm(margins[3])
    return doc


def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p
