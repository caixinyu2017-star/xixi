# -*- coding: utf-8 -*-
"""核对 Word 文档是否符合指定的字体字号规范。"""
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def fonts_of(run):
    rpr = run._element.find(qn('w:rPr'))
    ea = asc = None
    if rpr is not None:
        rf = rpr.find(qn('w:rFonts'))
        if rf is not None:
            ea = rf.get(qn('w:eastAsia')); asc = rf.get(qn('w:ascii'))
    return ea, asc

def check(path):
    d = Document(path)
    print("=" * 70); print(path); print("=" * 70)
    seen = {}
    bad = []
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        for r in p.runs:
            ea, asc = fonts_of(r)
            sz = r.font.size.pt if r.font.size else None
            key = (ea, sz, bool(r.font.bold))
            seen[key] = seen.get(key, 0) + 1
        pf = p.paragraph_format
        # 正文（宋体小四）必须无段前段后
        if any(fonts_of(r)[0] == '宋体' and r.font.size and abs(r.font.size.pt - 12) < .01 for r in p.runs):
            sb = pf.space_before.pt if pf.space_before is not None else None
            sa = pf.space_after.pt if pf.space_after is not None else None
            if sb not in (0, None) or sa not in (0, None):
                bad.append((p.text[:28], sb, sa))
    print(f"{'字体':<8}{'字号':>8}{'加粗':>6}{'出现次数':>10}   说明")
    label = {('黑体', 15.0, True): '一级标题 黑体小三加粗 ✓',
             ('黑体', 14.0, False): '二级标题 黑体四号 ✓',
             ('黑体', 12.0, False): '三级标题 黑体小四 ✓',
             ('宋体', 12.0, False): '正文 宋体小四 ✓',
             ('宋体', 12.0, True): '正文内加粗强调',
             ('黑体', 10.5, False): '图名/表名 五号黑体 ✓',
             ('黑体', 18.0, True): '论文标题',
             ('宋体', 10.5, False): '表格内容 五号宋体',
             ('宋体', 10.5, True): '表头 五号宋体加粗',
             ('宋体', 9.0, False): '表注 小五宋体'}
    for k, v in sorted(seen.items(), key=lambda t: (-t[1])):
        print(f"{str(k[0]):<8}{str(k[1]):>8}{str(k[2]):>6}{v:>10}   {label.get(k,'')}")
    if bad:
        print("\n[!] 正文段落存在非零段前/段后间距：")
        for t, sb, sa in bad[:10]:
            print("   ", t, sb, sa)
    else:
        print("\n[✓] 所有宋体小四正文段落的段前段后间距均为 0")
    # 三线表检查
    for i, t in enumerate(d.tables):
        xml = t._element.xml
        has_v = 'w:insideV' in xml and 'w:val="none"' not in xml.split('w:insideV')[1][:80]
        print(f"[表{i+1}] 行数 {len(t.rows)} 列数 {len(t.columns)}；"
              f"竖线={'有(✗)' if has_v else '无(✓)'}")

for p in sys.argv[1:]:
    check(p)
