"""Build the revised manuscript from the submitted .docx.

The original file is edited in place (styles, OMML equations and embedded
figures are preserved) and every insertion or modification is colour-coded by
the reviewer that prompted it:

    Reviewer 1 -> red        Reviewer 2 -> blue        Reviewer 3 -> green

Output: out/MSSBOA_revised_marked.docx
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, '..')
FIGS = os.path.join(ROOT, 'figs')
OUT = os.path.join(ROOT, 'out')
os.makedirs(OUT, exist_ok=True)

import revision_text as T                                          # noqa: E402
from docx_edit import (para_by, find_para, insert_para_after,       # noqa: E402
                       insert_paras_after, insert_table_after,
                       insert_picture_after, replace_in_para,
                       strike_and_replace, strike_paragraph, COLOR_NAMES)
import make_tables as MT                                           # noqa: E402


# ------------------------------------------------------------------ helpers
def body_style(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == 'MDPI_3.1_text':
            return p
    return doc.paragraphs[11]


def h2_style(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == 'MDPI_2.2_heading2':
            return p
    return None


def cap_style(doc, kind='table'):
    want = 'MDPI_4.1_table_caption' if kind == 'table' else 'MDPI_5.1_figure_caption'
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith(want[:16]):
            return p
    return None


def add_heading_after(doc, para, text, color):
    h = insert_para_after(para, text, color, template=h2_style(doc))
    return h


def replace_picture_before(doc, caption_para, png, width, color):
    """Swap the image in the paragraph that precedes `caption_para`.

    Falls back to inserting a new picture paragraph if no image is found.
    """
    from docx.text.paragraph import Paragraph
    from docx_edit import insert_para_after
    prev = caption_para._p.getprevious()
    while prev is not None and prev.tag != caption_para._p.tag:
        prev = prev.getprevious()
    if prev is not None:
        holder = Paragraph(prev, caption_para._parent)
        if 'blip' in prev.xml or 'imagedata' in prev.xml:
            for r in list(holder.runs):
                r._element.getparent().remove(r._element)
            holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            holder.add_run().add_picture(png, width=width)
            print('  replaced the Figure 2 image in place')
            return holder
    holder = insert_para_after(caption_para, '', color)
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.add_run().add_picture(png, width=width)
    holder._p.addprevious(caption_para._p)
    print('  ! inserted Figure 2 as a new paragraph (no existing image found)')
    return holder


def fix_algorithm1_threshold(doc):
    """Reviewer 1, comment 6: Algorithm 1 said 30%, the experiments used 0.2N."""
    from docx_edit import COLORS
    n = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if 'worst 30%' in r.text:
                            r.text = r.text.replace(
                                'worst 30%', 'worst N_min = 0.2N individuals')
                            r.font.color.rgb = COLORS['R1']
                            n += 1
    print(f'  Algorithm 1 threshold corrected in {n} run(s)')


def data_or_note(rows, note):
    """Return the table rows, or a one-row placeholder if the run is unfinished."""
    if rows:
        return rows
    return [['Note'], [note]]


# ------------------------------------------------------------------- build
def build():
    doc = Document(os.path.join(ROOT, 'src', 'original.docx'))
    body = body_style(doc)
    tcap = cap_style(doc, 'table')
    fcap = cap_style(doc, 'figure')

    # collected numbering for the new floats
    # Provisional numbers in the 100s so they cannot collide with the
    # existing Tables 1-10; renumber_tables() rewrites them, and every in-text
    # reference to them, into document order at the end of the build.
    nums = dict(novelty_table_no='101', k_table_no='102',
                factorial_table_no='103', variants_table_no='104',
                loss_table_no='105', template_table_no='106',
                weights_table_no='107', scale_table_no='108')

    # ============================================================ REVIEWER 2
    # -- 2.1  why SBOA was chosen (Introduction, after the SBOA-variants para)
    p = para_by(doc, 'The secretary bird optimization algorithm (SBOA) is a bio-inspired')
    insert_paras_after(p, T.R2_WHY_SBOA, 'R2', template=body)

    # ============================================================ REVIEWER 1
    # -- 1.1  novelty framing in the abstract and the contribution list
    p = doc.paragraphs[7]
    replace_in_para(p, T.R1_ABSTRACT_OLD, T.R1_ABSTRACT_NEW, 'R1')
    p = para_by(doc, '(1) An improved SBOA variant')
    strike_and_replace(p, T.R1_CONTRIB_OLD, T.R1_CONTRIB_NEW, 'R1')

    # -- 1.2  Section 3.2: rigorous derivation of Eq. (11)
    p = para_by(doc, 'Opposition-based learning evaluates a candidate solution')
    last = insert_paras_after(p, T.R1_LOBL_DERIVATION, 'R1', template=body)

    # The submitted paragraph misstated the effect of a large k and gave the
    # LOBL threshold as 30%; strike it and follow it with the corrected text.
    p = para_by(doc, 'A small  in the early stage produces a refracted solution')
    strike_paragraph(p, 'R1')
    p = insert_para_after(p, T.R1_LOBL_CORRECTED, 'R1', template=body)

    # justification of the k schedule, after the paragraph that describes it
    ktxt = [t.format(k_table_no=nums['k_table_no'],
                     k_conclusion=KCONC) for t in T.R1_LOBL_KJUSTIFY]
    insert_paras_after(p, ktxt, 'R1', template=body)

    # replace the (unrendered) Figure 2 with the new schematic
    fig2 = para_by(doc, 'Figure 2. Schematic of the lens opposition-based learning')
    for r in fig2.runs:
        r.text = ''
    from docx_edit import append_run
    append_run(fig2, 'Figure 2. Convex-lens imaging behind the lens '
                     'opposition-based learning strategy: (a) the similar-triangle '
                     'construction that yields Equation (11), in which k = h/h* is '
                     'the ratio of object height to image height; (b) the dynamic '
                     'scaling factor of Equation (12) for several exponent pairs; '
                     '(c) the resulting refraction radius about the interval '
                     'midpoint.', 'R1')
    # The submitted Figure 2 was an EMF object that did not render for the
    # reviewers; replace the picture in place rather than adding a second one.
    replace_picture_before(doc, fig2, os.path.join(FIGS, 'fig02_lens_imaging.png'),
                           Inches(6.3), 'R1')
    insert_table_after(doc, fig2, K_TABLE, 'R1',
                       caption=f'Table {nums["k_table_no"]}. Friedman ranking '
                               'of alternative exponent pairs in the lens '
                               'scaling schedule k(t) = (1 + (t/T)^nu)^mu, '
                               'including the constant k = 1 that reduces the '
                               'operator to plain opposition-based learning.',
                       template_style='Table Grid')

    # -- 1.1b  new Section 3.6 (before Section 4)
    anchor = para_by(doc, 'For MSSBOA, the good point set initialization has the same complexity')
    h = add_heading_after(doc, anchor, T.R1_NOVELTY_HEADING, 'R1')
    txt = [t.format(novelty_table_no=nums['novelty_table_no'])
           for t in T.R1_NOVELTY_TEXT]
    last = insert_paras_after(h, txt, 'R1', template=body)
    insert_table_after(doc, last,
                       T.NOVELTY_TABLE, 'R1',
                       caption=f'Table {nums["novelty_table_no"]}. How the three '
                               'operator families are implemented in MSSBOA '
                               'compared with their existing applications.',
                       template_style='Table Grid')

    # -- 1.4  balanced reading of the Wilcoxon results
    p = para_by(doc, 'As shown in Table 7 and Figure 11, MSSBOA significantly outperforms')
    for r in p.runs:
        r.text = ''
    txt = [t.format(loss_table_no=nums['loss_table_no'])
           for t in T.R1_WILCOXON_BALANCED]
    last = insert_paras_after(p, txt, 'R1', template=body)
    after_loss = insert_table_after(
        doc, last, LOSS_TABLE, 'R1',
        caption=f'Table {nums["loss_table_no"]}. Distribution of the cases in '
                'which MSSBOA is outranked, by CEC2017 function class, pooled '
                'over the four dimensions (116 cases per competitor).',
        template_style='Table Grid')

    # -- 2.3 / 3.3  comparison with SBOA variants, MS-TSA and I-CPA (Sec. 4.5)
    h = add_heading_after(doc, after_loss, T.R2_VARIANTS_HEADING, 'R2')
    last = insert_paras_after(h, T.R2_VARIANTS_INTRO, 'R2', template=body)
    last = insert_paras_after(last, VARIANTS_PARAS, 'R2', template=body)
    insert_table_after(doc, last, VARIANTS_TABLE, 'R2',
                       caption=f'Table {nums["variants_table_no"]}. MSSBOA '
                               'against three recent SBOA variants and two '
                               'recent multi-strategy algorithms on the CEC2017 '
                               'suite.',
                       template_style='Table Grid')

    # -- 1.3  colour model: space, variables, template derivation
    p = para_by(doc, 'Generating a harmonious yet expressive color palette')
    txt = [t.format(template_table_no=nums['template_table_no'])
           for t in T.R1_COLOR_SPACE]
    last = insert_paras_after(p, txt, 'R1', template=body)
    insert_table_after(doc, last, T.TEMPLATE_TABLE, 'R1',
                       caption=f'Table {nums["template_table_no"]}. The seven '
                               'chromatic harmonic templates of Matsuda in the '
                               'parameterization of Cohen-Or et al. [70,72]. '
                               'Offsets are relative to the rotation angle '
                               'alpha, which is the only free parameter.',
                       template_style='Table Grid')

    # complete the wording of Eqs. (18) and (19)
    p = para_by(doc, 'A purely harmonic palette may collapse into nearly identical colors')
    for r in p.runs:
        r.text = ''
    append_run(p, T.R1_COLOR_EQ_TEXT, 'R1')
    p = para_by(doc, 'The decision vector is the stacked palette')
    for r in p.runs:
        r.text = ''
    append_run(p, T.R1_COLOR_OBJ_TEXT, 'R1')

    # -- 1.6  aesthetic weights
    p = para_by(doc, 'As shown in Table 10 and Figure 14, MSSBOA achieves the best mean')
    txt = list(T.R1_WEIGHTS_TEXT) + [WEIGHTS_PARA]
    last = insert_paras_after(p, txt, 'R1', template=body)
    last = insert_table_after(
        doc, last, WEIGHTS_TABLE, 'R1',
        caption=f'Table {nums["weights_table_no"]}. Mean layout aesthetic score '
                'over the eight design problems under five weight '
                'configurations of Equation (22). Kendall tau compares the '
                'algorithm ordering with that obtained under W0.',
        template_style='Table Grid')

    # -- 1.12  scalability of the layout problem
    last = insert_para_after(last, '5.3. Scalability of the layout problem', 'R1',
                             template=h2_style(doc))
    last = insert_paras_after(last, SCALE_PARAS, 'R1', template=body)
    insert_table_after(doc, last, SCALE_TABLE, 'R1',
                       caption=f'Table {nums["scale_table_no"]}. Mean aesthetic '
                               'score (x100) on layout problems with 8 to 30 '
                               'elements, i.e. up to D = 60.',
                       template_style='Table Grid')

    # -- 1.1c  carry the integration-framework framing into the Conclusions
    p = para_by(doc, 'This paper proposed a multi-strategy secretary bird')
    strike_and_replace(p, T.R1_CONCL_FRAMING_OLD, T.R1_CONCL_FRAMING_NEW, 'R1')

    # -- 1.7  limitations, split into inherited and variant-specific
    p = para_by(doc, 'Nevertheless, MSSBOA has some limitations.')
    for r in p.runs:
        r.text = ''
    insert_paras_after(p, T.R1_LIMITATIONS, 'R1', template=body)

    # -- 1.9 / 3.5  code and data availability
    p = para_by(doc, 'Data Availability Statement:')
    for r in p.runs:
        r.text = ''
    append_run(p, T.R1_DATA_AVAILABILITY, 'R1')

    # ============================================================ REVIEWER 3
    # -- 3.1 / 3.2  full-factorial ablation and the synergy argument
    p = para_by(doc, 'To investigate the contribution of each improvement strategy')
    last = insert_paras_after(p, T.R3_ABLATION_INTRO, 'R3', template=body)
    p2 = para_by(doc, 'According to Table 5 and Figure 6, all p-values are far below')
    last = insert_paras_after(p2, FACTORIAL_PARAS, 'R3', template=body)
    insert_table_after(doc, last, FACTORIAL_TABLE, 'R3',
                       caption=f'Table {nums["factorial_table_no"]}. Full 2^3 '
                               'factorial ablation of the three strategies on '
                               'the CEC2017 suite.',
                       template_style='Table Grid')

    # -- 3.4  redrawn Figures 4 and 5
    for cap_needle, png, newcap in (
            ('Figure 4. Friedman rankings of MSSBOA with different',
             'fig04_nmin_ranking.png',
             'Figure 4. Friedman rankings of MSSBOA with different N_min '
             '(alpha = 0.05). The best setting is marked above the '
             'corresponding point.'),
            ('Figure 5. Friedman rankings of MSSBOA with different p',
             'fig05_pm_ranking.png',
             'Figure 5. Friedman rankings of MSSBOA with different p_m '
             '(alpha = 0.05). The best setting is marked above the '
             'corresponding point.')):
        i = find_para(doc, cap_needle)
        if i is None:
            continue
        cp = doc.paragraphs[i]
        for r in cp.runs:
            r.text = ''
        append_run(cp, newcap, 'R3')
        path = os.path.join(FIGS, png)
        if os.path.exists(path):
            holder = insert_para_after(cp, '', 'R3')
            holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            holder.add_run().add_picture(path, width=Inches(5.4))
            holder._p.addprevious(cp._p)

    # ------------------------------------------------- typography and notes
    fix_typography(doc)
    fix_algorithm1_threshold(doc)
    add_new_references(doc)
    add_colour_key(doc)
    renumber_tables(doc)

    out = os.path.join(OUT, 'MSSBOA_revised_marked.docx')
    doc.save(out)
    print('wrote', out)
    return out


# -------------------------------------------------------------- small fixes
EN = '–'

# Notation is normalized to the MDPI form "Cauchy-Gaussian" with an en dash
# (the submitted file mixed an en dash, a hyphen and the truncated
# "Cauchy-Gauss"), and any residual HTML artefact from the table conversion is
# stripped.
SUBS = [
    ('CauchyGaussian', f'Cauchy{EN}Gaussian'),
    ('Cauchy-Gaussian', f'Cauchy{EN}Gaussian'),
    ('Cauchy-Gauss ', f'Cauchy{EN}Gaussian '),
    (f'Cauchy{EN}Gauss ', f'Cauchy{EN}Gaussian '),
    ('Cauchy-Gauss mutation', f'Cauchy{EN}Gaussian mutation'),
    ('td rowspan', ''),
    ('</td>', ''),
    ('<td>', ''),
    ('rowspan=', ''),
]


def _fix_runs(paras, counter):
    for p in paras:
        for r in p.runs:
            t = r.text
            if not t:
                continue
            new = t
            for a, b in SUBS:
                if a in new:
                    new = new.replace(a, b)
            if new != t:
                r.text = new
                counter[0] += 1


def fix_typography(doc):
    """Reviewer 1, comment 8: notation and formatting consistency."""
    n = [0]
    _fix_runs(doc.paragraphs, n)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _fix_runs(cell.paragraphs, n)
    # make the header row of every appendix table repeat cleanly across pages
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rep = 0
    for tbl in doc.tables:
        if len(tbl.rows) < 20:
            continue
        tr = tbl.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            trPr.append(OxmlElement('w:tblHeader'))
            rep += 1
    print(f'  typography fixes: {n[0]} run(s); repeating headers set on {rep} table(s)')


def renumber_tables(doc):
    """Renumber every main-text table into document order.

    New tables are inserted with provisional numbers in the 100s so that they
    cannot collide with the existing Tables 1-10.  Here the captions are walked
    in document order, a bijection old -> new is built, and both the captions
    and every in-text "Table n" reference are rewritten.  Appendix tables
    (Table A1-A8) are left alone.
    """
    import re
    cap_re = re.compile(r'^Table\s+(\d+)\s*\.')
    caps = []
    for p in doc.paragraphs:
        m = cap_re.match(p.text.strip())
        if m:
            caps.append((p, int(m.group(1))))
    mapping = {old: i + 1 for i, (_, old) in enumerate(caps)}
    if mapping == {k: k for k in mapping}:
        print('  table numbering already sequential')
        return

    for p, old in caps:                      # rewrite the captions themselves
        new = mapping[old]
        for r in p.runs:
            if f'Table {old}.' in r.text:
                r.text = r.text.replace(f'Table {old}.', f'Table {new}.', 1)
                break

    ref_re = re.compile(r'\bTable\s+(\d+)\b')

    def fix(text):
        return ref_re.sub(
            lambda m: (f'Table {mapping[int(m.group(1))]}'
                       if int(m.group(1)) in mapping else m.group(0)), text)

    n = 0
    cap_paras = {id(p) for p, _ in caps}
    for p in doc.paragraphs:
        if id(p) in cap_paras:
            continue
        for r in p.runs:
            if 'Table' in r.text:
                new = fix(r.text)
                if new != r.text:
                    r.text = new
                    n += 1
    print(f'  renumbered {len(caps)} tables; updated {n} in-text reference run(s)')
    print('   mapping:', {k: v for k, v in sorted(mapping.items()) if k != v})


NEW_REFS = [
    '81. Safari, A.; Varaee, H. Opposition-Based Ideal Gas Molecular Movement '
    'Algorithm with Cauchy Mutation, Velocity Clamping, and Mirror Operator. '
    'Expert Syst. 2023, 40, e13306, doi:10.1111/exsy.13306.',
    '82. Beskirli, A.; Dag, I.; Kiran, M.S. A Tree Seed Algorithm with '
    'Multi-Strategy for Parameter Estimation of Solar Photovoltaic Models. '
    'Appl. Soft Comput. 2024, 167, 112220, doi:10.1016/j.asoc.2024.112220.',
    '83. Beskirli, A.; Dag, I. I-CPA: An Improved Carnivorous Plant Algorithm '
    'for Solar Photovoltaic Parameter Identification Problem. Biomimetics 2023, '
    '8, 569, doi:10.3390/biomimetics8080569.',
]


def add_new_references(doc):
    i = find_para(doc, '80. Jiang, Z.; Xia, J.')
    if i is None:
        print('  ! reference anchor not found')
        return
    p = doc.paragraphs[i]
    colors = ['R1', 'R3', 'R3']
    for ref, col in zip(NEW_REFS, colors):
        p = insert_para_after(p, ref, col, template=doc.paragraphs[i])
    print('  added 3 references')


def add_colour_key(doc):
    p = doc.paragraphs[8]           # after the keywords line
    key = insert_para_after(
        p, 'Note to the editor and reviewers: all revisions are marked by '
           'colour - text added or modified in response to Reviewer 1 is in '
           'red, to Reviewer 2 in blue, and to Reviewer 3 in green. Deleted '
           'wording is shown struck through in the colour of the reviewer who '
           'prompted its removal.', 'R1', template=doc.paragraphs[11])
    key.runs[0].font.italic = True
    return key


# =====================================================================
#  Data-dependent blocks, filled from the experiment tables
# =====================================================================
def _load_all():
    global LOSS_TABLE, VARIANTS_TABLE, VARIANTS_PARAS, FACTORIAL_TABLE
    global FACTORIAL_PARAS, WEIGHTS_TABLE, WEIGHTS_PARA, SCALE_TABLE
    global SCALE_PARAS, KCONC, K_TABLE, PARAM_TABLES

    import csv
    LOSS_TABLE = []
    lp = os.path.join(ROOT, 'results', 'loss_by_class.csv')
    if os.path.exists(lp):
        with open(lp, newline='') as fh:
            LOSS_TABLE = [r for r in csv.reader(fh)]
    if not LOSS_TABLE:
        LOSS_TABLE = [['(loss analysis not available)']]

    PENDING = ('[Experiment still running at the time this draft was '
               'generated - the table is produced by code/make_tables.py from '
               'results/ and will be filled before submission.]')

    v = MT.table_variants()
    if v:
        VARIANTS_TABLE, meta = v
        VARIANTS_PARAS = [
            f'Table 11 reports the Friedman ranking of every algorithm and the '
            f'Wilcoxon rank-sum outcome of MSSBOA against each of them over the '
            f'{meta["funcs"][meta["dims"][0]]} CEC2017 functions, together with '
            f'the win/tie/loss counts.']
    else:
        VARIANTS_TABLE = [['Status'], [PENDING]]
        VARIANTS_PARAS = [PENDING]

    f = MT.table_factorial()
    if f:
        FACTORIAL_TABLE, meta = f
        eff = meta['effects']
        FACTORIAL_PARAS = [
            f'Table 7 reports the eight configurations. Averaged over the other '
            f'two factors, the main effect of each strategy on the mean '
            f'Friedman rank is {eff["GPSI"]:+.3f} for GPSI, {eff["LOBL"]:+.3f} '
            f'for LOBL and {eff["ACGM"]:+.3f} for ACGM, where a negative value '
            f'means the strategy improves the ranking. The three-way '
            f'interaction is {meta["interaction"]:+.3f}.']
    else:
        FACTORIAL_TABLE = [['Status'], [PENDING]]
        FACTORIAL_PARAS = [PENDING]

    w = MT.table_weights()
    if w:
        WEIGHTS_TABLE, _ = w
        WEIGHTS_PARA = (
            'To test how far the conclusions depend on this choice, the eight '
            'design problems were re-optimized under five weight '
            'configurations: the Ngo-based setting W0 used in the paper, equal '
            'weights, and three configurations that each promote one term. '
            'Table 16 reports the mean score of each algorithm under each '
            'configuration together with Kendall\'s tau between the resulting '
            'algorithm ordering and the ordering under W0.')
    else:
        WEIGHTS_TABLE = [['Status'], [PENDING]]
        WEIGHTS_PARA = PENDING

    s = MT.table_scale()
    if s:
        SCALE_TABLE, meta = s
        SCALE_PARAS = [
            'The eight problems of Table 9 involve at most six elements, so the '
            'largest decision vector has 12 components. To establish whether '
            'the method scales, the layout problem was instantiated with 8, 12, '
            '20 and 30 elements, giving problems of 16 to 60 dimensions, with '
            'total element area held at about half the canvas so that the '
            'density term remains comparable. Table 17 reports the outcome.']
    else:
        SCALE_TABLE = [['Status'], [PENDING]]
        SCALE_PARAS = [
            'The eight problems of Table 9 involve at most six elements, so the '
            'largest decision vector has 12 components. To establish whether '
            'the method scales, the layout problem was instantiated with 8, 12, '
            '20 and 30 elements, giving problems of 16 to 60 dimensions. '
            + PENDING]

    K_TABLE = [['Status'], [PENDING]]
    PARAM_TABLES = []
    KCONC = ('the adopted pair is compared there with alternative roots and '
             'powers, including the constant k = 1 that reduces the operator to '
             'plain opposition-based learning.')
    pg = MT.table_params()
    if pg:
        PARAM_TABLES = pg
        for title, rows, meta in pg:
            if title.startswith('Lens scaling'):
                K_TABLE = rows
                best = min(rows[1:-1], key=lambda r: float(r[1]))
                KCONC = (f'the best mean rank over the '
                         f'{meta["nfuncs"]} functions is obtained by '
                         f'"{best[0]}", and the adopted pair is compared there '
                         f'with alternative roots and powers, including the '
                         f'constant k = 1 that reduces the operator to plain '
                         f'opposition-based learning.')


if __name__ == '__main__':
    _load_all()
    build()
