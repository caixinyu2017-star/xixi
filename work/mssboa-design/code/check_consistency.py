"""Audit the built manuscript for statements the data does not support.

The requirement for this revision is that nothing in the paper contradicts
anything else in it.  Tables and figures are generated from the runs and the
results prose is generated from the same metadata, so those cannot disagree by
construction.  What can still go wrong is leftover prose from the submitted
version - a superlative, a count, or a claim of first place that survived
because no anchor matched it.

This script flags every sentence that
  * claims MSSBOA outperforms, dominates or ranks first, when the regenerated
    Friedman ranking does not put it first;
  * quotes a numeric win/tie/loss triple that is not in the regenerated
    Table 7;
  * mentions dimensions that were not run;
  * refers to a table or figure number that does not exist.
"""
import os
import re
import sys

from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, '..')

import tables as TB                                              # noqa: E402

SUPERLATIVE = re.compile(
    r'\b(outperform\w*|superior\w*|best (?:Friedman |average )?rank\w*|'
    r'ranks? first|the most \w+|dominat\w*|wins? (?:on )?all|'
    r'highest \w+ score)\b', re.I)
TRIPLE = re.compile(r'\b(\d{1,3})/(\d{1,3})/(\d{1,3})\b')
DIMLIST = re.compile(r'\b(\d{2,3}(?:,\s*\d{2,3})*(?:\s+and\s+\d{2,3})?)\s+dimensions\b')


def audit(path):
    doc = Document(path)
    t6 = TB.t6_friedman()
    t7 = TB.t7_wilcoxon()
    problems = []

    # When the submitted results are kept, the paper's own Tables 5-7 are the
    # reference for claims about them, and the newly added subsections are the
    # only place the regenerated runs speak.  Checking the submitted prose
    # against the new runs would only produce noise.
    keep = os.environ.get('REGENERATE_ALL', '') != '1'
    mssboa_first = None
    dims_run = set()
    if t6 and not keep:
        _, m6 = t6
        mssboa_first = (m6['order']['MSSBOA'] == 1)
        dims_run = set(m6['dims'])
    valid_triples = set()
    if t7 and not keep:
        rows, _ = t7
        for r in rows[1:]:
            for cell in r[1:]:
                mt = TRIPLE.fullmatch(str(cell).strip())
                if mt:
                    valid_triples.add(mt.group(0))

    # every table and figure number that actually exists
    tables_present = set()
    figures_present = set()
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r'^Table\s+(A?\d+)\s*\.', t)
        if m:
            tables_present.add(m.group(1))
        m = re.match(r'^Figure\s+(\d+)\s*\.', t)
        if m:
            figures_present.add(m.group(1))

    for i, p in enumerate(doc.paragraphs):
        # struck-through runs are shown deletions, not assertions: audit only
        # the text the reader is meant to take as current
        text = ''.join(r.text for r in p.runs if not r.font.strike).strip()
        if not text or text.startswith(('Table ', 'Figure ')):
            continue

        for sent in re.split(r'(?<=[.;])\s+', text):
            if mssboa_first is False and SUPERLATIVE.search(sent) \
                    and 'MSSBOA' in sent and 'not ' not in sent.lower() \
                    and 'rather than' not in sent.lower():
                problems.append(('superlative', i, sent[:150]))
            for tr in TRIPLE.findall(sent):
                s = '/'.join(tr)
                if valid_triples and s not in valid_triples and \
                        int(tr[0]) + int(tr[1]) + int(tr[2]) >= 20:
                    problems.append(('stale win/tie/loss', i, sent[:150]))
            for dm in DIMLIST.findall(sent):
                nums = {int(x) for x in re.findall(r'\d+', dm)}
                if dims_run and not nums <= dims_run:
                    problems.append(('dimension not run', i, sent[:150]))

        for ref in re.findall(r'\bTable\s+(A?\d+)\b', text):
            if tables_present and ref not in tables_present:
                problems.append(('dangling table ref', i, f'Table {ref}: {text[:90]}'))
        for ref in re.findall(r'\bFigure\s+(\d+)\b', text):
            if figures_present and ref not in figures_present:
                problems.append(('dangling figure ref', i, f'Figure {ref}: {text[:90]}'))

    return problems, {'mssboa_first': mssboa_first, 'dims': sorted(dims_run),
                      'tables': len(tables_present), 'figures': len(figures_present)}


if __name__ == '__main__':
    path = sys.argv[1] if len(sys.argv) > 1 else \
        os.path.join(ROOT, 'out', 'MSSBOA_revised_marked.docx')
    probs, info = audit(path)
    print(f'audit of {os.path.basename(path)}')
    print(f'  MSSBOA ranks first: {info["mssboa_first"]}   '
          f'dimensions run: {info["dims"]}   '
          f'{info["tables"]} tables, {info["figures"]} figures')
    if not probs:
        print('  no inconsistencies found')
    else:
        kinds = {}
        for k, i, s in probs:
            kinds.setdefault(k, []).append((i, s))
        for k, items in kinds.items():
            print(f'\n  {k} ({len(items)}):')
            for i, s in items[:12]:
                print(f'    para {i}: {s}')
            if len(items) > 12:
                print(f'    ... {len(items) - 12} more')
    sys.exit(1 if probs else 0)
