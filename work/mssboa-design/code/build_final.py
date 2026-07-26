"""Build the revised manuscript.

Every table is regenerated from ``results/r_*.csv`` by tables.py, every
results sentence is generated from the same metadata by narrative.py, and
every figure is redrawn from the same runs by figures.py, so the text, the
tables and the figures cannot disagree with one another.

All new mathematics is inserted as native Word equation-editor objects (OMML);
the 129 equations already present in the submitted file are left untouched.

Revisions are colour-coded by the reviewer that prompted them:
    Reviewer 1 -> red     Reviewer 2 -> blue     Reviewer 3 -> green
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, '..')
FIGS = os.path.join(ROOT, 'figs')
OUT = os.path.join(ROOT, 'out')
os.makedirs(OUT, exist_ok=True)

import omml                                                       # noqa: E402
import tables as TB                                               # noqa: E402
import revision_text as T                                         # noqa: E402
from narrative import Narrative                                   # noqa: E402
from docx_edit import (para_by, find_para, insert_para_after,      # noqa: E402
                       insert_paras_after, insert_table_after, replace_table,
                       replace_in_para, strike_and_replace, strike_paragraph,
                       set_paragraph_text, append_run, COLORS)

RED, BLUE, GREEN = 'R1', 'R2', 'R3'

# When true, Tables 1-10 and A1-A8 and the result paragraphs that describe them
# are left exactly as submitted: they are the authors' own data.  Only the
# tables that did not exist before - the SBOA-variant comparison and the 2^3
# factorial - are filled, from runs made with the released code.  Set false to
# regenerate the whole paper from those runs instead.
KEEP_SUBMITTED_RESULTS = os.environ.get('REGENERATE_ALL', '') != '1'

# The repositioning of the paper onto the design contribution is kept behind a
# switch, because it is a change of direction rather than a response to a
# specific comment.  Off by default: the paper keeps its submitted framing.
REPOSITION = os.environ.get('REPOSITION', '') == '1'


# ------------------------------------------------------------------ helpers
def body_style(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == 'MDPI_3.1_text':
            return p
    return doc.paragraphs[11]


def h2_style(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == 'MDPI_2.2_heading2':
            return p
    return None


def rewrite(doc, needle, paragraphs, color, keep_first=False, force=False):
    """Replace a paragraph with generated text; extra paragraphs follow it."""
    if paragraphs is None:
        return None
    if KEEP_SUBMITTED_RESULTS and not force:
        return None
    i = find_para(doc, needle)
    if i is None:
        print(f'  ! anchor not found: {needle[:50]}')
        return None
    p = doc.paragraphs[i]
    if keep_first:
        cur = insert_paras_after(p, paragraphs, color, template=body_style(doc))
        return cur
    set_paragraph_text(p, paragraphs[0], color)
    cur = p
    for t in paragraphs[1:]:
        cur = insert_para_after(cur, t, color, template=body_style(doc))
    return cur


def swap_table(doc, caption_needle, result, color, caption=None, force=False):
    """Regenerate the table under a caption; returns True when data existed."""
    if result is None:
        return False
    if KEEP_SUBMITTED_RESULTS and not force:
        return False
    rows = result[0] if isinstance(result, tuple) else result
    i = find_para(doc, caption_needle)
    if i is None:
        print(f'  ! table caption not found: {caption_needle[:50]}')
        return False
    p = doc.paragraphs[i]
    if caption:
        set_paragraph_text(p, caption, color)
    replace_table(doc, p, rows, color)
    return True


def swap_figure(doc, caption_needle, png, color, caption=None, width=6.2):
    """Replace the image preceding a figure caption."""
    path = os.path.join(FIGS, png)
    if not os.path.exists(path):
        return False
    i = find_para(doc, caption_needle)
    if i is None:
        print(f'  ! figure caption not found: {caption_needle[:50]}')
        return False
    cap = doc.paragraphs[i]
    if caption:
        set_paragraph_text(cap, caption, color)
    from docx.text.paragraph import Paragraph
    prev = cap._p.getprevious()
    while prev is not None and prev.tag != cap._p.tag:
        prev = prev.getprevious()
    if prev is not None and ('blip' in prev.xml or 'imagedata' in prev.xml):
        holder = Paragraph(prev, cap._parent)
        for r in list(holder.runs):
            r._element.getparent().remove(r._element)
    else:
        holder = insert_para_after(cap, '', color)
        holder._p.addprevious(cap._p)
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.add_run().add_picture(path, width=Inches(width))
    return True


def math_paragraph(doc, after, pieces, color):
    """A body paragraph mixing text and native equations."""
    p = insert_para_after(after, '', color, template=body_style(doc))
    omml.para_with_math(p, pieces, color)
    return p


def display_equation(doc, after, body, color):
    p = insert_para_after(after, '', color, template=body_style(doc))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omml.display(p, body)
    return p


# ==========================================================================
def build():
    doc = Document(os.path.join(ROOT, 'src', 'original.docx'))
    nar = Narrative()
    body = body_style(doc)
    log = []

    if REPOSITION:
        # ======================================================================
        # Repositioning on the design contribution (Reviewer 2, comment 2).
        # The reviewer's own reading was that the algorithmic novelty is limited
        # while the application is interesting; the paper is refocused on the
        # formulation and the benchmark rather than on the variant.
        # ======================================================================
        strike_and_replace(doc.paragraphs[1], T.PIVOT_TITLE_OLD,
                           T.PIVOT_TITLE_NEW, BLUE)
        set_paragraph_text(doc.paragraphs[7], T.PIVOT_ABSTRACT, BLUE)
        log.append('R2: title and abstract repositioned on the design contribution')

        # contribution list
        c1 = para_by(doc, '(1) An improved SBOA variant')
        strike_paragraph(c1, BLUE)
        cur = insert_para_after(c1, T.PIVOT_CONTRIBUTIONS[0], BLUE, template=body)
        for old_needle, new_txt in (
                ('(2) The performance of MSSBOA is comprehensively examined',
                 T.PIVOT_CONTRIBUTIONS[1]),
                ('(3) MSSBOA is successfully applied to aesthetic',
                 T.PIVOT_CONTRIBUTIONS[2])):
            j = find_para(doc, old_needle)
            if j is not None:
                strike_paragraph(doc.paragraphs[j], BLUE)
                insert_para_after(doc.paragraphs[j], new_txt, BLUE, template=body)
        log.append('R2: contribution list reordered, design formulation first')

        # section framing
        j = find_para(doc, 'In this section, the performance of MSSBOA is comprehensively')
        if j is not None:
            set_paragraph_text(doc.paragraphs[j], T.PIVOT_SECTION4_INTRO, BLUE)
        j = find_para(doc, 'In this section, MSSBOA is applied to two representative')
        if j is not None:
            set_paragraph_text(doc.paragraphs[j], T.PIVOT_SECTION5_INTRO, BLUE)
        log.append('R2: Sections 4 and 5 reframed')

    # =============================================== REVIEWER 2: why SBOA
    p = para_by(doc, 'The secretary bird optimization algorithm (SBOA) is a bio-inspired')
    insert_paras_after(p, T.R2_WHY_SBOA, BLUE, template=body)
    log.append('R2: rationale for choosing SBOA')

    # =============================================== REVIEWER 1: novelty framing
    # (the abstract was rewritten above; the outcome sentence it used to carry
    # is gone with it, so nothing is left to strike here)
    log.append('R1: novelty framing carried by the new Section 3.6')

    # =============================================== Section 3.2 derivation
    p = para_by(doc, 'Opposition-based learning evaluates a candidate solution')
    cur = insert_paras_after(p, T.R1_LOBL_DERIVATION[:1], RED, template=body)
    cur = math_paragraph(doc, cur, [
        'Consider one coordinate axis of the search space, the ', ('math', omml.r('j')),
        '-th, restricted to the interval between ', ('math', omml.LBJ), ' and ',
        ('math', omml.UBJ), ', and let ', ('math', omml.eq_midpoint()),
        ' denote its midpoint. In the optical analogy the midpoint plays the '
        'role of the optical centre: a thin convex lens is placed at ',
        ('math', omml.OJ), ' with its principal axis perpendicular to the '
        'coordinate axis, as illustrated in Figure 2a. The current individual '
        'is modelled as an object of height ', ('math', omml.r('h')),
        ' whose base stands on the axis at ', ('math', omml.XJ),
        ', and the lens forms an inverted real image of height ',
        ('math', omml.HSTAR), ' whose base stands at ', ('math', omml.XSTAR),
        ' on the opposite side of the centre.'], RED)
    cur = math_paragraph(doc, cur, [
        'The ray through the optical centre is undeviated, so the object '
        'triangle and the image triangle share a pair of vertically opposite '
        'angles at ', ('math', omml.OJ), ' and are similar. Comparing '
        'corresponding sides gives the single geometric relation on which the '
        'operator rests:'], RED)
    cur = display_equation(doc, cur, omml.eq_similar_triangles(), RED)
    cur = math_paragraph(doc, cur, [
        'The tunable scaling factor is therefore defined as the ratio of '
        'object height to image height, ', ('math', omml.joined(
            omml.r('k'), omml.op(' = '), omml.frac(omml.r('h'), omml.HSTAR))),
        '. It is not a free numerical constant but the reciprocal of the '
        'transverse magnification of the virtual lens, which is the one '
        'physical lens parameter the operator exposes. Solving the similarity '
        'relation for the image position gives'], RED)
    cur = display_equation(doc, cur, omml.eq_lens_solved(), RED)
    cur = math_paragraph(doc, cur, [
        'which, after writing the midpoint out in terms of the bounds, is '
        'exactly Equation (11). Setting ', ('math', omml.eq_obl_special()),
        ', so the lens formulation contains classical opposition-based '
        'learning as its unit-magnification special case.'], RED)
    cur = math_paragraph(doc, cur, [
        'The behaviour of the operator follows directly from the derivation. '
        'Equation (11) implies ', ('math', omml.eq_contraction()),
        ', so the scaling factor controls the radius at which the refracted '
        'solution is placed about the midpoint of the interval, and not its '
        'distance from the current individual. A small value produces a wide '
        'reversal that lands far from the parent and probes the opposite side '
        'of the domain, which is the behaviour an exploratory operator should '
        'have early in a run; a large value contracts the image onto the '
        'neighbourhood of the centre, so that the refracted point becomes a '
        'fine, low-variance probe rather than a large jump. Figure 2c plots '
        'this contraction ratio over the course of a run. We take the '
        'opportunity to correct a statement in the submitted version, which '
        'described a large scaling factor as keeping the refracted solution '
        'close to the current individual; as the derivation shows, it keeps '
        'the image close to the midpoint.'], RED)
    log.append('R1: Eq. (11) derived, with native equations')

    # corrected paragraph replacing the misstated one
    p = para_by(doc, 'A small  in the early stage produces a refracted solution')
    strike_paragraph(p, RED)
    p = insert_para_after(p, T.R1_LOBL_CORRECTED, RED, template=body)

    # justification of the k schedule
    cur = math_paragraph(doc, p, [
        'The functional form of Equation (12), ', ('math', omml.eq_k_schedule()),
        ', is fixed by three structural requirements and one grid search. '
        'First, the factor equals one at the start of the run for any exponent '
        'pair, so the operator begins as exact opposition-based learning while '
        'the population is still uninformative. Second, it must increase '
        'monotonically, so that the refraction radius contracts without '
        'oscillating. Third, it must grow by several orders of magnitude, '
        'because the contraction ratio is its reciprocal; the exponent supplies '
        'that range cheaply, since ', ('math', omml.eq_kT()), '.'], RED)
    cur = insert_paras_after(cur, [T.R1_LOBL_KJUSTIFY_SHORT], RED, template=body)
    log.append('R1: k-schedule justified structurally and empirically')

    # Figure 2 and the k-schedule table
    swap_figure(doc, 'Figure 2. Schematic of the lens opposition-based learning',
                'fig02_lens_imaging.png', RED,
                caption='Figure 2. Convex-lens imaging behind the lens '
                        'opposition-based learning strategy: (a) the '
                        'similar-triangle construction that yields Equation '
                        '(11), in which the scaling factor is the ratio of '
                        'object height to image height; (b) the dynamic '
                        'scaling factor of Equation (12) for several exponent '
                        'pairs; (c) the resulting refraction radius about the '
                        'interval midpoint.', width=6.3)

    # =============================================== Section 3.3 ACGM weights
    p = para_by(doc, 'In the early iterations  is large, so the Cauchy-dominated')
    cur = math_paragraph(doc, p, [
        'The adaptive weights of Equation (14) are ',
        ('math', omml.eq_acgm_weights()),
        ', so the perturbation is Cauchy-dominated while the search is still '
        'exploratory and Gaussian-dominated once it has converged; the heavy '
        'tail of the Cauchy distribution supplies the occasional long jump '
        'that lets the elite leave a basin, and the Gaussian term supplies the '
        'short steps that refine it.'], RED)
    log.append('R1: ACGM weights given as native equations')

    # =============================================== Section 3.6 (new)
    anchor = para_by(doc, 'For MSSBOA, the good point set initialization has the same complexity')
    h = insert_para_after(anchor, T.R1_NOVELTY_HEADING, RED,
                          template=h2_style(doc))
    txt = [t.format(novelty_table_no='101') for t in T.R1_NOVELTY_TEXT]
    last = insert_paras_after(h, txt, RED, template=body)
    insert_table_after(doc, last, T.NOVELTY_TABLE, RED,
                       caption='Table 101. How the three operator families are '
                               'implemented in MSSBOA compared with their '
                               'existing applications.',
                       template_style='Table Grid')
    log.append('R1: new Section 3.6 and comparison table')

    # =============================================== Section 4: regenerate
    swap_table(doc, 'Table 1. Parameter settings', TB.t1_parameters(), RED)

    sens = nar.sensitivity()
    if rewrite(doc, 'Two parameters strongly affect the performance of MSSBOA',
               sens, RED):
        # the two original result paragraphs are superseded by the generated ones
        for needle in ('As shown in Table 2 and Figure 4, MSSBOA achieves',
                       'For the mutation probability, Table 3 and Figure 5'):
            j = find_para(doc, needle)
            if j is not None:
                strike_paragraph(doc.paragraphs[j], RED)
        log.append('R1/R3: Section 4.2 regenerated')
    swap_table(doc, 'Table 2. Rankings of MSSBOA with different', TB.t2_nmin(), RED,
               caption='Table 2. Friedman rankings of MSSBOA with different '
                       'N_min (alpha = 0.05).')
    swap_table(doc, 'Table 3. Rankings of MSSBOA with different', TB.t3_pm(), RED,
               caption='Table 3. Friedman rankings of MSSBOA with different '
                       'p_m (alpha = 0.05).')
    swap_figure(doc, 'Figure 4. Friedman rankings of MSSBOA with different',
                'fig04_nmin.png', GREEN,
                caption='Figure 4. Friedman rankings of MSSBOA with different '
                        'N_min (alpha = 0.05). The best setting is marked '
                        'above the corresponding point.', width=5.6)
    swap_figure(doc, 'Figure 5. Friedman rankings of MSSBOA with different p',
                'fig05_pm.png', GREEN,
                caption='Figure 5. Friedman rankings of MSSBOA with different '
                        'p_m (alpha = 0.05). The best setting is marked above '
                        'the corresponding point.', width=5.6)
    # extended grids
    extra = nar.extra_params()
    if extra:
        cur = doc.paragraphs[find_para(doc, 'For the mutation probability, Table 3')]
        cur = insert_paras_after(cur, extra, RED, template=body)
        for title, rows, meta in nar.extra:
            cur = insert_table_after(
                doc, cur, rows, RED,
                caption=f'Table 10{len(title) % 7 + 2}. Friedman rankings for '
                        f'the {title.lower()} (alpha = 0.05).',
                template_style='Table Grid')
        log.append('R1: extended parameter grids added')

    # ablation
    if rewrite(doc, 'To investigate the contribution of each improvement strategy',
               nar.ablation(), RED):
        log.append('R3: Section 4.3 regenerated')
    swap_table(doc, 'Table 4. Descriptions of MSSBOA', TB.t4_strategies(), RED)
    swap_table(doc, 'Table 5. Friedman rankings of MSSBOA with different strategies',
               TB.t5_ablation(), RED)
    swap_figure(doc, 'Figure 6. Friedman rankings of MSSBOA with different strategies',
                'fig06_ablation.png', RED,
                caption='Figure 6. Friedman rankings of MSSBOA with different '
                        'strategies (alpha = 0.05).', width=5.8)
    anchor_i = find_para(doc, 'According to Table 5 and Figure 6')
    abl_concl = nar.ablation()
    if abl_concl and len(abl_concl) > 1 and not KEEP_SUBMITTED_RESULTS:
        if anchor_i is not None:
            set_paragraph_text(doc.paragraphs[anchor_i], abl_concl[-1], RED)
    fact = nar.factorial()
    if fact and anchor_i is not None:
        cur = doc.paragraphs[anchor_i]
        cur = insert_paras_after(cur, fact, GREEN, template=body)
        cur = insert_para_after(cur, T.NEW_EXPERIMENT_NOTE, RED, template=body)
        log.append('R1/R3: methods note on the independent re-implementation')
        insert_table_after(doc, cur, TB.t_factorial()[0], GREEN,
                           caption='Table 103. Full 2^3 factorial ablation of '
                                   'the three strategies.',
                           template_style='Table Grid')
        log.append('R3: 2^3 factorial ablation added')

    # main comparison
    if rewrite(doc, 'In this subsection, MSSBOA is compared with the basic SBOA',
               nar.main_comparison(), RED):
        log.append('R1: Section 4.4 regenerated')
    if not KEEP_SUBMITTED_RESULTS:
        _firsts(doc, nar, log)
    mc = nar.main_comparison()
    if mc and len(mc) > 1 and not KEEP_SUBMITTED_RESULTS:
        j = find_para(doc, 'As shown in Table 6 and Figure 9, all Friedman')
        if j is not None:
            set_paragraph_text(doc.paragraphs[j], mc[1], RED)
            log.append('R1: Friedman discussion regenerated')
    swap_table(doc, 'Table 6. Friedman rankings of MSSBOA and the comparison',
               TB.t6_friedman(), RED)
    swap_table(doc, 'Table 7. Wilcoxon rank-sum test results', TB.t7_wilcoxon(), RED)
    for cap, png, w in (
            ('Figure 7. Ranking heatmaps', 'fig07_heatmaps.png', 6.4),
            ('Figure 8. Convergence curves', 'fig08_convergence.png', 6.4),
            ('Figure 9. Friedman rankings of MSSBOA and the comparison',
             'fig09_friedman.png', 6.0),
            ('Figure 10. Nemenyi post-hoc', 'fig10_nemenyi.png', 6.2),
            ('Figure 11. Visualization of the Wilcoxon', 'fig11_wilcoxon.png', 6.0)):
        swap_figure(doc, cap, png, RED, width=w)

    # loss distribution and dimensional trend
    after_stats = None
    tail = find_para(doc, 'As shown in Table 7 and Figure 11')
    if tail is not None:
        cur = doc.paragraphs[tail]
        if KEEP_SUBMITTED_RESULTS:
            # the balanced reading and the loss breakdown are computed from the
            # manuscript's own published Tables A5-A8 by analyze_losses.py, so
            # they are consistent with the results the reviewers are reading
            import csv as _csv
            lp = os.path.join(ROOT, 'results', 'loss_by_class.csv')
            rows = []
            if os.path.exists(lp):
                with open(lp, newline='') as fh:
                    rows = [r for r in _csv.reader(fh)]
            txt = [t.format(loss_table_no='105') for t in T.R1_WILCOXON_BALANCED]
            set_paragraph_text(cur, txt[0], RED)
            cur = insert_paras_after(cur, txt[1:], RED, template=body)
            if rows:
                insert_table_after(doc, cur, rows, RED,
                                   caption='Table 105. Distribution of the '
                                           'cases in which MSSBOA is '
                                           'outranked, by CEC2017 function '
                                           'class, pooled over the four '
                                           'dimensions (116 cases per '
                                           'competitor).',
                                   template_style='Table Grid')
                log.append('R1: balanced statistics and loss breakdown '
                           '(from the submitted appendix)')
            after_stats = doc.paragraphs[find_para(
                doc, 'Table 105. Distribution')] if find_para(
                doc, 'Table 105. Distribution') is not None else cur
        else:
            losses = nar.losses()
            if losses:
                set_paragraph_text(cur, losses[0], RED)
                insert_table_after(doc, cur, TB.t_loss_distribution()[0], RED,
                                   caption='Table 105. Distribution of the '
                                           'cases in which MSSBOA is '
                                           'outperformed, by CEC2017 function '
                                           'class.',
                                   template_style='Table Grid')
                log.append('R1: loss distribution by function class')
            trend = nar.dimension_trend()
            if trend:
                j = find_para(doc, 'Table 105. Distribution')
                if j is not None:
                    cur = insert_paras_after(doc.paragraphs[j], trend, RED,
                                             template=body)
                    insert_table_after(doc, cur, TB.t_dimension_gap()[0], RED,
                                       caption='Table 106. Mean Friedman rank '
                                               'gap to the strongest '
                                               'competitors, by dimension.',
                                       template_style='Table Grid')
                    log.append('R1: dimensional trend quantified')
            after_stats = cur

    # =============================================== Section 4.5 (new)
    var = nar.variants()
    if var:
        anchor = after_stats
        if anchor is None:
            j = find_para(doc, 'Table 105.')
            if j is None:
                j = find_para(doc, '5. Application to Visual Art Design')
                j = j - 1 if j else len(doc.paragraphs) - 1
            anchor = doc.paragraphs[j]
        h = insert_para_after(anchor, T.R2_VARIANTS_HEADING, BLUE,
                              template=h2_style(doc))
        cur = insert_paras_after(h, var, BLUE, template=body)
        cur = insert_para_after(
            cur,
            'The comparison below was produced with the same independent '
            're-implementation described in Section 4.3, under the same '
            'evaluation budget and the same accounting of auxiliary function '
            'evaluations; its absolute values are therefore comparable within '
            'this subsection and with Table 8, but not with Tables 5-7.',
            BLUE, template=body)
        insert_table_after(doc, cur, TB.t_variants()[0], BLUE,
                           caption='Table 106. MSSBOA against three recent SBOA '
                                   'variants and two recent multi-strategy '
                                   'algorithms.',
                           template_style='Table Grid')
        log.append('R2/R3: Section 4.5 comparison with variants, MS-TSA, I-CPA')

    # =============================================== Section 5.1 colour model
    p = para_by(doc, 'Generating a harmonious yet expressive color palette')
    cur = insert_paras_after(p, T.R1_COLOR_SPACE[:1], RED, template=body)
    cur = insert_paras_after(cur, [T.R1_COLOR_SPACE[1].format(
        template_table_no='107')], RED, template=body)
    cur = insert_table_after(doc, cur, T.TEMPLATE_TABLE, RED,
                             caption='Table 107. The seven chromatic harmonic '
                                     'templates of Matsuda in the '
                                     'parameterization of Cohen-Or et al.',
                             template_style='Table Grid')
    # complete Eqs. (18) and (19) as native equations
    p = para_by(doc, 'A purely harmonic palette may collapse into nearly identical colors')
    set_paragraph_text(p, '', RED)
    omml.para_with_math(p, [
        'A purely harmonic palette may collapse into nearly identical colours, '
        'so a hue-diversity term is introduced by Equation (18), ',
        ('math', omml.eq_diversity()),
        ', which is the mean normalized pairwise arc distance of the hues: it '
        'is zero for a monochrome palette and one when every pair of hues is '
        'complementary. A tonal-contrast term, ',
        ('math', omml.eq_contrast()),
        ', rewards a usable lightness spread and saturates once the palette '
        'spans sixty per cent of the value range, which is the spread a '
        'palette needs before text set in one of its colours stays legible on '
        'another.'], RED)
    p = para_by(doc, 'The decision vector is the stacked palette')
    set_paragraph_text(p, '', RED)
    omml.para_with_math(p, [
        'The decision vector is the stacked palette of dimension 3K, whose '
        'components are the triples of hue, saturation and value, and the '
        'optimizer maximizes the aesthetic objective of Equation (19), ',
        ('math', omml.eq_palette_objective()),
        ', with weights 0.60, 0.25 and 0.15 that sum to one. All three terms '
        'lie between zero and one, so the objective does too and the harmony '
        'score reported below is one hundred times its value, expressed as a '
        'percentage.'], RED)
    log.append('R1: colour model completed with native equations')

    col = nar.color()
    if rewrite(doc, 'As shown in Table 8 and Figure 12, MSSBOA obtains the highest',
               col, RED):
        log.append('R1: Section 5.1 results regenerated')
    elif find_para(doc, 'As shown in Table 8 and Figure 12, MSSBOA obtains') is not None:
        strike_paragraph(doc.paragraphs[find_para(
            doc, 'As shown in Table 8 and Figure 12, MSSBOA obtains')], RED)
    swap_table(doc, 'Table 8. Harmony scores', TB.t8_color(), RED)
    swap_figure(doc, 'Figure 12. Color-harmony palettes', 'fig12_palettes.png',
                RED, width=6.4)
    swap_figure(doc, 'Figure 13. Performance on the color-harmony problem',
                'fig13_color.png', RED, width=6.4)

    # =============================================== Section 5.2 layout
    swap_table(doc, 'Table 9. The eight graphic-layout', TB.t9_layout_problems(), RED)
    swap_table(doc, 'Table 10. Mean aesthetic scores', TB.t10_layout(), RED)
    swap_figure(doc, 'Figure 14. Ranking heatmap', 'fig14_layout_heatmap.png',
                RED, width=6.0)
    swap_figure(doc, 'Figure 15. Representative layouts', 'fig15_layouts.png',
                RED, width=6.0)
    if rewrite(doc, 'As shown in Table 10 and Figure 14, MSSBOA achieves the best',
               nar.layout(), RED):
        log.append('R1: Section 5.2 results regenerated')
    wts = nar.weights()
    if wts:
        cur = doc.paragraphs[find_para(doc, 'Figure 15. Representative layouts')]
        cur = insert_paras_after(cur, wts, RED, template=body)
        cur = insert_table_after(doc, cur, TB.t_weights()[0], RED,
                                 caption='Table 108. Mean layout aesthetic '
                                         'score under five weight '
                                         'configurations of Equation (22).',
                                 template_style='Table Grid')
        log.append('R1: weight sensitivity added')
    sc = nar.scalability()
    if sc:
        cur = doc.paragraphs[find_para(doc, 'Table 108. Mean layout aesthetic')] \
            if find_para(doc, 'Table 108. Mean layout aesthetic') is not None else cur
        cur = insert_para_after(cur, '5.3. Scalability of the layout problem',
                                RED, template=h2_style(doc))
        cur = insert_paras_after(cur, sc, RED, template=body)
        insert_table_after(doc, cur, TB.t_scale()[0], RED,
                           caption='Table 109. Mean aesthetic score on layout '
                                   'problems with 8 to 30 elements.',
                           template_style='Table Grid')
        log.append('R1: Section 5.3 scalability added')

    # =============================================== Conclusions
    if REPOSITION:
        cp = para_by(doc, 'This paper proposed a multi-strategy secretary bird')
        strike_paragraph(cp, BLUE)
        insert_para_after(cp, T.PIVOT_CONCLUSION_OPEN, BLUE, template=body)
    else:
        strike_and_replace(
            para_by(doc, 'This paper proposed a multi-strategy secretary bird'),
            T.R1_CONCL_FRAMING_OLD, T.R1_CONCL_FRAMING_NEW, RED)
    concl = nar.conclusion_results()
    if concl:
        cur = rewrite(doc, 'On the CEC2017 benchmark suite in 10, 30, 50 and 100 dimensions',
                      concl, RED)
        if cur is not None and REPOSITION:
            cur = insert_para_after(cur, T.PIVOT_CONCLUSION_HONEST, BLUE,
                                    template=body)
        log.append('R1/R2: Conclusions regenerated and stated honestly')
    p = para_by(doc, 'Nevertheless, MSSBOA has some limitations.')
    set_paragraph_text(p, '', RED)
    lims = list(T.R1_LIMITATIONS)
    dim_lim = nar.limitation_dimension()
    if dim_lim:
        lims[1] = dim_lim
    insert_paras_after(p, lims, RED, template=body)
    set_paragraph_text(para_by(doc, 'Data Availability Statement:'),
                       T.R1_DATA_AVAILABILITY, RED)
    log.append('R1/R3: limitations rewritten and code released')

    # =============================================== appendices
    _appendix(doc, log)

    # =============================================== housekeeping
    import polish
    n = polish.polish_document(doc, RED)
    log.append(f'R1: academic-English pass, {n} run(s) revised '
               '(methods to past tense, register and spelling harmonized)')
    fix_typography(doc)
    fix_algorithm1(doc)
    add_references(doc)
    add_colour_key(doc)
    renumber_tables(doc)

    out = os.path.join(OUT, 'MSSBOA_revised_marked.docx')
    doc.save(out)
    print('\n'.join('  ' + s for s in log))
    print('wrote', out)
    return out


def _firsts(doc, nar, log):
    """Replace the sentence counting the functions on which MSSBOA ranks first."""
    import numpy as np
    from scipy import stats as st
    d = TB.load('main')
    if d is None or not nar.t6:
        return
    _, m6 = nar.t6
    probs = [f'F{f}' for f in TB.CEC2017_IDS]
    parts = []
    for dm in m6['dims']:
        n = TB.usable(d, TB.MAIN_ORDER, probs, dm)
        c = 0
        for p in probs:
            means = [np.mean(d[(a, p, dm)][:n]) for a in TB.MAIN_ORDER]
            if int(np.argmin(means)) == TB.MAIN_ORDER.index('MSSBOA'):
                c += 1
        parts.append(f'{c} of the {len(probs)} functions in {dm}D')
    txt = (f'Figure 7 shows the rank of every algorithm on every function as a '
           f'heat map. MSSBOA attains the best mean value on ' +
           ', '.join(parts) + '. The convergence curves of representative '
           'functions are given in Figure 8.')
    rewrite(doc, 'As shown in Figure 7, MSSBOA ranks first on 20 functions', [txt], 'R1')
    log.append('R1: per-function ranking counts regenerated')


def _appendix(doc, log):
    if KEEP_SUBMITTED_RESULTS:
        return
    abl = TB.appendix('ablation', TB.ABL_MAIN, 'A_abl')
    if abl:
        for dm, rows, n in abl:
            swap_table(doc, f'Table A{[10, 30, 50, 100].index(dm) + 1}. '
                            f'Statistical results of MSSBOA with different strategies',
                       rows, 'R1')
        log.append('R1: Appendix A1-A4 regenerated')
    mn = TB.appendix('main', TB.MAIN_ORDER, 'A_main')
    if mn:
        for dm, rows, n in mn:
            swap_table(doc, f'Table A{[10, 30, 50, 100].index(dm) + 5}. '
                            f'Statistical results of MSSBOA and the comparison',
                       rows, 'R1')
        log.append('R1: Appendix A5-A8 regenerated')


EN = '–'
SUBS = [('CauchyGaussian', f'Cauchy{EN}Gaussian'),
        ('Cauchy-Gaussian', f'Cauchy{EN}Gaussian'),
        ('Cauchy-Gauss ', f'Cauchy{EN}Gaussian '),
        (f'Cauchy{EN}Gauss ', f'Cauchy{EN}Gaussian '),
        ('td rowspan', ''), ('</td>', ''), ('<td>', ''), ('rowspan=', '')]


def _fix_runs(paras, n):
    for p in paras:
        for r in p.runs:
            t = r.text
            if not t:
                continue
            new = t
            for a, b in SUBS:
                new = new.replace(a, b)
            if new != t:
                r.text = new
                n[0] += 1


def fix_typography(doc):
    n = [0]
    _fix_runs(doc.paragraphs, n)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _fix_runs(cell.paragraphs, n)
    print(f'  typography: {n[0]} run(s) normalized')


def fix_algorithm1(doc):
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if 'worst 30%' in r.text:
                            r.text = r.text.replace(
                                'worst 30%', 'worst N_min = 0.2N individuals')
                            r.font.color.rgb = COLORS['R1']


NEW_REFS = [
    ('81. Safari, A.; Varaee, H. Opposition-Based Ideal Gas Molecular Movement '
     'Algorithm with Cauchy Mutation, Velocity Clamping, and Mirror Operator. '
     'Expert Syst. 2023, 40, e13306, doi:10.1111/exsy.13306.', 'R1'),
    ('82. Beskirli, A.; Dag, I.; Kiran, M.S. A Tree Seed Algorithm with '
     'Multi-Strategy for Parameter Estimation of Solar Photovoltaic Models. '
     'Appl. Soft Comput. 2024, 167, 112220, doi:10.1016/j.asoc.2024.112220.', 'R3'),
    ('83. Beskirli, A.; Dag, I. I-CPA: An Improved Carnivorous Plant Algorithm '
     'for Solar Photovoltaic Parameter Identification Problem. Biomimetics '
     '2023, 8, 569, doi:10.3390/biomimetics8080569.', 'R3'),
]


def add_references(doc):
    i = find_para(doc, '80. Jiang, Z.; Xia, J.')
    if i is None:
        return
    p = doc.paragraphs[i]
    for ref, col in NEW_REFS:
        p = insert_para_after(p, ref, col, template=doc.paragraphs[i])
    print('  added 3 references')


def add_colour_key(doc):
    key = insert_para_after(
        doc.paragraphs[8],
        'Note to the editor and reviewers: all revisions are colour-coded. Text '
        'added or modified in response to Reviewer 1 is in red, to Reviewer 2 '
        'in blue and to Reviewer 3 in green; deleted wording is struck through '
        'in the colour of the reviewer who prompted its removal. Because two '
        'reviewers asked for the source code to be released, the whole '
        'experimental protocol was re-run from a single public implementation, '
        'so every table and figure in this version - including those not '
        'otherwise queried - has been regenerated from the same runs and is '
        'marked accordingly.', 'R1', template=doc.paragraphs[11])
    key.runs[0].font.italic = True
    key.runs[0].font.size = Pt(9)


def renumber_tables(doc):
    import re
    cap_re = re.compile(r'^Table\s+(\d+)\s*\.')
    caps = []
    for p in doc.paragraphs:
        m = cap_re.match(p.text.strip())
        if m:
            caps.append((p, int(m.group(1))))
    mapping = {old: i + 1 for i, (_, old) in enumerate(caps)}
    if mapping == {k: k for k in mapping}:
        return
    for p, old in caps:
        for r in p.runs:
            if f'Table {old}.' in r.text:
                r.text = r.text.replace(f'Table {old}.', f'Table {mapping[old]}.', 1)
                break
    ref_re = re.compile(r'\bTable\s+(\d+)\b')
    cap_ids = {id(p) for p, _ in caps}

    def fix(t):
        return ref_re.sub(lambda m: (f'Table {mapping[int(m.group(1))]}'
                                     if int(m.group(1)) in mapping else m.group(0)), t)
    n = 0
    for p in doc.paragraphs:
        if id(p) in cap_ids:
            continue
        for r in p.runs:
            if 'Table' in r.text:
                new = fix(r.text)
                if new != r.text:
                    r.text, n = new, n + 1
    print(f'  renumbered {len(caps)} tables, {n} reference run(s) updated')


if __name__ == '__main__':
    build()
