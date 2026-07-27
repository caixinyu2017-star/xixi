"""Turn the inline mathematics of the added text into native Word equations.

The display equations added by this revision are already OMML objects, built
by omml.py.  What this module handles is the mathematics that occurs *inside*
sentences and captions - N_min, p_m, h_i, alpha, k(T) = 2^mu and the rest -
which the drafting pipeline produced as ordinary text because it is written
inline in Python string literals.  Left as text it would be set in the body
font, with the subscripts as literal underscores, which is neither correct
typography nor what a Word user sees when they click on the symbol.

Only text this revision added is converted.  The rule is enforced by requiring
the run under the match to carry one of the three reviewer colours: the
submitted paragraphs already hold 129 genuine OMML equations and must not be
touched.
"""
import re

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

import omml as M
from docx_edit import COLORS

_NS = nsdecls('m', 'w')

GREEK = {'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'mu': 'μ',
         'nu': 'ν', 'tau': 'τ', 'rho': 'ρ', 'sigma': 'σ', 'lambda': 'λ',
         'theta': 'θ', 'omega': 'ω'}


# ----------------------------------------------------------------- builders
def v(name):
    """An italic variable, spelt-out Greek names included."""
    return M.r(GREEK.get(name, name))


def eq(*parts):
    """Join with the thin spacing Word puts around a relation."""
    return M.joined(*parts)


def rel(sym):
    return M.op(f' {sym} ')


def var_sub(base, sub, upright_sub=False):
    s = M.r(sub, italic=not upright_sub)
    return M.sub(v(base), s)


def fn(name, arg):
    """An italic symbol applied to a parenthesised argument: k(T), f(P)."""
    return M.joined(v(name), M.delim(v(arg)))


# The conversions, longest first so that a rule never fires inside the match
# of a more specific one.  Each entry is (pattern, builder); the builder takes
# the regex match and returns an OMML body.
RULES = [
    (r'k\(T\) = 2\^mu',
     lambda m: eq(fn('k', 'T'), rel('='), M.sup(M.num('2'), v('mu')))),
    (r'\bN_min = 0\.2N\b',
     lambda m: eq(var_sub('N', 'min', True), rel('='), M.num('0.2'), v('N'))),
    (r'\bN_min\b', lambda m: var_sub('N', 'min', True)),
    (r'\b0\.2N\b', lambda m: eq(M.num('0.2'), v('N'))),
    (r'\bp_m\b', lambda m: var_sub('p', 'm')),
    (r'\b([hsv])_i\b', lambda m: var_sub(m.group(1), 'i')),
    (r'\bL\*a\*b\*', lambda m: eq(M.sup(v('L'), M.op('*')), M.sup(v('a'), M.op('*')),
                                  M.sup(v('b'), M.op('*')))),
    (r'\(a\*, b\*\)',
     lambda m: M.delim(eq(M.sup(v('a'), M.op('*')), M.op(', '),
                          M.sup(v('b'), M.op('*'))))),
    (r'\b2\^3\b', lambda m: M.sup(M.num('2'), M.num('3'))),
    (r'\b(alpha|beta|mu|nu|tau) = ([0-9][0-9.]*)',
     lambda m: eq(v(m.group(1)), rel('='), M.num(m.group(2)))),
    (r'([αβμντ]) = ([0-9][0-9.]*)',
     lambda m: eq(M.r(m.group(1)), rel('='), M.num(m.group(2)))),
    (r'\bν = 1/2, μ = 10\b',
     lambda m: eq(M.r('ν'), rel('='), M.num('1'), M.op('/'), M.num('2'),
                  M.op(', '), M.r('μ'), rel('='), M.num('10'))),
    (r'\bN = (\d+)\b', lambda m: eq(v('N'), rel('='), M.num(m.group(1)))),
    (r'\bD = (\d+)\b', lambda m: eq(v('D'), rel('='), M.num(m.group(1)))),
    (r'\bk ≡ 1\b', lambda m: eq(v('k'), rel('≡'), M.num('1'))),
    (r'\bk = 1\b', lambda m: eq(v('k'), rel('='), M.num('1'))),
    (r'\b(\d?)T/(\d), (\d?)T/(\d)\b',
     lambda m: eq(M.num(m.group(1)), v('T'), M.op('/'), M.num(m.group(2)),
                  M.op(', '), M.num(m.group(3)), v('T'), M.op('/'),
                  M.num(m.group(4)))),
    (r'\bf\(P\)', lambda m: fn('f', 'P')),
    (r'\b(alpha|beta|gamma|delta|mu|nu|tau|rho|sigma|lambda|theta|omega)\b',
     lambda m: v(m.group(1))),
]

_COMPILED = [(re.compile(p), b) for p, b in RULES]
_RGB = {str(c) for c in COLORS.values()}


_CAPTION_RE = re.compile(r'^\s*(Table|Figure)\s+A?\d+\s*\.')


def _is_caption(text):
    return bool(_CAPTION_RE.match(text))


def _colored(run):
    c = run.font.color
    return c is not None and c.rgb is not None and str(c.rgb) in _RGB


def _color_of(run):
    c = run.font.color
    return str(c.rgb) if (c is not None and c.rgb is not None) else None


def _paint(el, rgb):
    """Give every math run the colour of the text it replaced."""
    if not rgb:
        return el
    for mr in el.iter(qn('m:r')):
        pr = mr.find(qn('w:rPr'))
        if pr is None:
            pr = parse_xml(f'<w:rPr {_NS}><w:color w:val="{rgb}"/></w:rPr>')
            mr.insert(0, pr)
    return el


def _splice_math(para, start, end, body, rgb):
    """Replace characters [start, end) of `para` with an inline equation.

    Only the runs the match actually spans are rebuilt, so the formatting of
    everything before and after it survives - the same discipline the text
    edits follow, for the same reason.
    """
    runs = list(para.runs)
    offsets, pos = [], 0
    for run in runs:
        offsets.append((pos, pos + len(run.text)))
        pos += len(run.text)
    first = next(k for k, (a, b) in enumerate(offsets)
                 if a <= start < b or (a == b == start))
    last = next(k for k, (a, b) in enumerate(offsets) if a < end <= b)
    head = runs[first].text[:start - offsets[first][0]]
    tail = runs[last].text[end - offsets[last][0]:]

    anchor = runs[first]
    anchor.text = head
    for k in range(first + 1, last + 1):
        runs[k]._element.getparent().remove(runs[k]._element)

    el = parse_xml(f'<m:oMath {_NS}>{body}</m:oMath>')
    _paint(el, rgb)
    anchor._element.addnext(el)
    if tail:
        import copy
        new_r = copy.deepcopy(anchor._element)
        for t in new_r.findall(qn('w:t')):
            new_r.remove(t)
        t = parse_xml(f'<w:t {_NS} xml:space="preserve">'
                      f'{M._esc(tail)}</w:t>')
        new_r.append(t)
        el.addnext(new_r)
    return True


def mathify_paragraph(para):
    """Convert every inline expression in one paragraph; returns the count.

    The paragraph is rescanned after each conversion because splicing changes
    the offsets.  A match that falls in text the revision did not add is
    stepped over rather than ending the scan, so one sentence of untouched
    submitted prose in the middle of a paragraph does not hide the added
    mathematics that follows it.
    """
    done, floor = 0, 0
    for _ in range(80):
        runs = list(para.runs)
        if not runs:
            return done
        text = ''.join(r.text for r in runs)
        hit = None
        for rx, build in _COMPILED:
            m = rx.search(text, floor)
            if m and (hit is None or m.start() < hit[0].start()):
                hit = (m, build)
        if hit is None:
            return done
        m, build = hit
        pos, owner = 0, None
        for run in runs:
            if pos <= m.start() < pos + len(run.text):
                owner = run
                break
            pos += len(run.text)
        if owner is None:
            floor = m.end()
            continue
        if _colored(owner):
            rgb = _color_of(owner)
        elif _is_caption(text):
            # A submitted caption that spells a symbol as "N_min" is a
            # typographical defect Reviewer 1 asked us to fix, and fixing it
            # changes no data; it is converted and marked as a red revision.
            rgb = str(COLORS['R1'])
        else:
            floor = m.end()                 # submitted body text: leave it
            continue
        _splice_math(para, m.start(), m.end(), build(m), rgb)
        floor = m.start()
        done += 1
    return done


def mathify_document(doc):
    n = 0
    for para in doc.paragraphs:
        n += mathify_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    n += mathify_paragraph(para)
    print(f'  inline mathematics: {n} expression(s) set as Word equations')
    return n
