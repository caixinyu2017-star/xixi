"""Helpers for editing the submitted manuscript in place.

The revision is produced by editing the original .docx rather than rebuilding
it, so the MDPI template styles, the native OMML equations and every embedded
figure survive untouched.  All inserted or modified text is colour-coded by
reviewer, as requested:

    Reviewer 1 -> red        Reviewer 2 -> blue        Reviewer 3 -> green
"""
import copy

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

COLORS = {
    'R1': RGBColor(0xC0, 0x00, 0x00),      # red
    'R2': RGBColor(0x00, 0x70, 0xC0),      # blue
    'R3': RGBColor(0x00, 0xA0, 0x33),      # green
}
COLOR_NAMES = {'R1': 'red', 'R2': 'blue', 'R3': 'green'}


# ------------------------------------------------------------------ lookup
def iter_paragraphs(doc):
    return doc.paragraphs


def find_para(doc, needle, start=0, exact=False):
    """Index of the first paragraph containing `needle` (or None)."""
    for i, p in enumerate(doc.paragraphs[start:], start):
        t = p.text.strip()
        if (t == needle) if exact else (needle in p.text):
            return i
    return None


def para_by(doc, needle, **kw):
    i = find_para(doc, needle, **kw)
    if i is None:
        raise LookupError(f'paragraph not found: {needle!r}')
    return doc.paragraphs[i]


# ------------------------------------------------------------- new content
def _style_run(run, color, italic=False, bold=False, size=None):
    run.font.color.rgb = COLORS[color]
    run.font.bold = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)


def _clone_format(src_para, new_para):
    """Copy paragraph properties (style, spacing, indent) from src to new."""
    if src_para.style is not None:
        try:
            new_para.style = src_para.style
        except Exception:
            pass
    src_pPr = src_para._p.find(qn('w:pPr'))
    if src_pPr is not None:
        new_pPr = new_para._p.find(qn('w:pPr'))
        if new_pPr is not None:
            new_para._p.remove(new_pPr)
        new_para._p.insert(0, copy.deepcopy(src_pPr))


def insert_para_after(para, text, color, template=None, bold=False,
                      italic=False, align=None, style=None):
    """Insert a new paragraph directly after `para`; return it."""
    new_p = copy.deepcopy(para._p)
    for child in list(new_p):
        if child.tag != qn('w:pPr'):
            new_p.remove(child)
    para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np_ = Paragraph(new_p, para._parent)
    if template is not None:
        _clone_format(template, np_)
    if style is not None:
        try:
            np_.style = style
        except Exception:
            pass
    if align is not None:
        np_.alignment = align
    if text:
        _style_run(np_.add_run(text), color, italic=italic, bold=bold)
    return np_


def insert_paras_after(para, texts, color, **kw):
    """Insert several paragraphs in order; return the last one."""
    cur = para
    for t in texts:
        cur = insert_para_after(cur, t, color, **kw)
    return cur


def append_run(para, text, color, bold=False, italic=False):
    _style_run(para.add_run(text), color, bold=bold, italic=italic)
    return para


# --------------------------------------------------------- text surgery
def _clone_run_after(run, text):
    """Duplicate `run` (keeping its formatting) with new text, right after it."""
    import copy as _copy
    new_r = _copy.deepcopy(run._element)
    run._element.addnext(new_r)
    from docx.text.run import Run
    r = Run(new_r, run._parent)
    r.text = text
    return r


def _splice(para, old, make_runs):
    """Replace the first occurrence of `old` without disturbing other runs.

    Only the runs that the match actually spans are rebuilt; everything before
    and after keeps its original formatting.  `make_runs(anchor)` is called to
    create the replacement runs immediately after `anchor`.
    """
    runs = list(para.runs)
    if not runs:
        return False
    offsets, pos = [], 0
    for r in runs:
        offsets.append((pos, pos + len(r.text)))
        pos += len(r.text)
    full = ''.join(r.text for r in runs)
    i = full.find(old)
    if i < 0:
        return False
    j = i + len(old)

    first = next(k for k, (a, b) in enumerate(offsets) if a <= i < b or (a == b == i))
    last = next(k for k, (a, b) in enumerate(offsets) if a < j <= b)

    head = runs[first].text[:i - offsets[first][0]]
    tail = runs[last].text[j - offsets[last][0]:]

    anchor = runs[first]
    anchor.text = head
    for k in range(first + 1, last + 1):
        runs[k]._element.getparent().remove(runs[k]._element)

    cur = anchor
    for r in make_runs(cur):
        cur = r
    if tail:
        _clone_run_after(cur, tail)
    return True


def strike_paragraph(para, color):
    """Mark a whole paragraph as deleted: struck through, in reviewer colour."""
    for r in para.runs:
        if r.text:
            r.font.strike = True
            r.font.color.rgb = COLORS[color]
    return para


def replace_in_para(para, old, new, color):
    """Replace `old` by `new`, colouring only the new text."""
    def mk(anchor):
        r = _clone_run_after(anchor, new)
        _style_run(r, color)
        r.font.strike = False
        return [r]
    return _splice(para, old, mk)


def strike_and_replace(para, old, new, color):
    """Show a deletion struck through, followed by the replacement text."""
    def mk(anchor):
        r_del = _clone_run_after(anchor, old)
        _style_run(r_del, color)
        r_del.font.strike = True
        r_new = _clone_run_after(r_del, new)
        _style_run(r_new, color)
        r_new.font.strike = False
        return [r_del, r_new]
    return _splice(para, old, mk)


# ------------------------------------------------------------------ tables
def insert_table_after(doc, para, rows, color, caption=None, template_style=None,
                       widths=None, header_bold=True):
    """Insert a caption paragraph plus a table right after `para`.

    Returns the empty paragraph that trails the table, so that further content
    can be chained after it in document order.
    """
    cur = para
    if caption:
        cur = insert_para_after(cur, caption, color, bold=False)
    ncols = max(len(r) for r in rows)
    rows = [list(r) + [''] * (ncols - len(r)) for r in rows]
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    try:
        tbl.style = template_style or 'Table Grid'
    except Exception:
        pass
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            _style_run(run, color, bold=(header_bold and i == 0), size=8.5)
    cur._p.addnext(tbl._tbl)
    trailer = insert_para_after(cur, '', color)
    trailer._p.getparent().remove(trailer._p)
    tbl._tbl.addnext(trailer._p)
    return trailer


def table_after(doc, para):
    """The first table element following a caption paragraph, or None."""
    el = para._p.getnext()
    while el is not None:
        if el.tag == qn('w:tbl'):
            from docx.table import Table
            return Table(el, para._parent)
        if el.tag == qn('w:p') and el.xpath('.//w:t'):
            return None                 # ran into the next paragraph of text
        el = el.getnext()
    return None


def replace_table(doc, para, rows, color, font_size=8.0, header_bold=True):
    """Replace the table following `para` with `rows`, keeping its look.

    The original table's properties (borders, width, look) are carried over so
    the regenerated table matches the MDPI template, and the cell text is
    written in the reviewer's colour so the change is visible.
    """
    old = table_after(doc, para)
    ncols = max(len(r) for r in rows)
    rows = [list(r) + [''] * (ncols - len(r)) for r in rows]
    new = doc.add_table(rows=len(rows), cols=ncols)
    if old is not None:
        old_pr = old._tbl.find(qn('w:tblPr'))
        if old_pr is not None:
            new_pr = new._tbl.find(qn('w:tblPr'))
            if new_pr is not None:
                new._tbl.remove(new_pr)
            new._tbl.insert(0, copy.deepcopy(old_pr))
        else:
            try:
                new.style = 'Table Grid'
            except Exception:
                pass
    else:
        try:
            new.style = 'Table Grid'
        except Exception:
            pass
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = new.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(str(val))
            _style_run(run, color, bold=(header_bold and i == 0), size=font_size)
    if old is not None:
        old._tbl.addnext(new._tbl)
        old._tbl.getparent().remove(old._tbl)
    else:
        para._p.addnext(new._tbl)
    # mark the first header row to repeat across page breaks
    tr = new.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        from docx.oxml import OxmlElement
        trPr.append(OxmlElement('w:tblHeader'))
    return new


def set_paragraph_text(para, text, color, keep_style=True):
    """Replace a paragraph's whole content with new coloured text."""
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for child in list(para._p):
        if child.tag in (qn('m:oMath'), qn('m:oMathPara')):
            para._p.remove(child)
    if text:
        _style_run(para.add_run(text), color)
    return para


def insert_picture_after(doc, para, path, color, caption, width):
    """Insert a centred picture plus its caption after `para`."""
    holder = insert_para_after(para, '', color)
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.add_run().add_picture(path, width=width)
    cap = insert_para_after(holder, caption, color)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap
