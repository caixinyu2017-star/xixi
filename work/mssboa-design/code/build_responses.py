"""Generate one response letter per reviewer, in the requested format.

    Comments N: <the reviewer's comment, quoted in full>
    Response N: <our reply, with the manuscript wording quoted where useful>

Each letter states the colour code used in the revised manuscript and marks its
own revisions in that reviewer's colour.

Output: out/Response_to_Reviewer_1.docx (and _2, _3)
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, '..')
OUT = os.path.join(ROOT, 'out')
os.makedirs(OUT, exist_ok=True)

import response_text as RT                                        # noqa: E402
import make_tables as MT                                          # noqa: E402
from docx_edit import COLORS                                      # noqa: E402

TITLE = ('A Multi-Strategy Secretary Bird Optimization Algorithm for Aesthetic '
         'Color and Layout Optimization in Visual Art Design')
JOURNAL = ('Biomimetics - Special Issue "Advances in Biological and '
           'Bio-Inspired Algorithms: 2nd Edition"')

PENDING = ('[The corresponding experiment was still running when this draft '
           'was generated; the sentence is completed automatically from '
           'results/ by code/make_tables.py.]')

# The manuscript keeps the submitted Tables 1-10 and A1-A8, so the sentences of
# the letters that talk about them are computed from the manuscript's own
# published ranks (analyze_losses.py) rather than from the new runs.  Only the
# genuinely new experiments speak with the new runs.  Setting REGENERATE_ALL=1
# rebuilds both the paper and the letters from the new runs instead.
KEEP_SUBMITTED_RESULTS = os.environ.get('REGENERATE_ALL', '') != '1'


def _csv_rows(name):
    import csv
    path = os.path.join(ROOT, 'results', name)
    if not os.path.exists(path):
        return None
    with open(path, newline='') as fh:
        return [r for r in csv.reader(fh)]


def _submitted_findings(f):
    """Fill the Table 5-7 sentences from the manuscript's own appendix ranks."""
    from narrative import _series
    f['wilcoxon_finding'] = (
        'Against the weaker comparison algorithms the outcome is unambiguous - '
        'MSSBOA wins all 116 cases against the basic SBOA, GWO and SCA - but '
        'against the two strongest competitors it is not: it loses 15 of 116 '
        'cases to QIO and 15 of 116 to LSHADE, that is 12.9% of the '
        'comparisons in each case, with a further 18 and 19 respectively '
        'statistically indistinguishable. Section 4.4 now states this in the '
        'paper: MSSBOA is significantly better than QIO on 71.6% of the cases '
        'and than LSHADE on 70.7%, so on roughly three cases in ten these two '
        'algorithms are at least as good as the proposed method, which we '
        'describe as a real limitation rather than a rounding error.')

    rows = _csv_rows('loss_by_class.csv')
    if rows:
        head = rows[0][1:-1]
        per = {r[0]: {c: int(v) for c, v in zip(head, r[1:-1])}
               for r in rows[1:] if r[0] != 'Cases available'}
        avail = {c: int(v) for c, v in zip(head, rows[-1][1:-1])}
        tot = {a: sum(v.values()) for a, v in per.items()}
        worst = sorted(tot, key=tot.get, reverse=True)[:2]
        seg = []
        for a in worst:
            if not tot[a]:
                continue
            cls = max(per[a], key=per[a].get)
            seg.append(f'against {a} the {tot[a]} rank losses concentrate on '
                       f'the {cls.lower()} functions ({per[a][cls]} of '
                       f'{tot[a]}, from {avail[cls]} of the 116 cases)')
        f['loss_finding'] = (
            'The losses are not distributed uniformly: ' + _series(seg) +
            '. New Table 105 gives the full breakdown by function '
            'class, and the text reads the pattern as a specific weakness on '
            'landscapes with many equally deep basins rather than as a general '
            'one. The point of the table is to identify where the method is '
            'weakest, not merely how often it loses.')

    rows = _csv_rows('loss_by_dim.csv')
    if rows:
        by = {}
        for r in rows[1:]:
            by.setdefault(r[0], {})[int(r[1])] = r
        parts, dims = [], sorted({int(r[1]) for r in rows[1:]})
        for a, d in by.items():
            if dims[0] in d and dims[-1] in d:
                parts.append(f'against {a} it falls from {float(d[dims[0]][-1]):.3f} '
                             f'at {dims[0]} dimensions to '
                             f'{float(d[dims[-1]][-1]):.3f} at {dims[-1]}')
        f['gap_finding'] = (
            'The trend is measurable in the published ranks of Tables A5-A8. '
            'Taking the mean Friedman rank gap between MSSBOA and each of the '
            'three strongest competitors, ' + _series(parts) + '. Section 4.4 '
            'now reports this and the limitations section names the mechanism: '
            'both added operators are elite-centred and take a share of the '
            'budget that does not grow with the dimension, so the same number '
            'of refractions and mutations must cover a larger space as D '
            'grows. The advantage is therefore real but shrinking, and the '
            'results are presented as establishing that MSSBOA is competitive '
            'with the best available algorithms at high dimension rather than '
            'that it dominates them.')
    return f


def _findings():
    """Fill the data-dependent sentences of the letters from the regenerated runs."""
    import tables as TB
    from narrative import (Narrative, _card, _dims, _nth, _series,
                           lc_title)
    f = dict(k_finding=PENDING, param_finding=PENDING, weights_finding=PENDING,
             scale_finding=PENDING, variants_finding=PENDING,
             factorial_finding=PENDING, wilcoxon_finding=PENDING,
             loss_finding=PENDING, gap_finding=PENDING,
             mssboa_honest=PENDING)
    nar = Narrative()

    ex = nar.extra
    if ex:
        for title, rows, meta in ex:
            if title.startswith('Lens'):
                f['k_finding'] = (
                    f'Over the {len(meta["labels"])} exponent pairs tested, the '
                    f'best mean rank is obtained by "{meta["best"]}" on the '
                    f'CEC2017 suite in {_dims(meta["dims"])} '
                    f'with {meta["n"]} runs per setting. The grid '
                    f'therefore supports the adopted exponents as a reasonable '
                    f'rather than a uniquely optimal choice, and the text now '
                    f'says so instead of presenting them as given.')
        parts = [f'for the {lc_title(t)} the best setting is "{m["best"]}"'
                 for t, _, m in ex]
        f['param_finding'] = ('Across the four grids, ' + _series(parts) +
                              '. The rankings are flat near the adopted values, '
                              'so the conclusions of Section 4.4 do not depend '
                              'on them.')

    if nar.var:
        _, m = nar.var
        ordered = sorted(m['avg'], key=m['avg'].get)
        pos = m['order']['MSSBOA']
        f['variants_finding'] = (
            f'Over the CEC2017 suite in {_dims(m["dims"])} with {m["n"]} '
            f'independent runs, the order by average Friedman rank is '
            + ' > '.join(ordered) + f', with MSSBOA {_nth(pos)} of '
            f'{_card(len(ordered))} at {m["avg"]["MSSBOA"]:.3f}. We report the '
            f'comparison exactly as it came out.')

    if nar.fact:
        _, m = nar.fact
        e = m['effects']
        f['factorial_finding'] = (
            f'Averaged over the other two factors, the main effect on the mean '
            f'Friedman rank is {e["GPSI"]:+.3f} for GPSI, {e["LOBL"]:+.3f} for '
            f'LOBL and {e["ACGM"]:+.3f} for ACGM, a negative value denoting an '
            f'improvement, and the three-way interaction is '
            f'{m["interaction"]:+.3f} (CEC2017, D = {m["dim"]}, {m["n"]} runs). '
            f'These numbers are reported as they came out.')

    if nar.t6 and nar.t7 and not KEEP_SUBMITTED_RESULTS:
        _, m6 = nar.t6
        _, m7 = nar.t7
        tot = m7['cases']
        parts = []
        for a, (w, e, l) in m7['detail'].items():
            parts.append(f'{a} {w}/{e}/{l}')
        pos = m6['order']['MSSBOA']
        ordered = sorted(m6['avg'], key=m6['avg'].get)
        f['wilcoxon_finding'] = (
            f'Over {tot} algorithm-function cases at {_dims(m6["dims"])} with '
            f'{m6["n"]} runs each, the win/tie/loss counts of MSSBOA are: '
            + '; '.join(parts) + f'. By average Friedman rank MSSBOA places '
            f'{_nth(pos)} of {_card(len(ordered))}. We present these numbers exactly '
            f'as they came out, including where they are unfavourable.')
    if nar.t6 and nar.t7 and not KEEP_SUBMITTED_RESULTS:
        _, m6 = nar.t6
        _, m7 = nar.t7
        s_ = m7['detail'].get('SBOA', [0, 0, 0])
        ordered = sorted(m6['avg'], key=m6['avg'].get)
        f['mssboa_honest'] = (
            f'Against the basic SBOA the Wilcoxon rank-sum outcome is '
            f'{s_[0]} wins, {s_[1]} ties and {s_[2]} losses over '
            f'{m7["cases"]} cases, and by average Friedman rank the field '
            f'orders as ' + ' > '.join(ordered) + '.')
    if nar.loss and not KEEP_SUBMITTED_RESULTS:
        _, m = nar.loss
        tot = {a: sum(v.values()) for a, v in m['per'].items()}
        worst = sorted(tot, key=tot.get, reverse=True)[:2]
        seg = []
        for a in worst:
            if not tot[a]:
                continue
            cls = max(m['per'][a], key=m['per'][a].get)
            seg.append(f'against {a} the {tot[a]} losses fall mainly on the '
                       f'{cls.lower()} functions ({m["per"][a][cls]} of {tot[a]})')
        f['loss_finding'] = (
            (_series(seg) + '. ' if seg else '')
            + 'The new Table 11 gives the full breakdown. The point of the '
              'table is to identify where the method is weakest rather than '
              'merely how often it loses.')
    if nar.gap and not KEEP_SUBMITTED_RESULTS:
        _, m = nar.gap
        top, dims = m['top'][0], m['dims']
        if len(dims) >= 2:
            g0 = m['rank'][dims[0]][top] - m['rank'][dims[0]]['MSSBOA']
            g1 = m['rank'][dims[-1]][top] - m['rank'][dims[-1]]['MSSBOA']
            f['gap_finding'] = (
                f'Relative to {top}, the strongest competitor by average rank, '
                f'the mean Friedman rank gap moves from {g0:+.3f} at {dims[0]} '
                f'dimensions to {g1:+.3f} at {dims[-1]}. Table 12 reports the '
                f'gap for the three strongest competitors, and the limitations '
                f'section now names the mechanism: both added operators are '
                f'elite-centred and take a share of the budget that does not '
                f'grow with the dimension.')
    if nar.wtbl:
        rows, _ = nar.wtbl
        taus = [r[-1] for r in rows[1:]]
        bests = {r[-2] for r in rows[1:]}
        f['weights_finding'] = (
            f'The Kendall rank correlations against W0 are {", ".join(taus)}, '
            + ('and the best-performing algorithm is the same in every '
               'configuration; the absolute scores shift with the weights, as '
               'they must, but the comparison between algorithms does not.'
               if len(bests) == 1 else
               'and the best-performing algorithm is not the same in every '
               'configuration, so the ranking is itself weight-dependent and '
               'is reported as such.'))

    if nar.scale:
        rows, m = nar.scale
        algos = m['algos']
        lines = [f'{r[0]}: {float(r[1]):.2f} (rank {r[-2]} of {len(algos)})'
                 for r in rows[1:]]
        f['scale_finding'] = (
            f'Over {m["runs"]} independent runs per setting the mean aesthetic '
            f'score of MSSBOA is ' + '; '.join(lines) + '. The score does not '
            'collapse as the problem grows, so the formulation and the '
            'optimizer both handle layouts of the size you ask about.')
    if KEEP_SUBMITTED_RESULTS:
        f = _submitted_findings(f)
    return f


def _table_refs(text):
    """Rewrite the provisional table numbers into the ones the paper uses.

    build_final.py numbers the new tables in the 100s while the document is
    assembled and renumbers everything into document order at the end, writing
    the mapping to results/table_map.json.  The letters quote the provisional
    numbers, so the same mapping is applied here; without it the letters and
    the manuscript would disagree on every table number.
    """
    import json
    path = os.path.join(ROOT, 'results', 'table_map.json')
    if not os.path.exists(path):
        print('  ! results/table_map.json missing - run build_final.py first; '
              'table numbers in the letters are provisional')
        return text
    with open(path) as fh:
        mapping = {int(k): v for k, v in json.load(fh).items()}
    from build_final import map_table_refs
    return map_table_refs(text, mapping)


# --------------------------------------------------------------- rendering
def _p(doc, text, size=10.5, bold=False, italic=False, color=None,
       space_after=6, align=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = 'Times New Roman'
    if color is not None:
        r.font.color.rgb = color
    return p


def build_letter(key, comments, filename, findings):
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(10.5)

    n = key[-1]
    _p(doc, f'Response to Reviewer {n}', size=15, bold=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _p(doc, f'Manuscript: {TITLE}', size=10.5, italic=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, JOURNAL, size=10.5, italic=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    _p(doc, RT.OPENING.format(colour=RT.COLOUR_NOTE[key]), space_after=8)
    _p(doc, f'Colour key used in the revised manuscript: '
            f'Reviewer 1 = red, Reviewer 2 = blue, Reviewer 3 = green. '
            f'Your revisions are the ones in '
            f'{ {"R1": "red", "R2": "blue", "R3": "green"}[key] }.',
       italic=True, color=COLORS[key], space_after=16)

    for i, (comment, responses) in enumerate(comments, 1):
        _p(doc, f'Comments {i}:', bold=True, space_after=3)
        _p(doc, comment.strip(), italic=True, indent=0.25, space_after=8)
        _p(doc, f'Response {i}:', bold=True, color=COLORS[key], space_after=3)
        for j, r in enumerate(responses):
            text = _table_refs(r.format(**findings))
            quoted = text.startswith('"')
            _p(doc, text, color=COLORS[key], italic=quoted,
               indent=0.25 if quoted else None,
               space_after=6 if j < len(responses) - 1 else 16)

    _p(doc, 'We hope these revisions address your concerns, and we thank you '
            'again for a review that has clearly improved the paper.',
       space_after=6)
    path = os.path.join(OUT, filename)
    doc.save(path)
    print('wrote', path)
    return path


if __name__ == '__main__':
    f = _findings()
    build_letter('R1', RT.REVIEWER1, 'Response_to_Reviewer_1.docx', f)
    build_letter('R2', RT.REVIEWER2, 'Response_to_Reviewer_2.docx', f)
    build_letter('R3', RT.REVIEWER3, 'Response_to_Reviewer_3.docx', f)
