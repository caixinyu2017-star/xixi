# -*- coding: utf-8 -*-
"""书稿组装器：约定式Markdown -> 规范Word文档。

格式规范：
- 章标题（# 一、…）：黑体 小三(15pt) 加粗 居中，段前后空行
- 节标题（## （一）…）：黑体 四号(14pt)
- 条标题（### 1．…）：黑体 小四(12pt)
- 正文：宋体 小四(12pt)，西文Times New Roman，1.5倍行距，首行缩进2字符
- 图名（图下方）、表名（表上方）：黑体 五号(10.5pt) 居中
- 表格：三线表，表内宋体五号
- 公式：OMML 居中，编号 (章-序) 右顶格（制表位实现）
- 参考文献：宋体 五号，悬挂缩进，[n]按正文首次引用顺序编码
"""
import os, re, json, sys
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
import latex2mathml.converter
import mathml2omml

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
CH_DIR = os.path.join(ROOT, 'chapters')
FIG_DIR = os.path.join(ROOT, 'figs')

HEI = '黑体'; SONG = '宋体'; TIMES = 'Times New Roman'
OMML_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" ' \
          'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

def set_run_font(run, east=SONG, ascii_f=TIMES, size=12, bold=False, color=None):
    run.font.name = ascii_f
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ascii_f); rFonts.set(qn('w:hAnsi'), ascii_f)
    rFonts.set(qn('w:eastAsia'), east)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def para(doc, text='', east=SONG, ascii_f=TIMES, size=12, bold=False, align=None,
         indent_first=True, space_before=0, space_after=0, line=1.5, outline=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None: p.alignment = align
    if indent_first:
        # 首行缩进2字符（用 firstLineChars 保证随字号缩放）
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind'); ind.set(qn('w:firstLineChars'), '200'); ind.set(qn('w:firstLine'), '480')
        pPr.append(ind)
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    if line == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif line:
        pf.line_spacing = line
    if outline is not None:
        pPr = p._p.get_or_add_pPr()
        ol = OxmlElement('w:outlineLvl'); ol.set(qn('w:val'), str(outline)); pPr.append(ol)
    if text:
        add_rich_text(p, text, east=east, ascii_f=ascii_f, size=size, bold=bold)
    return p

INLINE_MATH = re.compile(r'\$([^$]+)\$')

GROUPCHR_FIX = re.compile(r'<m:groupChrPr>((?:<m:\w+(?:\s[^>]*)?/>)*)</m:groupChr>')

def latex_to_omml_xml(latex):
    mml = latex2mathml.converter.convert(latex)
    om = mathml2omml.convert(mml)
    # 修复 mathml2omml 对 \bar 生成的 groupChrPr 未闭合缺陷
    om = GROUPCHR_FIX.sub(r'<m:groupChrPr>\1</m:groupChrPr>', om)
    om = om.replace('<m:oMath>', f'<m:oMath {OMML_NS}>', 1)
    return om

def append_omml(p, latex):
    try:
        xml = latex_to_omml_xml(latex)
        p._p.append(parse_xml(xml))
        return True
    except Exception as e:
        r = p.add_run('[公式转换失败:%s]' % latex[:40])
        set_run_font(r, size=10, color='FF0000')
        print('  ! OMML fail:', latex[:60], e)
        return False

BOLD_SEG = re.compile(r'\*\*([^*]+)\*\*')

def add_rich_text(p, text, east=SONG, ascii_f=TIMES, size=12, bold=False):
    """处理行内公式$...$与**加粗**。"""
    pos = 0
    for m in INLINE_MATH.finditer(text):
        seg = text[pos:m.start()]
        _add_bold_aware(p, seg, east, ascii_f, size, bold)
        append_omml(p, m.group(1))
        pos = m.end()
    _add_bold_aware(p, text[pos:], east, ascii_f, size, bold)

def _add_bold_aware(p, seg, east, ascii_f, size, bold):
    if not seg: return
    pos = 0
    for m in BOLD_SEG.finditer(seg):
        if m.start() > pos:
            r = p.add_run(seg[pos:m.start()]); set_run_font(r, east, ascii_f, size, bold)
        r = p.add_run(m.group(1)); set_run_font(r, east, ascii_f, size, True)
        pos = m.end()
    if pos < len(seg):
        r = p.add_run(seg[pos:]); set_run_font(r, east, ascii_f, size, bold)

def add_equation(doc, latex, number):
    """居中公式 + 编号右顶格：制表位 center@7.5cm, right@15cm。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    for pos_cm, val in ((7.5, 'center'), (15.0, 'right')):
        t = OxmlElement('w:tab'); t.set(qn('w:val'), val)
        t.set(qn('w:pos'), str(int(pos_cm * 567)))
        tabs.append(t)
    pPr.append(tabs)
    r = p.add_run('\t'); set_run_font(r, size=12)
    append_omml(p, latex)
    r2 = p.add_run('\t(' + number + ')'); set_run_font(r2, size=12)
    return p

def add_caption(doc, text, before=6, after=6):
    return para(doc, text, east=HEI, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER,
                indent_first=False, space_before=before, space_after=after, line=1.3)

def add_figure(doc, num):
    fname = os.path.join(FIG_DIR, 'fig_%s.png' % num.replace('.', '_'))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    if os.path.exists(fname):
        run = p.add_run()
        run.add_picture(fname, width=Cm(14.2))
        return True
    else:
        r = p.add_run('【图%s 占位：%s 未生成】' % (num, fname))
        set_run_font(r, size=10, color='FF0000')
        print('  ! missing figure', fname)
        return False

def set_cell_borders_none(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge); e.set(qn('w:val'), 'none'); borders.append(e)
    tblPr.append(borders)

def set_row_border(row, where, sz):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
        e = borders.find(qn('w:'+where))
        if e is None:
            e = OxmlElement('w:'+where); borders.append(e)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz)); e.set(qn('w:color'), '000000')

def add_table(doc, rows):
    """rows: list[list[str]]，第一行为表头。三线表。"""
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    set_cell_borders_none(tbl)
    for i, rdata in enumerate(rows):
        for j in range(ncol):
            cell = tbl.cell(i, j)
            txt = rdata[j] if j < len(rdata) else ''
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i == 0 or _is_numlike(txt)) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.15
            add_rich_text(p, txt, east=SONG, size=10.5, bold=(i == 0))
    set_row_border(tbl.rows[0], 'top', 12)      # 顶线 1.5pt
    set_row_border(tbl.rows[0], 'bottom', 6)    # 栏目线 0.75pt
    set_row_border(tbl.rows[-1], 'bottom', 12)  # 底线 1.5pt
    return tbl

def _is_numlike(s):
    s = s.strip().replace('−', '-').replace('%', '').replace('*', '')
    if not s: return True
    try:
        float(s.replace(',', '').replace('—', '0')); return True
    except ValueError:
        return False

def parse_md_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i

# ---------------- 引用编码 ----------------
CITE_RE = re.compile(r'（([^（）]{1,80}?)，((?:19|20)\d{2}[a-c]?)）')

def order_refs(all_text, refs):
    """按正文首次出现顺序为参考文献编码。refs: list of {key, cite, gb}。"""
    first_pos = {}
    for ref in refs:
        cites = ref.get('cite_variants') or [ref['cite']]
        best = None
        for c in cites:
            c_norm = c.strip('（）')
            for m in re.finditer(re.escape(c_norm), all_text):
                if best is None or m.start() < best:
                    best = m.start()
                break
        first_pos[ref['key']] = best if best is not None else 10**12
    ordered = sorted(refs, key=lambda r: (first_pos[r['key']], r['gb']))
    used = [r for r in ordered if first_pos[r['key']] < 10**12]
    unused = [r for r in ordered if first_pos[r['key']] >= 10**12]
    return used, unused

# ---------------- 主流程 ----------------

def build(md_files, out_path, refs_json, meta):
    doc = Document()
    # 页面：A4，边距
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)
    # 默认样式兜底
    st = doc.styles['Normal']
    st.font.name = TIMES; st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)

    # ---- 封面 ----
    for _ in range(3): para(doc, '', indent_first=False)
    para(doc, meta['project'], east=SONG, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, space_after=12)
    para(doc, meta['category'], east=SONG, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, space_after=60)
    para(doc, meta['title'], east=HEI, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, space_after=6)
    para(doc, meta['subtitle'], east=HEI, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, space_after=80)
    para(doc, meta['author'], east=SONG, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    doc.add_page_break()

    # ---- 目录（域，打开时自动更新）----
    para(doc, '目　　录', east=HEI, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, space_after=12)
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r' TOC \o "1-2" \h \z \u ')
    r = OxmlElement('w:r'); t = OxmlElement('w:t')
    t.text = '（目录将在Word中打开时自动更新；如未更新请全选后按F9）'
    r.append(t); fld.append(r); p._p.append(fld)
    doc.add_page_break()
    # settings: updateFields
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); settings.append(uf)

    stats = {'figures': 0, 'tables': 0, 'equations': 0, 'missing_figs': [], 'chars': 0}
    all_text_parts = []

    with open(refs_json, encoding='utf-8') as f:
        refs = json.load(f)

    # 先合并全文以确定参考文献顺序
    merged = []
    for fp in md_files:
        with open(fp, encoding='utf-8') as f:
            merged.append(f.read())
    used_refs, unused_refs = order_refs('\n'.join(merged), refs)

    for fp in md_files:
        with open(fp, encoding='utf-8') as f:
            lines = f.read().split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            s = line.strip()
            if not s:
                i += 1; continue
            if s.startswith('<!--'):
                i += 1; continue
            if s.startswith('# '):
                txt = s[2:].strip()
                doc.add_page_break() if stats.get('chapters', 0) else None
                para(doc, txt, east=HEI, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     indent_first=False, space_before=12, space_after=14, outline=0)
                stats['chapters'] = stats.get('chapters', 0) + 1
                all_text_parts.append(txt)
            elif s.startswith('## '):
                txt = s[3:].strip()
                para(doc, txt, east=HEI, size=14, bold=False, indent_first=False,
                     space_before=10, space_after=8, outline=1)
                all_text_parts.append(txt)
            elif s.startswith('### '):
                txt = s[4:].strip()
                para(doc, txt, east=HEI, size=12, bold=False, indent_first=False,
                     space_before=8, space_after=6, outline=2)
                all_text_parts.append(txt)
            elif s.startswith('[[FIG:'):
                num = s[6:s.index(']]')]
                ok = add_figure(doc, num)
                if not ok: stats['missing_figs'].append(num)
                stats['figures'] += 1
                # 下一行应为图名
                if i + 1 < len(lines) and lines[i+1].strip().startswith('图'):
                    add_caption(doc, lines[i+1].strip()); i += 1
            elif s.startswith('[[TAB:'):
                # 下一行表名，随后为markdown表
                if i + 1 < len(lines) and lines[i+1].strip().startswith('表'):
                    add_caption(doc, lines[i+1].strip(), before=8, after=4); i += 1
                j = i + 1
                while j < len(lines) and not lines[j].strip().startswith('|'):
                    j += 1
                rows, j2 = parse_md_table(lines, j)
                if rows:
                    add_table(doc, rows)
                    para(doc, '', indent_first=False, space_after=6)
                stats['tables'] += 1
                i = j2 - 1
            elif s.startswith('[[EQ:'):
                num = s[5:s.index(']]')]
                # 收集 $$...$$（可能多行）
                j = i + 1
                buf = []
                started = False
                while j < len(lines):
                    ls_ = lines[j].strip()
                    if ls_.startswith('$$'):
                        if started:
                            buf.append(ls_.strip('$')); break
                        started = True
                        core = ls_.strip('$')
                        if core: buf.append(core)
                        if ls_.endswith('$$') and len(ls_) > 4: break
                        j += 1; continue
                    if started:
                        buf.append(ls_)
                    j += 1
                latex = ' '.join(x for x in buf if x).strip()
                add_equation(doc, latex, num)
                stats['equations'] += 1
                i = j
            else:
                para(doc, s)
                all_text_parts.append(s)
            i += 1

    # ---- 参考文献 ----
    doc.add_page_break()
    para(doc, '参考文献', east=HEI, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         indent_first=False, space_after=12, outline=0)
    for n, ref in enumerate(used_refs, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.3
        pf.space_after = Pt(2)
        pf.left_indent = Cm(0.9); pf.first_line_indent = Cm(-0.9)
        r = p.add_run('[%d] %s' % (n, ref['gb']))
        set_run_font(r, east=SONG, size=10.5)
    print('refs used:', len(used_refs), 'unused:', len(unused_refs))
    if unused_refs:
        print('  unused keys:', [r['key'] for r in unused_refs][:40])

    doc.save(out_path)

    body = '\n'.join(all_text_parts)
    stats['chars'] = len(re.sub(r'\s', '', body))
    return stats, used_refs, unused_refs

if __name__ == '__main__':
    files = sorted(sys.argv[1:-1]) if len(sys.argv) > 2 else None
    print('run via build_all.py')
