# -*- coding: utf-8 -*-
"""对 pandoc 生成的 docx 做成稿排版后处理，规格见仓库根目录 STYLE.md。

要点：
  1. 字体字号：题名黑体小二加粗；一级标题黑体小三加粗；二级标题黑体四号；
     三级标题黑体小四；正文宋体小四且段前段后为 0；图名表名五号黑体。
  2. 公式：pandoc 已把 LaTeX 数学转为 Word 原生公式对象（OMML）。此处把
     "公式段 + 紧随其后的编号段"合并为一段，设居中制表位与右对齐制表位，
     使公式居中、编号右顶格。
  3. 表格：三线表（顶线底线 1.5 磅、栏目线 0.75 磅，无竖线），列宽按内容计算。
  4. 插图：按版心宽度等比缩放。
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

SRC = sys.argv[1] if len(sys.argv) > 1 else "paper.docx"
doc = Document(SRC)

# 字号（磅）：小二18 小三15 四号14 小四12 五号10.5 小五9
SIZE = dict(title=18, h1=15, h2=14, h3=12, body=12, caption=10.5,
            table=10.5, note=9, ref=10.5, abstract=10.5)
PAGE_W, MARGIN = Cm(21.0), Cm(2.6)
TEXT_W = Emu(int(PAGE_W - 2 * MARGIN))            # 版心宽 15.8 cm
TAB_MID, TAB_END = Emu(int(TEXT_W / 2)), TEXT_W   # 制表位须为整数 EMU

# ---------- 页面 ----------
for sec in doc.sections:
    sec.page_width, sec.page_height = PAGE_W, Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = MARGIN


def set_fonts(style, east, size, bold=None, ascii_font="Times New Roman"):
    f = style.font
    f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    f.italic = False
    f.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east)
    f.color.rgb = RGBColor(0, 0, 0)


def para_fmt(obj, align=None, first_indent=None, before=None, after=None,
             line=None):
    pf = obj.paragraph_format
    if align is not None:
        pf.alignment = align
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line


S = {s.name: s for s in doc.styles}
CENTER, LEFT, JUST = (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.JUSTIFY)

# ---------- 正文：宋体小四，段前段后为 0 ----------
BODY_STYLES = ("Normal", "Body Text", "First Paragraph",
               "Body Text First Indent", "Plain Text")
for name in BODY_STYLES:
    if name in S:
        set_fonts(S[name], "宋体", SIZE["body"])
        para_fmt(S[name], align=JUST, before=0, after=0, line=1.5)
for name in ("Body Text", "First Paragraph"):
    if name in S:
        para_fmt(S[name], first_indent=Cm(0.85))     # 小四字下的 2 字符缩进

# ---------- 标题层级 ----------
HEADINGS = [("Heading 1", "h1", True, 12, 6),
            ("Heading 2", "h2", False, 10, 5),
            ("Heading 3", "h3", False, 8, 4),
            ("Heading 4", "h3", False, 6, 3)]
for name, key, bold, before, after in HEADINGS:
    if name in S:
        set_fonts(S[name], "黑体", SIZE[key], bold=bold)
        para_fmt(S[name], align=LEFT, before=before, after=after, line=1.5,
                 first_indent=Cm(0))

# ---------- 自定义样式 ----------
custom = {
    "TitleCN": ("黑体", SIZE["title"], True, CENTER, 6, 12),
    "AuthorCN": ("仿宋", SIZE["abstract"], False, CENTER, 0, 3),
    "AffilCN": ("宋体", SIZE["note"], False, CENTER, 0, 10),
    "AbstractCN": ("楷体", SIZE["abstract"], False, JUST, 0, 4),
    "TableNote": ("宋体", SIZE["note"], False, LEFT, 2, 8),
    "RefItem": ("宋体", SIZE["ref"], False, JUST, 0, 2),
    "Table Caption": ("黑体", SIZE["caption"], False, CENTER, 8, 3),
    "Image Caption": ("黑体", SIZE["caption"], False, CENTER, 3, 8),
    "Captioned Figure": ("宋体", SIZE["body"], False, CENTER, 8, 0),
    "Compact": ("宋体", SIZE["table"], False, CENTER, 0, 0),
}
for name, (east, size, bold, align, before, after) in custom.items():
    if name in S:
        set_fonts(S[name], east, size, bold=bold)
        para_fmt(S[name], align=align, before=before, after=after, line=1.2,
                 first_indent=Cm(0))

# ============================================================
# 公式：居中排版，编号右顶格
# ============================================================
EQNUM = re.compile(r"^\(\s*\d+\s*[-–—]\s*\d+\s*\)$")


def _run(text=None, tab=False, size=SIZE["body"]):
    """构造一个 run，可含制表符与文本。"""
    r = doc.element.makeelement(qn("w:r"), {})
    rpr = r.makeelement(qn("w:rPr"), {})
    rpr.append(rpr.makeelement(qn("w:rFonts"),
                               {qn("w:ascii"): "Times New Roman",
                                qn("w:hAnsi"): "Times New Roman",
                                qn("w:eastAsia"): "宋体"}))
    rpr.append(rpr.makeelement(qn("w:sz"), {qn("w:val"): str(int(size * 2))}))
    r.append(rpr)
    if tab:
        r.append(r.makeelement(qn("w:tab"), {}))
    if text is not None:
        t = r.makeelement(qn("w:t"), {})
        t.text = text
        r.append(t)
    return r


def unwrap_math_para(p):
    """把 pandoc 生成的块级 m:oMathPara 拆成行内 m:oMath。

    m:oMathPara 在 Word 中独占一行，其前后的文本 run 会被挤到另起一行，
    无法与编号排在同一行；改为行内公式后才能用制表位实现"公式居中、
    编号右顶格"。
    """
    for mp in p._p.findall(qn("m:oMathPara")):
        for child in list(mp):
            if child.tag == qn("m:oMathParaPr"):
                continue
            mp.addprevious(child)
        mp.getparent().remove(mp)


paras = list(doc.paragraphs)
n_eq = 0
for i, p in enumerate(paras):
    # 仅处理独立成段的行间公式（pandoc 以 m:oMathPara 标记），行内公式不动
    if "<m:oMathPara" not in p._p.xml:
        continue
    unwrap_math_para(p)
    maths = p._p.findall(qn("m:oMath"))
    if not maths:
        continue
    para_fmt(p, align=LEFT, first_indent=Cm(0), before=6, after=6, line=1.5)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(TAB_MID, WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(TAB_END, WD_TAB_ALIGNMENT.RIGHT)
    maths[0].addprevious(_run(tab=True))           # 居中制表位前导
    nxt = paras[i + 1] if i + 1 < len(paras) else None
    if nxt is not None and EQNUM.match(nxt.text.strip()):
        maths[-1].addnext(_run(nxt.text.strip()))
        maths[-1].addnext(_run(tab=True))          # 右顶格制表位
        nxt._p.getparent().remove(nxt._p)          # 删除独立成段的编号
        n_eq += 1

# ============================================================
# 表格：三线表 + 按内容计算列宽
# ============================================================
TOP_BOTTOM_SZ, HEADER_SZ = "12", "6"            # 1.5 磅与 0.75 磅
TEXT_TWIPS = int(round(15.8 / 2.54 * 1440))
CELL_MARGIN, PAD, UNIT = 216, 120, 120          # 五号字下半角字符约 105 缇
MIN_DW = 10                                     # 每列至少容纳 5 个汉字的宽度
UNBREAK = re.compile(r"[0-9A-Za-z.,%*+\-−–—()（）:：/×±≤≥Ͱ-Ͽ"
                     r"⁰-₟]+")


def _border(parent, tag, val, sz=None):
    attrs = {qn("w:val"): val, qn("w:color"): "000000"}
    if sz is not None:
        attrs[qn("w:sz")] = sz
        attrs[qn("w:space")] = "0"
    parent.append(parent.makeelement(qn("w:" + tag), attrs))


def _dw(s):
    return sum(2 if ord(ch) > 0x2E80 else 1 for ch in s)


def demand(tbl):
    """逐列统计：full 为最长单元格的显示宽度，brk 为最长不可断串的显示宽度。"""
    ncol = len(tbl.columns)
    need, mini = [], []
    for j in range(ncol):
        full = brk = 0
        for row in tbl.rows:
            if j >= len(row.cells):
                continue
            t = row.cells[j].text.strip()
            full = max(full, _dw(t))
            for m in UNBREAK.findall(t):
                brk = max(brk, _dw(m))
        need.append(max(full, 3))
        # 中文串可以换行，但仍需保证不至于窄到每行只放一两个字：
        # 下限取"不可断串宽度"与"MIN_DW 或该列实际需求（取小者）"中的较大值。
        mini.append(max(brk, min(need[-1], MIN_DW), 3))
    return need, mini


def fit_size(tbl):
    """选取能让所有不可断串（如"（−0.3321）"）整体容纳于版心的最大字号。

    列数较多的回归结果表在五号字下常常放不下，若强行压缩列宽会把括号或负号
    甩到下一行；此处改为按 STYLE.md 允许的字号档位逐级下调，直至恰好放得下。
    """
    _, mini = demand(tbl)
    for fs in (10.5, 10, 9.5, 9, 8.5, 8):
        unit = UNIT * fs / SIZE["table"]
        if sum(m * unit + CELL_MARGIN + PAD for m in mini) <= TEXT_TWIPS:
            return fs
    return 8.0


def col_widths(tbl, fs):
    """先保证最长不可断串不被拆行，余量再按显示宽度需求比例分配。"""
    ncol = len(tbl.columns)
    need, mini = demand(tbl)
    unit = UNIT * fs / SIZE["table"]
    base = [m * unit + CELL_MARGIN + PAD for m in mini]
    if sum(base) > TEXT_TWIPS:
        k = TEXT_TWIPS / sum(base)
        return [int(b * k) for b in base]
    extra = TEXT_TWIPS - sum(base)
    gap = [max(need[j] - mini[j], 0) * unit for j in range(ncol)]
    tot = sum(gap)
    out = ([base[j] + int(extra * gap[j] / tot) for j in range(ncol)]
           if tot else [b + extra // ncol for b in base])
    cap = int(TEXT_TWIPS * (0.45 if ncol >= 3 else 0.62))
    surplus = sum(max(0, w - cap) for w in out)
    if surplus:
        free = [j for j, w in enumerate(out) if w <= cap]
        out = [min(w, cap) for w in out]
        for j in free:
            out[j] += surplus // len(free)
    out[-1] += TEXT_TWIPS - sum(out)
    return out


for tbl in doc.tables:
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout", "w:tblBorders", "w:tblStyle"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    tblPr.append(tblPr.makeelement(qn("w:tblW"),
                                   {qn("w:w"): str(TEXT_TWIPS),
                                    qn("w:type"): "dxa"}))
    tblPr.append(tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"}))
    bd = tblPr.makeelement(qn("w:tblBorders"), {})
    _border(bd, "top", "single", TOP_BOTTOM_SZ)
    _border(bd, "left", "none")
    _border(bd, "bottom", "single", TOP_BOTTOM_SZ)
    _border(bd, "right", "none")
    _border(bd, "insideH", "none")
    _border(bd, "insideV", "none")
    tblPr.append(bd)

    fs = fit_size(tbl)
    ws = col_widths(tbl, fs)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, wd in zip(grid.findall(qn("w:gridCol")), ws):
            gc.set(qn("w:w"), str(wd))
    nrow = len(tbl.rows)
    for ri, row in enumerate(tbl.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
        if ri == 0:
            trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            if j < len(ws):
                tcPr.append(tcPr.makeelement(qn("w:tcW"),
                                             {qn("w:w"): str(ws[j]),
                                              qn("w:type"): "dxa"}))
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tb = tcPr.makeelement(qn("w:tcBorders"), {})
            if ri == 0:
                _border(tb, "top", "single", TOP_BOTTOM_SZ)
                _border(tb, "bottom", "single", HEADER_SZ)
            elif ri == nrow - 1:
                _border(tb, "top", "none")
                _border(tb, "bottom", "single", TOP_BOTTOM_SZ)
            else:
                _border(tb, "top", "none")
                _border(tb, "bottom", "none")
            _border(tb, "left", "none")
            _border(tb, "right", "none")
            tcPr.append(tb)
            for para in cell.paragraphs:
                para.alignment = CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in para.runs:          # 落实自适应字号
                    run.font.size = Pt(fs)

# ---------- 插图：按版心宽度等比缩放 ----------
for shape in doc.inline_shapes:
    if shape.width and shape.width != TEXT_W:
        ratio = TEXT_W / shape.width
        shape.width = TEXT_W
        shape.height = int(shape.height * ratio)

# ---------- 清除正文段落上的直接间距设置 ----------
for p in doc.paragraphs:
    if p.style.name in BODY_STYLES and "<m:oMath" not in p._p.xml:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

doc.save(SRC)
print(f"finalize done: {SRC}（公式 {n_eq} 个已设编号右顶格）")
