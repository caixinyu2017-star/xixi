# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter and
the manuscript cannot disagree.
"""
import json
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"), encoding="utf-8") as fh:
    S = json.load(fh)
S["years"] = S["years"].replace("-", "\u2013")
GOVCUT = 100.0 * (1.0 - S["mod_gov_effect"] / S["mod_gov_b1"])
MED = S["mediation"]

FONT = "Times New Roman"
SIZE = 10.0

TITLE = ("The Narrowing Entry Port: Generative Artificial Intelligence "
         "Adoption, Task Routinisation and Youth Employment in Chinese "
         "Listed Firms")

BODY = [
    ("We would like to submit the enclosed manuscript entitled “%s” for "
     "consideration in the Special Issue “Systems Approaches to Generative AI: "
     "Workforce Development, Organisational Learning, and Economic "
     "Transformation” of Systems. The manuscript has not been published "
     "elsewhere and is not under consideration by any other journal. No "
     "conflict of interest exists in its submission, and it has been approved "
     "for publication by all authors." % TITLE),

    ("Generative artificial intelligence diffuses fastest through exactly the "
     "codified, routine-cognitive work that firms have traditionally allocated "
     "to new entrants, so the entry port of the professional career lies "
     "directly in the path of the technology. Whether that port narrows is an "
     "empirical question about firms, and it has not yet been answered with "
     "firm-level evidence. Using a panel of Chinese A-share listed firms over "
     f"{S['years']} ({S['n_obs']:,} firm-year observations for "
     f"{S['n_firms']:,} firms), a text-based measure of adoption and the age "
     "composition of the workforce disclosed in annual reports, the manuscript "
     "estimates how adoption reshapes the age structure of corporate "
     "employment and identifies the organisational conditions that govern the "
     "outcome."),

    "The contributions of the paper are as follows.",
]

POINTS = [
    ("(1) It provides direct firm-level evidence on a question that has so far "
     "been addressed with occupational or vacancy data. A one-unit increase in "
     "generative AI adoption intensity lowers the share of employees aged 30 "
     f"or below by {abs(S['beta_baseline']):.3f} percentage points, "
     f"{abs(S['econ_pct']):.1f}% of the sample mean for a one-standard-"
     "deviation increase. The estimate survives instrumental-variable, "
     "matched-sample, entropy-balanced, event-study and randomisation-based "
     "inference, and the parallel-trends assumption is not rejected."),

    ("(2) It opens the black box and finds two channels working against each "
     "other. Adoption compresses routine-cognitive task intensity, which "
     f"transmits {abs(MED['RTI']['share']):.1f}% of the total effect and is "
     "displacing; it simultaneously raises the human-capital requirement of "
     f"the remaining work, which transmits {abs(MED['HCU']['share']):.1f}% in "
     "the opposite direction and is complementary. The net figure that the "
     "literature reports is therefore a residual between two economically "
     "distinct forces, which is why estimates across settings disagree."),

    ("(3) It shows that the outcome is a property of the firm as a "
     "socio-technical system rather than of the technology. The displacement "
     f"effect is about {GOVCUT:.0f}% smaller in firms that disclose AI "
     "governance arrangements and is statistically indistinguishable from "
     "zero above the "
     f"{S['mod_olc_turn_pct']:.0f}th percentile of organisational learning "
     "capability. The same technology narrows the entry port in one firm and "
     "leaves it open in another, and the difference is organisational design."),

    ("(4) It is transparent and fully reproducible. Identification rests on a "
     "shift-share and a peer-diffusion instrument that pass the "
     "Kleibergen–Paap and overidentification tests, selection on unobservables "
     "is bounded using coefficient-stability methods, and every number in "
     "every table and figure is produced by the released estimation code."),
]

CLOSING = (
    "We believe the manuscript speaks directly to the aims of the Special "
    "Issue. It treats generative AI adoption as an organisational rather than "
    "a purely technological event, identifies the learning and governance "
    "mechanisms through which it reaches the labour market, and converts them "
    "into levers that firms and policymakers can actually pull. Youth "
    "employment is where the workforce-development consequences of the "
    "technology appear first, and the evidence here indicates that they are "
    "not predetermined. We deeply appreciate your consideration of our "
    "manuscript and look forward to receiving comments from the reviewers. "
    "Correspondence should be directed to the address below.")

SIGN = [
    "Name: Tiantian Mo",
    "Institution and address: College of Business, Jiaxing University; "
    "Jiaxing 314001, China",
    "Email: 00008227@zjxu.edu.cn",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(3.0)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def p(text, after=5, bold=False, italic=False, align=None, indent=0.0,
          size=SIZE):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        run = par.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        return par

    p("Dear Editors:", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    for t in BODY:
        p(t, after=5)
    for t in POINTS:
        p(t, after=4)
    p(CLOSING, after=7)

    p("Thanks very much for your attention to our paper.", after=8,
      align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in SIGN:
        p(line, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("Very sincerely yours,", after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p("Xinyu Cai and Tiantian Mo", after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
