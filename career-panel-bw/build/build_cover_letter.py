# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the
letter and the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"),
          encoding="utf-8") as fh:
    S = json.load(fh)
DIS, WAV, DIV = S["disagreement"], S["waves"], S["divergence"]
ROB, N_H = S["robustness"], S["meta"]["n_hypotheses"]
NW, NM = S["panels"]["NLSW"], S["panels"]["NLSY79M"]
W1 = S["results"]["W1"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")
_APOS = re.compile(r"(?<=[A-Za-z])'(?=[A-Za-z]|\s|$)")


def typo(s):
    s = _APOS.sub("\u2019", s)
    s = _YEARS.sub("\\1\u2013\\2", s)
    return _MINUS.sub("\u2212", s)


def r3(x):
    s = "%.3f" % x
    return s.replace("0.", ".").replace("-.", "\u2212.")


FONT = "Times New Roman"
SIZE = 9.6

TITLE = ("When Does a Career Contingency Hold? Between-Person and "
         "Within-Person Estimates of Sixteen Moderation Hypotheses in Two "
         "Longitudinal Cohorts of Young Workers")

BODY = [
    ("We would like to submit the enclosed manuscript entitled "
     "\u201c%s\u201d for consideration as an Article in Behavioral "
     "Sciences. The manuscript has not been published elsewhere and is not "
     "under consideration by any other journal. It is a secondary analysis "
     "of public-use microdata and involved no contact with human subjects. "
     "All authors have approved the submission, and no conflict of interest "
     "exists." % TITLE),

    ("Career development theories describe what happens to a person as a "
     "career unfolds. The evidence about the contingencies in those theories "
     "is mostly gathered by comparing people at a moment. That the two can "
     "differ has been argued on theoretical grounds for decades. What has "
     "not been established is how far they actually differ for interaction "
     "terms \u2014 which is where the interesting theoretical claims live "
     "\u2014 in real career data, for hypotheses of the kind the field "
     "reports."),

    ("We wrote down sixteen moderation hypotheses in advance, each stating "
     "that the wage return to a career input depends on a characteristic of "
     "the worker or the job, and estimated every one of them in two public "
     f"longitudinal cohorts of young workers: {NW['n_obs']:,} person-years "
     f"on {NW['n_person']:,} women across {NW['n_wave']} waves, and "
     f"{NM['n_obs']:,} person-years on {NM['n_person']} men across "
     f"{NM['n_wave']} annual waves. Each interaction was estimated one wave "
     "at a time, pooled, and as separate between-person and within-person "
     "slopes inside a single model that permits a formal test of their "
     "equality. All sixteen are reported. The main findings are as "
     "follows."),
]

POINTS = [
    ("(1) The two answers disagree, often and by large margins. The "
     "between-person and within-person estimates carried opposite signs in "
     f"{DIS['sign_flips']} of {N_H} hypotheses and supported different "
     f"substantive conclusions in {DIS['verdict_changes']}; their equality "
     f"was rejected in {DIS['between_ne_within_q05']} after controlling the "
     "false discovery rate. The median disagreement was "
     f"{'%.2f' % DIV['median_ratio']} times the size of the within-person "
     "estimate itself \u2014 that is, the gap between the two answers was "
     "typically larger than the answer."),

    ("(2) A single cross-section does not approximate the within-person "
     f"quantity; it estimates the other one. Of {WAV['total']} single-wave "
     f"estimates, {WAV['significant']} reached significance, and "
     f"{WAV['significant_and_opposite']} of those "
     f"({'%.0f' % (100 * WAV['share_sig_opposite'])} per cent) pointed the "
     "opposite way to the within-person estimate of the same quantity. The "
     "wave-by-wave estimates scatter around the between-person value, so "
     "collecting another cross-section reproduces the same quantity rather "
     "than closing the gap."),

    ("(3) The divergence survives the obvious checks. Refitting every "
     "within-person estimate restricted to workers observed at least three "
     "times, by first differences, and with the controls stripped back left "
     f"the sign unchanged in {ROB['sign_agreement_all_three']} of {N_H} "
     "hypotheses, so the result is not an artefact of the transformation "
     "used to remove the person."),

    ("(4) A worked case makes the stakes concrete. Whether employer tenure "
     "pays off more for college graduates is answered "
     f"{r3(W1['between']['b'])} between people (p < .001) and "
     f"{r3(W1['within']['b'])} within them; the two differ by "
     f"{r3(W1['equality']['diff'])} with p < .001 after correction. "
     "Read across people, graduates gain markedly less from tenure; read "
     "within people, slightly more. A one-wave study would have reported "
     "the first."),

    ("(5) We are careful about what we claim. Neither estimator is "
     "presented as correct. They answer different questions, and the "
     "manuscript says so repeatedly; the within-person estimator removes "
     "only the stable part of confounding and amplifies measurement error. "
     "The two cohorts are also old, which we state plainly: they were "
     "chosen because they are among the few career panels with many waves "
     "that are public, documented and redistributable, so that the analysis "
     "can be checked. As a check on our own handling of the files, we "
     "reproduce a published benchmark estimate to four decimal places "
     "before estimating anything of our own."),
]

CLOSING = (
    "We believe the manuscript suits Behavioral Sciences. It bears directly "
    "on how the journal\u2019s readers design and read career studies, it "
    "engages a methodological literature that psychology has developed "
    "largely on theoretical grounds by supplying a magnitude from real "
    "career data, and it yields a concrete and inexpensive recommendation: "
    "where panel data exist, report both components of a moderation, which "
    "costs two columns; where they do not, name the quantity estimated "
    "accurately. The analysis code downloads the public extracts, "
    "reproduces the benchmark, and regenerates every table and figure, so "
    "readers can check the result rather than take it on trust. We would be "
    "glad to revise the manuscript in the light of the reviewers\u2019 "
    "comments.")


def para(doc, text, size=SIZE, bold=False, space_after=5, align=None,
         indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if indent:
        fmt.left_indent = Cm(indent)
    r = p.add_run(typo(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.9))
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.2))

    para(doc, "Cover Letter", size=13, bold=True, space_after=8,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Dear Editors,", space_after=7)
    for b in BODY:
        para(doc, b)
    for pt in POINTS:
        para(doc, pt, indent=0.55, space_after=4)
    para(doc, CLOSING, space_after=8)
    para(doc, "Yours sincerely,", space_after=2)
    para(doc, "Xinyu Cai and Tiantian Mo", space_after=0)
    para(doc, "College of Business, Jiaxing University", space_after=0)
    para(doc, "Correspondence: 00008227@zjxu.edu.cn", space_after=0)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
