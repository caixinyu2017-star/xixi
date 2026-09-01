# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the
letter and the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"),
          encoding="utf-8") as fh:
    S = json.load(fh)
RF, EN = S["reference"], S["ensemble"]
TH_, PR, EU = S["threshold"], S["params"], S["eu"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")
_APOS = re.compile(r"(?<=[A-Za-z])'(?=[A-Za-z]|\s|$)")


def typo(s):
    s = _APOS.sub("\u2019", s)
    s = _YEARS.sub("\\1\u2013\\2", s)
    return _MINUS.sub("\u2212", s)


FONT = "Times New Roman"
SIZE = 9.8

TITLE = ("Siting Urban Green Infrastructure by Workplace Rather than "
         "Residence: Heat-Attributable Work Capacity Loss Among Young "
         "Workers under Exposure\u2013Response Uncertainty")

BODY = [
    ("We would like to submit the enclosed manuscript entitled "
     "\u201c%s\u201d for consideration in Sustainability, for the Special "
     "Issue \u201cThermal Mitigation Effects of Green and Blue "
     "Infrastructure and Urban Sustainability from an Interdisciplinary "
     "Perspective\u201d. The manuscript has not been published elsewhere "
     "and is not under consideration by any other journal. No conflict of "
     "interest exists in its submission, and it has been approved for "
     "publication by all authors." % TITLE),

    ("The Special Issue asks how the strategic implementation of green and "
     "blue infrastructure mitigates urban heat, and it lists the "
     "socioeconomic impacts of heatwaves, heat mitigation strategies, urban "
     "greenery, urban form and urban ventilation among its themes. The "
     "enclosed manuscript takes up a question that sits inside that list but "
     "has not been asked quantitatively: when a city decides where to plant, "
     "whose exposure should it count? Green infrastructure is sited almost "
     "everywhere by residential population and residential deprivation, "
     "while heat damages work where work happens \u2014 and entry-level "
     "outdoor employment, through which young Europeans enter the labour "
     "market, is concentrated in precisely the industrial and logistics "
     "districts where few people live."),

    "The main findings are as follows.",
]

POINTS = [
    ("(1) Siting rule dominates. Weighting a fixed planting budget by "
     "exposed workplaces rather than by residents protects "
     f"{RF['ratio_exposure_to_population_heat']:.2f} times as many hours of "
     "young workers\u2019 work capacity, against a control rule that targets "
     "the same heat but counts residents. Following residential population "
     "is worse than not targeting at all: the uniform rule outperforms it."),

    ("(2) The advantage is conditional, and the condition is measurable. It "
     "is absent where workplaces and homes coincide and reaches "
     f"{TH_['median_ratio_top_bin']:.1f}-fold where they are opposed, "
     "crossing into materiality once the correlation between workplace and "
     "residential density falls below roughly 0.25. Cities can compute that "
     "correlation from data they already hold, so the recommendation is "
     "checkable rather than universal."),

    ("(3) A negative finding we think the field needs. The magnitude of the "
     "benefit is not identified. Across "
     f"{EN['n']:d} draws over {EN['n_params']:d} parameters, five published "
     "exposure\u2013response functions, four climate settings and the "
     "geography of employment, the estimated loss varies by more than an "
     "order of magnitude and vanishes entirely in "
     f"{100 * EN['share_zero_loss']:.1f} per cent of draws, because the "
     "published response functions disagree about whether European summer "
     "conditions cause any loss at all. Point estimates in this genre are "
     "reporting a choice of function. Rankings survive; levels do not."),

    ("(4) Transparency about what the study is. This is a model-based "
     "assessment, and the manuscript says so in the abstract, the methods "
     "and the data availability statement. The city is synthetic, because "
     "the question \u2014 where cooling ought to go, as against where it "
     "goes \u2014 is counterfactual and no observational design answers it. "
     "Its residential density and its four climate settings are anchored to "
     f"the {EU['ucdb']['n']:d} EU-27 urban centres of the Global Human "
     "Settlement Layer and to Berkeley Earth country summer temperatures. "
     f"All {PR['n']:d} parameters are listed in an appendix with their "
     f"provenance: {PR['counts']['literature']:d} are taken from cited "
     f"sources and {PR['counts']['assumed']:d} are modelling choices, which "
     "are swept across declared intervals rather than defended."),
]

CLOSING = (
    "We believe the manuscript fits the Special Issue closely. It bridges "
    "microclimate, urban design and labour economics, which is the "
    "cross-sectoral perspective the call champions, and it delivers an "
    "actionable insight: the 3\u201330\u2013300 rule already names "
    "workplaces alongside homes and schools, but compliance has only ever "
    "been evaluated for dwellings, and our results indicate that the two "
    "definitions of the population served are not interchangeable. We hope "
    "the manuscript proves suitable for the Special Issue, and we would be "
    "pleased to revise it in the light of the reviewers\u2019 comments.")


def para(doc, text, size=SIZE, bold=False, space_after=5, align=None,
         indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if indent:
        fmt.left_indent = Cm(indent)
    r = p.add_run(typo(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.9))
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.2))

    para(doc, "Cover Letter", size=13, bold=True, space_after=8,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Dear Dr. Rus and Dr. Gaban, Guest Editors,",
         space_after=7)
    for b in BODY:
        para(doc, b)
    for pt in POINTS:
        para(doc, pt, indent=0.55, space_after=4)
    para(doc, CLOSING, space_after=8)
    para(doc, "Yours sincerely,", space_after=2)
    para(doc, "Xinyu Cai and Tiantian Mo", space_after=0)
    para(doc, "College of Business, Jiaxing University", space_after=0)
    para(doc, "Correspondence: 00008227@zjxu.edu.cn", space_after=0)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
