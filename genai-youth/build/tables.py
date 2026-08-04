# -*- coding: utf-8 -*-
"""Table and figure rendering for the manuscript.

Tables use the MDPI three-line format: a rule above the header, a rule below the
header, and a rule below the last row; no vertical rules.  Captions sit ABOVE the
table.  Figure captions and notes sit BELOW the image.
"""
import os

import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
TAB = os.path.join(HERE, '..', 'tables')
FIG = os.path.join(HERE, '..', 'figures')

# key -> (csv file, number, caption, note, column widths in cm or None, font pt)
TABLES = {
    'table1': ('table2_variable_definitions.csv', 1,
               'Definitions of the variables.',
               None, [1.9, 4.2, 10.2], 7.0),
    'table2': ('table3_descriptive_statistics.csv', 2,
               'Descriptive statistics.',
               'The sample comprises 31,015 firm-year observations for 3,240 Chinese '
               'A-share listed firms over 2014–2024. Continuous variables are winsorised '
               'at the 1st and 99th percentiles.', None, 7.5),
    'table3': ('table4_correlation.csv', 3,
               'Pearson correlation coefficients for the main variables.',
               'Superscripts ***, ** and * denote significance at the 1%, 5% and 10% '
               'levels, respectively.', None, 6.6),
    'table4': ('table5_baseline.csv', 4,
               'Effect of generative artificial intelligence exposure on the youth '
               'employment share: baseline estimates.',
               'The dependent variable is the youth employment share. Treat is the '
               'interaction of the post-2022 indicator with mean-centred entry-task '
               'exposure. Standard errors clustered at the firm level are reported in '
               'parentheses. All columns include firm fixed effects; columns (2) to (4) '
               'also include year fixed effects. Superscripts ***, ** and * denote '
               'significance at the 1%, 5% and 10% levels, respectively.', None, 7.5),
    'table5': ('table10_endogeneity.csv', 5,
               'Endogeneity: instrumental variables, propensity-score matching and '
               'entropy balancing.',
               'Column (1) instruments disclosed adoption intensity with a shift-share '
               'instrument and a historical-infrastructure instrument, estimated on '
               'two-way within-transformed data. Column (2) re-estimates the baseline on '
               'a 1:4 nearest-neighbour matched sample. Column (3) reweights the '
               'comparison group by entropy balancing on the first two moments of the '
               'pre-shock covariates. Standard errors clustered at the firm level are '
               'reported in parentheses. Superscripts ***, ** and * denote significance '
               'at the 1%, 5% and 10% levels, respectively.', None, 7.5),
    'table6': ('table12_robustness.csv', 6,
               'Robustness checks.',
               'Each row re-estimates the baseline specification of Equation (3) under '
               'the stated change. The reported coefficient is that on Treat, except in '
               'the binary-treatment row, where it is the coefficient on the interaction '
               'of the post-2022 indicator with an above-median exposure dummy. Standard '
               'errors are in parentheses. Superscripts ***, ** and * denote '
               'significance at the 1%, 5% and 10% levels, respectively.', None, 7.5),
    'table7': ('table7_mechanism.csv', 7,
               'Transmission channels: entry-task restructuring and training intensity.',
               'Columns (2) and (4) estimate Equation (5) with the stated mediator as '
               'the dependent variable; columns (3) and (5) estimate Equation (6). All '
               'columns include the full control set, firm fixed effects and year fixed '
               'effects. Standard errors clustered at the firm level are in parentheses. '
               'Superscripts ***, ** and * denote significance at the 1%, 5% and 10% '
               'levels, respectively.', None, 7.5),
    'table8': ('table8_bootstrap_mediation.csv', 8,
               'Bootstrap inference on the indirect effects.',
               'Indirect effects are the products of the coefficients from Equations (5) '
               'and (6). Confidence intervals are percentile intervals from a firm-level '
               'block bootstrap with 400 replications of 600 firms each. The share of '
               'the total effect is the indirect effect divided by the baseline '
               'coefficient of −0.1175.', None, 7.2),
    'table9': ('table9_moderation.csv', 9,
               'Moderating effects of organisational learning capacity and internal '
               'labour-market depth.',
               'Both moderators are standardised to mean zero and unit variance. All '
               'columns include the full control set, firm fixed effects and year fixed '
               'effects. Standard errors clustered at the firm level are in parentheses. '
               'Superscripts ***, ** and * denote significance at the 1%, 5% and 10% '
               'levels, respectively.', None, 7.5),
    'table10': ('table13_heterogeneity.csv', 10,
                'Heterogeneity by ownership, technology intensity, factor intensity and '
                'region.',
                'Each column re-estimates Equation (3) on the stated subsample with the '
                'full control set, firm fixed effects and year fixed effects. Standard '
                'errors clustered at the firm level are in parentheses. Tests of '
                'between-group differences are reported in Table A2. Superscripts ***, '
                '** and * denote significance at the 1%, 5% and 10% levels, '
                'respectively.', None, 6.8),
    'table11': ('table15_pipeline.csv', 11,
                'Downstream effect on the internal advancement rate.',
                'The dependent variable is the share of employees promoted into '
                'supervisory grades. All columns include the full control set, firm '
                'fixed effects and year fixed effects. Standard errors clustered at the '
                'firm level are in parentheses. Superscripts ***, ** and * denote '
                'significance at the 1%, 5% and 10% levels, respectively.', None, 7.5),
    'tableA1': ('tableA1_vif.csv', 'A1',
                'Variance inflation factors.',
                'Computed from a pooled regression of the youth employment share on the '
                'treatment, disclosed adoption intensity and the full control set.',
                None, 7.5),
    'tableA2': ('table14_group_differences.csv', 'A2',
                'Tests of differences between subsample coefficients.',
                'Differences are tested with a z statistic constructed from the two '
                'subsample coefficients and their clustered standard errors, under the '
                'assumption that the subsamples are independent.', None, 7.5),
}

# key -> (image file, number, caption, note, width in cm)
FIGURES = {
    'fig2': ('fig2_ladder.png',
             'Compression of the junior practice rung under generative artificial '
             'intelligence.',
             'Entry-level tasks are ordered by complexity on the horizontal axis. The '
             'model performs the interval below a, junior employees perform the interval '
             'between a and C, and experienced employees perform the remainder. An '
             'increase in the model-performed interval compresses the junior rung from '
             'below. The figure is a schematic representation of the task-assignment '
             'logic; the segment widths are illustrative and are not drawn to estimated '
             'values.', 13.5),
    'fig3': ('fig3_event_study.png',
             'Event-study estimates of the effect of entry-task exposure on the youth '
             'employment share.',
             'Coefficients are from Equation (4), with 2021 normalised to zero. Vertical '
             'bars are 95% confidence intervals constructed from standard errors '
             'clustered at the firm level. The vertical dashed line marks the timing of '
             'the generative artificial intelligence shock.', 11.5),
    'fig4': ('fig4_psm_balance.png',
             'Covariate balance before and after propensity-score matching.',
             'Standardised percentage bias is the difference in means between the '
             'treated and control groups divided by the pooled standard deviation. '
             'Matching is 1:4 nearest neighbour on the pre-shock means of the control '
             'set, with a caliper of 0.05 propensity-score standard deviations. The '
             'dotted lines mark the conventional ±10% threshold.', 11.5),
    'fig5': ('fig7_placebo.png',
             'Randomisation-inference distribution of placebo estimates.',
             'Entry-task exposure is randomly reassigned across firms and the baseline '
             'specification of Equation (3) re-estimated. The solid vertical line marks '
             'the estimate obtained with the true assignment of exposure.', 11.5),
    'fig6': ('fig5_marginal_effect.png',
             'Marginal effect of treatment across the range of organisational learning '
             'capacity.',
             'The marginal effect is π1 + π3M from Equation (7), with a 95% confidence '
             'band constructed from the full variance–covariance matrix of the two '
             'coefficients. The shaded histogram shows the sample distribution of the '
             'moderator.', 11.5),
    'fig7': ('fig6_heterogeneity.png',
             'Subsample estimates of the effect of entry-task exposure on the youth '
             'employment share.',
             'Each estimate comes from Equation (3) estimated on the stated subsample. '
             'Horizontal bars are 95% confidence intervals constructed from standard '
             'errors clustered at the firm level. Subsample sizes are reported to the '
             'right of each interval.', 11.5),
}


# ----------------------------------------------------------------------------------
TCPR_ORDER = ['cnfStyle', 'tcW', 'gridSpan', 'hMerge', 'vMerge', 'tcBorders', 'shd',
              'noWrap', 'tcMar', 'textDirection', 'tcFitText', 'vAlign', 'hideMark']


def insert_ordered(parent, el, order):
    """Insert el into parent at the position required by the OOXML schema."""
    tag = el.tag.split('}')[1]
    idx = order.index(tag)
    for child in parent:
        ctag = child.tag.split('}')[1]
        if ctag in order and order.index(ctag) > idx:
            child.addprevious(el)
            return el
    parent.append(el)
    return el


def _set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = insert_ordered(tcPr, OxmlElement('w:tcBorders'), TCPR_ORDER)
    for edge in ('top', 'bottom', 'left', 'right'):
        if edge in kwargs:
            spec = kwargs[edge]
            el = borders.find(qn('w:' + edge))
            if el is None:
                el = OxmlElement('w:' + edge)
                borders.append(el)
            for k, v in spec.items():
                el.set(qn('w:' + k), str(v))


NONE = {'val': 'nil'}
THICK = {'val': 'single', 'sz': 8, 'space': 0, 'color': '000000'}
THIN = {'val': 'single', 'sz': 4, 'space': 0, 'color': '000000'}


def render_table(doc, key, styles):
    fname, num, caption, note, widths, fs = TABLES[key]
    df = pd.read_csv(os.path.join(TAB, fname)).fillna('')
    df.columns = ['' if str(c).startswith('Unnamed') else str(c) for c in df.columns]

    # ---- caption ABOVE the table -------------------------------------------------
    p = doc.add_paragraph(style=styles['tabcap'])
    r = p.add_run(f'Table {num}. ')
    r.bold = True
    p.add_run(caption)

    n_rows, n_cols = df.shape[0] + 1, df.shape[1]
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    def put(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ''
        par = cell.paragraphs[0]
        par.style = styles['tabbody']
        par.alignment = align
        par.paragraph_format.space_before = Pt(1)
        par.paragraph_format.space_after = Pt(1)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = par.add_run(str(text))
        run.font.size = Pt(fs)
        run.font.bold = bold
        run.font.name = 'Palatino Linotype'
        run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Palatino Linotype')

    for j, c in enumerate(df.columns):
        put(tbl.cell(0, j), c, bold=True)
    for i in range(df.shape[0]):
        for j in range(n_cols):
            v = df.iloc[i, j]
            left = (j == 0 and n_cols > 2)
            put(tbl.cell(i + 1, j), v,
                align=WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER)

    if widths:
        for j, w in enumerate(widths[:n_cols]):
            for i in range(n_rows):
                tbl.cell(i, j).width = Cm(w)

    # ---- three-line rules ---------------------------------------------------------
    for j in range(n_cols):
        _set_cell_border(tbl.cell(0, j), top=THICK, bottom=THIN, left=NONE, right=NONE)
        for i in range(1, n_rows - 1):
            _set_cell_border(tbl.cell(i, j), top=NONE, bottom=NONE, left=NONE, right=NONE)
        _set_cell_border(tbl.cell(n_rows - 1, j), top=NONE, bottom=THICK,
                         left=NONE, right=NONE)

    # repeat the header row when the table breaks across pages
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    hdr = OxmlElement('w:tblHeader')
    hdr.set(qn('w:val'), 'true')
    trPr.append(hdr)

    if note:
        p = doc.add_paragraph(style=styles['tabnote'])
        r = p.add_run('Note: ')
        r.italic = True
        p.add_run(note)
    return tbl


FIG_NUM = {}


def figure_available(key):
    return os.path.exists(os.path.join(FIG, FIGURES[key][0]))


def assign_figure_numbers(keys):
    """Number the figures that actually have an image, in order of appearance."""
    FIG_NUM.clear()
    n = 0
    for k in keys:
        if figure_available(k):
            n += 1
            FIG_NUM[k] = n
    return FIG_NUM


def render_figure(doc, key, styles):
    fname, caption, note, width = FIGURES[key]
    num = FIG_NUM[key]
    path = os.path.join(FIG, fname)

    p = doc.add_paragraph(style=styles['figure'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width))

    # ---- caption BELOW the image --------------------------------------------------
    cap = doc.add_paragraph(style=styles['figcap'])
    r = cap.add_run(f'Figure {num}. ')
    r.bold = True
    cap.add_run(caption)

    # ---- note BELOW the caption, no indent, single spacing ------------------------
    if note:
        n = doc.add_paragraph(style=styles['fignote'])
        n.paragraph_format.first_line_indent = Cm(0)
        n.paragraph_format.left_indent = cap.paragraph_format.left_indent
        n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = n.add_run('Note: ')
        r.italic = True
        n.add_run(note)
