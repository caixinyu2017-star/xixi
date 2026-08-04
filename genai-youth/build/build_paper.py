# -*- coding: utf-8 -*-
"""Build the manuscript as a Word .docx on the official MDPI Systems template.

Guarantees
  * continuous line numbers in the left margin (inherited from the template)
  * every equation is a native Word equation-editor (OMML) object, centred, with a
    right-aligned number in a borderless two-column table, exactly as the MDPI
    template specifies
  * three-line tables with captions ABOVE
  * figure captions BELOW the image, notes below the caption, unindented, single spaced
  * MDPI reference style, numbered by order of first appearance, cross-referenced
"""
from __future__ import annotations

import os
import re
import shutil
import zipfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import content as C
import l2omml
import omml
import refs as R
import tables as T

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DOT = ('/root/.claude/uploads/60820523-945d-56c9-958a-936121511764/'
                'de763d36-systemstemplate.dot')
OUT = os.path.join(HERE, '..', 'The_Vanishing_First_Rung.docx')

AUTHORS = 'Xinyu Cai 1 and Tiantian Mo 1,*'
AFFIL = ['1  College of Business, Jiaxing University, Jiaxing 314001, China; '
         '00008227@zjxu.edu.cn']
CORRESP = '*  Correspondence: 00008227@zjxu.edu.cn'


# ==================================================================================
# 1. open the .dot template as a .docx and strip its placeholder body
# ==================================================================================
def open_template() -> Document:
    tmp = os.path.join(HERE, '_template.docx')
    zin = zipfile.ZipFile(TEMPLATE_DOT)
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename == '[Content_Types].xml':
                data = data.replace(b'wordprocessingml.template.main+xml',
                                    b'wordprocessingml.document.main+xml')
            zout.writestr(it, data)
    zin.close()
    doc = Document(tmp)
    body = doc.element.body
    kids = list(body.iterchildren())
    # keep only the trailing sectPr (page size, margins, line numbering)
    for k in kids:
        if k.tag != qn('w:sectPr'):
            body.remove(k)
    # the shipped template marks the section right-to-left, which reverses table
    # column order in conforming readers; the manuscript is left-to-right
    for sect in body.findall(qn('w:sectPr')):
        for tag in ('w:bidi', 'w:rtlGutter'):
            for el in sect.findall(qn(tag)):
                sect.remove(el)
    return doc


STYLE_MAP = {
    'artType': 'MDPI_1.1_article_type', 'title': 'MDPI_1.2_title',
    'authors': 'MDPI_1.3_authornames', 'affil': 'MDPI_1.6_affiliation',
    'abstract': 'MDPI_1.7_abstract', 'keywords': 'MDPI_1.8_keywords',
    'line': 'MDPI_1.9_line', 'h1': 'MDPI_2.1_heading1', 'h2': 'MDPI_2.2_heading2',
    'h3': 'MDPI_2.3_heading3', 'text': 'MDPI_3.1_text',
    'noindent': 'MDPI_3.2_text_no_indent', 'eq': 'MDPI_3.9_equation',
    'eqnum': 'MDPI_3.a_equation_number', 'tabcap': 'MDPI_4.1_table_caption',
    'tabbody': 'MDPI_4.2_table_body', 'tabnote': 'MDPI_4.3_table_footer',
    'figcap': 'MDPI_5.1_figure_caption', 'figure': 'MDPI_5.2_figure',
    'fignote': 'MDPI_4.3_table_footer', 'back': 'MDPI_6.2_back_matter',
    'refs': 'MDPI_8.1_references', 'theorem': 'MDPI_8.2_theorem',
}

TBLPR_ORDER = ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
               'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
               'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'tblCaption',
               'tblDescription']

BODY_INDENT = 2608          # left indent used by every MDPI body style, in twips
TEXT_WIDTH = 7858           # usable body column width, in twips


# ==================================================================================
# 2. reference numbering by order of first appearance
# ==================================================================================
CITE_RE = re.compile(r'\[\[([a-z0-9\-]+)\]\]')
TOKEN_RE = re.compile(r'(\$[^$]*\$|\[\[[a-z0-9\-]+\]\]|\{\{F:[a-z0-9]+\}\})')

_order: list[str] = []


def _scan_citations():
    def scan(text):
        for k in CITE_RE.findall(text):
            if k not in _order:
                if k not in R.REFS:
                    raise KeyError(f'unknown reference key: {k}')
                _order.append(k)
    for b in C.BLOCKS:
        for part in b[1:]:
            if isinstance(part, str):
                scan(part)
    scan(C.ABSTRACT)
    return _order


def cite_number(key) -> int:
    return _order.index(key) + 1


def fmt_citation(keys):
    nums = sorted({cite_number(k) for k in keys})
    parts, i = [], 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        parts.append(f'{nums[i]}–{nums[j]}' if j - i >= 1 else str(nums[i]))
        i = j + 1
    return '[' + ','.join(parts) + ']'


# ==================================================================================
# 3. rich text: inline maths and merged consecutive citations
# ==================================================================================
def add_rich(par, text):
    """Add text, rendering $latex$ as inline OMML and merging adjacent [[key]]."""
    segs = [s for s in TOKEN_RE.split(text) if s]
    i = 0
    while i < len(segs):
        s = segs[i]
        if s.startswith('[['):
            group = [s[2:-2]]
            j = i + 1
            while j < len(segs) and segs[j].startswith('[['):
                group.append(segs[j][2:-2])
                j += 1
            par.add_run(fmt_citation(group))
            i = j
        elif s.startswith('{{F:'):
            par.add_run(str(T.FIG_NUM.get(s[4:-2], '?')))
            i += 1
        elif s.startswith('$') and s.endswith('$') and len(s) > 2:
            omml.add_inline_math(par, l2omml.L(s[1:-1]))
            i += 1
        else:
            par.add_run(s)
            i += 1


def para(doc, text, style_key='text'):
    p = doc.add_paragraph(style=STYLE_MAP[style_key])
    add_rich(p, text)
    return p


# ==================================================================================
# 4. numbered equation in the MDPI two-column borderless table
# ==================================================================================
def add_equation(doc, latex, number):
    """Centred equation with a right-aligned number, laid out with tab stops so
    that both Word and other OOXML readers place them identically."""
    p = doc.add_paragraph(style=STYLE_MAP['eq'])
    pf = p.paragraph_format
    pf.left_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ppr = p._p.get_or_add_pPr()
    for old in ppr.findall(qn('w:tabs')):
        ppr.remove(old)
    tabs = OxmlElement('w:tabs')
    centre = OxmlElement('w:tab')
    centre.set(qn('w:val'), 'center')
    centre.set(qn('w:pos'), str(BODY_INDENT + TEXT_WIDTH // 2))
    right = OxmlElement('w:tab')
    right.set(qn('w:val'), 'right')
    right.set(qn('w:pos'), str(BODY_INDENT + TEXT_WIDTH))
    tabs.append(centre)
    tabs.append(right)
    PPR_ORDER = ['pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
                 'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd',
                 'tabs', 'suppressAutoHyphens', 'kinsoku', 'wordWrap',
                 'overflowPunct', 'topLinePunct', 'autoSpaceDE', 'autoSpaceDN',
                 'bidi', 'adjustRightInd', 'snapToGrid', 'spacing', 'ind',
                 'contextualSpacing', 'mirrorIndents', 'suppressOverlap', 'jc',
                 'textDirection', 'textAlignment', 'textboxTightWrap',
                 'outlineLvl', 'divId', 'cnfStyle', 'rPr', 'sectPr', 'pPrChange']
    T.insert_ordered(ppr, tabs, PPR_ORDER)

    p.add_run().add_tab()
    p._p.append(omml.omath(l2omml.L(latex)))
    r = p.add_run()
    r.add_tab()
    r.add_text(f'({number})')
    return p


# ==================================================================================
# 5. front matter
# ==================================================================================
def front_matter(doc):
    p = doc.add_paragraph(style=STYLE_MAP['artType']); p.add_run('Article')
    p = doc.add_paragraph(style=STYLE_MAP['title']); p.add_run(C.TITLE)

    p = doc.add_paragraph(style=STYLE_MAP['authors'])
    for chunk, sup in (('Xinyu Cai ', False), ('1', True), (' and Tiantian Mo ', False),
                       ('1,*', True)):
        r = p.add_run(chunk)
        if sup:
            r.font.superscript = True

    for a in AFFIL:
        doc.add_paragraph(style=STYLE_MAP['affil']).add_run(a)
    doc.add_paragraph(style=STYLE_MAP['affil']).add_run(CORRESP)

    p = doc.add_paragraph(style=STYLE_MAP['abstract'])
    r = p.add_run('Abstract: '); r.bold = True
    add_rich(p, C.ABSTRACT)

    p = doc.add_paragraph(style=STYLE_MAP['keywords'])
    r = p.add_run('Keywords: '); r.bold = True
    p.add_run(C.KEYWORDS)

    doc.add_paragraph(style=STYLE_MAP['line'])


# ==================================================================================
# 6. body
# ==================================================================================
def body(doc):
    for blk in C.BLOCKS:
        kind = blk[0]
        if kind in ('h1', 'h2', 'h3'):
            doc.add_paragraph(style=STYLE_MAP[kind]).add_run(blk[1])
        elif kind == 'p':
            para(doc, blk[1])
        elif kind == 'hyp':
            p = doc.add_paragraph(style=STYLE_MAP['noindent'])
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(blk[1] + ' '); r.bold = True; r.italic = True
            r2 = p.add_run(blk[2]); r2.italic = True
        elif kind == 'eq':
            add_equation(doc, blk[1], blk[2])
        elif kind == 'table':
            T.render_table(doc, blk[1], STYLE_MAP)
        elif kind == 'fig':
            if T.figure_available(blk[1]):
                T.render_figure(doc, blk[1], STYLE_MAP)
        else:
            raise ValueError(f'unknown block kind {kind}')


# ==================================================================================
# 7. back matter and references
# ==================================================================================
def back_matter(doc):
    for label, text in C.BACKMATTER:
        p = doc.add_paragraph(style=STYLE_MAP['back'])
        r = p.add_run(label + ': '); r.bold = True
        p.add_run(text)


def _ref_text(r):
    """Render one reference in MDPI style; returns a list of (text, italic) runs."""
    out = [(r['authors'] + ' ', False)]
    title = r['title'].rstrip('.')
    if r['kind'] == 'article':
        out.append((title + '. ', False))
        out.append((r['journal'] + ' ', True))
        out.append((str(r['year']), False))
        if r.get('volume'):
            out.append((', ', False))
            out.append((str(r['volume']), True))
            if r.get('pages'):
                out.append((', ' + r['pages'], False))
        out.append(('.', False))
    elif r['kind'] == 'book':
        out.append((title, True))
        if r.get('edition'):
            out.append((', ' + r['edition'], False))
        out.append(('; ' + r['publisher'] + ': ' + r['place'] + ', '
                    + str(r['year']) + '.', False))
    else:  # report / working paper
        out.append((title, True))
        if r.get('series'):
            out.append(('; ' + r['series'], False))
        out.append(('; ' + r['publisher'] + ': ' + r['place'] + ', '
                    + str(r['year']) + '.', False))
    if r.get('doi'):
        out.append((' [CrossRef]', False))
    return out


def references(doc):
    doc.add_paragraph(style=STYLE_MAP['h1']).add_run('References')
    for i, key in enumerate(_order, 1):
        r = R.REFS[key]
        p = doc.add_paragraph(style=STYLE_MAP['refs'])
        pf = p.paragraph_format
        pf.left_indent = Cm(BODY_INDENT / 566.93 + 0.60)
        pf.first_line_indent = Cm(-0.60)
        pf.space_after = Pt(1)
        # the template's reference style carries automatic list numbering; switch it
        # off at paragraph level (numId 0) so that only our own numbers appear
        ppr = p._p.get_or_add_pPr()
        for old in ppr.findall(qn('w:numPr')):
            ppr.remove(old)
        numpr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
        numid = OxmlElement('w:numId'); numid.set(qn('w:val'), '0')
        numpr.append(ilvl); numpr.append(numid)
        PPR_REF_ORDER = ['pStyle', 'keepNext', 'keepLines', 'pageBreakBefore',
                         'framePr', 'widowControl', 'numPr', 'suppressLineNumbers',
                         'pBdr', 'shd', 'tabs', 'suppressAutoHyphens', 'kinsoku',
                         'wordWrap', 'overflowPunct', 'topLinePunct', 'autoSpaceDE',
                         'autoSpaceDN', 'bidi', 'adjustRightInd', 'snapToGrid',
                         'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
                         'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
                         'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle',
                         'rPr', 'sectPr', 'pPrChange']
        T.insert_ordered(ppr, numpr, PPR_REF_ORDER)
        p.add_run(f'{i}. ')
        for text, ital in _ref_text(r):
            run = p.add_run(text)
            run.italic = ital
    return len(_order)


# ==================================================================================
def main():
    _scan_citations()
    fig_keys = [b[1] for b in C.BLOCKS if b[0] == 'fig']
    T.assign_figure_numbers(fig_keys)
    missing = [k for k in fig_keys if k not in T.FIG_NUM]
    if missing:
        print(f'figure image(s) unavailable, omitted and renumbered: {missing}')
    doc = open_template()
    front_matter(doc)
    body(doc)
    back_matter(doc)
    n = references(doc)
    doc.save(OUT)
    unused = sorted(set(R.REFS) - set(_order))
    print(f'saved {OUT}')
    print(f'{n} references cited, numbered by order of appearance')
    if unused:
        print(f'{len(unused)} reference(s) defined but not cited: {unused}')


if __name__ == '__main__':
    main()
