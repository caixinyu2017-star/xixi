# -*- coding: utf-8 -*-
"""Build the one-page A4 cover letter, matching the style of the author's
previous submission letter to Systems."""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, '..', 'Cover_Letter.docx')

FONT = 'Times New Roman'
SIZE = 9.0

BODY = [
    "Dear Editors:",

    "We would like to submit the enclosed manuscript entitled “The Vanishing First "
    "Rung: Generative Artificial Intelligence, Entry-Task Exposure and Youth Employment "
    "in Chinese Listed Firms” for consideration in Systems, for the Special Issue "
    "“Systems Approaches to Generative AI: Workforce Development, Organisational "
    "Learning, and Economic Transformation”. The manuscript has not been published "
    "elsewhere and is not under consideration by any other journal. No conflict of "
    "interest exists in its submission, and it has been approved for publication by all "
    "authors.",

    "Generative artificial intelligence is unusual among labour-saving technologies in "
    "that the work it performs best is the work through which people become qualified "
    "for the work it cannot perform. Drafting, reconciliation, summarising and "
    "first-pass analysis are simultaneously a firm’s production input and its training "
    "input. A firm that automates them captures the saving immediately and pays for the "
    "lost learning years later, in a currency — experienced staff — that cannot be "
    "bought back quickly. This manuscript asks whether that coincidence has already "
    "begun to close the entry rung of the firm-internal job ladder, and answers it with "
    "a design built around the late-2022 arrival of general-purpose language-model "
    "assistants: a shock that reached every firm at the same moment but bit in "
    "proportion to how much of its entry-level work lay inside the newly automatable "
    "frontier.",

    "The highlights of the paper are as follows.",

    "(1) A displacement that lands on one age margin and not the others. Using 31,015 "
    "firm-year observations for 3,240 Chinese A-share listed firms over 2014–2024, a "
    "one-standard-deviation increase in pre-shock entry-task exposure lowers the youth "
    "employment share by 1.50 percentage points, 5.42 per cent of its mean. The five "
    "pre-shock leads are individually and jointly insignificant (χ² = 4.73, p = 0.450), "
    "the effect appears in the shock year and persists, and it survives instrumental "
    "variables, propensity-score matching, entropy balancing and eight further "
    "specifications. The placebo is decisive: run on employees aged 41 to 50, whose "
    "tasks lie outside the exposed frontier, the coefficient is −0.0011 and "
    "insignificant, two orders of magnitude below the baseline.",

    "(2) The channel that works, and the channel that is assumed to work and does not. "
    "About 31 per cent of the effect travels through the observable contraction of "
    "entry-level routine cognitive positions (indirect effect −0.0368, bootstrap "
    "confidence interval [−0.0441, −0.0301]). Exposed firms do raise training "
    "expenditure, significantly so, and the offsetting indirect effect is nonetheless "
    "0.0026 with a confidence interval containing zero. The compensating mechanism that "
    "most policy discussion takes for granted is present in intention and absent in "
    "effect.",

    "(3) What protects the rung is capability, not incentive. Organisational learning "
    "capacity significantly attenuates the loss; internal labour-market depth, which "
    "gives a firm every reason to want its entry rung protected, does not. Firms that "
    "already know how to build competence deliberately can afford to lose the rung that "
    "used to build it for them; firms that merely need the competence cannot. The "
    "reduction in youth employment is followed by a fall in internal advancement, about "
    "half of which runs through the contracted entry rung itself.",

    "We believe the manuscript is well suited to this Special Issue and to Systems. Its "
    "contribution is not the decoration of an economic result with systems vocabulary "
    "but the identification of a specific, estimable instance of the classical "
    "socio-technical proposition that partial optimisation degrades the whole: "
    "generative artificial intelligence raises the productivity of the technical "
    "subsystem while quietly degrading the reproduction function of the social "
    "subsystem, because the tasks it absorbs belong to both. The practical payoff is "
    "that the diagnostic quantity — the share of entry-grade positions that still "
    "contain graded, supervised, consequential work — is observable now, from "
    "disclosures firms already file, whereas the capability consequence it foreshadows "
    "becomes observable only once it is expensive to reverse. The urgency is not "
    "hypothetical: youth unemployment remains far above the adult rate worldwide, and "
    "the earnings penalty for a poor labour-market entry is measured in decades. We "
    "deeply appreciate your consideration of our manuscript and look forward to "
    "receiving comments from the reviewers. Correspondence should be directed to the "
    "address below.",

    "Thanks very much for your attention to our paper.",
]

SIGN = ["Name: Tiantian Mo",
        "Institution and address: College of Business, Jiaxing University; "
        "Jiaxing 314001, China",
        "Email: 00008227@zjxu.edu.cn",
        "",
        "Very sincerely yours,",
        "Xinyu Cai and Tiantian Mo"]


def styled(doc, text, justify=True, space_after=3.0, bold=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(SIZE)
    r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), FONT)
    return p


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)
    sec.top_margin = Cm(1.1)
    sec.bottom_margin = Cm(1.1)

    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(SIZE)

    for i, para in enumerate(BODY):
        styled(doc, para, justify=(i > 0), space_after=3.0)
    for line in SIGN:
        styled(doc, line, justify=False, space_after=0.5)

    doc.save(OUT)
    print('saved', OUT)


if __name__ == '__main__':
    main()
