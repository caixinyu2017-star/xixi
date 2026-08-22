# -*- coding: utf-8 -*-
"""《经济研究》投稿体例的 Word 生成器。

要点：
- 版心 170mm×230mm（A4，左右边距 20mm、上下 33.5mm）
- 正文宋体小四、1.5 倍行距、首行缩进 2 字符；西文 Times New Roman
- 标题 一、／（一）／1.／(1)，末尾不加标点，标题行空两格
- 公式为 Word 原生 OMML 对象，居中，编号右对齐
- 表：三线表，表题在表上方；注在表下方
- 图：图题在图下方，其下依次为图例、数据来源、注（无缩进、单倍行距）
- 脚注：真正的 Word 页底脚注，圈码①②……每页重新编号
- 参考文献：按正文首次出现顺序编号 [1][2]，正文处上标标注
"""
import copy
import os
import re
import subprocess
import zipfile

import docx
from lxml import etree
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips

SONG, HEI, KAI, TNR = '宋体', '黑体', '楷体', 'Times New Roman'
SIZE = {'小三': 15, '四号': 14, '小四': 12, '五号': 10.5, '小五': 9, '六号': 7.5}
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# 粗体用 {b:...} 而非 **...**：注文中的显著性星号 ***、**、* 会与 ** 标记冲突
INLINE_RE = re.compile(r'(\$[^$]+\$|\{b:[^}]*\}|\{fn:[^}]*\}|\{c:[^}]*\})')


# --------------------------------------------------------------- LaTeX → OMML
def latex_to_omml(latex_list, workdir):
    """用 pandoc 批量把 LaTeX 公式转成 OMML 元素。"""
    os.makedirs(workdir, exist_ok=True)
    if not latex_list:
        return []
    md = os.path.join(workdir, '_eq.md')
    dx = os.path.join(workdir, '_eq.docx')
    SEP = 'ZQEQSEPZQ'
    with open(md, 'w', encoding='utf-8') as f:
        for i, s in enumerate(latex_list):
            f.write(f'{SEP}{i}\n\n$${s}$$\n\n')
    subprocess.run(['pandoc', md, '-o', dx], check=True,
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    xml = zipfile.ZipFile(dx).read('word/document.xml').decode('utf-8')
    root = etree.fromstring(xml.encode('utf-8'))
    out, cur = [], None
    for p in root.iter('{%s}p' % W_NS):
        txt = ''.join(t.text or '' for t in p.iter('{%s}t' % W_NS))
        m = re.match(r'^%s(\d+)$' % SEP, txt.strip())
        if m:
            cur = int(m.group(1))
            continue
        om = p.find('.//{%s}oMath' % M_NS)
        if om is not None and cur is not None:
            while len(out) <= cur:
                out.append(None)
            out[cur] = copy.deepcopy(om)
            cur = None
    for om in out:
        if om is None:
            continue
        # pandoc 会生成空的 m:sepChr（val=""），部分渲染器据此把定界符误绘为竖线，删去
        for dpr in om.iter('{%s}dPr' % M_NS):
            sep = dpr.find('{%s}sepChr' % M_NS)
            if sep is not None and not (sep.get('{%s}val' % M_NS) or ''):
                dpr.remove(sep)
        _flatten_simple_delims(om)
    return out


# 含高元素时须保留可伸缩定界符，否则括号包不住内容
_TALL = {'f', 'rad', 'nary', 'm', 'd', 'limLow', 'limUpp', 'box'}


def _flatten_simple_delims(om):
    """把简单圆括号的 m:d 换成字面量括号 run。

    m:d 的 begChr／endChr 在 Word 中正常，但部分渲染器（如 LibreOffice）会忽略这两个属性
    而一律画成竖线。内容不含分式、根式等高元素时，改用字面括号可在各渲染器下一致。
    """
    for d in list(om.iter('{%s}d' % M_NS)):
        dpr = d.find('{%s}dPr' % M_NS)
        if dpr is None:
            continue
        beg = dpr.find('{%s}begChr' % M_NS)
        end = dpr.find('{%s}endChr' % M_NS)
        b = beg.get('{%s}val' % M_NS) if beg is not None else '('
        e = end.get('{%s}val' % M_NS) if end is not None else ')'
        if (b, e) not in (('(', ')'), ('[', ']')):
            continue
        es = d.findall('{%s}e' % M_NS)
        if len(es) != 1:
            continue
        if any(ch.tag.split('}')[-1] in _TALL for ch in es[0].iter()):
            continue
        parent = d.getparent()
        if parent is None:
            continue
        idx = list(parent).index(d)

        def lit(ch):
            r = etree.SubElement(etree.Element('x'), '{%s}r' % M_NS)
            rpr = etree.SubElement(r, '{%s}rPr' % M_NS)
            sty = etree.SubElement(rpr, '{%s}sty' % M_NS)
            sty.set('{%s}val' % M_NS, 'p')
            t = etree.SubElement(r, '{%s}t' % M_NS)
            t.text = ch
            return r

        nodes = [lit(b)] + list(es[0]) + [lit(e)]
        for off, nd in enumerate(nodes):
            parent.insert(idx + off, nd)
        parent.remove(d)


# --------------------------------------------------------------- 脚注部件
FOOTNOTES_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<w:footnotes xmlns:w="%s">'
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" '
    'w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing '
    'w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/>'
    '</w:r></w:p></w:footnote>%%s</w:footnotes>' % W_NS
)


class Paper(object):
    def __init__(self, workdir, margins_mm=(33.5, 33.5, 20.0, 20.0)):
        self.doc = Document()
        self.workdir = workdir
        os.makedirs(workdir, exist_ok=True)
        self.eq_cache = {}
        self.footnotes = []          # [(id, text)]
        self._fn_next = 1
        self.refs_order = []         # 按首次出现排序的 refkey
        self.refmap = {}
        self._setup_page(margins_mm)
        self._setup_styles()

    # ---------------- 页面与样式
    def _setup_page(self, m):
        s = self.doc.sections[0]
        s.page_height, s.page_width = Cm(29.7), Cm(21.0)
        top, bottom, left, right = m
        s.top_margin, s.bottom_margin = Cm(top / 10), Cm(bottom / 10)
        s.left_margin, s.right_margin = Cm(left / 10), Cm(right / 10)
        self.text_w_cm = 21.0 - left / 10 - right / 10        # 版心宽 17.0 cm
        # 脚注编号：圈码①②……，每页重新编号
        sectPr = s._sectPr
        fpr = OxmlElement('w:footnotePr')
        for tag, val in (('w:numFmt', 'decimalEnclosedCircleChinese'),
                         ('w:numRestart', 'eachPage')):
            el = OxmlElement(tag)
            el.set(qn('w:val'), val)
            fpr.append(el)
        sectPr.insert(0, fpr)

    def _setup_styles(self):
        st = self.doc.styles['Normal']
        st.font.name = TNR
        st.font.size = Pt(SIZE['小四'])
        rpr = st.element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn('w:eastAsia'), SONG)
        rf.set(qn('w:ascii'), TNR)
        rf.set(qn('w:hAnsi'), TNR)
        # 引号“”、破折号、省略号属东亚宽度歧义类，须由 hint 判给中文字体
        rf.set(qn('w:hint'), 'eastAsia')

    @staticmethod
    def _font(run, cn=SONG, size=SIZE['小四'], bold=False, italic=False, western=None,
              sup=False, sub=False):
        run.font.size = Pt(size)
        run.font.bold = bold or None
        run.font.italic = italic or None
        if sup:
            run.font.superscript = True
        if sub:
            run.font.subscript = True
        west = western or TNR
        run.font.name = west
        rpr = run._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn('w:eastAsia'), cn)
        rf.set(qn('w:ascii'), west)
        rf.set(qn('w:hAnsi'), west)
        rf.set(qn('w:hint'), 'eastAsia')

    # ---------------- 脚注
    def _add_footnote(self, para, text):
        fid = self._fn_next
        self._fn_next += 1
        self.footnotes.append((fid, text, None))
        r = para.add_run()
        rpr = r._element.get_or_add_rPr()
        st = OxmlElement('w:vertAlign')
        st.set(qn('w:val'), 'superscript')
        rpr.append(st)
        ref = OxmlElement('w:footnoteReference')
        ref.set(qn('w:id'), str(fid))
        r._element.append(ref)

    def footnote_star(self, para, text):
        """给指定段落挂一个以 * 为标记的脚注（首页作者简介用）。"""
        fid = self._fn_next
        self._fn_next += 1
        self.footnotes.append((fid, text, '*'))
        r = para.add_run()
        rpr = r._element.get_or_add_rPr()
        va = OxmlElement('w:vertAlign')
        va.set(qn('w:val'), 'superscript')
        rpr.append(va)
        ref = OxmlElement('w:footnoteReference')
        ref.set(qn('w:customMarkFollows'), '1')
        ref.set(qn('w:id'), str(fid))
        r._element.append(ref)
        t = OxmlElement('w:t')
        t.text = '*'
        r._element.append(t)

    def _write_footnotes_part(self):
        if not self.footnotes:
            return
        body = []
        for fid, text, mark in self.footnotes:
            runs = ''.join(
                '<w:r><w:rPr><w:rFonts w:ascii="%s" w:hAnsi="%s" w:eastAsia="%s" '
                'w:hint="eastAsia"/><w:sz w:val="%d"/></w:rPr><w:t xml:space="preserve">%s'
                '</w:t></w:r>' % (TNR, TNR, SONG, int(SIZE['小五'] * 2), _esc(seg))
                for seg in [text])
            # 自定义标记的星号已由正文引用处的 w:t 提供，脚注体内不再重复，
            # 否则渲染出两个星号。
            head = '' if mark else (
                '<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                '<w:footnoteRef/></w:r>')
            body.append(
                '<w:footnote w:id="%d"><w:p><w:pPr><w:spacing w:after="0" w:line="240" '
                'w:lineRule="auto"/><w:ind w:firstLineChars="0" w:firstLine="0"/></w:pPr>'
                '%s%s</w:p></w:footnote>' % (fid, head, runs))
        xml = FOOTNOTES_XML % ''.join(body)

        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        ct = ('application/vnd.openxmlformats-officedocument.wordprocessingml.'
              'footnotes+xml')
        rt = ('http://schemas.openxmlformats.org/officeDocument/2006/relationships/'
              'footnotes')
        part = Part(PackURI('/word/footnotes.xml'), ct, xml.encode('utf-8'),
                    self.doc.part.package)
        self.doc.part.relate_to(part, rt)

    # ---------------- 引用
    def _cite(self, para, keys):
        nums = []
        for k in [x.strip() for x in keys.split(',') if x.strip()]:
            if k not in self.refmap:
                self.refs_order.append(k)
                self.refmap[k] = len(self.refs_order)
            nums.append(self.refmap[k])
        nums = sorted(set(nums))
        # 连续编号缩写为 a-b
        parts, i = [], 0
        while i < len(nums):
            j = i
            while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
                j += 1
            parts.append(str(nums[i]) if j == i else '%d-%d' % (nums[i], nums[j]))
            i = j + 1
        r = para.add_run('[%s]' % ','.join(parts))
        self._font(r, SONG, SIZE['小四'], sup=True)

    # ---------------- 富文本
    def _rich(self, para, text, cn=SONG, size=SIZE['小四'], bold=False):
        pos = 0
        for m in INLINE_RE.finditer(text):
            if m.start() > pos:
                r = para.add_run(_minus(text[pos:m.start()]))
                self._font(r, cn, size, bold)
            seg = m.group(0)
            if seg.startswith('$'):
                para._p.append(self._omml(seg[1:-1]))
            elif seg.startswith('{b:'):
                r = para.add_run(_minus(seg[3:-1]))
                self._font(r, cn, size, True)
            elif seg.startswith('{fn:'):
                self._add_footnote(para, seg[4:-1])
            elif seg.startswith('{c:'):
                self._cite(para, seg[3:-1])
            pos = m.end()
        if pos < len(text):
            r = para.add_run(_minus(text[pos:]))
            self._font(r, cn, size, bold)

    def _omml(self, latex):
        om = self.eq_cache.get(latex)
        if om is None:
            om = latex_to_omml([latex], self.workdir)[0]
            self.eq_cache[latex] = om
        return copy.deepcopy(om)

    def prepare_math(self, blocks):
        need = []
        for b in blocks:
            for v in b.values():
                if isinstance(v, str):
                    need += [m[1:-1] for m in re.findall(r'\$[^$]+\$', v)]
            if 'eq' in b:
                need.append(b['eq'])
        need = [x for x in dict.fromkeys(need) if x not in self.eq_cache]
        if need:
            for s, om in zip(need, latex_to_omml(need, self.workdir)):
                self.eq_cache[s] = om

    # ---------------- 段落构件
    def para(self, text, indent=True, align='justify', cn=SONG, size=SIZE['小四'],
             bold=False, spacing=1.5, space_after=0, first_line_chars=200):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        if spacing == 1.5:
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(space_after)
        p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                       'center': WD_ALIGN_PARAGRAPH.CENTER,
                       'left': WD_ALIGN_PARAGRAPH.LEFT,
                       'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
        if indent:
            ind = p._p.get_or_add_pPr().get_or_add_ind()
            ind.set(qn('w:firstLineChars'), str(first_line_chars))
            ind.set(qn('w:firstLine'), str(int(size * first_line_chars / 100 * 20)))
        self._rich(p, text, cn, size, bold)
        return p

    def heading(self, text, level):
        cfg = {1: (HEI, SIZE['小三'], 12, 8),
               2: (HEI, SIZE['四号'], 8, 6),
               3: (HEI, SIZE['小四'], 6, 4),
               4: (KAI, SIZE['小四'], 4, 3)}[level]
        cn, size, sb, sa = cfg
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(sb), Pt(sa)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.keep_with_next = True
        ind = p._p.get_or_add_pPr().get_or_add_ind()      # 标题行空两格
        ind.set(qn('w:firstLineChars'), '200')
        ind.set(qn('w:firstLine'), str(int(size * 2 * 20)))
        opr = OxmlElement('w:outlineLvl')
        opr.set(qn('w:val'), str(level - 1))
        p._p.get_or_add_pPr().append(opr)
        self._rich(p, text, cn, size, bold=(level == 1))
        return p

    def equation(self, latex, num):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before, pf.space_after = Pt(6), Pt(6)
        w = self.text_w_cm
        pf.tab_stops.add_tab_stop(Cm(w / 2.0), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(w), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run().add_tab()
        p._p.append(self._omml(latex))
        p.add_run().add_tab()
        r = p.add_run('(%s)' % num)
        self._font(r, SONG, SIZE['小四'])
        return p

    # ---------------- 图
    def figure(self, path, caption, legend=None, source=None, note=None, width_cm=None):
        from PIL import Image
        w = width_cm or min(14.0, self.text_w_cm)
        with Image.open(path) as im:
            ratio = im.height / im.width
        if w * ratio > 19.5:                       # 不超版心高度
            w = 19.5 / ratio
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(path, width=Cm(w))
        # 图题在图下方
        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self._rich(cp, caption, HEI, SIZE['五号'])
        # 图例、数据来源、注：均在图题下方，无缩进、单倍行距
        for txt in (legend, source, note):
            if not txt:
                continue
            np_ = self.doc.add_paragraph()
            np_.alignment = WD_ALIGN_PARAGRAPH.LEFT
            np_.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            np_.paragraph_format.space_after = Pt(2)
            self._rich(np_, txt, SONG, SIZE['小五'])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 表（三线表，表题在上）
    @staticmethod
    def _vlen(s):
        return sum(1.0 if ord(c) > 0x2E7F else 0.55 for c in str(s))

    def _fit_size(self, header, rows, total_cm, size):
        """列多时自动缩字号，避免把 −0.0292*** 这类单元格挤到折行。"""
        n = len(header)
        while size > 7.0:
            char_cm = size / 72.0 * 2.54
            need = []
            for j in range(n):
                body = [self._vlen(r[j]) if j < len(r) else 0 for r in rows]
                need.append(max([self._vlen(header[j]) * 1.15] + body + [1.0]))
            # 数值列须能整体容纳最长单元格，不得折行
            want = sum(min(x, 6.0) * char_cm + 0.24 for x in need)
            if want <= total_cm:
                return size
            size -= 0.5
        return size

    def table(self, caption, header, rows, note=None, size=SIZE['五号'], width_cm=None):
        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.keep_with_next = True
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self._rich(cp, caption, HEI, SIZE['五号'])

        ncol = len(header)
        total = width_cm or self.text_w_cm
        size = self._fit_size(header, rows, total, size)
        t = self.doc.add_table(rows=1 + len(rows), cols=ncol)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        self._tbl_width(t, total)
        for j, h in enumerate(header):
            self._cell(t.cell(0, j), str(h), size, bold=True)
        for i, row in enumerate(rows):
            for j, v in enumerate(row):
                self._cell(t.cell(i + 1, j), str(v), size)
        self._col_widths(t, header, rows, total, size)
        self._three_line(t)
        if note:
            np_ = self.doc.add_paragraph()
            np_.alignment = WD_ALIGN_PARAGRAPH.LEFT
            np_.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            np_.paragraph_format.space_after = Pt(8)
            self._rich(np_, note, SONG, SIZE['小五'])
        else:
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return t

    def _cell(self, cell, text, size, bold=False):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = 1.0
        pf.space_before = pf.space_after = Pt(0)
        pPr = pf.element.get_or_add_pPr()
        for tag, val in (('w:snapToGrid', '0'), ('w:wordWrap', '0')):
            el = OxmlElement(tag)
            el.set(qn('w:val'), val)
            pPr.append(el)
        self._rich(p, text, SONG, size)
        tcPr = cell._tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), 'center')
        tcPr.append(va)

    def _tbl_width(self, t, cm):
        tblPr = t._tbl.tblPr
        old = tblPr.find(qn('w:tblW'))
        if old is not None:
            tblPr.remove(old)
        el = OxmlElement('w:tblW')
        el.set(qn('w:type'), 'dxa')
        el.set(qn('w:w'), str(int(cm * 567)))
        tblPr.append(el)

    def _col_widths(self, t, header, rows, total_cm, size):
        import math
        n = len(header)
        char_cm = size / 72.0 * 2.54
        pad = 0.26
        need = []
        for j in range(n):
            body = [self._vlen(r[j]) if j < len(r) else 0 for r in rows]
            need.append(max([self._vlen(header[j]) * 1.15] + body + [1.0]))
        floors = [min(x, 6.0) * char_cm + pad for x in need]
        if sum(floors) >= total_cm:
            k = total_cm / sum(floors)
            widths = [f * k for f in floors]
        else:
            rest = total_cm - sum(floors)
            w = [x ** 0.72 for x in need]
            widths = [f + rest * wi / sum(w) for f, wi in zip(floors, w)]
        tblPr = t._tbl.tblPr
        for tag in ('w:tblCellMar', 'w:tblLayout'):
            old = tblPr.find(qn(tag))
            if old is not None:
                tblPr.remove(old)
        mar = OxmlElement('w:tblCellMar')
        for side, v in (('top', 0), ('left', 56), ('bottom', 0), ('right', 56)):
            e = OxmlElement('w:' + side)
            e.set(qn('w:w'), str(v))
            e.set(qn('w:type'), 'dxa')
            mar.append(e)
        tblPr.append(mar)
        lay = OxmlElement('w:tblLayout')
        lay.set(qn('w:type'), 'fixed')
        tblPr.append(lay)
        for j, wcm in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(wcm)
        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            for j, col in enumerate(grid.findall(qn('w:gridCol'))):
                if j < len(widths):
                    col.set(qn('w:w'), str(int(round(widths[j] * 567))))

    def _three_line(self, t):
        tblPr = t._tbl.tblPr
        b = OxmlElement('w:tblBorders')
        for name, val, sz in (('top', 'single', 12), ('left', 'none', None),
                              ('bottom', 'single', 12), ('right', 'none', None),
                              ('insideH', 'none', None), ('insideV', 'none', None)):
            e = OxmlElement('w:' + name)
            e.set(qn('w:val'), val)
            if sz:
                e.set(qn('w:sz'), str(sz))
                e.set(qn('w:color'), '000000')
            b.append(e)
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        after = [qn('w:' + n) for n in ('shd', 'tblLayout', 'tblCellMar', 'tblLook')]
        anchor = next((c for c in tblPr if c.tag in after), None)
        if anchor is not None:
            anchor.addprevious(b)
        else:
            tblPr.append(b)
        # 表头下的第二条线
        for cell in t.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tb = OxmlElement('w:tcBorders')
            e = OxmlElement('w:bottom')
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), '6')
            e.set(qn('w:color'), '000000')
            tb.append(e)
            tcPr.append(tb)

    # ---------------- 保存
    def save(self, path):
        self._write_footnotes_part()
        self.doc.save(path)


def _esc(s):
    return (s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))


_MINUS = re.compile(r'(?<![0-9A-Za-z\-])-(?=\d)')


def _minus(t):
    """正文负数用真正的减号 U+2212；页码范围、日期、带连字符专名不动。"""
    return _MINUS.sub('−', t)
