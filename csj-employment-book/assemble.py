# -*- coding: utf-8 -*-
"""组装完整专著 docx。按章节顺序登记引用，末尾按顺序编码制生成参考文献。"""
import importlib
import os

from docxbuilder import BookBuilder, count_chars, enforce_fig_after_mention, SIZE, SONG, HEI
from docx.enum.text import WD_ALIGN_PARAGRAPH
import references as refs

HERE = os.path.dirname(os.path.abspath(__file__))

CHAPTERS = [f'content_ch{i:02d}' for i in range(1, 15)]

TITLE = '长三角高校大学生就业和职业发展困境、成因及解决对策研究'
TITLE2 = ''
OUTNAME = '长三角高校大学生就业和职业发展困境、成因及解决对策研究.docx'


def title_page(bk):
    doc = bk.doc
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bk._append_rich(p, TITLE, HEI, SIZE['小二'], bold=True)
    if TITLE2:
        p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bk._append_rich(p1, TITLE2, HEI, SIZE['小二'], bold=True)
    for _ in range(6):
        doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bk._append_rich(p3, '蔡鑫宇  著', SONG, SIZE['四号'])
    bk.add_pagebreak()


def toc_page(bk):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    bk.add_heading('目  录', 1)
    p = bk.doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '（在 Word 中右键此处选择“更新域”生成目录）'
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    bk.add_pagebreak()


def load_blocks():
    all_blocks = []
    import content_front
    all_blocks += content_front.abstract_blocks()
    all_blocks.append({'pagebreak': True})
    all_blocks += content_front.preface_blocks()
    all_blocks.append({'pagebreak': True})
    counts = {}
    refs.reset()
    for mod in CHAPTERS:
        m = importlib.import_module(mod)
        blks = enforce_fig_after_mention(m.blocks())
        counts[mod] = count_chars(blks)
        all_blocks += blks
        all_blocks.append({'pagebreak': True})
    return all_blocks, counts


def main():
    bk = BookBuilder(os.path.join(HERE, 'build'))
    all_blocks, counts = load_blocks()
    total_chars = count_chars(all_blocks)
    title_page(bk)
    toc_page(bk)
    bk.build(all_blocks)
    bk.add_references(refs.gb_ordered())
    import content_front
    post = content_front.postscript_blocks()
    bk.add_pagebreak()
    bk.build(post)
    out = os.path.join(HERE, 'build', OUTNAME)
    bk.save(out)
    print('=== 字数统计 ===')
    for m in CHAPTERS:
        print(f'  {m}: {counts[m]:>7} 字')
    print(f'  正文合计（含前言摘要后记）: {total_chars + count_chars(post)} 字')
    print(f'  参考文献: {refs.n_used()} 条（已引用） / {len(refs.REFS)} 条（库）')
    print('  saved:', out)


if __name__ == '__main__':
    main()
