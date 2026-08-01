# -*- coding: utf-8 -*-
"""构建中文学术/方案类 Word 文档的排版辅助函数。

排版规范（由任务要求给定）：
    一级（章）标题：黑体、小三号（15 pt）、加粗
    二级（节）标题：黑体、四号（14 pt）
    三级（条）标题：黑体、小四号（12 pt）
    正文：宋体、小四号（12 pt），段前段后间距均为 0
    图题：置于图下方，五号（10.5 pt）黑体
    表题：置于表上方，五号（10.5 pt）黑体；表格采用三线表
"""
import re

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

# ---- 字号常量（磅） ---------------------------------------------------
XIAO_ER = 18.0     # 小二
SAN = 16.0         # 三号
XIAO_SAN = 15.0    # 小三
SI = 14.0          # 四号
XIAO_SI = 12.0     # 小四
WU = 10.5          # 五号
XIAO_WU = 9.0      # 小五

HEI = "黑体"
SONG = "宋体"
KAI = "楷体"
TIMES = "Times New Roman"


def style_run(run, cn_font=SONG, size=XIAO_SI, bold=False, en_font=None,
              color=None, italic=False):
    """设置 run 的中英文字体、字号与字重。"""
    en_font = en_font or (TIMES if cn_font == SONG else cn_font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:cs"), en_font)
    return run


# CT_PPrBase 中位于 w:ind 之后的元素，用于保证 XML 子元素顺序合法
_AFTER_IND = ("w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap",
              "w:jc", "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
              "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
              "w:pPrChange")
_AFTER_TABS = ("w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
               "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
               "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
               "w:spacing") + _AFTER_IND + ("w:ind",)
_AFTER_OUTLINE = ("w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange")


def set_indent_chars(paragraph, first_line_chars=0, left_chars=0, size=XIAO_SI):
    """按“字符”设置缩进（Word 中文排版惯例），同时写入磅值兜底。"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.insert_element_before(ind, *_AFTER_IND)
    if first_line_chars:
        ind.set(qn("w:firstLineChars"), str(int(first_line_chars * 100)))
        ind.set(qn("w:firstLine"), str(int(first_line_chars * size * 20)))
    if left_chars:
        ind.set(qn("w:leftChars"), str(int(left_chars * 100)))
        ind.set(qn("w:left"), str(int(left_chars * size * 20)))


_RICH = re.compile(r"(@[^@]+@)")


def emit_rich(paragraph, text, cn_font=SONG, size=XIAO_SI, bold=False):
    """把 @X_y@ / @X^y@ / @X@ 标记渲染为斜体变量与真上下标。

    @!X@ 表示基底保持正体（用于 CCI、R 等多字母或集合符号）。
    """
    for part in _RICH.split(text):
        if not part:
            continue
        if part.startswith("@") and part.endswith("@") and len(part) > 2:
            inner = part[1:-1]
            upright = inner.startswith("!")
            if upright:
                inner = inner[1:]
            m = re.match(r"^([^_^]+)(?:([_^])(.+))?$", inner)
            base, op, tail = m.group(1), m.group(2), m.group(3)
            r = paragraph.add_run(base)
            style_run(r, cn_font=SONG, size=size, bold=bold, en_font=TIMES,
                      italic=(not upright) and bool(re.fullmatch(r"[A-Za-z]", base)))
            if op:
                r2 = paragraph.add_run(tail)
                style_run(r2, cn_font=SONG, size=size, bold=bold, en_font=TIMES,
                          italic=bool(re.fullmatch(r"[A-Za-z]", tail)))
                if op == "_":
                    r2.font.subscript = True
                else:
                    r2.font.superscript = True
        else:
            style_run(paragraph.add_run(part), cn_font=cn_font, size=size,
                      bold=bold, en_font=TIMES if cn_font == SONG else None)


def para(doc, text="", cn_font=SONG, size=XIAO_SI, bold=False, align=None,
         first_line_chars=0, space_before=0, space_after=0, line_spacing=1.5,
         en_font=None, keep_with_next=False, color=None):
    """通用段落构造器。空 text 时返回空段落供后续手工填充 run。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    if first_line_chars:
        set_indent_chars(p, first_line_chars=first_line_chars, size=size)
    if text:
        if "@" in text:
            emit_rich(p, text, cn_font=cn_font, size=size, bold=bold)
        else:
            style_run(p.add_run(text), cn_font=cn_font, size=size, bold=bold,
                      en_font=en_font, color=color)
    return p


def body(doc, text):
    """正文：宋体、小四、首行缩进 2 字符、1.5 倍行距、段前段后为 0。"""
    return para(doc, text, cn_font=SONG, size=XIAO_SI, first_line_chars=2,
                space_before=0, space_after=0, line_spacing=1.5,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def h1(doc, text):
    """一级标题：黑体、小三、加粗、居中。"""
    p = para(doc, text, cn_font=HEI, size=XIAO_SAN, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=10,
             line_spacing=1.5, keep_with_next=True)
    _set_outline(p, 0)
    return p


def h2(doc, text):
    """二级标题：黑体、四号。"""
    p = para(doc, text, cn_font=HEI, size=SI, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=10, space_after=6, line_spacing=1.5,
             keep_with_next=True)
    _set_outline(p, 1)
    return p


def h3(doc, text):
    """三级标题：黑体、小四。"""
    p = para(doc, text, cn_font=HEI, size=XIAO_SI, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=6, space_after=4, line_spacing=1.5,
             keep_with_next=True)
    _set_outline(p, 2)
    return p


def h4(doc, text):
    """四级标题：宋体小四加粗，缩进 2 字符。"""
    return para(doc, text, cn_font=SONG, size=XIAO_SI, bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT, first_line_chars=2,
                space_before=4, space_after=2, line_spacing=1.5,
                keep_with_next=True)


def _set_outline(paragraph, level):
    """写入大纲级别，使自动目录可以收录自定义样式的标题。"""
    pPr = paragraph._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level))
    pPr.insert_element_before(ol, *_AFTER_OUTLINE)


def fig_caption(doc, text):
    """图题：置于图下方，五号黑体，居中。"""
    return para(doc, text, cn_font=HEI, size=WU, align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=4, space_after=10, line_spacing=1.0)


def tab_caption(doc, text):
    """表题：置于表上方，五号黑体，居中。"""
    return para(doc, text, cn_font=HEI, size=WU, align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=10, space_after=3, line_spacing=1.0,
                keep_with_next=True)


def fig_note(doc, text):
    """图/表下方的注释：小五宋体，居中。"""
    return para(doc, text, cn_font=SONG, size=XIAO_WU,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8,
                line_spacing=1.0)


def add_picture(doc, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.keep_with_next = True
    p.add_run().add_picture(path, width=Cm(width_cm))
    return p


def formula(doc, latex_plain, number):
    """公式段落：居中显示式子，右侧给出编号（用制表位对齐）。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    # 居中制表位 + 右对齐制表位，模拟公式居中、编号右对齐
    tabs = OxmlElement("w:tabs")
    for pos, val in ((4200, "center"), (8400, "right")):
        t = OxmlElement("w:tab")
        t.set(qn("w:val"), val)
        t.set(qn("w:pos"), str(pos))
        tabs.append(t)
    p._p.get_or_add_pPr().insert_element_before(tabs, *_AFTER_TABS)
    style_run(p.add_run("\t" + latex_plain), cn_font=SONG, size=XIAO_SI,
              en_font=TIMES, italic=False)
    style_run(p.add_run("\t（" + str(number) + "）"), cn_font=SONG,
              size=XIAO_SI, en_font=TIMES)
    return p


# ---- 三线表 -----------------------------------------------------------
def _border(el, edge, sz, val="single"):
    tag = "w:" + edge
    e = el.find(qn(tag))
    if e is None:
        e = OxmlElement(tag)
        el.append(e)
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), "000000")


def _cell_borders(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "nil")
        b.append(e)
    tcPr.insert_element_before(b, "w:shd", "w:noWrap", "w:tcMar",
                               "w:textDirection", "w:tcFitText", "w:vAlign",
                               "w:hideMark", "w:cellIns", "w:cellDel",
                               "w:cellMerge", "w:tcPrChange")
    if top:
        _border(b, "top", top)
    if bottom:
        _border(b, "bottom", bottom)


def _no_split(row):
    """禁止单行跨页拆分，避免同一行的文字被截为两半。"""
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:cantSplit")
    e.set(qn("w:val"), "true")
    trPr.insert_element_before(e, "w:trHeight", "w:tblHeader",
                               "w:tblCellSpacing", "w:jc", "w:hidden",
                               "w:ins", "w:del", "w:trPrChange")


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader")
    h.set(qn("w:val"), "true")
    trPr.insert_element_before(h, "w:tblCellSpacing", "w:jc", "w:hidden",
                               "w:ins", "w:del", "w:trPrChange")


def three_line_table(doc, header, rows, widths_cm, size=WU,
                     aligns=None, header_size=None):
    """三线表：顶线与底线 1.5 pt，表头下横线 0.75 pt，无竖线与内部横线。"""
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = doc.styles["Table Grid"]
    tblPr = table._tbl.tblPr
    # 去掉默认表格样式的全部边框，再由单元格逐个补三线
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.insert_element_before(borders, "w:shd", "w:tblLayout", "w:tblCellMar",
                                "w:tblLook", "w:tblCaption", "w:tblDescription",
                                "w:tblPrChange")

    aligns = aligns or ["center"] * ncol
    amap = {"center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT}

    hdr = table.rows[0]
    _no_split(hdr)
    _repeat_header(hdr)
    for j, txt in enumerate(header):
        cell = hdr.cells[j]
        cell.width = Cm(widths_cm[j])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        if "@" in txt:
            emit_rich(p, txt, cn_font=HEI, size=header_size or size)
        else:
            style_run(p.add_run(txt), cn_font=HEI, size=header_size or size)
        _cell_borders(cell, top=12, bottom=6)

    for i, r in enumerate(rows):
        row = table.add_row()
        _no_split(row)
        last = (i == len(rows) - 1)
        for j, txt in enumerate(r):
            cell = row.cells[j]
            cell.width = Cm(widths_cm[j])
            p = cell.paragraphs[0]
            p.alignment = amap[aligns[j]]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            txt = str(txt)
            if "@" in txt:
                emit_rich(p, txt, cn_font=SONG, size=size)
            else:
                style_run(p.add_run(txt), cn_font=SONG, size=size)
            _cell_borders(cell, bottom=12 if last else None)
    return table


def add_toc(doc):
    """插入自动目录域（打开文档时提示更新）。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "（请在 Word 中按 Ctrl+A、F9 更新目录）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        r._r.append(el)
    style_run(r, cn_font=SONG, size=XIAO_SI)
    return p


def enable_update_fields(doc):
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.insert_element_before(
        el, "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr", "w:compat",
        "w:docVars", "w:rsids", "m:mathPr", "w:attachedSchema",
        "w:themeFontLang", "w:clrSchemeMapping",
        "w:doNotIncludeSubdocsInStats", "w:doNotAutoCompressPictures",
        "w:forceUpgrade", "w:captions", "w:readModeInkLockDown",
        "w:smartTagType", "sl:schemaLibrary", "w:shapeDefaults",
        "w:doNotEmbedSmartTags", "w:decimalSymbol", "w:listSeparator")


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, i, e):
        r._r.append(el)
    style_run(r, cn_font=SONG, size=WU)
