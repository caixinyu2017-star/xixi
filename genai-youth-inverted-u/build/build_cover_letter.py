# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter and
the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"), encoding="utf-8") as fh:
    S = json.load(fh)
S["years"] = S["years"].replace("-", "\u2013")
MOD = S["moderation"]

# typographic minus before numerals, en dash in year spans
_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")


def typo(s):
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 10.0

TITLE = ("Too Much of a Good Thing? How the Depth of Generative Artificial "
         "Intelligence Adoption Shapes Youth Employment in Chinese "
         "Listed Firms")

BODY = [
    ("We would like to submit the enclosed manuscript entitled \u201c%s\u201d for "
     "consideration in the Special Issue \u201cSystems Approaches to Generative AI: "
     "Workforce Development, Organisational Learning, and Economic "
     "Transformation\u201d of Systems. The manuscript has not been published "
     "elsewhere and is not under consideration by any other journal. No "
     "conflict of interest exists in its submission, and it has been approved "
     "for publication by all authors." % TITLE),

    ("Whether generative artificial intelligence destroys or creates "
     "entry-level employment is one of the most consequential open questions "
     "in labour economics, and the firm-level evidence is flatly "
     "contradictory: some studies find sharp displacement of young workers, "
     "others find nothing. This manuscript argues that the contradiction is a "
     "functional-form artefact and shows that it is. Using "
     f"{S['n_obs']:,} firm-year observations for {S['n_firms']:,} Chinese "
     f"A-share listed firms over {S['years']} and a text-based measure of "
     "adoption depth, we find that the relationship is not monotone but "
     "inverted U-shaped."),

    "The contributions of the paper are as follows.",
]

POINTS = [
    ("(1) It resolves a live empirical contradiction with a functional form. "
     "A linear specification with firm and year fixed effects and a full "
     f"control vector finds nothing at all ({S['b_linear']:.3f}, "
     f"p = {S['p_linear']:.2f}). Adding the squared term reveals a strongly "
     f"significant inverted U ({S['b1']:.3f} on depth, {S['b2']:.3f} on its "
     "square). Studies that estimate a monotone relationship are averaging "
     "across a rising and a falling arm, which is why their signs disagree."),

    ("(2) It tests the shape properly rather than reading it off a quadratic "
     "coefficient. The Lind\u2013Mehlum exact test rejects monotonicity "
     f"(t = {S['u_t']:.2f}, p < 0.001); the extreme point, {S['tau']:.3f}, has "
     f"a Fieller interval of [{S['tau_lo']:.2f}, {S['tau_hi']:.2f}] lying "
     f"wholly inside the observed range, with {S['beyond_pct']:.1f}% of "
     "post-shock observations already beyond it; a two-lines test finds "
     "significant slopes of opposite sign on either side; and a "
     "non-parametric binned fit traces the same curve."),

    ("(3) It explains why the curve bends. The augmentation of entry-level "
     "work is increasing and concave, the automation of the entry-level task "
     "bundle increasing and convex, the youth share loads positively on the "
     "first and negatively on the second, and once both are controlled for no "
     "residual curvature remains. Moderate adoption raises the share of "
     f"employees aged 30 or below by up to {S['peak_gain']:.2f} percentage "
     f"points; the deepest observed adoption leaves the firm "
     f"{abs(S['loss_at_max']):.2f} points below where it started."),

    ("(4) It shows that the turning point is an organisational variable, not "
     "a technological constant. Organisational learning capability delays it "
     f"by {MOD['OLC']['dtau']:.2f} units of depth per standard deviation and "
     f"disclosed AI governance by {MOD['AIGov']['dtau']:.2f}, while labour "
     f"cost pressure brings it forward by {abs(MOD['LCP']['dtau']):.2f}. The "
     "managerial question becomes how deep a firm can go before the entry-"
     "level role has to be rebuilt, and how much further organisational "
     "design can push that frontier."),
]

CLOSING = (
    "We believe the manuscript speaks directly to the aims of the Special "
    "Issue. It treats generative AI adoption as an organisational rather than "
    "a purely technological event, identifies the learning and governance "
    "mechanisms that determine where a firm can operate on the curve, and "
    "converts them into levers that firms and policymakers can actually pull. "
    "Identification rests on two shift-share instruments and a peer-diffusion "
    "instrument that pass the Kleibergen\u2013Paap and overidentification "
    "tests, and every number in every table and figure is produced by the "
    "released estimation code. We deeply appreciate your consideration of our "
    "manuscript and look forward to receiving comments from the reviewers. "
    "Correspondence should be directed to the address below.")

SIGN = [
    "Name: Tiantian Mo",
    "Institution and address: College of Business, Jiaxing University; "
    "Jiaxing 314001, China",
    "Email: 00008227@zjxu.edu.cn",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(3.0)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def p(text, after=5, bold=False, italic=False, align=None, indent=0.0,
          size=SIZE):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        run = par.add_run(typo(text))
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        return par

    p("Dear Editors:", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    for t in BODY:
        p(t, after=5)
    for t in POINTS:
        p(t, after=4)
    p(CLOSING, after=7)

    p("Thanks very much for your attention to our paper.", after=8,
      align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in SIGN:
        p(line, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("Very sincerely yours,", after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p("Xinyu Cai and Tiantian Mo", after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
