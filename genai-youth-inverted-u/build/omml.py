"""Minimal OMML (Word native equation) builder for python-docx.

Produces <m:oMath> element trees that Word opens in the built-in equation editor.
Supports runs, subscripts, superscripts, sub-sup, fractions, radicals, grouping,
delimiters, and accents (prime).  Numbered display equations are laid out with a
centered equation and a right-aligned number via paragraph tab stops.
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Twips

M = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _m(tag):
    return OxmlElement("m:" + tag)


def run(text, sty=None):
    """A math run.  sty: 'p'=plain upright, 'i'=italic, 'bi'=bold-italic, None=default."""
    r = _m("r")
    if sty is not None:
        rpr = _m("rPr")
        s = _m("sty")
        s.set(qn("m:val"), sty)
        rpr.append(s)
        r.append(rpr)
    t = _m("t")
    # preserve spaces
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _wrap_e(children):
    e = _m("e")
    for c in children:
        e.append(c)
    return e


def _slot(tag, children):
    node = _m(tag)
    for c in children:
        node.append(c)
    return node


def _as_list(x):
    if isinstance(x, list):
        return x
    return [x]


def sub(base, subscript):
    node = _m("sSub")
    node.append(_slot("e", _as_list(base)))
    node.append(_slot("sub", _as_list(subscript)))
    return node


def sup(base, supscript):
    node = _m("sSup")
    node.append(_slot("e", _as_list(base)))
    node.append(_slot("sup", _as_list(supscript)))
    return node


def subsup(base, subscript, supscript):
    node = _m("sSubSup")
    node.append(_slot("e", _as_list(base)))
    node.append(_slot("sub", _as_list(subscript)))
    node.append(_slot("sup", _as_list(supscript)))
    return node


def frac(num, den):
    f = _m("f")
    f.append(_slot("num", _as_list(num)))
    f.append(_slot("den", _as_list(den)))
    return f


def rad(radicand, degree=None):
    r = _m("rad")
    pr = _m("radPr")
    deghide = _m("degHide")
    if degree is None:
        deghide.set(qn("m:val"), "1")
    else:
        deghide.set(qn("m:val"), "0")
    pr.append(deghide)
    r.append(pr)
    if degree is None:
        r.append(_slot("deg", []))
    else:
        r.append(_slot("deg", _as_list(degree)))
    r.append(_slot("e", _as_list(radicand)))
    return r


def delim(children, left="(", right=")"):
    """Delimiter (auto-sizing brackets)."""
    d = _m("d")
    pr = _m("dPr")
    lp = _m("begChr"); lp.set(qn("m:val"), left)
    rp = _m("endChr"); rp.set(qn("m:val"), right)
    pr.append(lp); pr.append(rp)
    d.append(pr)
    d.append(_slot("e", _as_list(children)))
    return d


def acc(base, char="́"):
    """Accent over base; default combining acute (prime-like). Use char='^' etc."""
    a = _m("acc")
    pr = _m("accPr")
    ch = _m("chr"); ch.set(qn("m:val"), char)
    pr.append(ch)
    a.append(pr)
    a.append(_slot("e", _as_list(base)))
    return a


def bar(base):
    b = _m("bar")
    pr = _m("barPr")
    pos = _m("pos"); pos.set(qn("m:val"), "top")
    pr.append(pos)
    b.append(pr)
    b.append(_slot("e", _as_list(base)))
    return b


def omath(children):
    """Wrap a list of math elements into an <m:oMath> (inline)."""
    o = _m("oMath")
    for c in _as_list(children):
        o.append(c)
    return o


def _set_tabs(paragraph, right_twips, center_twips):
    ppr = paragraph._p.get_or_add_pPr()
    tabs = _m_w("tabs")
    tc = _m_w("tab"); tc.set(qn("w:val"), "center"); tc.set(qn("w:pos"), str(center_twips))
    tr = _m_w("tab"); tr.set(qn("w:val"), "right"); tr.set(qn("w:pos"), str(right_twips))
    tabs.append(tc); tabs.append(tr)
    ppr.append(tabs)


def _m_w(tag):
    return OxmlElement("w:" + tag)


def add_numbered_equation(doc, math_children, number, right_twips=9072, center_twips=4536,
                          style=None):
    """Add a display equation centered with a right-aligned number '(n)'.

    right_twips: right margin position (default 6.3in text width).
    """
    p = doc.add_paragraph(style=style)
    _set_tabs(p, right_twips, center_twips)
    # tab -> equation -> tab -> number
    p.add_run().add_tab()
    p._p.append(omath(math_children))
    r2 = p.add_run()
    r2.add_tab()
    r2.add_text("(%s)" % number)
    return p


def add_inline_math(paragraph, math_children):
    """Insert inline math into an existing paragraph run position (appends)."""
    paragraph._p.append(omath(math_children))


def make_inline_para(doc, prefix_runs, math_children, suffix_text=""):
    """Convenience: paragraph with text, inline math, text."""
    p = doc.add_paragraph()
    if prefix_runs:
        p.add_run(prefix_runs)
    p._p.append(omath(math_children))
    if suffix_text:
        p.add_run(suffix_text)
    return p
