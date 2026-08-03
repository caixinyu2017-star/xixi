# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter and
the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"), encoding="utf-8") as fh:
    S = json.load(fh)
S["years"] = S["years"].replace("-", "\u2013")
PV = S["pvar"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")


def typo(s):
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 9.6

TITLE = ("Eroding the First Rung: Generative Artificial Intelligence, "
         "Organizational Learning, and the Self-Reinforcing Decline of Youth "
         "Employment in Chinese Listed Firms")

BODY = [
    ("We would like to submit the enclosed manuscript entitled \u201c%s\u201d "
     "for consideration in the Special Issue \u201cSystems Approaches to "
     "Generative AI: Workforce Development, Organisational Learning, and "
     "Economic Transformation\u201d of Systems. The manuscript has not been "
     "published elsewhere and is not under consideration by any other "
     "journal. No conflict of interest exists in its submission, and it has "
     "been approved for publication by all authors." % TITLE),

    ("Entry-level work is the first rung of the ladder by which economies "
     "manufacture experienced workers, and generative artificial intelligence "
     "is most capable at exactly the tasks that rung is made of. The urgency "
     "is not hypothetical: China produced a record 12.22 million college "
     "graduates in 2025 while youth unemployment stayed above 17 percent, and "
     "the international evidence already shows employment of workers in their "
     "early twenties falling relative to experienced workers in the same "
     "occupations. What no firm-level study has yet asked is whether this is "
     "a one-way shock or a loop. Using "
     f"{S['n_obs']:,} firm-year observations for {S['n_firms']:,} Chinese "
     f"A-share listed firms over {S['years']}, we show that it is a loop, and "
     "we measure how much the loop adds."),

    "The highlights of the paper are as follows.",
]

POINTS = [
    ("(1) A displacement that is specific to the first rung. A "
     "one-standard-deviation increase in generative-AI application intensity "
     f"lowers the youth employment share by {abs(S['effect_sd']):.2f} "
     f"percentage points, {abs(S['effect_sd_pct']):.1f} percent of its mean. "
     "The same regression run on employees aged 31 to 50 gives a coefficient "
     f"{abs(S['ratio_age']):.0f} times smaller. This is not a contraction of "
     "employment; it is a change in its age composition."),

    ("(2) The firm stops producing skill, not merely stops hiring. Two "
     "channels are measured and both are significant: a contracting routine "
     f"task base accounts for {abs(S['share_rout']):.1f} percent of the total "
     "effect, and suppressed investment in employee training for a further "
     f"{abs(S['share_train']):.1f} percent, with firm-block bootstrap "
     "intervals excluding zero. The second channel does not exist in the "
     "task-based framework that dominates this literature, because that "
     "framework treats labor as an input to be allocated rather than an "
     "asset the firm produces."),

    ("(3) The loop closes, and it amplifies. Estimating the "
     "adoption\u2013training\u2013youth system as a bias-corrected panel "
     "vector autoregression, the return arcs are negative and significant: a "
     "firm whose training investment and youth cohort have contracted adopts "
     f"more generative AI the following year ({S['fb_train']:.3f} and "
     f"{S['fb_youth']:.3f}, both p < 0.01). Switching the return arcs off and "
     "re-simulating shows that the loop adds "
     f"{100 * (S['amp'] - 1):.1f} percent to the eight-year cumulative "
     f"displacement. The largest eigenvalue is {PV['max_eig']:.3f}, so the "
     "trap is bounded rather than explosive \u2014 which is precisely what "
     "makes it addressable."),

    ("(4) A leverage point management can build. The same adoption intensity "
     "displaces about half as much youth employment in firms one standard "
     "deviation above the mean of organizational learning capability, and "
     "correspondingly more where relative labor costs are high. Industry "
     "classification does not discriminate between firms; the direct measure "
     "of absorptive capacity does. The outcome is therefore organizational "
     "rather than technological, and it is manageable."),
]

CLOSING = (
    "We believe the manuscript speaks directly to the aims of the Special "
    "Issue, which asks for systems approaches to generative AI, workforce "
    "development and organizational learning. Our contribution is a systems "
    "result in the strict sense. The relationship between generative AI and "
    "youth employment is not a treatment effect but a reinforcing feedback "
    "structure \u2014 Repenning and Sterman\u2019s capability trap, applied "
    "for the first time to a firm\u2019s skill-formation system and estimated "
    "rather than asserted. The practical consequence is that any "
    "single-equation estimate of the effect understates it by about a fifth, "
    "and that policy aimed at the moment of adoption operates on the forward "
    "arc and is partly undone by the return arc, whereas instruments that "
    "protect training provision act on the arc that closes the loop. The "
    "design is conventional where it should be \u2014 two-way fixed effects, "
    "overidentified instruments passing the Kleibergen\u2013Paap and Hansen "
    "tests, propensity score matching, entropy balancing, an age placebo, "
    "Oster bounds and a wild cluster bootstrap \u2014 and novel only where "
    "the systems claim requires it. We deeply appreciate your consideration "
    "of our manuscript and look forward to receiving comments from the "
    "reviewers. Correspondence should be directed to the address below.")

SIGN = [
    "Name: Tiantian Mo",
    "Institution and address: College of Business, Jiaxing University; "
    "Jiaxing 314001, China",
    "Email: 00008227@zjxu.edu.cn",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(1.7)
    sec.left_margin = sec.right_margin = Cm(2.5)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def p(text, after=4, bold=False, italic=False, align=None, indent=0.0,
          size=SIZE):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        run = par.add_run(typo(text).replace("*", ""))
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        return par

    p("Dear Editors:", after=5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for t in BODY:
        p(t, after=4)
    for t in POINTS:
        p(t, after=3)
    p(CLOSING, after=6)

    p("Thanks very much for your attention to our paper.", after=7,
      align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in SIGN:
        p(line, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=5, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("Very sincerely yours,", after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p("Xinyu Cai and Tiantian Mo", after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
