# -*- coding: utf-8 -*-
"""将带 #标记# 与 [@key] 引用占位符的论文源文件转换为规范排版的 Word 文档。
引用编号按全篇首次出现顺序自动分配；相邻编号自动合并为 [n-m]；
文末参考文献表由 citations.json 数据库自动生成，保证正文与文献严格交叉对应。
"""
import json, re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CIT_RE = re.compile(r'\[@([a-zA-Z0-9_]+)\]')

def set_font(run, cn_font='宋体', en_font='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_para(doc, text, cn_font='宋体', size=12, bold=False, align=None,
             indent_chars=0, space_before=0, space_after=0, line_spacing=1.5,
             citations=None, sup_ok=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent_chars:
        # 首行缩进2字符
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLineChars'), str(indent_chars * 100))
        p._p.get_or_add_pPr().append(ind)
    # 拆分引用占位符 -> 上标
    parts = re.split(r'(\[\d+(?:[-,，]\d+)*\])', text)
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        set_font(r, cn_font=cn_font, size=size, bold=bold)
        if sup_ok and re.fullmatch(r'\[\d+(?:[-,，]\d+)*\]', part):
            r.font.superscript = True
            r.font.size = Pt(max(size - 3, 8))
    return p

def resolve_citations(lines, citdb):
    """按首次出现顺序为 [@key] 分配编号，返回(处理后的行, 有序key列表)。"""
    order = []
    def key_num(k):
        if k not in order:
            order.append(k)
        return order.index(k) + 1
    out = []
    for ln in lines:
        # 先收集本行所有 key（保持出现顺序编号）
        def repl_group(m):
            return m.group(0)
        # 逐个替换：连续的 [@a][@b][@c] 合并
        def merge_repl(match):
            keys = CIT_RE.findall(match.group(0))
            nums = sorted(key_num(k) for k in keys)
            # 合并连续区间
            groups, start, prev = [], nums[0], nums[0]
            for n in nums[1:]:
                if n == prev + 1:
                    prev = n
                else:
                    groups.append((start, prev)); start = prev = n
            groups.append((start, prev))
            segs = []
            for a, b in groups:
                if a == b: segs.append(str(a))
                elif b == a + 1: segs.append(f'{a},{b}')
                else: segs.append(f'{a}-{b}')
            return '[' + ','.join(segs) + ']'
        ln2 = re.sub(r'(?:\[@[a-zA-Z0-9_]+\])+', merge_repl, ln)
        out.append(ln2)
    missing = [k for k in order if k not in citdb]
    if missing:
        raise SystemExit(f'citations.json 缺少条目: {missing}')
    return out, order

def build(src_path, cit_path, out_path):
    citdb = json.load(open(cit_path, encoding='utf-8'))
    raw = open(src_path, encoding='utf-8').read().splitlines()
    lines, order = resolve_citations(raw, citdb)

    doc = Document()
    # 页面设置 A4
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.0)

    meta = {}
    body = []
    for ln in lines:
        m = re.match(r'^#(TITLE|SUBTITLE|AUTHOR|AFFIL|ABSTRACT|KEYWORDS|CLC|ENTITLE|ENABSTRACT|ENKEYWORDS)#\s*(.*)$', ln)
        if m:
            meta[m.group(1)] = m.group(2).strip()
        else:
            body.append(ln)

    # 标题区
    add_para(doc, meta['TITLE'], cn_font='黑体', size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.3)
    if meta.get('SUBTITLE'):
        add_para(doc, meta['SUBTITLE'], cn_font='黑体', size=14, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line_spacing=1.3)
    add_para(doc, meta['AUTHOR'], cn_font='楷体', size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.3)
    add_para(doc, meta['AFFIL'], cn_font='楷体', size=10.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line_spacing=1.3)
    for k, fs in (('ABSTRACT', 10.5), ('KEYWORDS', 10.5), ('CLC', 10.5)):
        if meta.get(k):
            add_para(doc, meta[k], cn_font='宋体', size=fs, space_after=4, line_spacing=1.3)
    for k in ('ENTITLE', 'ENABSTRACT', 'ENKEYWORDS'):
        if meta.get(k):
            add_para(doc, meta[k], cn_font='宋体', size=9,
                     space_after=4, line_spacing=1.15,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    i = 0
    while i < len(body):
        ln = body[i].strip()
        if not ln:
            i += 1; continue
        if ln.startswith('#H1#'):
            add_para(doc, ln[4:].strip(), cn_font='黑体', size=14, bold=True,
                     space_before=12, space_after=6, line_spacing=1.3)
        elif ln.startswith('#H2#'):
            add_para(doc, ln[4:].strip(), cn_font='黑体', size=12, bold=True,
                     space_before=8, space_after=4, line_spacing=1.3)
        elif ln.startswith('#TABLE#'):
            caption = ln[7:].strip()
            add_para(doc, caption, cn_font='黑体', size=10.5, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4, line_spacing=1.2)
            rows = []
            i += 1
            while i < len(body) and not body[i].strip().startswith('#ENDTABLE#'):
                if body[i].strip():
                    rows.append([c.strip() for c in body[i].strip().split('|')])
                i += 1
            ncol = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncol)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci in range(ncol):
                    cell = tbl.cell(ri, ci)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    txt = row[ci] if ci < len(row) else ''
                    parts = re.split(r'(\[\d+(?:[-,，]\d+)*\])', txt)
                    for part in parts:
                        if not part: continue
                        r = p.add_run(part)
                        set_font(r, cn_font='宋体', size=9, bold=(ri == 0))
                        if re.fullmatch(r'\[\d+(?:[-,，]\d+)*\]', part):
                            r.font.superscript = True; r.font.size = Pt(7)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif ln.startswith('#REFS#'):
            add_para(doc, '参考文献：', cn_font='黑体', size=12, bold=True,
                     space_before=14, space_after=6, line_spacing=1.3)
            for idx, key in enumerate(order, 1):
                add_para(doc, f'［{idx}］ {citdb[key]}', cn_font='宋体', size=9,
                         space_after=2, line_spacing=1.2, sup_ok=False)
        else:
            add_para(doc, ln, cn_font='宋体', size=12, indent_chars=2,
                     space_after=0, line_spacing=1.5,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        i += 1

    doc.save(out_path)
    print(f'OK {out_path}  引用文献 {len(order)} 条')
    return order

if __name__ == '__main__':
    src, cit, out = sys.argv[1], sys.argv[2], sys.argv[3]
    build(src, cit, out)
