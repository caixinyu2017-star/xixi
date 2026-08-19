# -*- coding: utf-8 -*-
"""
生成《"双优"工程"优秀办学"与"优秀服务"关键要素梳理（附来源）》Word 文档。

用法：  python3 build_docx.py
输出：  双优工程关键要素梳理.docx
"""
import os
import sys

from docx import Document

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from content_core import YOUZHI_BANXUE, YOUXIU_FUWU  # noqa: E402
from content_web import (  # noqa: E402
    COVER_SUBTITLE, EDITORIAL_NOTE, POLICY_ORIGIN, THREE_PROJECTS_TABLE,
    SCHOLAR_FRAMEWORKS, SEMINAR_VIEWS, BENCHMARK_SYSTEMS, JIAXING_SECTION,
    REFERENCES, SOURCE_TIERS, WEB_SUPPLEMENT,
)

HEI = "黑体"
SONG = "宋体"
FANGSONG = "仿宋"
KAI = "楷体"
EN = "Times New Roman"

ACCENT = RGBColor(0x1F, 0x3B, 0x73)      # 深蓝
ACCENT_LIGHT = "DCE3F0"
GREY_LIGHT = "F2F4F8"
BAND_A = "1F3B73"
BAND_B = "2E6B4F"


# --------------------------------------------------------------------------
# OOXML 元素顺序：pPr / tblPr / tcPr 的子元素必须按 schema 顺序排列，
# 直接 append 会生成 LibreOffice 等严格解析器无法打开的文档。
# --------------------------------------------------------------------------
_PPR_AFTER_PBDR = (
    "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)
_PPR_AFTER_OUTLINE = ("w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange")
_TBLPR_AFTER_BORDERS = (
    "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook", "w:tblCaption",
    "w:tblDescription", "w:tblPrChange",
)
_TCPR_AFTER_SHD = (
    "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign",
    "w:hideMark", "w:headers", "w:cellIns", "w:cellDel", "w:cellMerge",
    "w:tcPrChange",
)


def _insert_ordered(parent, el, successors):
    """按 OOXML schema 顺序把 el 插入 parent。"""
    parent.insert_element_before(el, *successors)
    return el


# --------------------------------------------------------------------------
# 底层排版工具
# --------------------------------------------------------------------------
def set_run(run, *, cn=SONG, en=EN, size=10.5, bold=False, color=None,
            italic=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    return run


def para(doc, text="", *, cn=SONG, en=EN, size=10.5, bold=False, color=None,
         align=None, first_indent=True, space_before=0, space_after=4,
         line=18, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(line)
    if align is not None:
        p.alignment = align
    if first_indent:
        pf.first_line_indent = Pt(size * 2)
    if text:
        set_run(p.add_run(text), cn=cn, en=en, size=size, bold=bold,
                color=color, italic=italic)
    return p


def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    _insert_ordered(tcpr, shd, _TCPR_AFTER_SHD)


def cell_text(cell, text, *, cn=SONG, size=9, bold=False, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, line=13):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(line)
    set_run(p.add_run(text), cn=cn, size=size, bold=bold, color=color)
    return cell


def set_table_borders(table, color="9AA7BF", sz=4):
    tbl = table._tbl
    tblpr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    _insert_ordered(tblpr, borders, _TBLPR_AFTER_BORDERS)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    _insert_ordered(trpr, el, ("w:cantSplit", "w:trHeight", "w:tblCellSpacing",
                               "w:jc", "w:hidden", "w:ins", "w:del",
                               "w:trPrChange"))


def heading(doc, text, level=1, *, page_break=False):
    if page_break:
        doc.add_page_break()
    sizes = {1: 16, 2: 13, 3: 11.5}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt({1: 18, 2: 14, 3: 10}[level])
    pf.space_after = Pt({1: 10, 2: 7, 3: 5}[level])
    pf.line_spacing = Pt(sizes[level] * 1.5)
    pf.keep_with_next = True
    set_run(p.add_run(text), cn=HEI, size=sizes[level], bold=True,
            color=ACCENT if level <= 2 else None)
    # 设置大纲级别，使 Word 的 TOC 域能够抓取（不改变外观）
    ppr = p._p.get_or_add_pPr()
    olvl = OxmlElement("w:outlineLvl")
    olvl.set(qn("w:val"), str(level - 1))
    _insert_ordered(ppr, olvl, _PPR_AFTER_OUTLINE)
    if level == 1:
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), BAND_A)
        pbdr.append(bottom)
        _insert_ordered(p._p.get_or_add_pPr(), pbdr, _PPR_AFTER_PBDR)
    return p


def source_note(doc, text):
    """来源标注（灰色小字，无首行缩进）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(8)
    pf.line_spacing = Pt(13)
    pf.left_indent = Pt(12)
    set_run(p.add_run("【来源】" + text), cn=KAI, size=8.5,
            color=RGBColor(0x55, 0x5F, 0x70))
    return p


def callout(doc, title, body, fill=GREY_LIGHT):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t, color="C3CCE0", sz=4)
    c = t.rows[0].cells[0]
    shade(c, fill)
    c.text = ""
    p0 = c.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    p0.paragraph_format.line_spacing = Pt(15)
    set_run(p0.add_run(title), cn=HEI, size=10, bold=True, color=ACCENT)
    for line_ in body:
        p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(15)
        set_run(p.add_run(line_), cn=SONG, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "（在 Word 中按 Ctrl+A 后按 F9 可更新目录）"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for e in (fld, instr, sep, txt, end):
        r._r.append(e)
    return p


# --------------------------------------------------------------------------
# 页眉页脚
# --------------------------------------------------------------------------
def add_page_number(paragraph):
    run = paragraph.add_run()
    for instr_text in ("PAGE",):
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr_text
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for e in (fld, it, end):
            run._r.append(e)
    set_run(run, cn=SONG, size=9, color=RGBColor(0x66, 0x6F, 0x80))


def setup_section(sec):
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)
    sec.header_distance = Cm(1.4)
    sec.footer_distance = Cm(1.4)

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(hp.add_run("“双优”工程“优秀办学”与“优秀服务”关键要素梳理"),
            cn=KAI, size=9, color=RGBColor(0x66, 0x6F, 0x80))
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), "AAB4C8")
    pbdr.append(b)
    _insert_ordered(hp._p.get_or_add_pPr(), pbdr, _PPR_AFTER_PBDR)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


# --------------------------------------------------------------------------
# 文档构建
# --------------------------------------------------------------------------
def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = EN
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

    setup_section(doc.sections[0])

    # ---------------- 封面 ----------------
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run("“双优”工程"), cn=HEI, size=30, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("“优秀办学”与“优秀服务”关键要素梳理"),
            cn=HEI, size=22, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    set_run(p.add_run("——政策依据、学理框架、指标清单与来源核验"),
            cn=KAI, size=14, color=RGBColor(0x44, 0x4F, 0x63))

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t, color="1F3B73", sz=8)
    c = t.rows[0].cells[0]
    shade(c, ACCENT_LIGHT)
    c.text = ""
    for i, line_ in enumerate(COVER_SUBTITLE):
        pp = c.paragraphs[0] if i == 0 else c.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_after = Pt(3)
        pp.paragraph_format.line_spacing = Pt(17)
        set_run(pp.add_run(line_), cn=KAI, size=10.5,
                color=RGBColor(0x1F, 0x3B, 0x73))

    doc.add_page_break()

    # ---------------- 目录 ----------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run(p.add_run("目　　录"), cn=HEI, size=18, bold=True, color=ACCENT)
    add_toc(doc)
    doc.add_page_break()

    # ---------------- 编制说明 ----------------
    heading(doc, "编制说明", 1)
    for t_ in EDITORIAL_NOTE:
        para(doc, t_)

    heading(doc, "资料的权威性分层（阅读本文件前请务必先看）", 2)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_borders(tbl)
    hdr = tbl.rows[0]
    repeat_header(hdr)
    for i, h in enumerate(("层级", "内容", "可引用性")):
        shade(hdr.cells[i], BAND_A)
        cell_text(hdr.cells[i], h, cn=HEI, size=9.5, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for k, (tier, cont, usab) in enumerate(SOURCE_TIERS):
        row = tbl.add_row()
        fill = "FFFFFF" if k % 2 == 0 else GREY_LIGHT
        for i, v in enumerate((tier, cont, usab)):
            shade(row.cells[i], fill)
            cell_text(row.cells[i], v, size=9,
                      bold=(i == 0), cn=HEI if i == 0 else SONG)
    for w, col in zip((Cm(2.6), Cm(9.4), Cm(4.0)), tbl.columns):
        for cell in col.cells:
            cell.width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---------------- 第一章 ----------------
    heading(doc, "一、“双优”工程的政策由来与核心内涵", 1, page_break=True)
    for h2, paras, src in POLICY_ORIGIN:
        heading(doc, h2, 2)
        for t_ in paras:
            if t_.startswith("§QUOTE§"):
                q = doc.add_paragraph()
                qf = q.paragraph_format
                qf.left_indent = Pt(20)
                qf.right_indent = Pt(10)
                qf.space_before = Pt(4)
                qf.space_after = Pt(6)
                qf.line_spacing = Pt(17)
                set_run(q.add_run(t_[7:]), cn=KAI, size=10,
                        color=RGBColor(0x2A, 0x33, 0x45))
                pbdr = OxmlElement("w:pBdr")
                left = OxmlElement("w:left")
                left.set(qn("w:val"), "single")
                left.set(qn("w:sz"), "18")
                left.set(qn("w:space"), "8")
                left.set(qn("w:color"), BAND_A)
                pbdr.append(left)
                _insert_ordered(q._p.get_or_add_pPr(), pbdr,
                                _PPR_AFTER_PBDR)
            else:
                para(doc, t_)
        if src:
            source_note(doc, src)

    heading(doc, "（五）“双一流”“双高”“新双高”与“双优”的定位差异", 2)
    para(doc, "厘清三大工程的差异，是理解“双优”指标为何这样设计的前提。下表据陈文博（2026）"
              "对三大工程政策文本的比较研究整理。")
    cols = ["比较维度", "“双一流”", "“双高”", "“双优”"]
    tbl = doc.add_table(rows=1, cols=4)
    set_table_borders(tbl)
    hdr = tbl.rows[0]
    repeat_header(hdr)
    for i, h in enumerate(cols):
        shade(hdr.cells[i], BAND_A)
        cell_text(hdr.cells[i], h, cn=HEI, size=9.5, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for k, row_data in enumerate(THREE_PROJECTS_TABLE):
        row = tbl.add_row()
        fill = "FFFFFF" if k % 2 == 0 else GREY_LIGHT
        for i, v in enumerate(row_data):
            shade(row.cells[i], fill)
            cell_text(row.cells[i], v, size=8.5, bold=(i == 0),
                      cn=HEI if i == 0 else SONG)
    for w, col in zip((Cm(2.4), Cm(4.5), Cm(4.5), Cm(4.6)), tbl.columns):
        for cell in col.cells:
            cell.width = w
    source_note(doc, "陈文博. 因果耦合视角下“双优工程”的定位与评价转向——基于与“双一流”"
                     "“双高”的比较研究[J]. 应用型高等教育研究, 2026, 11(2): 1-9. 表1及正文。")

    # ---------------- 第二章：优质办学 ----------------
    heading(doc, "二、“优秀办学”（优质办学）十大关键要素", 1, page_break=True)
    callout(doc, "术语说明", [
        "论文原文的表述是“优质办学”，青塔等媒体制作的图表标注为“优秀办学”，二者所指同一组要素。",
        "本文件正文以论文原文“优质办学”为准，标题保留用户熟悉的“优秀办学”说法。",
        "十大要素的定位：既是“双优”工程遴选的要素，也是动态监测评价的要素，"
        "还是应用型高校建设的基本考量。",
    ])
    para(doc, "优质办学涵盖办学方向、师资队伍、人才培养、学科布局、科技创新、国际交流、治理体系、"
              "资源配置、毕业生发展和社会声誉等十个方面的关键要素。抓好这十个关键要素，就能够夯实"
              "应用型高校发展的基本盘，摆脱“夹心层”困境，推进学校的高质量发展。")
    source_note(doc, "瞿振元, 蔺跟荣（2026），原文第11页。")

    _elements_table(doc, YOUZHI_BANXUE, BAND_A)
    heading(doc, "（二）十大要素逐条释义（原文）", 2)
    _elements_detail(doc, YOUZHI_BANXUE, "优质办学")

    # ---------------- 第三章：优秀服务 ----------------
    heading(doc, "三、“优秀服务”十大关键要素", 1, page_break=True)
    para(doc, "应用型高校的优秀服务应该包含毕业生留用、技术合作、成果转化、平台共建、智库服务、"
              "社会培训、人才引育、产业支持度、资源开放、专家派驻等十个主要方面。")
    source_note(doc, "瞿振元, 蔺跟荣（2026），原文第12—13页。")

    _elements_table(doc, YOUXIU_FUWU, BAND_B)
    heading(doc, "（二）十大要素逐条释义（原文）", 2)
    _elements_detail(doc, YOUXIU_FUWU, "优秀服务")

    # ---------------- 第四章：转化与统一 ----------------
    heading(doc, "四、二十个关键要素的转化与统一", 1, page_break=True)
    heading(doc, "（一）互为因果的逻辑整体与“三个统一”", 2)
    para(doc, "综观上述二十个关键要素，优质办学与优秀服务构成了一个相互支撑、互为因果的逻辑整体。"
              "优质办学是优秀服务的基础，没有高素质师资队伍、科学的学科布局、完善的治理体系，"
              "就难以产生高质量的技术服务和人才供给；优秀服务是优质办学的价值体现，毕业生留用率、"
              "技术合同成交额、成果转化收益等外部贡献指标，最终检验了办学质量的实际成效。")
    para(doc, "优质办学与优秀服务的二十个关键要素，还体现了“双优”工程内外的统一性，"
              "即以下“三个统一”。这三个统一，是应用型高校评价体系去“研究型大学依附”的关键所在，"
              "也是“双优”工程区别于传统教育工程的根本特征。")
    unify = [
        ("“投入导向”与“产出导向”的统一",
         "以区域需求配置优质资源，以产出质量调整专业结构，以服务贡献检验办学成效。"),
        ("“内部评价”与“外部评价”的统一",
         "在内部质量评价的基础上，重视用人单位反馈、第三方评价等外部评价。"),
        ("“学术标准”与“应用标准”的统一",
         "将学术指标聚焦于解决产业实际问题的应用研究能力，以技术攻关实效、工艺改进水平、"
         "解决复杂工程问题能力作为衡量学术深度与广度的新标准，实现学术与应用的同生共长。"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    set_table_borders(tbl)
    hdr = tbl.rows[0]
    repeat_header(hdr)
    for i, h in enumerate(("“三个统一”", "内涵（原文）")):
        shade(hdr.cells[i], BAND_A)
        cell_text(hdr.cells[i], h, cn=HEI, size=9.5, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for k, (a, b) in enumerate(unify):
        row = tbl.add_row()
        fill = "FFFFFF" if k % 2 == 0 else GREY_LIGHT
        shade(row.cells[0], fill)
        shade(row.cells[1], fill)
        cell_text(row.cells[0], a, cn=HEI, size=9, bold=True)
        cell_text(row.cells[1], b, size=9)
    for w, col in zip((Cm(4.6), Cm(11.4)), tbl.columns):
        for cell in col.cells:
            cell.width = w
    source_note(doc, "瞿振元, 蔺跟荣（2026），原文第13—14页。")

    heading(doc, "（二）配套推进的六个重要事项", 2)
    para(doc, "论文同时提出，当前推进“双优”工程建设应当抓好六件事，可视为二十个要素落地的“施工图”：")
    six = [
        ("明确地方性应用型办学定位",
         "实现四个“契合”：学科专业与区域经济社会发展需求契合、人才培养与区域行业产业人才需求契合、"
         "人才高地建设与区域创新建设契合、社会服务与区域文化和特色发展契合。"),
        ("构建全过程产教协同人才培养模式",
         "将企业真实项目转化为学生学习与研究课题，企业技术难题转化为教师科研选题；"
         "产教协同贯穿招生、培养、就业全周期。现代产业学院是重要探索。"),
        ("加强“双师双能型”教师队伍建设",
         "建立认定条件和评价标准；实施“教师企业实践必修制度”（青年教师每3～5年累计不少于6个月"
         "企业实践）；设立“产业教授”岗位；对艺术、体育、医药卫生等不同学科应有不同标准，"
         "不能“一刀切”。"),
        ("以多元合作平台深化区域合作",
         "建设技术转移中心、大学科技园/众创空间、教师教育培训、医疗卫生服务、农业科技小院、"
         "新型智库、地方研究院等平台；推行“揭榜挂帅”，形成“政府征题、企业出题、高校接题、"
         "师生解题”；鼓励教师任“科技副总”。"),
        ("加强应用型高校的研究生教育",
         "着力加强专业硕士、专业博士学位点建设；坚持实践导向，以项目成果、技术专利、解决方案"
         "作为学位评定重要依据；实行“校内导师+企业导师”双导师制。"),
        ("构建新型治理体系，理顺政校关系",
         "构建“国家主导、省级主责、学校主体”的三级联动治理模式：中央出标准、出政策、出资金；"
         "省市明晰区域教育布局与分类发展定位；学校主动争取发展条件。"),
    ]
    for i, (a, b) in enumerate(six, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(17)
        pf.left_indent = Pt(21)
        pf.first_line_indent = Pt(-21)
        set_run(p.add_run(f"{i}. {a}　"), cn=HEI, size=10, bold=True,
                color=ACCENT)
        set_run(p.add_run(b), cn=SONG, size=10)
    source_note(doc, "瞿振元, 蔺跟荣（2026），原文第14—17页。")

    # ---------------- 第五章：学界补充框架 ----------------
    heading(doc, "五、学界提出的补充性评价框架与指标", 1, page_break=True)
    para(doc, "除瞿振元、蔺跟荣提出的二十个关键要素外，2026年已有多位学者从不同角度提出“双优”"
              "评价的维度与指标。这些属于学术研究成果，尚未进入官方文件，但对理解未来指标"
              "走向具有参考价值。")
    for h2, lead, items, src in SCHOLAR_FRAMEWORKS:
        heading(doc, h2, 2)
        if lead:
            para(doc, lead)
        if items:
            _bullet_table(doc, items)
        source_note(doc, src)

    heading(doc, "（三）学术研讨会上专家提出的评价要点", 2)
    para(doc, "2026年4月29日，“双优”工程与高水平应用型大学建设学术研讨会在广东技术师范大学举行，"
              "18位国内高等教育领域知名专家出席。会上提出的若干评价要点对指标设计具有直接参考价值：")
    for who, what in SEMINAR_VIEWS:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(17)
        pf.left_indent = Pt(21)
        pf.first_line_indent = Pt(-21)
        set_run(p.add_run("· "), cn=SONG, size=10)
        set_run(p.add_run(who + "："), cn=HEI, size=10, bold=True)
        set_run(p.add_run(what), cn=SONG, size=10)
    source_note(doc, "姚宇华, 陈伟. 以“双优”工程驱动高水平应用型大学建设——“双优”工程与"
                     "高水平应用型大学建设学术研讨会综述[J]. 高教探索, 2026(6): 126-128.")

    # ---------------- 第六章：可对标体系 ----------------
    heading(doc, "六、可对标的现行评价体系", 1, page_break=True)
    para(doc, "在“双优”工程官方指标体系正式发布之前，下列现行评价体系是判断“双优”指标"
              "可能形态最可靠的现实参照。它们均有正式发布的官方文件或公开的方法说明。")
    for name, issuer, dims, src in BENCHMARK_SYSTEMS:
        heading(doc, name, 2)
        para(doc, issuer)
        if dims:
            _bullet_table(doc, dims)
        source_note(doc, src)

    # ---------------- 第七章：网络检索补充 ----------------
    if WEB_SUPPLEMENT:
        heading(doc, "七、网络检索补充与核验结论", 1, page_break=True)
        for h2, paras, src in WEB_SUPPLEMENT:
            heading(doc, h2, 2)
            for t_ in paras:
                para(doc, t_)
            if src:
                source_note(doc, src)

    # ---------------- 第八章：嘉兴大学对标 ----------------
    heading(doc, "八、对嘉兴大学的对标建议与自检清单", 1, page_break=True)
    for h2, paras, items, src in JIAXING_SECTION:
        if h2:
            heading(doc, h2, 2)
        for t_ in paras:
            para(doc, t_)
        if items:
            _bullet_table(doc, items)
        if src:
            source_note(doc, src)

    # ---------------- 参考文献 ----------------
    heading(doc, "参考文献与资料来源", 1, page_break=True)
    for group, refs in REFERENCES:
        heading(doc, group, 2)
        for i, r in enumerate(refs, 1):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(4)
            pf.line_spacing = Pt(16)
            pf.left_indent = Pt(24)
            pf.first_line_indent = Pt(-24)
            set_run(p.add_run(f"[{i}] "), cn=SONG, size=9.5)
            set_run(p.add_run(r), cn=SONG, size=9.5)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "“双优”工程关键要素梳理.docx")
    doc.save(out)
    return out


def _elements_table(doc, elements, band):
    heading(doc, "（一）要素总表", 2)
    tbl = doc.add_table(rows=1, cols=4)
    set_table_borders(tbl)
    hdr = tbl.rows[0]
    repeat_header(hdr)
    for i, h in enumerate(("序号", "关键要素", "核心界定（论文原文要点）", "可观测的量化观测点（整理）")):
        shade(hdr.cells[i], band)
        cell_text(hdr.cells[i], h, cn=HEI, size=9, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for k, (no, name, brief, _full, obs) in enumerate(elements):
        row = tbl.add_row()
        fill = "FFFFFF" if k % 2 == 0 else GREY_LIGHT
        for i, v in enumerate((no, name, brief, obs)):
            shade(row.cells[i], fill)
            cell_text(row.cells[i], v, size=8.5,
                      cn=HEI if i == 1 else SONG,
                      bold=(i == 1),
                      align=WD_ALIGN_PARAGRAPH.CENTER if i == 0
                      else WD_ALIGN_PARAGRAPH.LEFT)
    for w, col in zip((Cm(1.0), Cm(2.6), Cm(6.2), Cm(6.2)), tbl.columns):
        for cell in col.cells:
            cell.width = w
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("注：第1—3列为论文原文表述；第4列“量化观测点”为本文件依据论文"
                      "释义及现行高校统计口径整理，供内部对标自检使用，非官方指标。"),
            cn=KAI, size=8.5, color=RGBColor(0x55, 0x5F, 0x70))


def _elements_detail(doc, elements, tag):
    order = "一二三四五六七八九十"
    for (no, name, _brief, full, _obs) in elements:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(7)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(17)
        pf.keep_with_next = True
        set_run(p.add_run(f"{order[int(no) - 1]}是{name}。"), cn=HEI,
                size=10.5, bold=True, color=ACCENT)
        para(doc, full, size=10.5, line=18, space_after=2)


def _bullet_table(doc, items):
    for a, b in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(17)
        pf.left_indent = Pt(21)
        pf.first_line_indent = Pt(-21)
        set_run(p.add_run("· "), cn=SONG, size=10)
        if a:
            set_run(p.add_run(a + "　"), cn=HEI, size=10, bold=True)
        set_run(p.add_run(b), cn=SONG, size=10)


if __name__ == "__main__":
    print(build())
