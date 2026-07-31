# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission."""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.abspath(os.path.join(HERE, "..", "Cover_Letter.docx"))

FONT = "Times New Roman"
SIZE = 10.0

TITLE = ("The Vanishing First Rung? Generative AI, Organisational Learning and "
         "Youth Entry Employment in a Socio-Technical System Dynamics Model")

BODY = [
    ("We would like to submit the enclosed manuscript entitled “%s” for "
     "consideration in the Special Issue “Systems Approaches to Generative AI: "
     "Workforce Development, Organisational Learning, and Economic "
     "Transformation” of Systems. The manuscript has not been published "
     "elsewhere and is not under consideration by any other journal. No "
     "conflict of interest exists in its submission, and it has been approved "
     "for publication by the author." % TITLE),

    ("Generative AI is diffusing fastest through exactly the codified, "
     "routine-cognitive tasks that have traditionally formed the entry port of "
     "a professional career, and the first high-frequency evidence shows "
     "employment among workers aged 22–25 in exposed occupations falling by "
     "13–16% while employment of their experienced colleagues is unchanged. "
     "The manuscript argues that this is not a technological verdict but a "
     "system outcome, and it builds and validates the socio-technical system "
     "dynamics model needed to establish that claim. The entry port is "
     "simultaneously a factor of production and the intake stage of the firm's "
     "own capability pipeline, so its contraction feeds back — with a delay of "
     "the order of the time to competence — on organisational learning, on the "
     "supply of experienced staff and on the firm's capacity to absorb the "
     "technology at all."),

    "The contributions of the paper are as follows.",
]

POINTS = [
    ("(1) It reframes the effect of generative AI on youth employment as a "
     "question of loop dominance rather than of technology. A closed-form net "
     "entry-port elasticity decomposes the marginal effect of AI capability "
     "into displacement, reinstatement and headcount-saving components and "
     "makes loop dominance directly observable. Over the resulting design "
     "space the 25-year effect on entry employment ranges from −70.0% to "
     "+19.2% and turns positive in 17.1% of configurations, so the sign is set "
     "by two organisational design parameters, not by the model weights."),

    ("(2) It identifies and quantifies a capability trap in workforce entry. "
     "Thinning the entry port depletes both the learning stock and the "
     "promotion flow, and because the sector's external market for experienced "
     "staff is itself fed by that promotion flow, the shortfall cannot be "
     "bought back. An automation-first strategy delivers the fastest unit-cost "
     "reduction of any scenario over the first three years yet ends the horizon "
     "with "
     "output 15.7% below the counterfactual in which the technology was never "
     "adopted — and the damage is invisible for the first five years."),

    ("(3) It provides a direct numerical test of joint optimisation, the "
     "founding claim of socio-technical design. In a pre-specified 2 × 2 "
     "factorial, work redesign (technical subsystem) and investment in "
     "human–AI collaboration routines (social subsystem) interact positively "
     "by 2.74 percentage points, while a four-lever decomposition is close to "
     "additive because instruments acting on the same bottleneck substitute "
     "for one another. Complementarity in socio-technical design is a property "
     "of subsystems, not of instruments."),

    ("(4) It is fully transparent and reproducible. The model's diffusion and "
     "augmentation parameters are calibrated to four independent published "
     "benchmarks, and the model passes the standard battery of structure, "
     "equilibrium, extreme-condition, integration-error and "
     "behaviour-reproduction tests, with a Latin hypercube sensitivity "
     "analysis over fifteen parameters. Every number in every table and figure "
     "is produced by the released code."),
]

CLOSING = (
    "We believe the manuscript speaks directly to the aims of the Special "
    "Issue: it treats generative AI as a socio-technical system rather than a "
    "tool, models the feedback loops, delays and emergent consequences that "
    "connect workforce development to organisational learning and economic "
    "capability, and turns the resulting insight into actionable design and "
    "policy levers. We deeply appreciate your consideration of our manuscript "
    "and look forward to receiving comments from the reviewers. Correspondence "
    "should be directed to the address below.")

SIGN = [
    "Name: Xinyu Cai",
    "Degree: Ph.D.",
    "Institution and address: College of Business, Jiaxing University; "
    "Jiaxing 314001, China",
    "Email: caixinyu@zjxu.edu.cn",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.3)
    sec.left_margin = sec.right_margin = Cm(3.0)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def p(text, after=5, bold=False, italic=False, align=None, indent=0.0,
          size=SIZE):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        run = par.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        return par

    p("Dear Editors:", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    for t in BODY:
        p(t, after=5)
    for t in POINTS:
        p(t, after=4, indent=0.0)
    p(CLOSING, after=7)

    p("Thanks very much for your attention to our paper.", after=8,
      align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in SIGN:
        p(line, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("Very sincerely yours,", after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p("Xinyu Cai", after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
