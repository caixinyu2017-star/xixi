# -*- coding: utf-8 -*-
"""Assemble the manuscript as a Word document based on the official MDPI
*Systems* template.

The template supplies the page geometry, the continuous left-hand line numbers
and every named MDPI style, so the output matches the journal layout exactly.
Equations are native Word (OMML) objects, centred, with right-aligned numbers in
the two-column layout used by the template.  Citations are inserted as Word REF
fields pointing at bookmarks on the reference list, so they remain live
cross-references.
"""
from __future__ import annotations

import os
import re
import shutil
import zipfile

from docx import Document
from docx.shared import Cm, Pt, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import omml
import l2omml
import content as C
import refs as R
import tables_spec as TS

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
FIGDIR = os.path.join(ROOT, "figures")
TEMPLATE = os.path.join(HERE, "systemstemplate.dot")
OUT = os.path.join(ROOT, "Vanishing_First_Rung_Systems_manuscript.docx")

BODY_INDENT = 2608          # twips, MDPI body zone
BODY_W = 7859               # twips, width of the body zone
FULL_W = 10467              # twips, full text width
NUM_W = 431                 # twips, equation-number column


# ---------------------------------------------------------------------------
# template handling
# ---------------------------------------------------------------------------
DROP = {"word/customizations.xml"}


def template_to_docx(dot_path, docx_path):
    """Rewrite the .dot package as a .docx.

    The document part changes content type and the Word-only key-map
    customisation part (which is not valid inside a .docx) is removed together
    with its relationship and content-type override.
    """
    with zipfile.ZipFile(dot_path) as zin, \
            zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in DROP:
                continue
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml")
                data = re.sub(
                    br'<Override PartName="/word/customizations\.xml"[^/]*/>',
                    b"", data)
            elif item.filename == "word/_rels/document.xml.rels":
                data = re.sub(
                    br'<Relationship[^>]*keyMapCustomizations[^>]*/>', b"", data)
            elif item.filename == "word/settings.xml":
                data = re.sub(br'<w:attachedTemplate[^/]*/>', b"", data)
            zout.writestr(item, data)


def clear_body(doc):
    """Empty the template body, keeping the section properties.

    The MDPI template marks the section as right-to-left (<w:bidi/>), which
    mirrors table column order in strict renderers and would place equation
    numbers on the wrong side; it is removed here.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            for sub in list(child):
                if sub.tag == qn("w:bidi"):
                    child.remove(sub)
            continue
        body.remove(child)


# ---------------------------------------------------------------------------
# low-level helpers
# ---------------------------------------------------------------------------
_BID = [1000]


def _next_bid():
    _BID[0] += 1
    return _BID[0]


def set_ind(p, left=None, first=None, right=None):
    """Indentation via the python-docx API so that pPr child order stays valid."""
    pf = p.paragraph_format
    if left is not None:
        pf.left_indent = Twips(left)
    if first is not None:
        pf.first_line_indent = Twips(first)
    if right is not None:
        pf.right_indent = Twips(right)


def add_bookmark(paragraph, name, text, size=9, bold=False):
    bid = _next_bid()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.append(end)
    return run


def add_ref_field(paragraph, bookmark, shown, size=None):
    """Insert { REF <bookmark> \\h } with a cached result."""
    def _r():
        r = OxmlElement("w:r")
        if size is not None:
            rpr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(size * 2)))
            rpr.append(sz)
            r.append(rpr)
        return r

    r1 = _r()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r1.append(fc)
    r2 = _r()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " REF %s \\h " % bookmark
    r2.append(it)
    r3 = _r()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "separate")
    r3.append(fc)
    r4 = _r()
    t = OxmlElement("w:t")
    t.text = shown
    r4.append(t)
    r5 = _r()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "end")
    r5.append(fc)
    for r in (r1, r2, r3, r4, r5):
        paragraph._p.append(r)


# ---------------------------------------------------------------------------
# citation bookkeeping
# ---------------------------------------------------------------------------
CIT_ORDER = []          # keys in order of first appearance
CIT_NUM = {}


def cite_number(key):
    if key not in CIT_NUM:
        CIT_ORDER.append(key)
        CIT_NUM[key] = len(CIT_ORDER)
    return CIT_NUM[key]


SUBSUP = re.compile(r"([_^]\{[^}]*\})")
TOKEN = re.compile(r"(\$[^$]*\$|(?:\[\[[a-zA-Z0-9_\-]+\]\](?:,)?)+)")
CITE = re.compile(r"\[\[([a-zA-Z0-9_\-]+)\]\]")


def _fmt_group(nums):
    """MDPI citation group: collapse runs of three or more into a range."""
    nums = sorted(set(nums))
    out, i = [], 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j - i >= 2:
            out.append((nums[i], nums[j]))
        else:
            out.extend((n, n) for n in nums[i:j + 1])
        i = j + 1
    return out


def add_rich(paragraph, text, size=None, italic=False, bold=False):
    """Add text with inline math ($...$) and citations ([[key]])."""
    for seg in TOKEN.split(text):
        if not seg:
            continue
        if seg.startswith("$") and seg.endswith("$") and len(seg) > 1:
            omml.add_inline_math(paragraph, l2omml.L(seg[1:-1]))
        elif seg.startswith("[["):
            keys = CITE.findall(seg)
            nums = [cite_number(k) for k in keys]
            groups = _fmt_group(nums)
            r = paragraph.add_run("[")
            if size:
                r.font.size = Pt(size)
            for gi, (a, b) in enumerate(groups):
                if gi:
                    r = paragraph.add_run(",")
                    if size:
                        r.font.size = Pt(size)
                add_ref_field(paragraph, "_Ref_%d" % a, str(a), size)
                if b != a:
                    r = paragraph.add_run("–")
                    if size:
                        r.font.size = Pt(size)
                    add_ref_field(paragraph, "_Ref_%d" % b, str(b), size)
            r = paragraph.add_run("]")
            if size:
                r.font.size = Pt(size)
        else:
            r = paragraph.add_run(seg)
            r.font.italic = italic
            r.font.bold = bold
            if size:
                r.font.size = Pt(size)


# ---------------------------------------------------------------------------
# builders
# ---------------------------------------------------------------------------
def para(doc, style, text, size=None, first_line=None):
    p = doc.add_paragraph(style=style)
    if first_line is not None:
        set_ind(p, first=first_line)
    add_rich(p, text, size=size)
    return p


TBLPR_ORDER = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
               "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
               "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
               "tblCellMar", "tblLook", "tblCaption", "tblDescription"]


def configure_table(tbl, twips, indent=None, center=False, borders=None,
                    cell_left=0, cell_right=0):
    """Rebuild w:tblPr with children in the order required by the schema.

    borders: mapping edge -> "single"/"none"; omitted edges default to none.
    """
    tblPr = tbl._tbl.tblPr
    keep = {}
    for child in list(tblPr):
        tag = child.tag.split("}")[1]
        if tag in ("tblStyle", "tblLook"):
            keep[tag] = child
        tblPr.remove(child)

    parts = {}
    if "tblStyle" in keep:
        parts["tblStyle"] = keep["tblStyle"]

    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(twips))
    w.set(qn("w:type"), "dxa")
    parts["tblW"] = w

    if center:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        parts["jc"] = jc
    if indent is not None:
        ind = OxmlElement("w:tblInd")
        ind.set(qn("w:w"), str(indent))
        ind.set(qn("w:type"), "dxa")
        parts["tblInd"] = ind

    b = OxmlElement("w:tblBorders")
    borders = borders or {}
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        val = borders.get(edge, "none")
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), "8" if val == "single" else "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        b.append(e)
    parts["tblBorders"] = b

    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    parts["tblLayout"] = lay

    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 0), ("left", cell_left), ("bottom", 0),
                      ("right", cell_right)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    parts["tblCellMar"] = mar

    if "tblLook" in keep:
        parts["tblLook"] = keep["tblLook"]

    for tag in TBLPR_ORDER:
        if tag in parts:
            tblPr.append(parts[tag])


def add_equation(doc, latex, number):
    wide = len(latex) > 74
    total = FULL_W if wide else BODY_W
    eq_w = total - NUM_W
    tbl = doc.add_table(rows=1, cols=2)
    configure_table(tbl, total, indent=None if wide else BODY_INDENT,
                    center=wide)
    tbl.columns[0].width = Emu(int(eq_w * 635))
    tbl.columns[1].width = Emu(int(NUM_W * 635))

    c0, c1 = tbl.rows[0].cells
    c0.width = Emu(int(eq_w * 635))
    c1.width = Emu(int(NUM_W * 635))

    p = c0.paragraphs[0]
    p.style = doc.styles["MDPI_3.9_equation"]
    set_ind(p, left=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(omml.omath(l2omml.L(latex)))

    p = c1.paragraphs[0]
    p.style = doc.styles["MDPI_3.a_equation_number"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("(%d)" % number)
    return tbl


def add_figure(doc, key):
    spec = TS.FIGURES[key]
    p = doc.add_paragraph(style="MDPI_5.2_figure")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(os.path.join(FIGDIR, spec["file"]),
                            width=Cm(spec["width_cm"]))
    cap = doc.add_paragraph(style="MDPI_5.1_figure_caption")
    cap.paragraph_format.space_after = Pt(2)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run("Figure %d. " % spec["number"])
    r.font.bold = True
    r.font.size = Pt(9)
    add_rich(cap, spec["caption"], size=9)
    if spec.get("note"):
        n = doc.add_paragraph(style="MDPI_5.1_figure_caption")
        set_ind(n, left=BODY_INDENT, first=0)
        n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        n.paragraph_format.space_before = Pt(0)
        n.paragraph_format.space_after = Pt(10)
        r = n.add_run("Note: ")
        r.font.italic = True
        r.font.size = Pt(8)
        add_rich(n, spec["note"], size=8)
    return p


def _shade_header(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    e = OxmlElement("w:bottom")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "8")
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), "auto")
    b.append(e)
    tcPr.append(b)


def add_table(doc, key):
    spec = TS.TABLES[key]()
    ncol = len(spec["header"])
    total = FULL_W if spec["wide"] else BODY_W
    ws = spec["widths"]
    s = sum(ws)
    widths = [int(total * w / s) for w in ws]

    cap = doc.add_paragraph(style="MDPI_4.1_table_caption")
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run("Table %d. " % spec["number"])
    r.font.bold = True
    r.font.size = Pt(9)
    add_rich(cap, spec["caption"], size=9)

    body_rows = spec["rows"]
    tbl = doc.add_table(rows=1 + len(body_rows), cols=ncol)
    tbl.style = doc.styles["MDPI_4.1_three_line_table"]
    configure_table(tbl, total, indent=None if spec["wide"] else BODY_INDENT,
                    center=spec["wide"],
                    borders={"top": "single", "bottom": "single"},
                    cell_left=57, cell_right=57)

    fs = 8.0 if ncol > 8 else 9.0

    def rich_cell(p, text, bold, italic, size):
        """Render X_{sub} and X^{sup} with real Word sub/superscripts."""
        for part in SUBSUP.split(text):
            if not part:
                continue
            if part.startswith("_{") or part.startswith("^{"):
                run = p.add_run(part[2:-1])
                run.font.subscript = part[0] == "_"
                run.font.superscript = part[0] == "^"
            else:
                run = p.add_run(part)
            run.font.bold = bold
            run.font.italic = italic
            run.font.size = Pt(size)

    def fill(cell, text, w, bold=False, italic=False, align="c", size=fs):
        cell.width = Emu(int(w * 635))
        p = cell.paragraphs[0]
        p.style = doc.styles["MDPI_4.2_table_body"]
        p.alignment = {"l": WD_ALIGN_PARAGRAPH.LEFT,
                       "c": WD_ALIGN_PARAGRAPH.CENTER,
                       "r": WD_ALIGN_PARAGRAPH.RIGHT}[align]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(1.5)
        p.paragraph_format.space_after = Pt(1.5)
        if text:
            rich_cell(p, text, bold, italic, size)

    for j, w in enumerate(widths):
        tbl.columns[j].width = Emu(int(w * 635))

    hdr = tbl.rows[0]
    for j, h in enumerate(spec["header"]):
        fill(hdr.cells[j], h, widths[j], bold=True, align="c")
        _shade_header(hdr.cells[j])

    for i, item in enumerate(body_rows, start=1):
        row = tbl.rows[i]
        if item[0] == "sec":
            merged = row.cells[0]
            for j in range(1, ncol):
                merged = merged.merge(row.cells[j])
            merged.width = Emu(int(total * 635))
            p = merged.paragraphs[0]
            p.style = doc.styles["MDPI_4.2_table_body"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1.5)
            rich_cell(p, item[1], True, True, fs)
        else:
            for j, cellval in enumerate(item[1]):
                fill(row.cells[j], cellval, widths[j],
                     italic=(spec.get("italic_col") == j),
                     align=spec["align"][j])

    if spec.get("note"):
        n = doc.add_paragraph(style="MDPI_4.3_table_footer")
        set_ind(n, left=BODY_INDENT, first=0)
        n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        n.paragraph_format.space_before = Pt(2)
        n.paragraph_format.space_after = Pt(10)
        r = n.add_run("Note: ")
        r.font.italic = True
        r.font.size = Pt(8)
        add_rich(n, spec["note"], size=8)
    return tbl


# ---------------------------------------------------------------------------
def render_reference(p, key, number):
    ref = R.REFS[key]
    add_bookmark(p, "_Ref_%d" % number, str(number), size=8)
    r = p.add_run(". ")
    r.font.size = Pt(8)

    def txt(s, italic=False, bold=False):
        run = p.add_run(s)
        run.font.size = Pt(8)
        run.font.italic = italic
        run.font.bold = bold

    kind = ref["kind"]
    txt(ref["authors"] + " ")
    if kind in ("article", "preprint"):
        title = ref["title"].rstrip(".")
        txt(title + ". ")
        txt(ref["journal"] + " ", italic=True)
        txt(str(ref["year"]), bold=True)
        txt(", ")
        if ref.get("volume"):
            txt(str(ref["volume"]), italic=True)
            txt(", " + str(ref["pages"]) + ".")
        else:
            txt(str(ref["pages"]) + ".")
        if ref.get("doi"):
            txt(" [CrossRef]")
    elif kind == "book":
        txt(ref["title"], italic=True)
        txt((", " + ref["edition"] if ref.get("edition") else "") + "; ")
        txt("%s: %s, %s, %d." % (ref["publisher"], ref["city"],
                                 ref["country"], ref["year"]))
    elif kind == "report":
        txt(ref["title"] + "; ", italic=True)
        if ref.get("series"):
            txt(ref["series"] + "; ")
        txt("%s: %s, %s, %d." % (ref["publisher"], ref["city"],
                                 ref["country"], ref["year"]))


# ---------------------------------------------------------------------------
def build():
    tmp = os.path.join(HERE, "_template.docx")
    template_to_docx(TEMPLATE, tmp)
    doc = Document(tmp)
    clear_body(doc)

    # ---- front matter --------------------------------------------------
    p = doc.add_paragraph(style="MDPI_1.1_article_type")
    p.add_run("Article")

    p = doc.add_paragraph(style="MDPI_1.2_title")
    p.add_run(C.TITLE)

    p = doc.add_paragraph(style="MDPI_1.3_authornames")
    p.add_run("Firstname Lastname ")
    r = p.add_run("1,*")
    r.font.superscript = True

    for line, sid in (("Academic Editor: Firstname Lastname",
                       "MDPI_1.5_academic_editor"),
                      ("Received: date", "MDPI_1.4_history"),
                      ("Revised: date", "MDPI_1.4_history"),
                      ("Accepted: date", "MDPI_1.4_history"),
                      ("Published: date", "MDPI_1.4_history")):
        doc.add_paragraph(style=sid).add_run(line)

    doc.add_paragraph(style="MDPI_6.1_citation").add_run(
        "Citation: To be added by editorial staff during production.")
    doc.add_paragraph(style="MDPI_7.2_copyright").add_run(
        "Copyright: © 2026 by the author. Submitted for possible open access "
        "publication under the terms and conditions of the Creative Commons "
        "Attribution (CC BY) license (https://creativecommons.org/licenses/by/"
        "4.0/).")

    p = doc.add_paragraph(style="MDPI_1.6_affiliation")
    r = p.add_run("1")
    r.font.superscript = True
    p.add_run(" Affiliation; e-mail@e-mail.com")
    p = doc.add_paragraph(style="MDPI_1.6_affiliation")
    p.add_run("* Correspondence: e-mail@e-mail.com")

    p = doc.add_paragraph(style="MDPI_1.7_abstract")
    r = p.add_run("Abstract")
    r.font.bold = True
    p = doc.add_paragraph(style="MDPI_1.7_abstract")
    add_rich(p, C.ABSTRACT)

    p = doc.add_paragraph(style="MDPI_1.8_keywords")
    r = p.add_run("Keywords: ")
    r.font.bold = True
    p.add_run(C.KEYWORDS)

    doc.add_paragraph(style="MDPI_1.9_line")

    # ---- body ----------------------------------------------------------
    prev = None
    for block in C.BLOCKS:
        kind = block[0]
        if kind == "h1":
            doc.add_paragraph(style="MDPI_2.1_heading1").add_run(block[1])
        elif kind == "h1b":
            p = doc.add_paragraph(style="MDPI_6.2_back_matter")
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(block[1] + ": ")
            r.font.bold = True
            r.font.italic = True
            prev = p
        elif kind == "h2":
            doc.add_paragraph(style="MDPI_2.2_heading2").add_run(block[1])
        elif kind == "h3":
            doc.add_paragraph(style="MDPI_2.3_heading3").add_run(block[1])
        elif kind == "p":
            style = ("MDPI_3.2_text_no_indent"
                     if prev in ("eq", "h1", "h2", "h3") else "MDPI_3.1_text")
            para(doc, style, block[1])
        elif kind == "prop":
            p = doc.add_paragraph(style="MDPI_3.2_text_no_indent")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_ind(p, left=BODY_INDENT + 340, first=0)
            r = p.add_run(block[1] + " ")
            r.font.bold = True
            add_rich(p, block[2], italic=True)
        elif kind == "eq":
            add_equation(doc, block[1], block[2])
        elif kind == "fig":
            add_figure(doc, block[1])
        elif kind == "table":
            add_table(doc, block[1])
        elif kind == "stmt":
            if prev == "h1b":
                add_rich(doc.paragraphs[-1], block[2], size=9)
            else:
                p = doc.add_paragraph(style="MDPI_6.2_back_matter")
                add_rich(p, block[2], size=9)
        prev = kind

    # ---- references ----------------------------------------------------
    doc.add_paragraph(style="MDPI_2.1_heading1").add_run("References")
    for n, key in enumerate(CIT_ORDER, start=1):
        p = doc.add_paragraph(style="MDPI_3.2_text_no_indent")
        set_ind(p, left=BODY_INDENT + 340, first=-340)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(2)
        render_reference(p, key, n)

    doc.save(OUT)
    os.remove(tmp)

    unused = sorted(set(R.REFS) - set(CIT_ORDER))
    print("saved:", OUT)
    print("references cited:", len(CIT_ORDER), "/", len(R.REFS))
    if unused:
        print("UNCITED:", ", ".join(unused))
    missing = [k for k in CIT_ORDER if k not in R.REFS]
    if missing:
        print("MISSING FROM refs.py:", missing)


if __name__ == "__main__":
    build()
