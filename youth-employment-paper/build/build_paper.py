# -*- coding: utf-8 -*-
"""Build the youth-employment + AI manuscript as a Word .docx with native OMML equations,
right-aligned equation numbers, MDPI-style tables and references."""
import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import omml
import l2omml
import content as C
import tables as T
import refs as R  # provides REFERENCES (list) and CIT (dict key->list of ref numbers)

BODY_FONT = "Palatino Linotype"
BODY_SIZE = 9.5
TEXT_W_TWIPS = 9200      # ~ right margin position for equation number (A4, 2.0cm margins)
CENTER_TWIPS = 4600

doc = Document()

# ---- page setup: A4, MDPI-ish margins ----
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(2.0)

# ---- base style ----
style = doc.styles['Normal']
style.font.name = BODY_FONT
style.font.size = Pt(BODY_SIZE)
style._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_font(run, size=BODY_SIZE, bold=False, italic=False, color=None, name=BODY_FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), name)


def fmt_citation(nums):
    """Format a list of reference numbers as MDPI bracketed citation, collapsing runs."""
    nums = sorted(set(nums))
    parts = []
    i = 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j - i >= 1:
            parts.append(f"{nums[i]}–{nums[j]}")  # en dash range
        else:
            parts.append(str(nums[i]))
        i = j + 1
    return "[" + ",".join(parts) + "]"


TOKEN_RE = re.compile(r'(\$[^$]*\$|\[\[[a-z0-9\-]+\]\])')


def add_rich_runs(paragraph, text, size=BODY_SIZE, bold=False, italic=False):
    """Add text to paragraph, interpreting $latex$ math and [[key]] citations."""
    for seg in TOKEN_RE.split(text):
        if not seg:
            continue
        if seg.startswith('$') and seg.endswith('$'):
            omml.add_inline_math(paragraph, l2omml.L(seg[1:-1]))
        elif seg.startswith('[[') and seg.endswith(']]'):
            key = seg[2:-2]
            nums = R.CIT.get(key)
            if not nums:
                run = paragraph.add_run('[?]')
                set_font(run, size, color=RGBColor(0xC0, 0x00, 0x00))
            else:
                run = paragraph.add_run(fmt_citation(nums))
                set_font(run, size, bold=bold, italic=italic)
        else:
            run = paragraph.add_run(seg)
            set_font(run, size, bold=bold, italic=italic)


def add_paragraph(text, size=BODY_SIZE, space_after=6, first_indent=0.0, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    add_rich_runs(p, text, size=size)
    return p


def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_font(run, 11, bold=True)
    elif level == 2:
        set_font(run, 10, bold=True)
    else:
        set_font(run, 9.5, bold=True, italic=True)
    return p


# =====================================================================================
# FRONT MATTER
# =====================================================================================
def build_front_matter():
    # article type
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Article"); set_font(run, 10, bold=True, italic=True)

    # title
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(C.TITLE); set_font(run, 15, bold=True)

    # authors
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Author Name "); set_font(run, 10, bold=True)
    run = p.add_run("1,*"); set_font(run, 7); run.font.superscript = True

    # affiliations
    aff = ("1  School of Economics / Labor Economics, University, City, Country; "
           "author@university.edu")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(aff); set_font(run, 8)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    run = p.add_run("*  Correspondence: author@university.edu"); set_font(run, 8)

    # Abstract
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Abstract: "); set_font(run, 9, bold=True)
    add_rich_runs(p, C.ABSTRACT, size=9)

    # Keywords
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Keywords: "); set_font(run, 9, bold=True)
    run = p.add_run(C.KEYWORDS); set_font(run, 9, italic=False)


# =====================================================================================
# BODY BLOCKS
# =====================================================================================
def build_body():
    for block in C.BLOCKS:
        kind = block[0]
        if kind == 'h1':
            add_heading(block[1], 1)
        elif kind == 'h2':
            add_heading(block[1], 2)
        elif kind == 'h3':
            add_heading(block[1], 3)
        elif kind == 'p':
            add_paragraph(block[1])
        elif kind == 'hyp':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run(block[1] + " "); set_font(run, 9.5, bold=True)
            add_rich_runs(p, block[2], size=9.5, italic=True)
        elif kind == 'eq':
            latex, num = block[1], block[2]
            p = omml.add_numbered_equation(doc, l2omml.L(latex), num,
                                           right_twips=TEXT_W_TWIPS, center_twips=CENTER_TWIPS)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                set_font(r, BODY_SIZE)
        elif kind == 'fig':
            T.render_figure(doc, block[1], set_font)
        elif kind == 'stmt':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(3)
            run = p.add_run(block[1] + " "); set_font(run, 8.5, bold=True, italic=True)
            add_rich_runs(p, block[2], size=8.5)
        elif kind == 'table':
            T.render_table(doc, block[1], set_font, add_rich_runs)


def build_references():
    add_heading("References", 1)
    for i, ref in enumerate(R.REFERENCES, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.55)
        run = p.add_run(f"{i}. "); set_font(run, 8)
        add_ref_runs(p, ref)


def add_ref_runs(paragraph, ref):
    """Render one MDPI reference: authors + title + italic journal + year;vol;pages + doi."""
    # authors
    run = paragraph.add_run(ref['authors'] + " "); set_font(run, 8)
    # title
    title = ref['title'].strip()
    if not title.endswith('.'):
        title += '.'
    run = paragraph.add_run(title + " "); set_font(run, 8)
    # journal (italic)
    run = paragraph.add_run(ref['journal_abbrev'] + " "); set_font(run, 8, italic=True)
    # year bold, vol italic, pages
    yr = str(ref.get('year', ''))
    run = paragraph.add_run(yr + ", "); set_font(run, 8, bold=True)
    vol = str(ref.get('volume', '') or '')
    if vol:
        run = paragraph.add_run(vol); set_font(run, 8, italic=True)
        pages = str(ref.get('pages', '') or '')
        if pages:
            run = paragraph.add_run(", " + pages); set_font(run, 8)
        run = paragraph.add_run("."); set_font(run, 8)
    else:
        pages = str(ref.get('pages', '') or '')
        if pages:
            run = paragraph.add_run(pages + "."); set_font(run, 8)
    doi = ref.get('doi', '')
    if doi:
        run = paragraph.add_run(" [CrossRef]"); set_font(run, 8)


build_front_matter()
build_body()
build_references()
doc.save("youth_ai_employment.docx")
print("Saved youth_ai_employment.docx with", len(C.BLOCKS), "blocks and", len(R.REFERENCES), "references")
