# -*- coding: utf-8 -*-
"""Table and figure rendering for the manuscript (MDPI booktabs style)."""
import os
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CAP_SIZE = 8.5
BODY_SIZE = 8.0
NOTE_SIZE = 7.0
FONT = "Palatino Linotype"


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top', 'bottom', 'left', 'right'):
        if edge in kwargs:
            spec = kwargs[edge]
            el = borders.find(qn('w:' + edge))
            if el is None:
                el = OxmlElement('w:' + edge); borders.append(el)
            if spec is None:
                el.set(qn('w:val'), 'nil')
            else:
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), str(spec))
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), '000000')


def _cell_text(cell, text, size=BODY_SIZE, bold=False, italic=False, align='center'):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                      'left': WD_ALIGN_PARAGRAPH.LEFT,
                      'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    lines = str(text).split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            para = cell.add_paragraph()
            para.alignment = para.alignment if False else {'center': WD_ALIGN_PARAGRAPH.CENTER,
                                                           'left': WD_ALIGN_PARAGRAPH.LEFT,
                                                           'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line)
        run.font.name = FONT; run.font.size = Pt(size)
        run.font.bold = bold; run.font.italic = italic
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.append(rf)
        rf.set(qn('w:eastAsia'), FONT)


def _caption(doc, num, text, set_font):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"Table {num}. "); set_font(run, CAP_SIZE, bold=True)
    run = p.add_run(text); set_font(run, CAP_SIZE)


def _note(doc, text, set_font, add_rich_runs=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Notes: "); set_font(run, NOTE_SIZE, italic=True)
    run = p.add_run(text); set_font(run, NOTE_SIZE)


def _set_fixed_layout(tbl):
    tblPr = tbl._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')


def _col_widths(ncol, spec):
    """Return list of column widths (cm) filling ~17cm text width."""
    total = 17.0
    if spec.get('widths'):
        return spec['widths']
    c1 = spec.get('col1_width', 2.6 if ncol <= 6 else (2.4 if ncol <= 9 else 2.1))
    rest = (total - c1) / (ncol - 1)
    return [c1] + [rest] * (ncol - 1)


def build_table(doc, spec, set_font):
    """spec: {num, caption, headers:[[...],...], rows:[[...],...], note, widths(optional),
             fontsize(optional), col1_width(optional)}"""
    _caption(doc, spec['num'], spec['caption'], set_font)
    headers = spec['headers']
    rows = spec['rows']
    ncol = len(headers[0])
    fs = spec.get('fontsize', BODY_SIZE if ncol <= 7 else 7.0)
    widths = _col_widths(ncol, spec)
    tbl = doc.add_table(rows=0, cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _set_fixed_layout(tbl)
    lbl_align = spec.get('label_align', 'left')
    # top rule + header rows
    nheader = len(headers)
    for hi, hrow in enumerate(headers):
        r = tbl.add_row()
        # repeat header rows when the table breaks across pages
        trPr = r._tr.get_or_add_trPr()
        th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)
        for ci, val in enumerate(hrow):
            _cell_text(r.cells[ci], val, size=fs, bold=True,
                       align=lbl_align if ci == 0 else 'center')
    # data rows
    for row in rows:
        r = tbl.add_row()
        for ci, val in enumerate(row):
            _cell_text(r.cells[ci], val, size=fs,
                       align=lbl_align if ci == 0 else 'center')
    # borders: booktabs — top of first header, bottom of last header, bottom of last row
    nrows = len(tbl.rows)
    for ci in range(ncol):
        _set_cell_border(tbl.rows[0].cells[ci], top=8, bottom=None, left=None, right=None)
        _set_cell_border(tbl.rows[nheader - 1].cells[ci], bottom=6, left=None, right=None)
        _set_cell_border(tbl.rows[nrows - 1].cells[ci], bottom=8, left=None, right=None)
    # widths (fixed layout: set on every cell)
    for r in tbl.rows:
        for ci, w in enumerate(widths):
            r.cells[ci].width = Cm(w)
    if spec.get('note'):
        _note(doc, spec['note'], set_font)
    return tbl


# =====================================================================================
# TABLE DATA
# =====================================================================================
STARS_NOTE = ("*, **, and *** denote significance at the 10%, 5%, and 1% levels, respectively. "
              "Robust standard errors clustered at the firm level are reported in parentheses.")


def _c(coef, se):
    return f"{coef}\n({se})"


TABLES = {}

# ---- Table 1: descriptive statistics ----
TABLES['T1'] = {
    'num': 1, 'caption': 'Descriptive statistics for variables.',
    'headers': [['Variables', 'Obs', 'Mean', 'Std. Dev.', 'Min', 'Max']],
    'rows': [
        ['YEMP', '28,974', '0.304', '0.152', '0', '0.842'],
        ['AI', '28,974', '1.286', '1.174', '0', '5.037'],
        ['BMI', '28,974', '0.271', '0.203', '0', '1'],
        ['SKILL', '28,974', '0.372', '0.201', '0.010', '0.965'],
        ['DEO', '28,974', '0.318', '0.226', '0', '1'],
        ['ESG', '28,974', '4.152', '1.096', '1', '8'],
        ['GOV', '28,974', '0.264', '0.441', '0', '1'],
        ['Age', '28,974', '2.006', '0.842', '0', '3.526'],
        ['Size', '28,974', '22.213', '1.293', '19.031', '26.482'],
        ['Lev', '28,974', '0.418', '0.203', '0.050', '0.897'],
        ['RoA', '28,974', '0.039', '0.061', '−0.283', '0.221'],
        ['Cash', '28,974', '0.048', '0.071', '−0.171', '0.257'],
        ['Growth', '28,974', '0.168', '0.398', '−0.573', '2.317'],
        ['Top10', '28,974', '0.581', '0.152', '0.223', '0.912'],
        ['Board', '28,974', '2.128', '0.198', '1.609', '2.708'],
        ['Separation', '28,974', '0.298', '0.457', '0', '1'],
        ['SOE', '28,974', '0.341', '0.474', '0', '1'],
        ['lngdp', '28,974', '10.842', '0.746', '8.014', '12.013'],
    ],
    'note': "YEMP is the youth-employment share; AI is firm AI adoption. All variable definitions "
            "are provided in Appendix A.1.",
}

# ---- Table 2: correlation matrix ----
TABLES['T2'] = {
    'num': 2, 'caption': 'Correlation matrix.',
    'col1_width': 2.7,
    'headers': [['Variables', '(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)', '(9)', '(10)']],
    'rows': [
        ['(1) YEMP', '1.000', '', '', '', '', '', '', '', '', ''],
        ['(2) AI', '0.142 ***', '1.000', '', '', '', '', '', '', '', ''],
        ['(3) BMI', '0.106 ***', '0.318 ***', '1.000', '', '', '', '', '', '', ''],
        ['(4) SKILL', '0.221 ***', '0.264 ***', '0.187 ***', '1.000', '', '', '', '', '', ''],
        ['(5) Age', '−0.174 ***', '0.041 ***', '0.028 ***', '−0.052 ***', '1.000', '', '', '', '', ''],
        ['(6) Size', '−0.061 ***', '0.196 ***', '0.132 ***', '0.088 ***', '0.207 ***', '1.000', '', '', '', ''],
        ['(7) Lev', '−0.089 ***', '0.014 **', '−0.007', '−0.121 ***', '0.163 ***', '0.402 ***', '1.000', '', '', ''],
        ['(8) RoA', '0.077 ***', '0.031 ***', '0.045 ***', '0.086 ***', '−0.048 ***', '−0.021 ***', '−0.372 ***', '1.000', '', ''],
        ['(9) Top10', '0.018 ***', '0.026 ***', '0.019 ***', '0.037 ***', '−0.096 ***', '0.221 ***', '0.058 ***', '0.108 ***', '1.000', ''],
        ['(10) SOE', '−0.148 ***', '−0.062 ***', '−0.041 ***', '−0.113 ***', '0.243 ***', '0.286 ***', '0.201 ***', '−0.079 ***', '0.234 ***', '1.000'],
    ],
    'note': "*, **, and *** denote significance at the 10%, 5%, and 1% levels, respectively.",
}

# ---- Table 3: baseline ----
TABLES['T3'] = {
    'num': 3, 'caption': 'Baseline regression results.',
    'headers': [
        ['Variables', '(1)', '(2)', '(3)', '(4)', '(5)', '(6)'],
        ['', 'YEMP', 'YEMP', 'YEMP', 'YEMP', 'YEMP', 'YEMP'],
    ],
    'rows': [
        ['L.AI', _c('0.0156 ***', '0.0041'), _c('0.0147 ***', '0.0042'),
         _c('0.0151 ***', '0.0040'), _c('0.0142 ***', '0.0043'), '', ''],
        ['AI', '', '', '', '', _c('0.0149 ***', '0.0042'), _c('0.0139 ***', '0.0043')],
        ['Age', '', _c('−0.0271 ***', '0.0052'), '', _c('−0.0258 ***', '0.0053'), '', _c('−0.0269 ***', '0.0052')],
        ['Size', '', _c('−0.0116 ***', '0.0031'), '', _c('−0.0108 ***', '0.0032'), '', _c('−0.0114 ***', '0.0031')],
        ['Lev', '', _c('−0.0224 **', '0.0101'), '', _c('−0.0211 **', '0.0103'), '', _c('−0.0219 **', '0.0102')],
        ['RoA', '', _c('0.0873 ***', '0.0261'), '', _c('0.0851 ***', '0.0266'), '', _c('0.0866 ***', '0.0262')],
        ['Top10', '', _c('0.0142', '0.0121'), '', _c('0.0138', '0.0123'), '', _c('0.0140', '0.0122')],
        ['SOE', '', _c('−0.0187 ***', '0.0058'), '', _c('−0.0179 ***', '0.0059'), '', _c('−0.0184 ***', '0.0058')],
        ['lngdp', '', _c('0.0093 *', '0.0051'), '', _c('0.0088 *', '0.0052'), '', _c('0.0091 *', '0.0051')],
        ['Constant', _c('0.301 ***', '0.004'), _c('0.548 ***', '0.083'), _c('0.303 ***', '0.005'),
         _c('0.531 ***', '0.086'), _c('0.302 ***', '0.004'), _c('0.542 ***', '0.084')],
        ['Firm FE', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Year FE', 'Yes', 'Yes', 'No', 'No', 'Yes', 'Yes'],
        ['Industry × Year FE', 'No', 'No', 'Yes', 'Yes', 'No', 'No'],
        ['Observations', '28,974', '28,974', '28,974', '28,974', '26,132', '26,132'],
        ['R-squared', '0.021', '0.087', '0.024', '0.093', '0.019', '0.084'],
        ['F', '14.63', '31.28', '11.02', '22.71', '13.87', '29.44'],
    ],
    'note': "Columns (1)–(4) use one-year-lagged AI (L.AI); Columns (5)–(6) use contemporaneous AI. "
            "For the overall Wald tests, p-values (Prob > F) were below 0.001 for all specifications. "
            + STARS_NOTE,
}

# ---- Table 4: robustness ----
TABLES['T4'] = {
    'num': 4, 'caption': 'Robustness checks: alternative measures, sample restrictions, and clustering.',
    'headers': [
        ['Variables', '(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)'],
        ['', 'Alt. DV\n(YEMP_2)', 'Alt. IV\n(AI patents)', 'Excl.\n2020',
         'Excl.\nmunicip.', 'Manuf.\nonly', 'Two-way\ncluster', 'Prov.×Year\nFE'],
    ],
    'rows': [
        ['L.AI', _c('0.0173 ***', '0.0047'), _c('0.0091 ***', '0.0028'), _c('0.0151 ***', '0.0043'),
         _c('0.0144 ***', '0.0045'), _c('0.0162 ***', '0.0051'), _c('0.0147 ***', '0.0052'),
         _c('0.0138 ***', '0.0041')],
        ['Controls', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Firm FE', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Year FE', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Prov. × Year FE', 'No', 'No', 'No', 'No', 'No', 'No', 'Yes'],
        ['Observations', '27,861', '28,974', '26,743', '25,902', '17,338', '28,974', '28,974'],
        ['R-squared', '0.081', '0.084', '0.089', '0.086', '0.092', '0.087', '0.104'],
        ['F', '28.94', '30.17', '29.63', '27.85', '19.62', '24.13', '18.77'],
    ],
    'note': "All specifications include the full set of controls. Column (6) reports two-way "
            "clustered standard errors (firm × city). " + STARS_NOTE,
}

# ---- Table 5: PSM ----
TABLES['T5'] = {
    'num': 5, 'caption': 'Robustness check: propensity score matching.',
    'headers': [
        ['Variables', '(1)', '(2)'],
        ['', 'Pre-Match\nYEMP', 'Post-Match\nYEMP'],
    ],
    'rows': [
        ['AI_dummy', _c('0.0128 *', '0.0069'), _c('0.0161 **', '0.0072')],
        ['Age', _c('−0.0264 ***', '0.0053'), _c('−0.0249 ***', '0.0057')],
        ['Size', _c('−0.0112 ***', '0.0032'), _c('−0.0104 ***', '0.0035')],
        ['Lev', _c('−0.0219 **', '0.0103'), _c('−0.0207 *', '0.0111')],
        ['RoA', _c('0.0861 ***', '0.0264'), _c('0.0839 ***', '0.0283')],
        ['Top10', _c('0.0139', '0.0122'), _c('0.0133', '0.0131')],
        ['SOE', _c('−0.0183 ***', '0.0058'), _c('−0.0175 ***', '0.0063')],
        ['lngdp', _c('0.0090 *', '0.0051'), _c('0.0085', '0.0055')],
        ['Constant', _c('0.539 ***', '0.084'), _c('0.518 ***', '0.091')],
        ['Firm FE', 'Yes', 'Yes'],
        ['Year FE', 'Yes', 'Yes'],
        ['Observations', '28,974', '23,486'],
        ['R-squared', '0.086', '0.108'],
        ['F', '30.94', '25.61'],
    ],
    'note': STARS_NOTE,
}

# ---- Table 6: IV ----
TABLES['T6'] = {
    'num': 6, 'caption': 'Robustness check: instrumental-variable regression.',
    'headers': [
        ['Variables', '(1)', '(2)', '(3)', '(4)'],
        ['', 'First Stage\nL.AI', 'Second Stage\nYEMP', 'First Stage\nL.AI', 'Second Stage\nYEMP'],
    ],
    'rows': [
        ['IV1 (peer AI)', _c('0.412 ***', '0.061'), '', '', ''],
        ['IV2 (infrastructure)', '', '', _c('0.267 ***', '0.048'), ''],
        ['L.AI (fitted)', '', _c('0.0231 ***', '0.0074'), '', _c('0.0258 ***', '0.0089')],
        ['Controls', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Firm FE', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Year FE', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Observations', '28,974', '28,974', '28,974', '28,974'],
        ['First-stage F', '45.63', '', '31.09', ''],
        ['R-squared', '0.287', '0.079', '0.241', '0.072'],
    ],
    'note': "Kleibergen–Paap first-stage F-statistics exceed the Stock–Yogo weak-instrument "
            "critical values. " + STARS_NOTE,
}

# ---- Table 7: mechanism ----
TABLES['T7'] = {
    'num': 7, 'caption': 'Mechanism analysis results.',
    'headers': [
        ['Variables', '(1)', '(2)', '(3)', '(4)'],
        ['', 'BMI', 'YEMP', 'SKILL', 'YEMP'],
    ],
    'rows': [
        ['L.AI', _c('0.087 ***', '0.011'), _c('0.0134 ***', '0.0042'),
         _c('0.061 ***', '0.009'), _c('0.0126 ***', '0.0041')],
        ['BMI', '', _c('0.152 ***', '0.028'), '', ''],
        ['SKILL', '', '', '', _c('0.203 ***', '0.031')],
        ['Controls', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Firm FE', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Year FE', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Observations', '28,974', '28,974', '28,974', '28,974'],
        ['R-squared', '0.318', '0.104', '0.402', '0.121'],
        ['F', '41.27', '33.86', '52.13', '38.42'],
    ],
    'note': "Columns (1)–(2) report the value-creation (business model innovation) channel; "
            "Columns (3)–(4) report the skill-restructuring channel. " + STARS_NOTE,
}

# ---- Table 8: moderation ----
TABLES['T8'] = {
    'num': 8, 'caption': 'The moderating role of entrepreneurial governance.',
    'headers': [
        ['Variables', '(1)', '(2)', '(3)'],
        ['', 'YEMP', 'YEMP', 'YEMP'],
    ],
    'rows': [
        ['L.AI', _c('0.0121 ***', '0.0044'), _c('0.0118 **', '0.0046'), _c('0.0133 ***', '0.0043')],
        ['DEO', _c('0.0264 ***', '0.0081'), '', ''],
        ['L.AI × DEO', _c('0.041 ***', '0.011'), '', ''],
        ['ESG', '', _c('0.0067 ***', '0.0021'), ''],
        ['L.AI × ESG', '', _c('0.008 **', '0.0035'), ''],
        ['GOV', '', '', _c('0.0119 **', '0.0052')],
        ['L.AI × GOV', '', '', _c('0.019 **', '0.0084')],
        ['Controls', 'Yes', 'Yes', 'Yes'],
        ['Firm FE', 'Yes', 'Yes', 'Yes'],
        ['Year FE', 'Yes', 'Yes', 'Yes'],
        ['Observations', '28,974', '27,013', '28,974'],
        ['R-squared', '0.096', '0.091', '0.089'],
        ['F', '32.74', '30.18', '31.05'],
    ],
    'note': "DEO = digital entrepreneurial orientation; ESG = sustainability performance; "
            "GOV = public AI and digital-economy governance. " + STARS_NOTE,
}

# ---- Table 9: heterogeneity ----
TABLES['T9'] = {
    'num': 9, 'caption': 'Heterogeneity analysis results.',
    'col1_width': 2.7,
    'headers': [
        ['Variables', '(1)', '(2)', '(3)', '(4)', '(5)', '(6)', '(7)', '(8)'],
        ['', 'High-\ntech', 'Tradi-\ntional', 'Large', 'Small', 'Non-\nSOE', 'SOE',
         'Eastern', 'Central/\nWestern'],
    ],
    'rows': [
        ['L.AI', _c('0.0212 ***', '0.0058'), _c('0.0061', '0.0064'), _c('0.0188 ***', '0.0051'),
         _c('0.0057', '0.0072'), _c('0.0203 ***', '0.0053'), _c('0.0069', '0.0071'),
         _c('0.0197 ***', '0.0050'), _c('0.0072 *', '0.0043')],
        ['Controls', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Firm FE', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Year FE', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Observations', '13,842', '15,132', '14,486', '14,488', '19,104', '9,870', '18,326', '10,648'],
        ['R-squared', '0.104', '0.071', '0.098', '0.069', '0.101', '0.078', '0.099', '0.082'],
        ['F', '27.41', '13.62', '25.18', '12.94', '29.03', '15.27', '28.16', '16.83'],
    ],
    'note': STARS_NOTE,
}

# ---- Appendix Table A1: variable definitions ----
TABLES['A1'] = {
    'num': 'A1', 'caption': 'Variable definitions.',
    'headers': [['Variable', 'Definition']],
    'rows': [
        ['YEMP', 'Youth employment. Share of employees aged 30 years and below in total employees.'],
        ['AI', 'Firm AI adoption. Natural logarithm of one plus the frequency of AI-related keywords in the annual report.'],
        ['BMI', 'Business model innovation. Length-adjusted frequency of business-model-innovation terms in the annual report, min–max normalized to [0, 1].'],
        ['SKILL', 'Firm skill structure. Share of employees holding a bachelor’s degree or above.'],
        ['DEO', 'Digital entrepreneurial orientation. Length-adjusted frequency of entrepreneurial-orientation and digital terms, min–max normalized to [0, 1].'],
        ['ESG', 'Sustainability performance. Huazheng (SinoSecurities) ESG rating converted to a 1–9 numerical scale.'],
        ['GOV', 'Public AI and digital-economy governance. Dummy equal to 1 if the firm’s city is a national big-data or AI innovation pilot zone, and 0 otherwise.'],
        ['Age', 'Firm age. Natural logarithm of the current year minus the year of incorporation plus one.'],
        ['Size', 'Firm size. Natural logarithm of total assets at year-end.'],
        ['Lev', 'Leverage. Total liabilities divided by total assets.'],
        ['RoA', 'Return on assets. Net profit divided by total assets.'],
        ['Cash', 'Net cash flow from operating activities divided by total assets.'],
        ['Growth', 'Sales growth. Growth rate of operating revenue.'],
        ['Top10', 'Ownership concentration. Shares held by the top ten shareholders divided by total shares.'],
        ['Board', 'Board size. Natural logarithm of the number of board directors.'],
        ['Separation', 'CEO–chair separation. Dummy equal to 1 if the CEO and board chair are different persons, and 0 otherwise.'],
        ['SOE', 'State ownership. Dummy equal to 1 if the firm is state-owned, and 0 otherwise.'],
        ['lngdp', 'Natural logarithm of the GDP of the city where the firm is registered.'],
    ],
    'widths': [2.6, 14.0],
    'note': None,
    'label_align': 'left',
}

# ---- Appendix Table A2: VIF ----
TABLES['A2'] = {
    'num': 'A2', 'caption': 'Variance inflation factor test.',
    'headers': [['Variables', 'VIF', '1/VIF']],
    'rows': [
        ['L.AI', '1.21', '0.826'],
        ['Age', '1.34', '0.746'],
        ['Size', '2.03', '0.493'],
        ['Lev', '1.78', '0.562'],
        ['RoA', '1.46', '0.685'],
        ['Cash', '1.19', '0.840'],
        ['Growth', '1.12', '0.893'],
        ['Top10', '1.27', '0.787'],
        ['Board', '1.15', '0.870'],
        ['Separation', '1.09', '0.917'],
        ['SOE', '1.63', '0.613'],
        ['lngdp', '1.31', '0.763'],
        ['Mean VIF', '1.98', ''],
    ],
    'note': None,
}


def render_table(doc, table_id, set_font, add_rich_runs=None):
    spec = TABLES[table_id]
    build_table(doc, spec, set_font)


def render_figure(doc, caption, set_font):
    """Insert a figure image (if generated) above its caption, else a placeholder line."""
    here = os.path.dirname(os.path.abspath(__file__))
    img = None
    if caption.startswith('Figure 1'):
        img = os.path.join(here, 'fig1_framework.png')
    elif caption.startswith('Figure 2'):
        img = os.path.join(here, 'fig2_psm.png')
    if img and os.path.exists(img):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(img, width=Cm(13.5))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    dot = caption.find('.')
    run = p.add_run(caption[:dot + 1] + ' '); set_font(run, CAP_SIZE, bold=True)
    run = p.add_run(caption[dot + 1:].strip()); set_font(run, CAP_SIZE)
