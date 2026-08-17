# -*- coding: utf-8 -*-
"""Assemble the manuscript as a Word document based on the official MDPI
*Systems* template.

The template supplies the page geometry, the continuous left-hand line numbers
and every named MDPI style, so the output matches the journal layout exactly.
Equations are native Word (OMML) objects, centred in the body zone with
right-aligned numbers.  Citations are inserted as Word REF fields pointing at
bookmarks on the reference list, so they remain live cross-references and
renumber automatically if the order of the text changes.

The layout follows the two published *Systems* articles used as the format
reference: body-zone text indented by 2608 twips, table and figure captions at
the body-zone left edge, table and figure notes at the same edge with no
first-line indent and single line spacing, back-matter labels in bold, and the
reference list set across the full text width with a 425-twip hanging indent.
"""
from __future__ import annotations

import os
import re
import zipfile

from docx import Document
from docx.shared import Cm, Pt, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import omml
import l2omml
import content as C
import refs as R
import tables_spec as TS

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
FIGDIR = os.path.join(ROOT, "figures")
TEMPLATE = os.path.join(HERE, "systemstemplate.dot")
OUT = os.path.join(ROOT, "AI_Readiness_Youth_Employment_EU_Systems_manuscript.docx")

BODY_INDENT = 2608          # twips, left edge of the MDPI body zone
BODY_W = 7858               # twips, width of the body zone
FULL_W = 10466              # twips, full text width
NUM_W = 431                 # twips, equation-number column
REF_HANG = 425              # twips, hanging indent of the reference list


# ---------------------------------------------------------------------------
# template handling
# ---------------------------------------------------------------------------
DROP = {"word/customizations.xml"}


def template_to_docx(dot_path, docx_path):
    """Rewrite the .dot package as a .docx.

    The document part changes content type and the Word-only key-map
    customisation part (which is not valid inside a .docx) is removed together
    with its relationship and content-type override.
    """
    with zipfile.ZipFile(dot_path) as zin, \
            zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in DROP:
                continue
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml")
                data = re.sub(
                    br'<Override PartName="/word/customizations\.xml"[^>]*/>',
                    b"", data)
            elif item.filename == "word/_rels/document.xml.rels":
                data = re.sub(
                    br'<Relationship[^>]*keyMapCustomizations[^>]*/>', b"", data)
            elif item.filename == "word/settings.xml":
                data = re.sub(br'<w:attachedTemplate[^>]*/>', b"", data)
            zout.writestr(item, data)


def clear_body(doc):
    """Empty the template body, keeping the section properties.

    The MDPI template marks the section as right-to-left (<w:bidi/>), which
    mirrors table column order in strict renderers and would place equation
    numbers on the wrong side; it is removed here.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            for sub in list(child):
                if sub.tag == qn("w:bidi"):
                    child.remove(sub)
            continue
        body.remove(child)


# ---------------------------------------------------------------------------
# low-level helpers
# ---------------------------------------------------------------------------
_BID = [1000]


def _next_bid():
    _BID[0] += 1
    return _BID[0]


def set_ind(p, left=None, first=None, right=None):
    """Indentation via the python-docx API so that pPr child order stays valid."""
    pf = p.paragraph_format
    if left is not None:
        pf.left_indent = Twips(left)
    if first is not None:
        pf.first_line_indent = Twips(first)
    if right is not None:
        pf.right_indent = Twips(right)


def single_spaced(p, before=0, after=0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_bookmark(paragraph, name, text, size=9, bold=False):
    bid = _next_bid()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.append(end)
    return run


def add_ref_field(paragraph, bookmark, shown, size=None):
    """Insert { REF <bookmark> \\h } with a cached result."""
    def _r():
        r = OxmlElement("w:r")
        if size is not None:
            rpr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(size * 2)))
            rpr.append(sz)
            r.append(rpr)
        return r

    r1 = _r()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r1.append(fc)
    r2 = _r()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " REF %s \\h " % bookmark
    r2.append(it)
    r3 = _r()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "separate")
    r3.append(fc)
    r4 = _r()
    t = OxmlElement("w:t")
    t.text = shown
    r4.append(t)
    r5 = _r()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "end")
    r5.append(fc)
    for r in (r1, r2, r3, r4, r5):
        paragraph._p.append(r)


# ---------------------------------------------------------------------------
# citation bookkeeping
# ---------------------------------------------------------------------------
CIT_ORDER = []          # keys in order of first appearance
CIT_NUM = {}


def cite_number(key):
    if key not in CIT_NUM:
        CIT_ORDER.append(key)
        CIT_NUM[key] = len(CIT_ORDER)
    return CIT_NUM[key]


SUBSUP = re.compile(r"([_^]\{[^}]*\})")
# a citation group is [[a]] or [[a]],[[b]],[[c]] — the comma is only part
# of the group when another citation follows it, so sentence commas survive
_K = r"[a-zA-Z0-9_\-]+"
_GRP = r"\[\[" + _K + r"(?:\],\[" + _K + r")*\]\]"
TOKEN = re.compile(r"(\$[^$]*\$|" + _GRP + r"(?:," + _GRP + r")*)")
CITE = re.compile(_K)


def _fmt_group(nums):
    """MDPI citation group: collapse runs of three or more into a range."""
    nums = sorted(set(nums))
    out, i = [], 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j - i >= 2:
            out.append((nums[i], nums[j]))
        else:
            out.extend((n, n) for n in nums[i:j + 1])
        i = j + 1
    return out


MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
YEARSPAN = re.compile(r"\b(19|20)(\d{2})-((19|20)\d{2})\b")


APOS = re.compile(r"(?<=[A-Za-z])'(?=[A-Za-z]|\s|$)")


def _typo(s):
    """Typographic minus, en dash in year spans and a curly apostrophe."""
    s = APOS.sub("\u2019", s)
    return MINUS.sub("\u2212",
                     YEARSPAN.sub(lambda m: m.group(0).replace("-", "\u2013"),
                                  s))


def add_rich(paragraph, text, size=None, italic=False, bold=False):
    """Add text with inline math ($...$) and citations ([[key]])."""
    for seg in TOKEN.split(text):
        if not seg:
            continue
        if seg.startswith("$") and seg.endswith("$") and len(seg) > 1:
            omml.add_inline_math(paragraph, math(seg[1:-1]))
        elif seg.startswith("[["):
            keys = CITE.findall(seg)
            nums = [cite_number(k) for k in keys]
            groups = _fmt_group(nums)
            r = paragraph.add_run("[")
            if size:
                r.font.size = Pt(size)
            for gi, (a, b) in enumerate(groups):
                if gi:
                    r = paragraph.add_run(",")
                    if size:
                        r.font.size = Pt(size)
                add_ref_field(paragraph, "_Ref_%d" % a, str(a), size)
                if b != a:
                    r = paragraph.add_run("–")
                    if size:
                        r.font.size = Pt(size)
                    add_ref_field(paragraph, "_Ref_%d" % b, str(b), size)
            r = paragraph.add_run("]")
            if size:
                r.font.size = Pt(size)
        else:
            r = paragraph.add_run(_typo(seg))
            r.font.italic = italic
            r.font.bold = bold
            if size:
                r.font.size = Pt(size)


M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _split_scripts(el):
    """Give every character of a sub- or superscript its own run.

    Semantically identical OMML, but it stops renderers that translate OMML
    into their own math grammar from reading a multi-letter script such as the
    panel index "it" as a single unknown token.
    """
    for tag in ("sub", "sup"):
        for node in el.iter("{%s}%s" % (M_NS, tag)):
            for r in list(node):
                if r.tag != "{%s}r" % M_NS:
                    continue
                t = r.find("{%s}t" % M_NS)
                if t is None:
                    continue
                # the ASCII star is read as an operator by some renderers
                t.text = (t.text or "").replace("*", "\u2217")
                if len(t.text) < 2:
                    continue
                idx = list(node).index(r)
                node.remove(r)
                for j, ch in enumerate(t.text):
                    nr = OxmlElement("m:r")
                    nt = OxmlElement("m:t")
                    nt.set(qn("xml:space"), "preserve")
                    nt.text = ch
                    nr.append(nt)
                    node.insert(idx + j, nr)
    return el


def math(latex):
    return [_split_scripts(e) for e in l2omml.L(latex)]


# ---------------------------------------------------------------------------
# builders
# ---------------------------------------------------------------------------
def para(doc, style, text, size=None, first_line=None):
    p = doc.add_paragraph(style=style)
    if first_line is not None:
        set_ind(p, first=first_line)
    add_rich(p, text, size=size)
    return p


TBLPR_ORDER = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
               "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
               "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
               "tblCellMar", "tblLook", "tblCaption", "tblDescription"]


def configure_table(tbl, twips, indent=None, center=False, borders=None,
                    cell_left=0, cell_right=0):
    """Rebuild w:tblPr with children in the order required by the schema.

    borders: mapping edge -> "single"/"none"; omitted edges default to none.
    """
    tblPr = tbl._tbl.tblPr
    keep = {}
    for child in list(tblPr):
        tag = child.tag.split("}")[1]
        if tag in ("tblStyle", "tblLook"):
            keep[tag] = child
        tblPr.remove(child)

    parts = {}
    if "tblStyle" in keep:
        parts["tblStyle"] = keep["tblStyle"]

    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(twips))
    w.set(qn("w:type"), "dxa")
    parts["tblW"] = w

    if center:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        parts["jc"] = jc
    if indent is not None:
        ind = OxmlElement("w:tblInd")
        ind.set(qn("w:w"), str(indent))
        ind.set(qn("w:type"), "dxa")
        parts["tblInd"] = ind

    b = OxmlElement("w:tblBorders")
    borders = borders or {}
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        val = borders.get(edge, "none")
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), "8" if val == "single" else "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        b.append(e)
    parts["tblBorders"] = b

    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    parts["tblLayout"] = lay

    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 0), ("left", cell_left), ("bottom", 0),
                      ("right", cell_right)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    parts["tblCellMar"] = mar

    if "tblLook" in keep:
        parts["tblLook"] = keep["tblLook"]

    for tag in TBLPR_ORDER:
        if tag in parts:
            tblPr.append(parts[tag])


def _no_split(tbl):
    """Keep a row from breaking across pages."""
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        e = OxmlElement("w:cantSplit")
        trPr.append(e)


def add_equation(doc, latex, number):
    """A borderless two-column table: equation centred, number right-aligned.

    Short equations sit in the body zone; long ones are allowed the full text
    width so that they are not broken by the 2608-twip body indent.
    """
    wide = len(latex) > 74
    total = FULL_W if wide else BODY_W
    eq_w = total - NUM_W
    tbl = doc.add_table(rows=1, cols=2)
    configure_table(tbl, total, indent=None if wide else BODY_INDENT,
                    center=wide)
    tbl.columns[0].width = Emu(int(eq_w * 635))
    tbl.columns[1].width = Emu(int(NUM_W * 635))
    _no_split(tbl)

    c0, c1 = tbl.rows[0].cells
    c0.width = Emu(int(eq_w * 635))
    c1.width = Emu(int(NUM_W * 635))
    # the number is centred against the height of the equation, which matters
    # for two-line expressions such as fractions
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = c0.paragraphs[0]
    p.style = doc.styles["MDPI_3.9_equation"]
    set_ind(p, left=0, right=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(omml.omath(math(latex)))

    p = c1.paragraphs[0]
    p.style = doc.styles["MDPI_3.a_equation_number"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("(%d)" % number)
    return tbl


def add_note(doc, style, text, prefix="Note: ", size=8):
    """Caption note: body-zone left edge, no indent, single line spacing."""
    n = doc.add_paragraph(style=style)
    set_ind(n, left=BODY_INDENT, first=0)
    single_spaced(n, before=2, after=10)
    if prefix:
        r = n.add_run(prefix)
        r.font.size = Pt(size)
    add_rich(n, text, size=size)
    return n


def add_figure(doc, key):
    spec = TS.FIGURES[key]
    p = doc.add_paragraph(style="MDPI_5.2_figure")
    # MDPI_5.2_figure centres across the full text width; align the image with
    # the body zone instead so that it sits flush with the text column
    set_ind(p, left=BODY_INDENT, first=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(os.path.join(FIGDIR, spec["file"]),
                            width=Cm(spec["width_cm"]))
    cap = doc.add_paragraph(style="MDPI_5.1_figure_caption")
    cap.paragraph_format.space_after = Pt(2 if spec.get("note") else 10)
    r = cap.add_run("Figure %d. " % spec["number"])
    r.font.bold = True
    r.font.size = Pt(9)
    add_rich(cap, spec["caption"], size=9)
    if spec.get("note"):
        add_note(doc, "MDPI_5.1_figure_caption", spec["note"])
    return p


def _bottom_rule(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    e = OxmlElement("w:bottom")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "8")
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), "auto")
    b.append(e)
    tcPr.append(b)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    trPr.append(e)


def add_table(doc, key):
    spec = TS.TABLES[key]()
    ncol = len(spec["header"])
    total = FULL_W if spec["wide"] else BODY_W
    ws = spec["widths"]
    s = sum(ws)
    widths = [int(total * w / s) for w in ws]

    cap = doc.add_paragraph(style="MDPI_4.1_table_caption")
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run("Table %s. " % spec["number"])
    r.font.bold = True
    r.font.size = Pt(9)
    add_rich(cap, spec["caption"], size=9)

    body_rows = spec["rows"]
    tbl = doc.add_table(rows=1 + len(body_rows), cols=ncol)
    tbl.style = doc.styles["MDPI_4.1_three_line_table"]
    configure_table(tbl, total, indent=None if spec["wide"] else BODY_INDENT,
                    center=spec["wide"],
                    borders={"top": "single", "bottom": "single"},
                    cell_left=57, cell_right=57)
    _no_split(tbl)

    fs = spec.get("fs") or (7.5 if ncol > 9 else (8.0 if ncol > 6 else 9.0))

    def rich_cell(p, text, bold, italic, size):
        """Render X_{sub} and X^{sup} with real Word sub/superscripts."""
        for part in SUBSUP.split(text):
            if not part:
                continue
            if part.startswith("_{") or part.startswith("^{"):
                run = p.add_run(part[2:-1])
                run.font.subscript = part[0] == "_"
                run.font.superscript = part[0] == "^"
            else:
                run = p.add_run(part)
            run.font.bold = bold
            run.font.italic = italic
            run.font.size = Pt(size)

    def fill(cell, text, w, bold=False, italic=False, align="c", size=fs):
        cell.width = Emu(int(w * 635))
        p = cell.paragraphs[0]
        p.style = doc.styles["MDPI_4.2_table_body"]
        p.alignment = {"l": WD_ALIGN_PARAGRAPH.LEFT,
                       "c": WD_ALIGN_PARAGRAPH.CENTER,
                       "r": WD_ALIGN_PARAGRAPH.RIGHT}[align]
        single_spaced(p, before=1.5, after=1.5)
        if text:
            rich_cell(p, text, bold, italic, size)

    for j, w in enumerate(widths):
        tbl.columns[j].width = Emu(int(w * 635))

    hdr = tbl.rows[0]
    for j, h in enumerate(spec["header"]):
        fill(hdr.cells[j], h, widths[j], bold=True,
             align=spec.get("header_align", ["c"] * ncol)[j])
        _bottom_rule(hdr.cells[j])
    _repeat_header(hdr)

    for i, item in enumerate(body_rows, start=1):
        row = tbl.rows[i]
        if item[0] == "sec":
            merged = row.cells[0]
            for j in range(1, ncol):
                merged = merged.merge(row.cells[j])
            merged.width = Emu(int(total * 635))
            p = merged.paragraphs[0]
            p.style = doc.styles["MDPI_4.2_table_body"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            single_spaced(p, before=3, after=1.5)
            rich_cell(p, item[1], True, True, fs)
        else:
            for j, cellval in enumerate(item[1]):
                fill(row.cells[j], cellval, widths[j],
                     italic=(spec.get("italic_col") == j),
                     align=spec["align"][j])
        if item[0] == "rule":
            for c in row.cells:
                _bottom_rule(c)

    # keep the table on a single page where it fits: Word breaks it anyway if
    # it cannot, but this stops a one-row orphan following the header
    keep_rows = tbl.rows if spec.get("note") else tbl.rows[:-1]
    for row in keep_rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True

    if spec.get("note"):
        add_note(doc, "MDPI_4.3_table_footer", spec["note"])
    return tbl


# ---------------------------------------------------------------------------
def render_reference(p, key, number):
    ref = R.REFS[key]
    # the bookmark spans the bare number only: the REF fields in the text copy
    # the bookmarked text, so anything else inside it would be copied too
    add_bookmark(p, "_Ref_%d" % number, str(number), size=9)
    r = p.add_run(".\t")
    r.font.size = Pt(9)

    def txt(s, italic=False, bold=False):
        run = p.add_run(s)
        run.font.size = Pt(9)
        run.font.italic = italic
        run.font.bold = bold

    def title_of(r):
        """Title with a single terminal full stop, keeping ? and !."""
        t = _typo(r["title"]).rstrip(". ")
        return t + (" " if t[-1] in "?!" else ". ")

    kind = ref["kind"]
    au = _typo(ref["authors"])
    txt(au + ("" if au.endswith(".") else ".") + " ")
    if kind in ("article", "preprint"):
        txt(title_of(ref))
        txt(ref["journal"] + " ", italic=True)
        txt(str(ref["year"]), bold=True)
        txt(", ")
        if ref.get("volume"):
            txt(str(ref["volume"]), italic=True)
            txt(", " + str(ref["pages"]) + ".")
        else:
            txt(str(ref["pages"]) + ".")
        if ref.get("doi"):
            txt(" [CrossRef]")
    elif kind == "book":
        txt(_typo(ref["title"]), italic=True)
        txt((", " + ref["edition"] if ref.get("edition") else "") + "; ")
        txt("%s: %s, %s, %d." % (ref["publisher"], ref["city"],
                                 ref["country"], ref["year"]))
    elif kind == "chapter":
        txt(title_of(ref))
        txt("In ")
        txt(ref["booktitle"], italic=True)
        txt("; " + ref["editors"] + "; ")
        txt("%s: %s, %s, %d; %s." % (ref["publisher"], ref["city"],
                                     ref["country"], ref["year"],
                                     ref["pages"]))
    elif kind == "report":
        txt(_typo(ref["title"]) + "; ", italic=True)
        if ref.get("series"):
            txt(ref["series"] + "; ")
        txt("%s: %s, %s, %d." % (ref["publisher"], ref["city"],
                                 ref["country"], ref["year"]))
    elif kind == "web":
        txt(title_of(ref))
        txt("Available online: " + ref["url"] +
            " (accessed on %s)." % ref["accessed"])


# ---------------------------------------------------------------------------
def front_matter(doc):
    p = doc.add_paragraph(style="MDPI_1.1_article_type")
    p.add_run("Article")

    p = doc.add_paragraph(style="MDPI_1.2_title")
    p.add_run(C.TITLE)

    p = doc.add_paragraph(style="MDPI_1.3_authornames")
    for text, sup in C.AUTHORS:
        r = p.add_run(text)
        r.font.superscript = sup

    for line, sid in (("Academic Editor: Firstname Lastname", "MDPI_1.5_academic_editor"),
                      ("Received: date", "MDPI_1.4_history"),
                      ("Revised: date", "MDPI_1.4_history"),
                      ("Accepted: date", "MDPI_1.4_history"),
                      ("Published: date", "MDPI_1.4_history")):
        doc.add_paragraph(style=sid).add_run(line)

    doc.add_paragraph(style="MDPI_6.1_citation").add_run(
        "Citation: To be added by editorial staff during production.")
    doc.add_paragraph(style="MDPI_7.2_copyright").add_run(
        "Copyright: © 2025 by the authors. Submitted for possible open "
        "access publication under the terms and conditions of the Creative "
        "Commons Attribution (CC BY) license (https://creativecommons.org/"
        "licenses/by/4.0/).")

    for sup, rest in C.AFFILIATIONS:
        p = doc.add_paragraph(style="MDPI_1.6_affiliation")
        r = p.add_run(sup)
        r.font.superscript = sup.isdigit()
        r.font.bold = not sup.isdigit()
        p.add_run("\t" + rest)

    # Highlights, in the journal's current front-matter format: a bold
    # section word, then each prompt in bold followed by its bullets
    if getattr(C, "HIGHLIGHTS", None):
        p = doc.add_paragraph(style="MDPI_1.7_abstract")
        r = p.add_run("Highlights")
        r.font.bold = True
        for prompt, bullets in C.HIGHLIGHTS:
            p = doc.add_paragraph(style="MDPI_1.7_abstract")
            r = p.add_run(prompt)
            r.font.bold = True
            for b in bullets:
                p = doc.add_paragraph(style="MDPI_1.7_abstract")
                set_ind(p, left=BODY_INDENT + 280, first=-200)
                p.add_run("•	")
                add_rich(p, b)

    p = doc.add_paragraph(style="MDPI_1.7_abstract")
    r = p.add_run("Abstract")
    r.font.bold = True
    p = doc.add_paragraph(style="MDPI_1.7_abstract")
    add_rich(p, C.ABSTRACT)

    p = doc.add_paragraph(style="MDPI_1.8_keywords")
    r = p.add_run("Keywords: ")
    r.font.bold = True
    p.add_run(C.KEYWORDS)

    doc.add_paragraph(style="MDPI_1.9_line")


def body(doc):
    prev = None
    for block in C.BLOCKS:
        kind = block[0]
        if kind == "h1":
            doc.add_paragraph(style="MDPI_2.1_heading1").add_run(block[1])
        elif kind == "h1b":
            p = doc.add_paragraph(style="MDPI_6.2_back_matter")
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(block[1] + ": ")
            r.font.bold = True
            r.font.size = Pt(9)
        elif kind == "h2":
            doc.add_paragraph(style="MDPI_2.2_heading2").add_run(block[1])
        elif kind == "h3":
            doc.add_paragraph(style="MDPI_2.3_heading3").add_run(block[1])
        elif kind == "p":
            style = ("MDPI_3.2_text_no_indent"
                     if prev in ("h1", "h2", "h3") else "MDPI_3.1_text")
            para(doc, style, block[1])
        elif kind == "hyp":
            p = doc.add_paragraph(style="MDPI_3.2_text_no_indent")
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            set_ind(p, left=BODY_INDENT + 340, first=0)
            r = p.add_run(block[1] + " ")
            r.font.bold = True
            add_rich(p, block[2], italic=True)
        elif kind == "eq":
            add_equation(doc, block[1], block[2])
        elif kind == "fig":
            add_figure(doc, block[1])
        elif kind == "table":
            add_table(doc, block[1])
        elif kind == "stmt":
            # an in-text statement with a bold lead-in
            p = doc.add_paragraph(style="MDPI_3.2_text_no_indent")
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            set_ind(p, left=BODY_INDENT + 340, first=0)
            r = p.add_run(block[1] + " ")
            r.font.bold = True
            r.font.size = Pt(9)
            add_rich(p, block[2], size=9)
        elif kind == "stmt2":
            # continuation of the preceding bold back-matter label
            add_rich(doc.paragraphs[-1], block[2], size=9)
        prev = kind


def references(doc):
    h = doc.add_paragraph(style="MDPI_2.1_heading1")
    set_ind(h, left=0, first=0)
    h.add_run("References")
    for n, key in enumerate(CIT_ORDER, start=1):
        p = doc.add_paragraph(style="MDPI_3.2_text_no_indent")
        set_ind(p, left=REF_HANG, first=-REF_HANG)
        single_spaced(p, before=0, after=2)
        # a tab stop at the hanging indent keeps the entries aligned
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), str(REF_HANG))
        tabs.append(tab)
        p._p.get_or_add_pPr().append(tabs)
        render_reference(p, key, n)


def _suppress_header_footer_linenos(doc):
    """Line numbering must run over the body only. LibreOffice otherwise
    numbers the header and footer frames as a second sequence that prints
    beside the body numbers."""
    for sec in doc.sections:
        parts = [sec.header, sec.footer, sec.first_page_header,
                 sec.first_page_footer, sec.even_page_header,
                 sec.even_page_footer]
        for part in parts:
            if part is None:
                continue
            for para in part.paragraphs:
                pr = para._p.get_or_add_pPr()
                if pr.find(qn("w:suppressLineNumbers")) is None:
                    el = OxmlElement("w:suppressLineNumbers")
                    pr.append(el)
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            pr = para._p.get_or_add_pPr()
                            if pr.find(qn("w:suppressLineNumbers")) is None:
                                el = OxmlElement("w:suppressLineNumbers")
                                pr.append(el)


def build():
    tmp = os.path.join(HERE, "_template.docx")
    template_to_docx(TEMPLATE, tmp)
    doc = Document(tmp)
    clear_body(doc)

    front_matter(doc)
    body(doc)
    references(doc)

    _suppress_header_footer_linenos(doc)
    doc.save(OUT)
    os.remove(tmp)

    unused = sorted(set(R.REFS) - set(CIT_ORDER))
    print("saved:", OUT)
    print("references cited:", len(CIT_ORDER), "/", len(R.REFS))
    if unused:
        print("UNCITED:", ", ".join(unused))
    missing = [k for k in CIT_ORDER if k not in R.REFS]
    if missing:
        print("MISSING FROM refs.py:", missing)


if __name__ == "__main__":
    build()
