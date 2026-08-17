# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter
and the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"),
          encoding="utf-8") as fh:
    S = json.load(fh)
M = S["models"]
FR = S["factreg"]
C1, C2 = S["centers"]["cluster1"], S["centers"]["cluster2"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")
_APOS = re.compile(r"(?<=[A-Za-z])'(?=[A-Za-z]|\s|$)")


def typo(s):
    s = _APOS.sub("\u2019", s)
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 9.8

TITLE = ("Artificial Intelligence Readiness and Youth Employment in the "
         "European Union: Multivariate Modeling and Cluster Typologies "
         "of Youth Disengagement and Unemployment")

BODY = [
    ("We would like to submit the enclosed manuscript entitled "
     "\u201c%s\u201d for consideration in Systems, as a contribution to "
     "the Topic \u201cRecent Applications of Artificial Intelligence in "
     "Economy and Society\u201d. The manuscript has not been published "
     "elsewhere and is not under consideration by any other journal. No "
     "conflict of interest exists in its submission, and it has been "
     "approved for publication by all authors." % TITLE),

    ("The topic could hardly be more timely. Artificial intelligence is "
     "most capable at precisely the entry-level analytical and "
     "administrative tasks from which early careers have always been "
     "built, and Europe monitors its young people through two headline "
     "indicators \u2014 the NEET rate and the youth unemployment rate "
     "\u2014 without knowing which of them, if either, moves with the "
     "diffusion of the technology. Using harmonised Eurostat data for "
     "the 27 member states, the manuscript measures national AI "
     "readiness on four dimensions (enterprise adoption, generative AI "
     "use among young individuals, youth digital skills, and ICT "
     "specialist employment) and relates them to both youth indicators "
     "through multivariate general linear models, an exploratory factor "
     "analysis, and a cluster typology \u2014 the comparative, "
     "systems-oriented register in which the Topic's recent articles "
     "are written."),

    "The highlights of the paper are as follows.",
]

POINTS = [
    ("(1) A clean and policy-relevant asymmetry. All four readiness "
     "dimensions are significantly and negatively associated with the "
     "youth NEET rate (partial eta squared between "
     f"{M['ICTS']['uni']['NEET']['eta2']:.3f} and "
     f"{M['GAIY']['uni']['NEET']['eta2']:.3f}), while none is "
     "significantly associated with the youth unemployment rate. The "
     "statistical relationship between AI and youth labour markets "
     "runs through engagement, not through measured unemployment "
     "\u2014 a distinction the literature has not previously drawn at "
     "this level."),

    ("(2) A latent readiness dimension. Principal Axis Factoring "
     f"(KMO = {S['efa']['kmo']:.3f}) condenses the four indicators "
     "into a single factor explaining "
     f"{S['efa']['var_extracted']:.1f}% of their variance, with "
     "generative AI use among the young cohort as its purest "
     f"expression (loading "
     f"{S['efa']['communalities']['GAIY']['loading']:.3f}). One "
     "standard deviation of this factor is associated with a NEET "
     f"rate lower by {abs(FR['NEET']['b1']):.2f} percentage points "
     f"(R\u00b2 = {FR['NEET']['r2']:.3f})."),

    ("(3) A two-regime map of Europe. Hierarchical and k-means "
     f"clustering agree on a higher-readiness group of {C1['n']} "
     f"economies (average NEET {C1['NEET']:.1f}%) and a "
     f"lower-readiness group of {C2['n']} economies (average NEET "
     f"{C2['NEET']:.1f}%), with France's below-average enterprise "
     "adoption placing it in the second group. The typology separates "
     "engaged from disengaged youth labour markets \u2014 not "
     "low-unemployment from high-unemployment ones \u2014 and "
     "motivates differentiated policy for the two regimes."),

    ("(4) Disciplined, transparent inference. All conclusions are "
     "presented as exploratory cross-sectional associations: "
     "diagnostics (Shapiro\u2013Wilk, Cook's distance), an "
     "influential-country re-estimation, and supplementary "
     "correlations with tertiary attainment, income, and R&D "
     "intensity are reported, and the insignificant "
     "unemployment-margin results are stated as such rather than "
     "explained away."),
]

CLOSING = (
    "We believe the manuscript fits the Topic precisely: it addresses "
    "the call's interest in AI's influence on workforce dynamics, "
    "employment trends, and labor markets with comparative European "
    "evidence, methodological transparency, and directly actionable "
    "implications \u2014 applied AI skills inside Youth Guarantee "
    "offers, protected learning content in AI-augmented entry-level "
    "jobs, and infrastructure-first strategies for the lower-readiness "
    "cluster. We hope the manuscript proves suitable for the Topic, "
    "and we would be pleased to revise it in light of the reviewers' "
    "comments.")


def para(doc, text, size=SIZE, bold=False, space_after=5, align=None,
         indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if indent:
        fmt.left_indent = Cm(indent)
    r = p.add_run(typo(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.9))
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.2))

    para(doc, "Cover Letter", size=13, bold=True, space_after=8,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Dear Topic Editors,", space_after=7)
    for b in BODY:
        para(doc, b)
    for pt in POINTS:
        para(doc, pt, indent=0.55, space_after=4)
    para(doc, CLOSING, space_after=8)
    para(doc, "Yours sincerely,", space_after=2)
    para(doc, "Xinyu Cai and Tiantian Mo", space_after=0)
    para(doc, "College of Business, Jiaxing University", space_after=0)
    para(doc, "Correspondence: 00008227@zjxu.edu.cn", space_after=0)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
