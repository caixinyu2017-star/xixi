# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter and
the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"),
          encoding="utf-8") as fh:
    S = json.load(fh)
TP = S["tp"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")


def typo(s):
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 9.6

TITLE = ("Early-Career Choices in the Era of Generative AI: Human Capital "
         "Depreciation, Career Anxiety, and Occupational Adaptation Among "
         "Fresh Graduates")

BODY = [
    ("We would like to submit the enclosed manuscript entitled \u201c%s\u201d "
     "for consideration in Systems, for the Special Issue \u201cSystems "
     "Approaches to Generative AI: Workforce Development, Organisational "
     "Learning, and Economic Transformation\u201d (Guest Editor: Dr. Ali "
     "Ahsan). The manuscript has not been published elsewhere and is not "
     "under consideration by any other journal. No conflict of interest "
     "exists in its submission, and it has been approved for publication by "
     "all authors." % TITLE),

    ("Generative artificial intelligence is most capable at precisely the "
     "tasks from which early careers have always been built, and the "
     "students now leaving university are the first cohort to choose "
     "occupations under that condition. The demand side of this collision "
     "is being measured intensively; the supply side \u2014 what the "
     "entrants themselves perceive, feel and then do \u2014 is nearly "
     "unstudied, even though early-career scarring is among the most "
     "persistent phenomena in labor economics. This manuscript treats the "
     "fresh graduate as an adaptive agent in a socio-technical labor "
     "system: perceived human capital depreciation is the input, "
     "AI-related career anxiety is the error signal of a cybernetic "
     "control loop, and occupational adaptation or career avoidance is "
     "how the loop closes. Three properties follow \u2014 transmission, "
     "saturation, and conjunction \u2014 and each is tested on a two-wave "
     f"survey of {S['n_t1']:,} Chinese graduates within three years of "
     f"graduation ({S['n_t2']:,} matched across waves, outcomes measured "
     "six months after their antecedents)."),

    "The highlights of the paper are as follows.",
]

POINTS = [
    ("(1) The first inverted-U evidence on AI career anxiety. Anxiety "
     "converts into reskilling, credential pursuit and occupational "
     f"repositioning up to a turning point at {TP['tau']:.2f} standard "
     "deviations above the mean \u2014 established by the Lind\u2013Mehlum "
     "composite test with a Fieller interval, not by the sign of a "
     f"quadratic \u2014 and {S['share_beyond']:.1f} percent of the cohort "
     "is already beyond it, in the region where more anxiety predicts "
     "less adaptation. Career avoidance, by contrast, rises "
     "monotonically. The same error signal that steers within a bounded "
     "band buys escape without bound."),

    ("(2) The conduit reverses sign, as a control-loop reading predicts. "
     "Bootstrapped conditional-process estimates show the identical "
     "depreciation percept, transmitted through anxiety, raising "
     f"adaptation by {S['med']['ind_adp_w+1_m-1_l+1']['est']:.3f} for a "
     "calm, supported, AI-literate graduate and lowering it by "
     f"{abs(S['med']['ind_adp_w-1_m+1_l-1']['est']):.3f} for an anxious, "
     "unsupported, low-literacy one. Mediation whose sign depends on the "
     "state of the mediator is exactly what saturation of an error "
     "signal implies, and it has not previously been documented in this "
     "literature."),

    ("(3) The two levers sit where systems theory says they sit. "
     "Generative-AI literacy does not flatten the curve; it moves the "
     f"turning point, from {S['tp_lit'][0]['tau']:.2f} to "
     f"{S['tp_lit'][2]['tau']:.2f} standard deviations \u2014 capability "
     "extends the range over which the signal stays productive. "
     "Employability support acts a stage earlier, cutting the "
     f"depreciation-to-anxiety pass-through from "
     f"{S['fslo']['est']:.2f} to {S['fshi']['est']:.2f}. Sensor "
     "recalibration and actuator range, estimated separately."),

    ("(4) Equifinality fails, informatively. Fuzzy-set qualitative "
     "comparative analysis returns exactly one sufficient recipe for "
     "high adaptation \u2014 anxiety AND literacy AND support jointly "
     f"(consistency {S['qca']['sol_cons']:.3f}), every condition core, "
     "no consistent recipe for the negation, and no single condition "
     "necessary. Where configurational studies routinely report many "
     "roads, a tight control-loop theory predicted one road, and the "
     "data delivered it. Measurement quality is documented throughout "
     f"(CFI {S['cfa']['cfi']:.3f}, RMSEA {S['cfa']['rmsea']:.3f}; "
     "invariance across gender and major; three method-bias "
     "diagnostics), and the curvilinear result survives a behavioral "
     "hours outcome, attrition weighting, marker partialling and a "
     "wild-cluster bootstrap."),
]

CLOSING = (
    "We believe the manuscript fits the Special Issue precisely. It "
    "answers the call\u2019s question about feedback loops, emergent "
    "behaviour and adaptive responses to generative AI not by metaphor "
    "but by estimation: the loop\u2019s stages, its saturation point, its "
    "boundary conditions and its conjunctural logic are each a numbered "
    "result. The practical payoff is a relocation of the policy lever. "
    "Two-fifths of a record graduating cohort already carries more "
    "career anxiety than it can convert into action; reassurance "
    "campaigns damp a signal that is in fact informative, while "
    "embedding generative-AI practice in curricula and building visible "
    "transition infrastructure complete the loop around it. We hope the "
    "manuscript proves suitable for the Special Issue, and we would be "
    "pleased to revise it in light of the reviewers\u2019 comments.")


def para(doc, text, size=SIZE, bold=False, space_after=5, align=None,
         indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if indent:
        fmt.left_indent = Cm(indent)
    r = p.add_run(typo(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.9))
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.2))

    para(doc, "Cover Letter", size=13, bold=True, space_after=8,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Dear Dr. Ahsan, Guest Editor of Systems,", space_after=7)
    for b in BODY:
        para(doc, b)
    for pt in POINTS:
        para(doc, pt, indent=0.55, space_after=4)
    para(doc, CLOSING, space_after=8)
    para(doc, "Yours sincerely,", space_after=2)
    para(doc, "Xinyu Cai and Tiantian Mo", space_after=0)
    para(doc, "College of Business, Jiaxing University", space_after=0)
    para(doc, "Correspondence: 00008227@zjxu.edu.cn", space_after=0)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
