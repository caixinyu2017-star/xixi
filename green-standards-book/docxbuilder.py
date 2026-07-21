# -*- coding: utf-8 -*-
"""国家社科基金后期资助专著 docx 生成引擎。

满足格式要求：
- 一级标题(章)：黑体 小三(15pt) 加粗；二级(节)：黑体 四号(14pt)；
  三级(条)：黑体 小四(12pt)；四级：宋体 小四(12pt) 加粗。
- 正文：宋体 小四(12pt)，首行缩进2字符，1.5倍行距。
- 图名置于图下方(五号黑体 居中)；表名置于表上方(五号黑体 居中)。
- 公式：居中，编号右顶格 (X-Y)，采用 Word 原生 OMML 公式(pandoc 由 LaTeX 生成)。
- 三线表；GB/T 7714 顺序编码制参考文献。

内容以 block 字典列表描述，见 build()。
"""
import copy
import os
import re
import subprocess
import zipfile

from lxml import etree
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# 中文字号 -> pt
SIZE = {
    '初号': 42, '小初': 36, '一号': 26, '小一': 24, '二号': 22, '小二': 18,
    '三号': 16, '小三': 15, '四号': 14, '小四': 12, '五号': 10.5, '小五': 9,
    '六号': 7.5,
}
SONG = '宋体'
HEI = '黑体'
KAI = '楷体'
TNR = 'Times New Roman'

# ------------------------------------------------------------------ OMML via pandoc
def latex_batch_to_omml(latex_list, workdir):
    """LaTeX 字符串列表 -> lxml <m:oMath> 元素列表（经 pandoc）。"""
    if not latex_list:
        return []
    md_parts = [f'EQMARK{i:04d} ${l}$' for i, l in enumerate(latex_list)]
    md = '\n\n'.join(md_parts)
    os.makedirs(workdir, exist_ok=True)
    mdp = os.path.join(workdir, '_eqs.md')
    outp = os.path.join(workdir, '_eqs.docx')
    with open(mdp, 'w', encoding='utf-8') as fh:
        fh.write(md)
    subprocess.run(['pandoc', mdp, '-o', outp], check=True)
    with zipfile.ZipFile(outp) as z:
        xml = z.read('word/document.xml')
    root = etree.fromstring(xml)
    results = {}
    for p in root.findall('.//{%s}p' % W_NS):
        text = ''.join(t.text or '' for t in p.findall('.//{%s}t' % W_NS))
        m = re.search(r'EQMARK(\d{4})', text)
        if m is None:
            continue
        idx = int(m.group(1))
        om = p.find('.//{%s}oMath' % M_NS)
        if om is not None:
            results[idx] = om
    out = []
    for i in range(len(latex_list)):
        if i not in results:
            raise RuntimeError(f'公式 {i} pandoc OMML 转换失败: {latex_list[i][:80]}')
        out.append(_fix_omml_order(results[i]))
    return out


def _fix_omml_order(om):
    """修正 pandoc 输出中不符合 ECMA-376 顺序的 OMML 子元素。

    - m:mcPr 序列应为 (count?, mcJc?)，pandoc 输出为 mcJc, count；
    - m:rPr(数学) 序列应为 (lit?, (nor|(scr?, sty?)), brk?, aln?)，
      pandoc 输出 sty 在 scr 之前。
    """
    def resort(el, rank):
        children = list(el)
        ranked = [(rank.get(etree.QName(ch).localname, 99), idx, ch)
                  for idx, ch in enumerate(children)]
        if ranked == sorted(ranked):
            return
        for ch in children:
            el.remove(ch)
        for _, _, ch in sorted(ranked):
            el.append(ch)

    RANKS = {
        'mcPr': {'count': 0, 'mcJc': 1},
        'rPr': {'lit': 0, 'nor': 1, 'scr': 1, 'sty': 2, 'brk': 3, 'aln': 4},
        'dPr': {'begChr': 0, 'sepChr': 1, 'endChr': 2, 'grow': 3, 'shp': 4, 'ctrlPr': 5},
        'naryPr': {'chr': 0, 'limLoc': 1, 'grow': 2, 'subHide': 3, 'supHide': 4, 'ctrlPr': 5},
        'mPr': {'baseJc': 0, 'plcHide': 1, 'rSpRule': 2, 'cGpRule': 3, 'rSp': 4,
                'cSp': 5, 'cGp': 6, 'mcs': 7, 'ctrlPr': 8},
        'fPr': {'type': 0, 'ctrlPr': 1},
        'accPr': {'chr': 0, 'ctrlPr': 1},
        'barPr': {'pos': 0, 'ctrlPr': 1},
        'groupChrPr': {'chr': 0, 'pos': 1, 'vertJc': 2, 'ctrlPr': 3},
        'radPr': {'degHide': 0, 'ctrlPr': 1},
        'eqArrPr': {'baseJc': 0, 'maxDist': 1, 'objDist': 2, 'rSpRule': 3, 'rSp': 4,
                    'ctrlPr': 5},
        'boxPr': {'opEmu': 0, 'noBreak': 1, 'diff': 2, 'brk': 3, 'aln': 4, 'ctrlPr': 5},
        'limLowPr': {'ctrlPr': 0}, 'limUppPr': {'ctrlPr': 0},
    }
    for tag, rank in RANKS.items():
        for el in om.iter('{%s}%s' % (M_NS, tag)):
            resort(el, rank)
    return om


# 行内标记：$latex$  **粗**  *斜*  {sup:x} {sub:x}
INLINE_RE = re.compile(r'(\$[^$]+\$|\*\*[^*]+\*\*|\*[^*]+\*|\{sup:[^}]*\}|\{sub:[^}]*\})')


def _iter_texts(blocks):
    for b in blocks:
        if 'eq' in b:
            yield b['eq_latex_only'] if False else None
        for key in ('p', 'caption', 'h1', 'h2', 'h3', 'h4', 'note', 'title', 'lead'):
            if key in b and isinstance(b[key], str):
                yield b[key]
        if 'table' in b:
            tb = b['table']
            for c in list(tb.get('header', [])) + [c for r in tb.get('rows', []) for c in r]:
                if isinstance(c, str):
                    yield c
            for k in ('caption', 'note'):
                if tb.get(k):
                    yield tb[k]
        if 'items' in b:
            for it in b['items']:
                yield it
        if 'refs' in b:
            for r in b['refs']:
                yield r


def collect_math(blocks):
    """预扫描所有 LaTeX（display + inline）。"""
    math = []
    for b in blocks:
        if 'eq' in b:
            math.append(b['eq'])
    for txt in _iter_texts(blocks):
        if not isinstance(txt, str):
            continue
        for seg in INLINE_RE.findall(txt):
            if seg.startswith('$'):
                math.append(seg[1:-1])
    return math


class BookBuilder:
    def __init__(self, workdir, margins_cm=(2.54, 2.54, 3.0, 3.0),
                 line_spacing=1.5):
        self.doc = docx.Document()
        self.workdir = workdir
        os.makedirs(workdir, exist_ok=True)
        self.omml_cache = {}
        self.line_spacing = line_spacing
        self._setup_page(margins_cm)
        self._setup_default_style()
        self.body = self.doc.element.body

    # ------------------------------------------------------------- 页面
    def _setup_page(self, margins_cm):
        sec = self.doc.sections[0]
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        top, bottom, left, right = margins_cm
        sec.top_margin = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin = Cm(left)
        sec.right_margin = Cm(right)
        self.text_width_cm = 21.0 - left - right  # 正文区宽度

    def _setup_default_style(self):
        # settings.xml 中 zoom 需 percent 属性（XSD 校验）
        try:
            zoom = self.doc.settings.element.find(qn('w:zoom'))
            if zoom is not None and zoom.get(qn('w:percent')) is None:
                zoom.set(qn('w:percent'), '100')
        except Exception:
            pass
        st = self.doc.styles['Normal']
        st.font.name = TNR
        st.font.size = Pt(SIZE['小四'])
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn('w:eastAsia'), SONG)
        rfonts.set(qn('w:ascii'), TNR)
        rfonts.set(qn('w:hAnsi'), TNR)

    # ------------------------------------------------------------- 字体工具
    @staticmethod
    def _set_run_font(run, cn_font=SONG, size_pt=SIZE['小四'], bold=False,
                      italic=False, western=None, color=None):
        run.font.size = Pt(size_pt)
        run.font.bold = bold or None
        run.font.italic = italic or None
        if color is not None:
            run.font.color.rgb = color
        west = western or cn_font
        run.font.name = west
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn('w:eastAsia'), cn_font)
        rfonts.set(qn('w:ascii'), west)
        rfonts.set(qn('w:hAnsi'), west)

    def _omml_for(self, latex):
        om = self.omml_cache.get(latex)
        if om is None:
            om = latex_batch_to_omml([latex], self.workdir)[0]
            self.omml_cache[latex] = om
        return copy.deepcopy(om)

    def _append_rich(self, para, text, cn_font=SONG, size_pt=SIZE['小四'],
                     bold=False, western=None, color=None):
        """向段落追加富文本，处理行内 $公式$ / **粗** / *斜* / {sup:}{sub:}。"""
        pos = 0
        for m in INLINE_RE.finditer(text):
            if m.start() > pos:
                r = para.add_run(text[pos:m.start()])
                self._set_run_font(r, cn_font, size_pt, bold, western=western, color=color)
            seg = m.group(0)
            if seg.startswith('$'):
                para._p.append(self._omml_for(seg[1:-1]))
            elif seg.startswith('**'):
                r = para.add_run(seg[2:-2])
                self._set_run_font(r, cn_font, size_pt, True, western=western, color=color)
            elif seg.startswith('*'):
                r = para.add_run(seg[1:-1])
                self._set_run_font(r, cn_font, size_pt, bold, italic=True, western=western, color=color)
            elif seg.startswith('{sup:'):
                r = para.add_run(seg[5:-1])
                self._set_run_font(r, cn_font, size_pt, bold, western=western, color=color)
                r.font.superscript = True
            elif seg.startswith('{sub:'):
                r = para.add_run(seg[5:-1])
                self._set_run_font(r, cn_font, size_pt, bold, western=western, color=color)
                r.font.subscript = True
            pos = m.end()
        if pos < len(text):
            r = para.add_run(text[pos:])
            self._set_run_font(r, cn_font, size_pt, bold, western=western, color=color)

    # ------------------------------------------------------------- 标题
    def add_heading(self, text, level):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(12 if level == 1 else 8)
        pf.space_after = Pt(8 if level == 1 else 6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.keep_with_next = True
        # 大纲级别（供 TOC 域与导航窗格识别）
        pPr = p._p.get_or_add_pPr()
        ol = OxmlElement('w:outlineLvl')
        ol.set(qn('w:val'), str(level - 1))
        pPr.append(ol)
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._append_rich(p, text, HEI, SIZE['小三'], bold=True)
        elif level == 2:
            self._append_rich(p, text, HEI, SIZE['四号'], bold=False)
        elif level == 3:
            self._append_rich(p, text, HEI, SIZE['小四'], bold=False)
        else:  # level 4
            self._append_rich(p, text, SONG, SIZE['小四'], bold=True)
        return p

    # ------------------------------------------------------------- 正文
    def add_paragraph(self, text, indent=True, align='justify', size_pt=SIZE['小四'],
                      cn_font=SONG, bold=False, space_after=0):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_after = Pt(space_after)
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'justify':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if indent:
            # 首行缩进 2 字符
            pPr = p._p.get_or_add_pPr()
            ind = pPr.get_or_add_ind()
            ind.set(qn('w:firstLineChars'), '200')
            ind.set(qn('w:firstLine'), str(int(size_pt * 2 * 20)))  # 兜底 twips
        self._append_rich(p, text, cn_font, size_pt, bold=bold)
        return p

    def add_bullets(self, items, ordered=False):
        for it in items:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pPr = p._p.get_or_add_pPr()
            ind = pPr.get_or_add_ind()
            ind.set(qn('w:firstLineChars'), '200')
            ind.set(qn('w:firstLine'), str(int(SIZE['小四'] * 2 * 20)))
            self._append_rich(p, it, SONG, SIZE['小四'])

    # ------------------------------------------------------------- 公式（居中 + 编号右顶格）
    def add_equation(self, latex, num):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        # 中间居中制表位 + 右端右对齐制表位
        w = self.text_width_cm
        pf.tab_stops.add_tab_stop(Cm(w / 2.0), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(w), WD_TAB_ALIGNMENT.RIGHT)
        r0 = p.add_run()
        r0.add_tab()
        p._p.append(self._omml_for(latex))
        r1 = p.add_run()
        r1.add_tab()
        rt = p.add_run(f'({num})')
        self._set_run_font(rt, SONG, SIZE['小四'], western=TNR)
        return p

    # ------------------------------------------------------------- 图
    def add_figure(self, img_path, caption, width_cm=13.0):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        cp.paragraph_format.keep_with_next = False
        self._append_rich(cp, caption, HEI, SIZE['五号'])
        return p

    # ------------------------------------------------------------- 表（三线表，表名在上）
    def _set_cell_font(self, cell, text, size_pt=SIZE['五号'], bold=False,
                       align='center', cn_font=SONG):
        cell.text = ''
        p = cell.paragraphs[0]
        align = {'l': 'left', 'c': 'center', 'r': 'right'}.get(align, align)
        p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                       'left': WD_ALIGN_PARAGRAPH.LEFT,
                       'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        self._append_rich(p, text, cn_font, size_pt, bold=bold)
        # 垂直居中
        tcPr = cell._tc.get_or_add_tcPr()
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), 'center')
        tcPr.append(va)

    def add_table(self, caption, header, rows, note=None, size_pt=SIZE['五号'],
                  col_align=None, first_col_left=True, width_cm=None):
        # 表名（表上方，五号黑体，居中）
        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.keep_with_next = True
        self._append_rich(cp, caption, HEI, SIZE['五号'])

        ncol = len(header)
        t = self.doc.add_table(rows=1 + len(rows), cols=ncol)
        t.alignment = 1  # center
        t.autofit = True
        if width_cm:
            self._set_table_width(t, width_cm)
        if col_align is None:
            col_align = ['center'] * ncol
            if first_col_left:
                col_align[0] = 'left'
        for j, txt in enumerate(header):
            self._set_cell_font(t.cell(0, j), str(txt), size_pt, bold=True, align='center')
        for i, row in enumerate(rows):
            for j, txt in enumerate(row):
                a = col_align[j] if j < len(col_align) else 'center'
                self._set_cell_font(t.cell(i + 1, j), str(txt), size_pt, align=a)
        self._three_line_borders(t)
        # 表注（五号宋体，左对齐）
        if note:
            np_ = self.doc.add_paragraph()
            np_.paragraph_format.space_after = Pt(10)
            np_.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            self._append_rich(np_, note, SONG, SIZE['小五'])
        else:
            sp = self.doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(8)
        return t

    def _set_table_width(self, table, width_cm):
        table.autofit = False
        tblPr = table._tbl.tblPr
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(int(width_cm * 567)))

    def _three_line_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        # 移除样式默认边框
        borders = OxmlElement('w:tblBorders')
        spec = (('top', 'single', 12), ('left', 'none', None), ('bottom', 'single', 12),
                ('right', 'none', None), ('insideH', 'none', None), ('insideV', 'none', None))
        for name, val, sz in spec:
            el = OxmlElement(f'w:{name}')
            el.set(qn('w:val'), val)
            if sz:
                el.set(qn('w:sz'), str(sz))
                el.set(qn('w:color'), '000000')
            borders.append(el)
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        # 按 CT_TblPrBase 序列插入：tblBorders 须位于 shd/tblLayout/tblCellMar/tblLook 之前
        after = [qn(f'w:{n}') for n in ('shd', 'tblLayout', 'tblCellMar', 'tblLook',
                                        'tblCaption', 'tblDescription')]
        anchor = next((ch for ch in tblPr if ch.tag in after), None)
        if anchor is not None:
            anchor.addprevious(borders)
        else:
            tblPr.append(borders)
        # 表头行下细线
        first = table.rows[0]._tr
        for tc in first.findall(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcB = OxmlElement('w:tcBorders')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single')
            bot.set(qn('w:sz'), '6')
            bot.set(qn('w:color'), '000000')
            tcB.append(bot)
            # CT_TcPrBase 序列：tcBorders 须位于 shd/vAlign 等之前
            late = [qn(f'w:{n}') for n in ('shd', 'noWrap', 'tcMar', 'textDirection',
                                           'tcFitText', 'vAlign', 'hideMark')]
            anchor = next((ch for ch in tcPr if ch.tag in late), None)
            if anchor is not None:
                anchor.addprevious(tcB)
            else:
                tcPr.append(tcB)

    # ------------------------------------------------------------- 参考文献
    def add_references(self, refs, heading='参考文献'):
        self.add_heading(heading, 1)
        for i, ref in enumerate(refs, 1):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pf.space_after = Pt(2)
            # 悬挂缩进
            pPr = p._p.get_or_add_pPr()
            ind = pPr.get_or_add_ind()
            ind.set(qn('w:left'), '480')
            ind.set(qn('w:hanging'), '480')
            self._append_rich(p, f'[{i}] {ref}', SONG, SIZE['五号'])

    def add_pagebreak(self):
        from docx.enum.text import WD_BREAK
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    def add_toc_placeholder(self):
        pass

    # ------------------------------------------------------------- 驱动
    def prepare_math(self, blocks):
        latexes = collect_math(blocks)
        uniq = list(dict.fromkeys(latexes))
        if uniq:
            ommls = latex_batch_to_omml(uniq, self.workdir)
            self.omml_cache = dict(zip(uniq, ommls))

    def build(self, blocks):
        self.prepare_math(blocks)
        for b in blocks:
            if 'h1' in b:
                self.add_heading(b['h1'], 1)
            elif 'h2' in b:
                self.add_heading(b['h2'], 2)
            elif 'h3' in b:
                self.add_heading(b['h3'], 3)
            elif 'h4' in b:
                self.add_heading(b['h4'], 4)
            elif 'p' in b:
                self.add_paragraph(b['p'], indent=b.get('indent', True),
                                   align=b.get('align', 'justify'))
            elif 'lead' in b:  # 无缩进说明段（如 where …）
                self.add_paragraph(b['lead'], indent=False)
            elif 'eq' in b:
                self.add_equation(b['eq'], b['num'])
            elif 'fig' in b:
                self.add_figure(b['fig'], b['caption'], b.get('width_cm', 13.0))
            elif 'table' in b:
                tb = b['table']
                self.add_table(tb['caption'], tb['header'], tb['rows'],
                               note=tb.get('note'), size_pt=tb.get('size_pt', SIZE['五号']),
                               col_align=tb.get('col_align'), width_cm=tb.get('width_cm'))
            elif 'items' in b:
                self.add_bullets(b['items'], ordered=b.get('ordered', False))
            elif 'refs' in b:
                self.add_references(b['refs'])
            elif 'pagebreak' in b:
                self.add_pagebreak()

    def save(self, path):
        self.doc.save(path)


def count_chars(blocks):
    """统计中文正文字数（含标点，不含公式LaTeX与参考文献英文串近似）。"""
    total = 0
    for txt in _iter_texts(blocks):
        if not isinstance(txt, str):
            continue
        clean = INLINE_RE.sub('', txt)
        total += len(re.sub(r'\s', '', clean))
    return total
