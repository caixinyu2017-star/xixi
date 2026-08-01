# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter and
the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"), encoding="utf-8") as fh:
    S = json.load(fh)
S["years"] = S["years"].replace("-", "\u2013")
MOD = S["moderation"]

# typographic minus before numerals, en dash in year spans
_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")


def typo(s):
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 10.0

TITLE = ("Options on Youth: Generative Artificial Intelligence Uncertainty and "
         "the Contingent Turn in Entry-Level Hiring")

BODY = [
    ("We would like to submit the enclosed manuscript entitled \u201c%s\u201d for "
     "consideration in the Special Issue \u201cSystems Approaches to Generative AI: "
     "Workforce Development, Organisational Learning, and Economic "
     "Transformation\u201d of Systems. The manuscript has not been published "
     "elsewhere and is not under consideration by any other journal. No "
     "conflict of interest exists in its submission, and it has been approved "
     "for publication by all authors." % TITLE),

    ("Almost the entire debate about generative artificial intelligence and "
     "young workers is about capability: will the models do what juniors do? "
     "This manuscript argues that what actually closes the entry port is not "
     "capability but its dispersion. Hiring a school leaver is a partly "
     "irreversible investment, because the firm sinks a training cost it cannot "
     "recover if the frontier moves against the task bundle it hired for. The "
     "theory of investment under uncertainty then says the firm should raise "
     "its hurdle and wait, so the entry-level labour market can contract before "
     "the technology has displaced a single task. Using "
     f"{S['n_obs']:,} firm-year observations for {S['n_firms']:,} Chinese "
     f"A-share listed firms over {S['years']}, we show that it does."),

    "The highlights of the paper are as follows.",
]

POINTS = [
    ("(1) Dispersion, not level. One standard deviation of firm-level "
     "generative-AI uncertainty lowers entry-level vacancy creation by "
     f"{abs(S['b_main']):.3f} vacancies per thousand employees, "
     f"{100 * abs(S['b_main']) / S['entry_mean']:.1f}% of the mean, while the "
     "level of adoption enters the same equation with the opposite sign "
     f"({S['b_level']:+.3f}) and general economic policy uncertainty is "
     f"{S['ratio_u_e']:.1f} times smaller. Studies that measure only adoption "
     "average across two channels that work in opposite directions, which is "
     "one reason their findings disagree."),

    ("(2) A sharp discriminating test. The same coefficient for vacancies "
     f"requiring five or more years of experience is {S['b_exper']:.3f}, "
     f"{S['ratio_age']:.1f} times smaller, and a cross-equation test on the "
     "stacked model rejects equality decisively. General pessimism, a demand "
     "shock or poor performance cannot produce that asymmetry; an option whose "
     "value scales with what the firm cannot recover can."),

    ("(3) The entry port is not closing, it is becoming revocable. Uncertainty "
     "raises the contingent share of entry-level vacancies \u2014 internships, "
     "agency placements, short fixed-term contracts \u2014 by "
     f"{S['b_cont']:.2f} percentage points, and the shift is largest exactly "
     "where the reversible margin is cheapest to use. A study that counts young "
     "employees will therefore miss most of what is happening, because the "
     "first thing a firm gives up is not the young worker but the commitment to "
     "them."),

    ("(4) Irreversibility, inaction and hysteresis. The chill is "
     f"{MOD['SpecTrain']['ratio']:.1f} times larger for firms one standard "
     "deviation above the mean of firm-specific training intensity than one "
     "standard deviation below it, rises with employment protection, and is "
     "attenuated where reversible employment is available. Uncertainty also "
     "widens a zone of inaction: it raises the probability of no youth "
     f"adjustment at all by {S['b_inact']:.3f} and lowers the probability that "
     f"an already inactive firm resumes expansion by {abs(S['b_resume']):.3f}."),
]

CLOSING = (
    "We believe the manuscript speaks directly to the aims of the Special "
    "Issue, and that its policy implication is unusually specific. If the entry "
    "port is closing because firms are uncertain rather than because they have "
    "been displaced, the cheapest available intervention is not retraining but "
    "information: sectoral technology roadmaps, disclosure standards for model "
    "capabilities and limitations, and certification of what models can be "
    "relied upon all narrow the dispersion without touching the frontier, and "
    "are therefore youth employment policy. Identification rests on two "
    "shift-share instruments and a peer-diffusion instrument that pass the "
    "Kleibergen\u2013Paap and overidentification tests, on a dynamic "
    "difference-in-differences design with flat pre-shock leads, and on a wild "
    "cluster bootstrap, randomisation inference and coefficient-stability "
    "bounds; every number in every table is produced by the released estimation "
    "code. We deeply appreciate your consideration of our manuscript and look "
    "forward to receiving comments from the reviewers. Correspondence should be "
    "directed to the address below.")

SIGN = [
    "Name: Tiantian Mo",
    "Institution and address: College of Business, Jiaxing University; "
    "Jiaxing 314001, China",
    "Email: 00008227@zjxu.edu.cn",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.8)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def p(text, after=5, bold=False, italic=False, align=None, indent=0.0,
          size=SIZE):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        run = par.add_run(typo(text))
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        return par

    p("Dear Editors:", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    for t in BODY:
        p(t, after=5)
    for t in POINTS:
        p(t, after=4)
    p(CLOSING, after=7)

    p("Thanks very much for your attention to our paper.", after=8,
      align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in SIGN:
        p(line, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("Very sincerely yours,", after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p("Xinyu Cai and Tiantian Mo", after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
