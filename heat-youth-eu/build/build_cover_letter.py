# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the
letter and the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"),
          encoding="utf-8") as fh:
    S = json.load(fh)
B = S["base"]
ME = S["me_at"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")
_APOS = re.compile(r"(?<=[A-Za-z])'(?=[A-Za-z]|\s|$)")


def typo(s):
    s = _APOS.sub("\u2019", s)
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 9.8

TITLE = ("Summer Heat Stress, Urban Green Infrastructure, and Youth "
         "Labour Markets in the European Union: Evidence from a System "
         "GMM Approach")

BODY = [
    ("We would like to submit the enclosed manuscript entitled "
     "\u201c%s\u201d for consideration in Sustainability, for the "
     "Special Issue \u201cThermal Mitigation Effects of Green and Blue "
     "Infrastructure and Urban Sustainability from an "
     "Interdisciplinary Perspective\u201d. The manuscript has not been "
     "published elsewhere and is not under consideration by any other "
     "journal. No conflict of interest exists in its submission, and "
     "it has been approved for publication by all authors." % TITLE),

    ("The Special Issue asks how green and blue infrastructure "
     "mitigates urban heat and enhances human well-being, and it "
     "explicitly lists the socioeconomic impacts of heatwaves among "
     "its themes. The enclosed manuscript answers on the ground where "
     "those impacts will be felt longest: the labour-market entry of "
     "the young generation. The sectors through which young Europeans "
     "start working \u2014 tourism, hospitality, agriculture, "
     "construction \u2014 are the most heat-exposed in the economy, "
     "and conditions at entry scar careers for a decade. Using a "
     "balanced panel of the 27 EU Member States (2010\u20132024), "
     "cooling degree days, and the green infrastructure endowment of "
     "functional urban areas from the Copernicus Urban Atlas, we "
     "estimate a dynamic panel model of the youth NEET rate with a "
     "two-step System GMM estimator \u2014 the methodological register "
     "of the recent EU-panel literature in Sustainability."),

    "The highlights of the paper are as follows.",
]

POINTS = [
    ("(1) Summer heat measurably raises youth disengagement. An "
     "additional hundred cooling degree days raises the NEET rate by "
     f"{B['coef']['CDD100']:.2f} percentage points at a zero green "
     "endowment ({}), with persistence implying substantially larger "
     "long-run effects.".format(
         "p < 0.001" if B['p']['CDD100'] < 0.001
         else "p = %.3f" % B['p']['CDD100'])),

    ("(2) Urban greenery buffers the shock \u2014 the Special "
     "Issue's thermal-mitigation thesis, read in labour-market "
     "units. The heat effect declines by "
     f"{abs(B['coef']['CDD100 × GRN']):.4f} percentage points per "
     "green-share point (p = "
     f"{B['p']['CDD100 × GRN']:.3f}), falling from "
     f"{ME['15']['me']:.2f} at a 15% endowment to statistical "
     f"insignificance beyond roughly {S['me_cutoff']:.0f}% and to "
     "essentially zero in the greenest urban systems. Thermal "
     "mitigation is, at the macro level, social infrastructure."),

    ("(3) The effect sits on the engagement margin. Heat moves the "
     "NEET rate but is undetectable in the youth unemployment rate, "
     "which conditions on active search: hot summers push marginal "
     "young people out of structured activity altogether. "
     "Climate-adaptation dashboards that watch unemployment alone "
     "will miss this."),

    ("(4) Disciplined dynamic-panel inference. Two-step System GMM "
     "with Windmeijer-corrected standard errors, a deliberately "
     f"collapsed instrument set ({B['instruments']} instruments for "
     f"{B['groups']} countries), clean Arellano\u2013Bond and Hansen "
     f"diagnostics (AR(2) p = {B['ar2_p']:.3f}; Hansen p = "
     f"{B['hansen_p']:.3f}), a pre-pandemic robustness estimation, "
     "and period demeaning equivalent to full time dummies."),
]

CLOSING = (
    "We believe the manuscript fits the Special Issue precisely: it "
    "bridges climate science, urban design, and labour economics "
    "\u2014 the cross-sectoral perspective the call champions \u2014 "
    "and delivers actionable insights: urban greening under the "
    "Nature Restoration Regulation yields youth-employment "
    "co-benefits concentrated in the heat-exposed South, and the "
    "Youth Guarantee should treat recurring summer shocks as a "
    "predictable disengagement risk. We hope the manuscript proves "
    "suitable for the Special Issue, and we would be pleased to "
    "revise it in light of the reviewers' comments.")


def para(doc, text, size=SIZE, bold=False, space_after=5, align=None,
         indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if indent:
        fmt.left_indent = Cm(indent)
    r = p.add_run(typo(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.9))
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.2))

    para(doc, "Cover Letter", size=13, bold=True, space_after=8,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Dear Dr. Rus and Dr. Gaban, Guest Editors,",
         space_after=7)
    for b in BODY:
        para(doc, b)
    for pt in POINTS:
        para(doc, pt, indent=0.55, space_after=4)
    para(doc, CLOSING, space_after=8)
    para(doc, "Yours sincerely,", space_after=2)
    para(doc, "Xinyu Cai and Tiantian Mo", space_after=0)
    para(doc, "College of Business, Jiaxing University", space_after=0)
    para(doc, "Correspondence: 00008227@zjxu.edu.cn", space_after=0)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
