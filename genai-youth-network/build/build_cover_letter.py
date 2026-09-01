# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter and
the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"), encoding="utf-8") as fh:
    S = json.load(fh)
S["years"] = S["years"].replace("-", "\u2013")
AGE = S["agespec"]
PLA = S["placebo"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")


def typo(s):
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 10.0

TITLE = ("The Sum Is Not the Parts: Generative Artificial Intelligence, "
         "Production Networks, and the System-Level Displacement of Youth "
         "Employment")

BODY = [
    ("We would like to submit the enclosed manuscript entitled \u201c%s\u201d for "
     "consideration in the Special Issue \u201cSystems Approaches to Generative AI: "
     "Workforce Development, Organisational Learning, and Economic "
     "Transformation\u201d of Systems. The manuscript has not been published "
     "elsewhere and is not under consideration by any other journal. No "
     "conflict of interest exists in its submission, and it has been approved "
     "for publication by all authors." % TITLE),

    ("Every firm-level study of generative artificial intelligence and youth "
     "employment estimates the same object: what happens inside the firm that "
     "adopts. This manuscript argues that the object is the wrong one. Firms "
     "are not independent units; they buy from one another, sell to one "
     "another, and hire from the same local pools of graduates. When one of "
     "them adopts a technology that changes what junior labour is for, the "
     "consequences travel. We measure where they go. Using "
     f"{S['n_obs']:,} firm-year observations for {S['n_firms']:,} Chinese "
     f"A-share listed firms over {S['years']}, their disclosed supplier and "
     f"customer links, and the {S['n_cities']} local labour markets they "
     "share, we estimate a network autoregressive model of the youth "
     "employment share and decompose its impacts."),

    "The highlights of the paper are as follows.",
]

POINTS = [
    ("(1) Three transmissions, two signs. Own adoption lowers the youth "
     f"employment share by {abs(S['b_partial']):.3f} percentage points per "
     "standard deviation. Adoption by the firm's customers lowers it by a "
     f"further {abs(S['b_down']):.3f} \u2014 {S['ratio_down_own']:.2f} times "
     "the own-firm effect, obtained from firms that may have adopted nothing "
     f"themselves \u2014 and by suppliers by {abs(S['b_up']):.3f}. But "
     "adoption by other firms in the same city *raises* it, by "
     f"{S['b_peer']:.3f}: young workers one firm does not hire are hired by "
     "the firm next door."),

    ("(2) A network multiplier of "
     f"{S['multiplier']:.2f}. In the network autoregressive model, estimated "
     "by generalised spatial two-stage least squares with second-order network "
     "lags as excluded instruments, the direct impact is "
     f"{S['direct']:.3f}, the indirect impact {S['indirect']:.3f} and the "
     f"total impact {S['total']:.3f}. That is, "
     f"{S['indirect_share']:.0f}% of the system-level effect of adoption on "
     "youth employment occurs in firms other than the adopter."),

    ("(3) Half of the firm-level effect is reallocation, not destruction. "
     "Estimated on city aggregates, the same coefficient is "
     f"{S['b_city']:.3f}, only {100 - S['realloc_share']:.0f}% of the "
     "firm-level estimate. The firm-level coefficient is therefore neither an "
     "upper nor a lower bound on anything policy cares about: it overstates "
     "the local damage by about a factor of two and understates the "
     f"system-level damage by about a factor of {S['total_over_firm']:.1f}. "
     "Studies that find large effects and studies that find none may both be "
     "right about their own estimand."),

    ("(4) The signature is a network, not a common shock. Propagation is "
     "concentrated in relationship-specific customer links "
     f"({S['b_spec']:.3f} against {S['b_gen']:.3f} for generic inputs), is "
     "indifferent to whether the customer is in the same city "
     f"(difference {S['diff_near']:.3f}, p = {S['p_diff_near']:.3f}), is "
     f"{AGE['DownAI']['ratio']:.1f} times larger for young workers than for "
     "experienced ones, and vanishes when the network is randomly rewired "
     "while every firm keeps its own adoption and characteristics "
     f"(randomisation p = {PLA['DownAI']['p']:.3f})."),
]

CLOSING = (
    "We believe the manuscript speaks directly to the aims of the Special "
    "Issue, which asks for systems approaches to generative AI and workforce "
    "development. Our result is a systems result in the strict sense: the "
    "effect of the technology on youth employment is a property of the system "
    "of firms and not of any firm in it, and it cannot be estimated, or "
    "governed, one firm at a time. A young person's chances are being changed "
    "by technology decisions taken in firms they will never apply to. "
    "Identification rests on the intransitive triads of the disclosed "
    "supplier\u2013customer network, which make partners-of-partners adoption "
    "an excluded instrument for the endogenous network lag; the instruments "
    "pass the Kleibergen\u2013Paap and overidentification tests, the estimates "
    "survive a restricted wild cluster bootstrap and coefficient-stability "
    "bounds, and every number in every table is produced by the released "
    "estimation code. We deeply appreciate your consideration of our "
    "manuscript and look forward to receiving comments from the reviewers. "
    "Correspondence should be directed to the address below.")

SIGN = [
    "Name: Tiantian Mo",
    "Institution and address: College of Business, Jiaxing University; "
    "Jiaxing 314001, China",
    "Email: 00008227@zjxu.edu.cn",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.8)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def p(text, after=5, bold=False, italic=False, align=None, indent=0.0,
          size=SIZE):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        run = par.add_run(typo(text).replace("*", ""))
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        return par

    p("Dear Editors:", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    for t in BODY:
        p(t, after=5)
    for t in POINTS:
        p(t, after=4)
    p(CLOSING, after=7)

    p("Thanks very much for your attention to our paper.", after=8,
      align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in SIGN:
        p(line, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("Very sincerely yours,", after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p("Xinyu Cai and Tiantian Mo", after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
