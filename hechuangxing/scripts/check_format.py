# -*- coding: utf-8 -*-
"""排版合规自检：逐段核对字体、字号、对齐、缩进、段前段后、行距。"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document                      # noqa: E402
from docx.oxml.ns import qn                    # noqa: E402

SPEC = {
    "一级标题/标题": dict(font="黑体", size=15.0, align="center", indent=0, bold=False),
    "二级标题":     dict(font="黑体", size=14.0, align="left", indent=200, bold=False),
    "三级标题":     dict(font="宋体", size=12.0, align="left", indent=200, bold=True),
    "正文":         dict(font="宋体", size=12.0, align="both", indent=200, bold=False),
    "表格":         dict(font="宋体", size=10.5, align="center", indent=0),
}


def para_info(p):
    ppr = p._p.find(qn("w:pPr"))
    jc = ind = None
    before = after = line = None
    if ppr is not None:
        e = ppr.find(qn("w:jc"));   jc = e.get(qn("w:val")) if e is not None else None
        e = ppr.find(qn("w:ind"))
        if e is not None:
            ind = e.get(qn("w:firstLineChars"))
        e = ppr.find(qn("w:spacing"))
        if e is not None:
            before, after, line = e.get(qn("w:before")), e.get(qn("w:after")), e.get(qn("w:line"))
    runs = [r for r in p.runs if r.text.strip()]
    if not runs:
        return None
    r = runs[0]
    rpr = r._element.find(qn("w:rPr"))
    fonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
    return dict(
        text=p.text.strip(),
        ea=fonts.get(qn("w:eastAsia")) if fonts is not None else None,
        ascii=fonts.get(qn("w:ascii")) if fonts is not None else None,
        size=r.font.size.pt if r.font.size else None,
        bold=bool(r.font.bold),
        align=jc, indent=ind, before=before, after=after, line=line,
    )


def classify(info):
    if info["ea"] == "黑体" and info["size"] == 15.0:
        return "一级标题/标题"
    if info["ea"] == "黑体" and info["size"] == 14.0:
        return "二级标题"
    if info["ea"] == "宋体" and info["size"] == 12.0 and info["bold"]:
        return "三级标题"
    if info["ea"] == "宋体" and info["size"] == 12.0:
        return "正文"
    return "未知"


def check(path):
    doc = Document(str(path))
    errs, counts = [], {}
    for p in doc.paragraphs:
        info = para_info(p)
        if not info:
            continue
        kind = classify(info)
        counts[kind] = counts.get(kind, 0) + 1
        if kind == "未知":
            errs.append(f"未识别样式：{info['text'][:24]} ea={info['ea']} size={info['size']}")
            continue
        s = SPEC[kind]
        if info["ascii"] != "Times New Roman":
            errs.append(f"[{kind}] 西文字体不是 Times New Roman：{info['text'][:20]}")
        if info["align"] != s["align"]:
            errs.append(f"[{kind}] 对齐应为 {s['align']}，实为 {info['align']}：{info['text'][:20]}")
        want = "0" if s["indent"] == 0 else str(s["indent"])
        if (info["indent"] or "0") != want:
            errs.append(f"[{kind}] 首行缩进应为 {want}，实为 {info['indent']}：{info['text'][:20]}")
        if info["before"] not in (None, "0") or info["after"] not in (None, "0"):
            errs.append(f"[{kind}] 段前段后应为 0：before={info['before']} after={info['after']}")
        if info["bold"] != s["bold"]:
            errs.append(f"[{kind}] 加粗状态应为 {s['bold']}：{info['text'][:20]}")

    tbl_bad = 0
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    info = para_info(p)
                    if not info:
                        continue
                    counts["表格"] = counts.get("表格", 0) + 1
                    ok = (info["ea"] == "宋体" and info["ascii"] == "Times New Roman"
                          and info["align"] == "center" and (info["indent"] or "0") == "0"
                          and info["before"] in (None, "0") and info["after"] in (None, "0")
                          and info["line"] == "240")     # 240 二十分之一磅 = 单倍行距
                    if not ok:
                        tbl_bad += 1
                        if tbl_bad <= 3:
                            errs.append(f"[表格] 单元格不合规：{info}")
    print(f"\n=== {path.name} ===")
    print("  段落统计：" + "，".join(f"{k} {v}" for k, v in sorted(counts.items())))
    print(f"  表格数：{len(doc.tables)}")
    if errs:
        print(f"  ✗ 发现 {len(errs)} 处问题：")
        for e in errs[:12]:
            print("    -", e)
    else:
        print("  ✓ 全部符合排版规范")
    return len(errs)


total = 0
for f in sorted((ROOT / "samples").glob("*.docx")):
    total += check(f)
print(f"\n合计问题数：{total}")
sys.exit(1 if total else 0)
