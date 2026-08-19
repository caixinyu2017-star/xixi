# -*- coding: utf-8 -*-
"""
禾创星 · Word 排版引擎
====================
把结构化内容块渲染成符合下列排版规范的 .docx：

    一级标题：黑体，小三（15pt），居中，无首行缩进
    二级标题：黑体，四号（14pt），左对齐，首行缩进 2 字符
    三级标题：宋体加粗，小四（12pt），左对齐，首行缩进 2 字符
    正  文：中文宋体 / 西文 Times New Roman，无段前段后间距，首行缩进 2 字符
    表  格：段前段后 0 磅，单倍行距，居中

首行缩进统一使用 Word 的"字符"单位（w:firstLineChars="200"），
因此换字号不会导致缩进量走样。
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

# ---------------------------------------------------------------- 可调参数 ---
# 正文行距。规范只对表格明确要求"单倍行距"，正文行距未作规定，
# 这里按中文商务文书的通行做法取 1.5 倍；若需改成单倍，把下面一行改成 1.0 即可。
BODY_LINE_SPACING = 1.5
TABLE_LINE_SPACING = 1.0          # 表格：单倍行距（规范明确要求）

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_WEST = "Times New Roman"

SIZE_H1 = Pt(15)                  # 小三
SIZE_H2 = Pt(14)                  # 四号
SIZE_H3 = Pt(12)                  # 小四
SIZE_BODY = Pt(12)                # 小四
SIZE_TABLE = Pt(10.5)             # 五号
SIZE_FOOTER = Pt(9)               # 小五


# ------------------------------------------------------------ 底层小工具 ---
def _set_run_font(run, east_asia: str, west: str, size: Pt, bold: bool = False):
    """同时设定中文字体、西文字体与复杂文种字体，避免 Word 回退到默认字体。"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), west)
    # hint 必须挂在 w:rFonts 上：它决定引号、括号、顿号等中西文共用的标点
    # 走哪套字体。设为 eastAsia 后，引号跟随宋体（或黑体），不会被西文字体接管。
    rfonts.set(qn("w:hint"), "eastAsia")


def _set_indent_chars(paragraph, chars: int):
    """用字符数设置首行缩进；chars=0 表示取消缩进。"""
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if chars <= 0:
        ind.set(qn("w:firstLineChars"), "0")
        ind.set(qn("w:firstLine"), "0")
    else:
        ind.set(qn("w:firstLineChars"), str(chars * 100))
        # 同时给一个磅值兜底，兼容不读 firstLineChars 的阅读器（如 WPS 旧版、部分预览器）
        ind.set(qn("w:firstLine"), str(int(SIZE_BODY.pt * chars * 20)))
    ind.set(qn("w:leftChars"), "0")
    ind.set(qn("w:left"), "0")


def _set_spacing(paragraph, line_spacing: float, before=0, after=0):
    """段前段后一律 0，行距按倍数设置。"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    # 关闭"如果定义了文档网格，则对齐到网格"，否则行距会被网格吃掉
    ppr = paragraph._p.get_or_add_pPr()
    snap = ppr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        ppr.append(snap)
    snap.set(qn("w:val"), "0")


def _no_widow_control(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    kn = OxmlElement("w:keepNext")
    ppr.append(kn)


# --------------------------------------------------------------- 文档骨架 ---
def new_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_WEST
    normal.font.size = SIZE_BODY
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    normal.element.rPr.rFonts.set(qn("w:ascii"), FONT_WEST)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_WEST)
    _add_page_number_footer(sec)
    return doc


def _add_page_number_footer(section):
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(para, 1.0)
    _set_indent_chars(para, 0)
    run = para.add_run()
    _set_run_font(run, FONT_SONG, FONT_WEST, SIZE_FOOTER)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# --------------------------------------------------------------- 段落工厂 ---
def add_title(doc, text: str):
    """文档标题：与一级标题同规格（黑体 小三 居中 无缩进）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_indent_chars(p, 0)
    _set_spacing(p, BODY_LINE_SPACING)
    run = p.add_run(text)
    _set_run_font(run, FONT_HEI, FONT_WEST, SIZE_H1)
    _no_widow_control(p)
    return p


def add_h1(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_indent_chars(p, 0)
    _set_spacing(p, BODY_LINE_SPACING)
    run = p.add_run(text)
    _set_run_font(run, FONT_HEI, FONT_WEST, SIZE_H1)
    _no_widow_control(p)
    return p


def add_h2(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_indent_chars(p, 2)
    _set_spacing(p, BODY_LINE_SPACING)
    run = p.add_run(text)
    _set_run_font(run, FONT_HEI, FONT_WEST, SIZE_H2)
    _no_widow_control(p)
    return p


def add_h3(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_indent_chars(p, 2)
    _set_spacing(p, BODY_LINE_SPACING)
    run = p.add_run(text)
    _set_run_font(run, FONT_SONG, FONT_WEST, SIZE_H3, bold=True)
    _no_widow_control(p)
    return p


def add_body(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_indent_chars(p, 2)
    _set_spacing(p, BODY_LINE_SPACING)
    run = p.add_run(text)
    _set_run_font(run, FONT_SONG, FONT_WEST, SIZE_BODY)
    return p


# ----------------------------------------------------------------- 表  格 ---
def _style_cell(cell, text: str, bold: bool = False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_indent_chars(para, 0)
    _set_spacing(para, TABLE_LINE_SPACING, before=0, after=0)
    run = para.add_run(text)
    _set_run_font(run, FONT_SONG, FONT_WEST, SIZE_TABLE, bold=bold)
    # 单元格内容垂直居中
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "center")
    tc_pr.append(v_align)


def _set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def add_table(doc, headers, rows, col_widths_cm=None):
    """headers: list[str]；rows: list[list[str]]；col_widths_cm: list[float] 可选。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = col_widths_cm is None
    _set_table_borders(table)

    for idx, head in enumerate(headers):
        _style_cell(table.rows[0].cells[idx], str(head), bold=True)
    # 表头跨页重复
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)

    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            _style_cell(cells[idx], "" if val is None else str(val))

    if col_widths_cm:
        for row in table.rows:
            for idx, width in enumerate(col_widths_cm):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    return table


# ------------------------------------------------------------ 渲染入口 -----
BLOCK_RENDERERS = {
    "title": lambda doc, b: add_title(doc, b["text"]),
    "h1": lambda doc, b: add_h1(doc, b["text"]),
    "h2": lambda doc, b: add_h2(doc, b["text"]),
    "h3": lambda doc, b: add_h3(doc, b["text"]),
    "p": lambda doc, b: add_body(doc, b["text"]),
    "table": lambda doc, b: add_table(
        doc, b["headers"], b["rows"], b.get("col_widths_cm")
    ),
    "pagebreak": lambda doc, b: doc.add_page_break(),
}


def render(blocks, out_path: str) -> str:
    """把内容块列表渲染成 Word 文件，返回文件路径。

    块的形态：
        {"type": "title" | "h1" | "h2" | "h3" | "p", "text": "..."}
        {"type": "table", "headers": [...], "rows": [[...], ...]}
        {"type": "pagebreak"}
    """
    doc = new_document()
    for block in blocks:
        kind = block.get("type", "p")
        renderer = BLOCK_RENDERERS.get(kind)
        if renderer is None:
            add_body(doc, str(block.get("text", "")))
        else:
            renderer(doc, block)
    doc.save(out_path)
    return out_path
