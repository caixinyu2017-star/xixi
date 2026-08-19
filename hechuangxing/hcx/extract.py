# -*- coding: utf-8 -*-
"""附件解析：把上传的创业计划书变成纯文本。"""
import base64
import io
import mimetypes
from pathlib import Path

TEXT_EXT = {".txt", ".md", ".markdown", ".csv", ".json", ".log"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
MAX_TEXT_CHARS = 60000


def extract(path: Path):
    """返回 (text, image_block)。image_block 用于走 Claude 的视觉输入。"""
    ext = path.suffix.lower()
    if ext == ".docx":
        return _from_docx(path), None
    if ext == ".pdf":
        return _from_pdf(path), None
    if ext in TEXT_EXT:
        return path.read_text(encoding="utf-8", errors="ignore")[:MAX_TEXT_CHARS], None
    if ext in IMAGE_EXT:
        return "", _image_block(path)
    if ext == ".doc":
        return ("[提示] 检测到 .doc 旧版格式，禾创星无法直接读取。"
                "请在 Word 里另存为 .docx 后重新上传，或把正文粘贴到对话框。"), None
    # 兜底：按文本猜一次
    try:
        return path.read_text(encoding="utf-8", errors="ignore")[:MAX_TEXT_CHARS], None
    except Exception:
        return f"[提示] 暂不支持的文件类型：{ext}", None


def _from_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _from_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n".join(parts).strip()
    if not text:
        return "[提示] 这份 PDF 没有可提取的文字层（可能是扫描件）。请上传 Word 版本或把正文粘贴到对话框。"
    return text[:MAX_TEXT_CHARS]


def _image_block(path: Path):
    media_type = mimetypes.guess_type(str(path))[0] or "image/png"
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return {
        "type": "image",
        "source": {"type": "base64", "media_type": media_type, "data": data},
    }
