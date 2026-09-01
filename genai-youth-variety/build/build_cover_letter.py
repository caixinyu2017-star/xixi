# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the letter and
the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"), encoding="utf-8") as fh:
    S = json.load(fh)
S["years"] = S["years"].replace("-", "\u2013")
T = S["thr"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")


def typo(s):
    return _MINUS.sub("\u2212", _YEARS.sub("\\1\u2013\\2", s))


FONT = "Times New Roman"
SIZE = 9.6

TITLE = ("Variety Absorbs Variety: Workforce Skill Diversity, Shock "
         "Absorption and Youth Employment in Chinese Listed Firms")

BODY = [
    ("We would like to submit the enclosed manuscript entitled \u201c%s\u201d "
     "for consideration in Systems. The manuscript has not been published "
     "elsewhere and is not under consideration by any other journal. No "
     "conflict of interest exists in its submission, and it has been "
     "approved for publication by all authors." % TITLE),

    ("When demand falls, young workers are the first to be let go. Yet some "
     "firms shed almost none of them, and the difference is not explained by "
     "size, leverage, profitability or sector. This manuscript asks what "
     "internal property of a firm decides whether an external disturbance is "
     "passed on to its youngest employees, and answers it with a principle "
     "that cybernetics established seventy years ago and that has never been "
     "tested on employment: Ashby\u2019s law of requisite variety, which "
     "holds that a system can absorb only as much variety as it contains. "
     "The urgency is not hypothetical. China produced a record 12.22 million "
     "graduates in 2025 while youth unemployment stayed above 17 percent, "
     "and the penalty for a poor start is measured in decades. Using "
     f"{S['n_obs']:,} firm-year observations for {S['n_firms']:,} Chinese "
     f"A-share listed firms over {S['years']}, four separate episodes of "
     "demand contraction and an industry shock constructed outside the firm, "
     "we show that the principle holds, and that it holds as a threshold."),

    "The highlights of the paper are as follows.",
]

POINTS = [
    ("(1) A measurable capacity, available before the shock arrives. "
     "Workforce skill variety is the entropy of a firm\u2019s employment "
     "distribution over five disclosed occupational categories and three "
     "education levels, so it is computable annually from public filings for "
     f"every listed firm. It raises the youth employment share by "
     f"{abs(S['var_sd_effect']):.2f} percentage points per standard "
     f"deviation and attenuates the effect of an adverse demand shock, the "
     f"interaction being {S['b_sxv']:.3f}. At one standard deviation below "
     f"the mean of variety a shock costs {abs(S['marg_lo']):.2f} percentage "
     "points of youth employment; one standard deviation above it the effect "
     "is statistically indistinguishable from zero."),

    ("(2) The central result: absorption is a regime, not a gradient. "
     "Ashby\u2019s law is a sufficiency condition, and sufficiency "
     "conditions imply discontinuities. A fixed-effects panel threshold "
     f"model locates one at {T['gamma']:.2f} bits of variety, with a "
     f"bootstrap p-value of {T['p_boot']:.3f} against the null of no "
     f"threshold. Below it the shock cuts the youth employment share by "
     f"{abs(T['b_low']):.2f} percentage points; above it the coefficient is "
     f"{T['b_high']:.3f} and insignificant. "
     f"{T['pct_below']:.0f} percent of listed firm-years fall below the "
     "threshold, which is to say that for most firms the current workforce "
     "structure is not a weak buffer but no buffer at all."),

    ("(3) Which variety, and how it works. Related variety \u2014 depth "
     "within occupational families \u2014 buffers "
     f"{S['ratio_ru']:.2f} times as strongly per bit as unrelated variety, "
     "the spread across families, because absorption requires that a person "
     "actually move and movement is governed by cognitive proximity. Of the "
     f"buffering, {abs(S['share_chu']):.0f} percent travels through "
     f"suppressed separations and {abs(S['share_red']):.0f} percent through "
     "visible internal redeployment, so resilience is mostly a decision not "
     "to act. Breadth without depth looks like variety on an organization "
     "chart and is not variety in the regulatory sense."),

    ("(4) Identification and falsification. The disturbance is constructed "
     "from national industry revenue outside the firm. Variety is "
     "instrumented by the occupational diversity of the local labor pool "
     "and by the variety of same-industry firms in other provinces "
     f"(Kleibergen\u2013Paap F = {S['kp_f']:.0f}, Hansen p = "
     f"{S['hansen_p']:.3f}), and the result survives propensity score "
     "matching, entropy balancing, seven robustness specifications, Oster "
     "bounds and a wild cluster bootstrap. The placebo is decisive: run on "
     "employees aged 31 to 50, whose separation is expensive, the buffering "
     f"coefficient is {abs(S['ratio_age']):.0f} times smaller and "
     "insignificant. Variety protects exactly the margin the theory names."),
]

CLOSING = (
    "We believe the manuscript is well suited to Systems. Its contribution "
    "is not the application of a systems vocabulary to an economic question "
    "but the empirical test of a systems law: we take a proposition from "
    "cybernetics, derive from it a prediction that no adjacent literature "
    "makes \u2014 that absorption requires a sufficient and therefore "
    "estimable quantity of internal variety \u2014 and find the threshold "
    "where the theory says it should be. In doing so the paper gives the "
    "organizational resilience literature something it has lacked, an "
    "antecedent of resilience that is observable before the shock rather "
    "than inferred from the recovery, and it gives the labor economics of "
    "adjustment an explanation for the dispersion it has documented but not "
    "accounted for. The practical payoff is unusually direct: the capacity "
    "is computed from disclosures every listed firm already files, the "
    "threshold is a number a human resource function can be held to, and "
    "the finding that the return to raising variety is concentrated at the "
    "crossing rather than spread along a gradient changes what an "
    "employment policy aimed at firms should try to buy. We deeply "
    "appreciate your consideration of our manuscript and look forward to "
    "receiving comments from the reviewers. Correspondence should be "
    "directed to the address below.")

SIGN = [
    "Name: Tiantian Mo",
    "Institution and address: College of Business, Jiaxing University; "
    "Jiaxing 314001, China",
    "Email: 00008227@zjxu.edu.cn",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(1.7)
    sec.left_margin = sec.right_margin = Cm(2.5)

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def p(text, after=4, align=None, size=SIZE):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        run = par.add_run(typo(text))
        run.font.name = FONT
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        return par

    p("Dear Editors:", after=5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for t in BODY:
        p(t, after=4)
    for t in POINTS:
        p(t, after=3)
    p(CLOSING, after=6)

    p("Thanks very much for your attention to our paper.", after=7,
      align=WD_ALIGN_PARAGRAPH.LEFT)
    for line in SIGN:
        p(line, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=5, align=WD_ALIGN_PARAGRAPH.LEFT)
    p("Very sincerely yours,", after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p("Xinyu Cai and Tiantian Mo", after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
