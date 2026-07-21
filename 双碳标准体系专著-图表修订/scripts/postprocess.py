# -*- coding: utf-8 -*-
"""docx 后处理：
1) 在指定图题后插入图注（与表注同格式：宋体9磅、无缩进），已存在注则替换文字；
2) 全部表格统一三线表（顶线/栏目线/底线）、字号五号(10.5pt)、单倍行距。
"""
import sys, copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_NOTES = {
    "图4.2": "注：数量为截至2025年累计有效标准数，数据来源于标准文本数据库。",
    "图9.2": "注：权重由德尔菲法两轮征询（26位专家，Kendall W=0.73）确定；判断矩阵一致性检验中准则层CR=0.032、方案层CR=0.047，均小于0.1。",
    "图10.1": "注：＋表示同向变动，－表示反向变动；R为增强回路，B为平衡回路；（延迟）表示存量累积滞后效应。",
    "图10.2": "注：矩形为存量，蝶形阀为流量，云朵为系统边界外的源与汇；蓝色箭头为信息链，红色虚线箭头为动态更新反馈。",
    "图13.3": "注：括号内数字为编码参考点频次；理论饱和检验显示第44份访谈后无新范畴出现，补访5份确认饱和。",
}

def make_note_para(doc, text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '240'); spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), '0')
    pPr.append(ind)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')      # 9pt
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '18')
    rPr.append(sz); rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p

def insert_fig_notes(doc):
    paras = doc.paragraphs
    done = {}
    for i, p in enumerate(paras):
        txt = p.text.strip().replace('　', ' ')
        for key, note in FIG_NOTES.items():
            if done.get(key):
                continue
            if txt.startswith(key + ' ') or txt.startswith(key + '  ') or txt == key:
                # 检查下一段是否已是注
                nxt = paras[i+1].text.strip() if i+1 < len(paras) else ''
                if nxt.startswith('注：'):
                    for r in list(paras[i+1].runs):
                        r.text = ''
                    if paras[i+1].runs:
                        paras[i+1].runs[0].text = note
                    else:
                        paras[i+1].add_run(note)
                    for r in paras[i+1].runs:
                        r.font.size = Pt(9); r.font.name = 'Times New Roman'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    p._p.addnext(make_note_para(doc, note))
                done[key] = True
    return done

def set_border(el, name, val, sz):
    b = OxmlElement(f'w:{name}')
    b.set(qn('w:val'), val)
    if val != 'nil':
        b.set(qn('w:sz'), str(sz)); b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
    el.append(b)

def three_line_table(tbl):
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    set_border(borders, 'top', 'single', 12)
    set_border(borders, 'bottom', 'single', 12)
    set_border(borders, 'left', 'nil', 0)
    set_border(borders, 'right', 'nil', 0)
    set_border(borders, 'insideH', 'nil', 0)
    set_border(borders, 'insideV', 'nil', 0)
    tblPr.append(borders)
    # 清除单元格级边框与底纹
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcBorders')) + tcPr.findall(qn('w:shd')):
                tcPr.remove(old)
    # 栏目线：第一行底边
    for cell in tbl.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tb = OxmlElement('w:tcBorders')
        set_border(tb, 'bottom', 'single', 6)
        tcPr.append(tb)
    # 字号五号 + 单倍行距
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                for old in pPr.findall(qn('w:spacing')):
                    pPr.remove(old)
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:line'), '240'); spacing.set(qn('w:lineRule'), 'auto')
                pPr.append(spacing)
                for r in p.runs:
                    r.font.size = Pt(10.5)
                    rPr = r._element.get_or_add_rPr()
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is None:
                        szCs = OxmlElement('w:szCs'); rPr.append(szCs)
                    szCs.set(qn('w:val'), '21')

def main(src, dst):
    doc = Document(src)
    done = insert_fig_notes(doc)
    print('fig notes inserted:', sorted(done))
    n = 0
    for tbl in doc.tables:
        three_line_table(tbl); n += 1
    print('tables restyled:', n)
    doc.save(dst)

if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
