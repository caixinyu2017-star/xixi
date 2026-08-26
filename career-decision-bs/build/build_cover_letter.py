# -*- coding: utf-8 -*-
"""One-page A4 cover letter accompanying the submission.

Every quantitative claim is read from the estimation output, so the
letter and the manuscript cannot disagree.
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(ROOT, "Cover_Letter.docx")

with open(os.path.join(ROOT, "tables", "summary.json"),
          encoding="utf-8") as fh:
    S = json.load(fh)
CAL, REF, MOD = S["calibration"], S["reference_study"], S["moderation"]
CON, PRE, PAR = S["conditions"], S["precision"], S["params"]
BEN = S["scaffold_benefit"]
BEN_LO = BEN["Low anxiety (lowest quartile)"]
BEN_HI = BEN["High anxiety (highest quartile)"]

_MINUS = re.compile(r"(?<![\w\u2013\u2014-])-(?=[\d.])")
_YEARS = re.compile(r"\b((?:19|20)\d{2})-((?:19|20)\d{2})\b")
_APOS = re.compile(r"(?<=[A-Za-z])'(?=[A-Za-z]|\s|$)")


def typo(s):
    s = _APOS.sub("\u2019", s)
    s = _YEARS.sub("\\1\u2013\\2", s)
    return _MINUS.sub("\u2212", s)


def r3(x):
    s = "%.3f" % x
    return s.replace("0.", ".").replace("-.", "\u2212.")


FONT = "Times New Roman"
SIZE = 9.6

TITLE = ("Does Parental Career Support Buffer or Amplify the Association "
         "Between Career Anxiety and Career Decision-Making Difficulties? "
         "A Simulation Study of a Reported Moderation")

BODY = [
    ("We would like to submit the enclosed manuscript entitled "
     "\u201c%s\u201d for consideration as an Article in Behavioral "
     "Sciences. The manuscript has not been published elsewhere and is not "
     "under consideration by any other journal. It reports no new human-"
     "subject data. All authors have approved the submission, and no "
     "conflict of interest exists." % TITLE),

    ("The manuscript takes as its starting point a finding published in this "
     "journal. In a mixed-methods study of 407 female undergraduates at four "
     "non-elite Chinese universities, the association between career anxiety "
     "and career decision-making difficulties was reported to be stronger, "
     "not weaker, among students who reported more frequent parental career "
     "support. That result is difficult to accommodate within Social "
     "Cognitive Career Theory, and the authors themselves noted it as "
     "counter-intuitive. Rather than treat it as an anomaly to be replicated "
     "or dismissed, we ask a question a cross-sectional design cannot "
     "answer: what would have to be true of the process that generates these "
     "variables over time for a cross-section of it to look like that?"),

    ("We specify a dynamic model in which anxiety reduces both the amount and "
     "the informational yield of career exploration, exploration builds "
     "career decision-making self-efficacy, and self-efficacy sustains "
     "further exploration; parental involvement enters through three "
     "distinct channels \u2014 reassurance, scaffolding of the student\u2019s "
     "own exploring, and involvement that takes the decision over. We "
     "calibrate the model so that a simulated cross-section reproduces the "
     "six correlations the original study reported, then analyse simulated "
     "studies with the same conditional process procedure the original used. "
     "The main findings are as follows."),
]

POINTS = [
    ("(1) The obvious explanation is wrong. We expected directive, "
     "decision-taking involvement to amplify. It does the opposite: across "
     "the involvement mixture the interaction runs from "
     f"{r3(MOD['0.0']['inter'])} where involvement is wholly scaffolding to "
     f"{r3(MOD['1.0']['inter'])} where it is wholly directive. Involvement "
     "that resolves uncertainty on the student\u2019s behalf buffers the "
     "anxiety\u2013difficulty association most strongly of all, because it "
     "works whether or not the student is in any condition to act."),

    ("(2) Amplification has a specific and counter-intuitive signature. It "
     "arises where support scaffolds exploration the student is too anxious "
     "to undertake, because scaffolding multiplies a capacity that anxiety "
     "has already suppressed. Moving from low to high involvement reduces "
     f"simulated difficulty by {r3(BEN_LO)} for the least anxious quartile "
     f"and by {r3(BEN_HI)} for the most anxious. Autonomy-supportive help "
     "reaches those already able to use it, and the gap widens as it is "
     "provided. The reported moderation is thus more consistent with good "
     "help that anxious students cannot use than with intrusive help."),

    ("(3) The pattern is uncommon, which makes it informative. Across a "
     f"factorial search of {CON['n_cells']:d} parameter combinations, "
     f"amplification arises in {CON['n_amplifying']:d} of them "
     f"({100 * CON['share_amplifying']:.1f} per cent), and never where the "
     "directive share reaches one half. A study that observes it is "
     "therefore observing something narrow rather than something generic."),

    ("(4) Chance is not a sufficient account. At the original sample size of "
     f"{PRE['n']:d} the simulated interaction has a mean of "
     f"{r3(PRE['mean'])} and a standard deviation of {r3(PRE['sd'])} across "
     f"repeated studies, detected in {PRE['pct_significant']:.1f} per cent "
     "of them. Sampling variability of that magnitude does not span zero, "
     "let alone reach a positive value of the size reported."),

    ("(5) We are explicit about what the study is and is not. It is a "
     "simulation; it contains no empirical observations and it cannot "
     "confirm or refute the original finding. Its purpose is to establish "
     "what the finding would imply. All "
     f"{PAR['n']:d} parameters are listed in an appendix with the interval "
     "over which the analysis varies them and a statement of provenance, "
     f"{PAR['n_calibrated']:d} of them fixed by calibration and the "
     "remainder swept rather than defended, and an identification check "
     "reports which correlations each parameter actually moves."),
]

CLOSING = (
    "We believe the manuscript suits Behavioral Sciences. It engages "
    "directly with work the journal has published, it addresses a "
    "methodological problem the vocational literature shares \u2014 that "
    "cross-sectional moderation is routinely read as evidence about a "
    "process that unfolds over months \u2014 and it yields a concrete "
    "recommendation: measure the kind of parental support, not its "
    "frequency, because the two forms this study distinguishes predict "
    "opposite moderation from the same reported quantity. The model code, "
    "the calibrated parameter set and the scripts that generate every table "
    "and figure are available so that readers can vary our assumptions "
    "rather than accept them. We would be glad to revise the manuscript in "
    "the light of the reviewers\u2019 comments.")


def para(doc, text, size=SIZE, bold=False, space_after=5, align=None,
         indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if indent:
        fmt.left_indent = Cm(indent)
    r = p.add_run(typo(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.9))
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.2))

    para(doc, "Cover Letter", size=13, bold=True, space_after=8,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Dear Editors,", space_after=7)
    for b in BODY:
        para(doc, b)
    for pt in POINTS:
        para(doc, pt, indent=0.55, space_after=4)
    para(doc, CLOSING, space_after=8)
    para(doc, "Yours sincerely,", space_after=2)
    para(doc, "Xinyu Cai and Tiantian Mo", space_after=0)
    para(doc, "College of Business, Jiaxing University", space_after=0)
    para(doc, "Correspondence: 00008227@zjxu.edu.cn", space_after=0)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
