# -*- coding: utf-8 -*-
"""Word 文档构建工具：中文学术排版、三线表、OMML 公式、红色修订标记。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, parse_xml
import copy, re, os

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

SONG = '宋体'
HEI = '黑体'
TIMES = 'Times New Roman'

M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSDECL = 'xmlns:m="%s" xmlns:w="%s"' % (M_NS, W_NS)


# ---------------------------------------------------------------- 字体
def set_run_font(run, cn=SONG, en=TIMES, size=12, bold=False, color=BLACK, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
    rF.set(qn('w:ascii'), en); rF.set(qn('w:hAnsi'), en); rF.set(qn('w:eastAsia'), cn)
    return run


def para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_chars=2.0,
             line=1.5, before=0, after=0, keep_next=False):
    pf = p.paragraph_format
    p.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pPr = p._element.get_or_add_pPr()
    ind = pPr.get_or_add_ind()
    if indent_chars:
        ind.set(qn('w:firstLineChars'), str(int(indent_chars * 100)))
        ind.set(qn('w:firstLine'), str(int(indent_chars * 240)))
    else:
        for k in ('w:firstLineChars', 'w:firstLine'):
            if ind.get(qn(k)) is not None:
                ind.attrib.pop(qn(k))
    if keep_next:
        pPr.get_or_add_keepNext()
    wc = pPr.get_or_add_widowControl()
    wc.set(qn('w:val'), '0')
    return p


# ------------------------------------------------- 富文本标记解析 «红» 【粗】
TOKEN = re.compile(r'(«|»|【|】)')


def add_rich(p, text, cn=SONG, en=TIMES, size=12, base_color=BLACK, base_bold=False):
    """解析 «…» 标红、【…】加粗（可嵌套），逐段写入 run。"""
    red = 0; bold = 0
    buf = ''
    def flush():
        nonlocal buf
        if buf:
            r = p.add_run(buf)
            set_run_font(r, cn=cn, en=en, size=size,
                         bold=base_bold or bool(bold),
                         color=RED if red else base_color)
            buf = ''
    for tok in TOKEN.split(text):
        if tok == '«':
            flush(); red += 1
        elif tok == '»':
            flush(); red = max(0, red - 1)
        elif tok == '【':
            flush(); bold += 1
        elif tok == '】':
            flush(); bold = max(0, bold - 1)
        else:
            buf += tok
    flush()
    return p


# ---------------------------------------------------------------- 三线表
def _set_cell_border(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kw:
            spec = kw[edge]
            tag = 'w:{}'.format(edge)
            el = borders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag); borders.append(el)
            for k, v in spec.items():
                el.set(qn('w:{}'.format(k)), str(v))


NONE_B = {'val': 'none', 'sz': 0, 'space': 0, 'color': 'auto'}
THICK = {'val': 'single', 'sz': 12, 'space': 0, 'color': '000000'}   # 1.5pt
THIN = {'val': 'single', 'sz': 6, 'space': 0, 'color': '000000'}     # 0.75pt


def add_three_line_table(doc, headers, rows, widths=None, font_size=10.5,
                         header_cn=HEI, body_cn=SONG, red_all=False,
                         align_first_left=True, header_rows=1):
    ncol = len(headers[0]) if isinstance(headers[0], (list, tuple)) else len(headers)
    hdr_rows = headers if isinstance(headers[0], (list, tuple)) else [headers]
    t = doc.add_table(rows=len(hdr_rows) + len(rows), cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tblPr = t._tbl.tblPr
    lay = tblPr.get_or_add_tblLayout()
    lay.set(qn('w:type'), 'fixed')
    if widths:
        total = sum(widths)
        for r in t.rows:
            for i, c in enumerate(r.cells):
                c.width = Cm(widths[i])
        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            for i, gc in enumerate(grid.findall(qn('w:gridCol'))):
                gc.set(qn('w:w'), str(int(widths[i] * 567)))
    data = [list(h) for h in hdr_rows] + [list(r) for r in rows]
    nh = len(hdr_rows)
    for i, rowvals in enumerate(data):
        for j, val in enumerate(rowvals):
            cell = t.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            is_h = i < nh
            para_fmt(p, align=(WD_ALIGN_PARAGRAPH.LEFT if (not is_h and j == 0 and align_first_left)
                               else WD_ALIGN_PARAGRAPH.CENTER) if not is_h else WD_ALIGN_PARAGRAPH.CENTER,
                     indent_chars=0, line=1.0, before=1.5, after=1.5)
            txt = str(val)
            if red_all and '«' not in txt:
                txt = '«' + txt + '»'
            add_rich(p, txt, cn=header_cn if is_h else body_cn, en=TIMES,
                     size=font_size, base_bold=False)
            # 边框
            b = {'left': NONE_B, 'right': NONE_B, 'insideV': NONE_B, 'insideH': NONE_B,
                 'top': NONE_B, 'bottom': NONE_B}
            if i == 0:
                b['top'] = THICK
            if i == nh - 1:
                b['bottom'] = THIN
            if i == len(data) - 1:
                b['bottom'] = THICK
            _set_cell_border(cell, **b)
            tcPr = cell._tc.get_or_add_tcPr()
            va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)
    # 表头跨页重复；单行不跨页断开
    for i, row in enumerate(t.rows):
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit'); trPr.append(cs)
        if i < nh:
            th = OxmlElement('w:tblHeader'); trPr.append(th)
    return t


# ---------------------------------------------------------------- OMML
def _m(tag, *children, **attrs):
    a = ''.join(' m:%s="%s"' % (k, v) for k, v in attrs.items())
    inner = ''.join(children)
    return '<m:%s%s>%s</m:%s>' % (tag, a, inner, tag)


def mr(text, sty='i'):
    """数学 run。sty: 'i' 斜体变量, 'p' 正体, 'bi' 粗斜"""
    esc = (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
    rpr = '<m:rPr><m:sty m:val="%s"/></m:rPr>' % sty
    wpr = ('<w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" '
           'w:eastAsia="宋体"/><w:sz w:val="24"/></w:rPr>')
    return '<m:r>%s%s<m:t xml:space="preserve">%s</m:t></m:r>' % (rpr, wpr, esc)


def msup(base, sup):
    return _m('sSup', _m('e', base), _m('sup', sup))


def msub(base, sub):
    return _m('sSub', _m('e', base), _m('sub', sub))


def mfrac(num, den):
    return _m('f', '<m:fPr><m:type m:val="bar"/></m:fPr>', _m('num', num), _m('den', den))


def mrad(deg, e):
    pr = '<m:radPr>%s</m:radPr>' % ('' if deg else '<m:degHide m:val="1"/>')
    return _m('rad', pr, _m('deg', deg), _m('e', e))


def mdelim(inner, beg='(', end=')'):
    if beg == '(' and end == ')':
        pr = ''          # 默认即为圆括号，省略 dPr 可避免部分渲染器解析异常
    else:
        pr = ('<m:dPr><m:begChr m:val="%s"/><m:endChr m:val="%s"/></m:dPr>' % (beg, end))
    return _m('d', pr, _m('e', inner))


def mnary(chr_, sub, sup, e):
    pr = ('<m:naryPr><m:chr m:val="%s"/><m:limLoc m:val="subSup"/>'
          '<m:subHide m:val="%d"/><m:supHide m:val="%d"/></m:naryPr>'
          % (chr_, 0 if sub else 1, 0 if sup else 1))
    return _m('nary', pr, _m('sub', sub), _m('sup', sup), _m('e', e))


def add_equation(doc, omml_body, number=None, usable_cm=14.66, red=True):
    """居中公式 + 行末右顶格编号。"""
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, indent_chars=0, line=1.5, before=6, after=6)
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    t1 = OxmlElement('w:tab'); t1.set(qn('w:val'), 'center')
    t1.set(qn('w:pos'), str(int(usable_cm / 2 * 567)))
    t2 = OxmlElement('w:tab'); t2.set(qn('w:val'), 'right')
    t2.set(qn('w:pos'), str(int(usable_cm * 567)))
    tabs.append(t1); tabs.append(t2)
    old = pPr.find(qn('w:tabs'))
    if old is not None:
        pPr.remove(old)
    pPr._insert_tabs(tabs)

    r0 = p.add_run(); r0._element.append(OxmlElement('w:tab'))
    set_run_font(r0, size=12, color=RED if red else BLACK)
    xml = '<m:oMath %s>%s</m:oMath>' % (NSDECL, omml_body)
    p._element.append(parse_xml(xml))
    if number:
        r1 = p.add_run(); r1._element.append(OxmlElement('w:tab'))
        set_run_font(r1, size=12, color=RED if red else BLACK)
        r2 = p.add_run('（%s）' % number)
        set_run_font(r2, cn=SONG, en=TIMES, size=12, color=RED if red else BLACK)
    return p


# ---------------------------------------------------------------- 域代码（目录）
def add_field(p, instr):
    r = p.add_run()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
    r._element.append(fc)
    r2 = p.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    r2._element.append(it)
    r3 = p.add_run()
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
    r3._element.append(fc2)
    r4 = p.add_run('请在 Word 中右键“更新域”以生成目录。')
    set_run_font(r4, cn=SONG, size=12, color=BLACK)
    r5 = p.add_run()
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')
    r5._element.append(fc3)
    for rr in (r, r2, r3, r5):
        set_run_font(rr, cn=SONG, size=12)
    return p


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r._element.append(fc)
    r2 = p.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    r2._element.append(it)
    r3 = p.add_run()
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end'); r3._element.append(fc2)
    for rr in (r, r2, r3):
        set_run_font(rr, cn=SONG, en=TIMES, size=10.5)


def msubsup(base, sub, sup):
    return _m('sSubSup', _m('e', base), _m('sub', sub), _m('sup', sup))


def mgroup(*parts):
    return ''.join(parts)
