# -*- coding: utf-8 -*-
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docbuild import (set_run_font, para_fmt, add_rich, add_three_line_table,
                      add_equation, add_field, add_page_number_footer,
                      RED, BLACK, SONG, HEI, TIMES)

from content_front import FRONT, TITLE, SUBTITLE
from content_ch1 import CH1
from content_ch2 import CH2
from content_ch3 import CH3
from content_ch4 import CH4
from content_ch5 import CH5, APPENDIX
from content_refs import REFDB


# ---------------- 引用编号：顺序编码制 ----------------
import re as _re
_CIT = _re.compile(r'\[(@[A-Za-z0-9_]+(?:,\s*@[A-Za-z0-9_]+)*)\]')
_order = []


def _fmt(nums):
    nums = sorted(set(nums))
    out, i = [], 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j - i >= 1:
            out.append('%d-%d' % (nums[i], nums[j]))
        else:
            out.append(str(nums[i]))
        i = j + 1
    return '[' + ','.join(out) + ']'


def _sub(m):
    keys = [k.strip().lstrip('@') for k in m.group(1).split(',')]
    nums = []
    for k in keys:
        if k not in REFDB:
            raise KeyError('未知参考文献键：' + k)
        if k not in _order:
            _order.append(k)
        nums.append(_order.index(k) + 1)
    return _fmt(nums)


def resolve_citations(blocks):
    out = []
    for blk in blocks:
        blk = list(blk)
        for i, item in enumerate(blk):
            if isinstance(item, str):
                blk[i] = _CIT.sub(_sub, item)
            elif isinstance(item, list):
                blk[i] = [[_CIT.sub(_sub, c) if isinstance(c, str) else c for c in row]
                          if isinstance(row, list) else
                          (_CIT.sub(_sub, row) if isinstance(row, str) else row)
                          for row in item]
        out.append(tuple(blk))
    return out


BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, 'fig')
USABLE = 14.66

doc = Document()

# ---------------- 页面设置 ----------------
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.17)
sec.right_margin = Cm(3.17)
add_page_number_footer(sec)

# ---------------- 默认样式 ----------------
st = doc.styles['Normal']
st.font.name = TIMES
st.font.size = Pt(12)
st._element.rPr.rFonts.set(qn('w:eastAsia'), SONG)
st.paragraph_format.space_before = Pt(0)
st.paragraph_format.space_after = Pt(0)

HEAD_SPEC = {
    1: dict(cn=HEI, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=14),
    2: dict(cn=HEI, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=8),
    3: dict(cn=HEI, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6),
    4: dict(cn=SONG, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=4),
}
for lv, spec in HEAD_SPEC.items():
    s = doc.styles['Heading %d' % lv]
    s.font.name = TIMES
    s.font.size = Pt(spec['size'])
    s.font.bold = spec['bold']
    s.font.color.rgb = BLACK
    s._element.rPr.rFonts.set(qn('w:eastAsia'), spec['cn'])
    s._element.rPr.rFonts.set(qn('w:ascii'), TIMES)
    s._element.rPr.rFonts.set(qn('w:hAnsi'), TIMES)


def heading(text, level):
    spec = HEAD_SPEC[level]
    p = doc.add_paragraph(style='Heading %d' % level)
    para_fmt(p, align=spec['align'], indent_chars=0, line=1.5,
             before=spec['before'], after=spec['after'], keep_next=True)
    add_rich(p, text, cn=spec['cn'], en=TIMES, size=spec['size'],
             base_bold=spec['bold'])
    return p


def body(text, indent=2.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, cn=SONG):
    p = doc.add_paragraph()
    para_fmt(p, align=align, indent_chars=indent, line=1.5, before=0, after=0)
    add_rich(p, text, cn=cn, en=TIMES, size=size)
    return p


def figure(fname, caption, width_cm):
    for ext in ('', '.png', '.jpeg', '.jpg'):
        path = os.path.join(FIG, fname + ext)
        if os.path.exists(path):
            break
    else:
        raise FileNotFoundError(fname)
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, line=1.0,
             before=8, after=3, keep_next=True)
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    cp = doc.add_paragraph()
    para_fmt(cp, align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, line=1.0,
             before=0, after=8)
    add_rich(cp, '«' + caption + '»', cn=HEI, en=TIMES, size=10.5)
    return cp


def table_caption(text):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, line=1.0,
             before=8, after=3, keep_next=True)
    add_rich(p, '«' + text + '»', cn=HEI, en=TIMES, size=10.5)
    return p


def table_note(text):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, indent_chars=0, line=1.0,
             before=3, after=8)
    add_rich(p, '«' + text + '»', cn=SONG, en=TIMES, size=10.5)
    return p


def render(blocks):
    for b in blocks:
        k = b[0]
        if k == 'h1':
            heading(b[1], 1)
        elif k == 'h2':
            heading(b[1], 2)
        elif k == 'h3':
            heading(b[1], 3)
        elif k == 'h4':
            heading(b[1], 4)
        elif k == 'p':
            body(b[1])
        elif k == 'pn':
            body(b[1], indent=0)
        elif k == 'fig':
            figure(b[1], b[2], b[3])
        elif k == 'tcap':
            table_caption(b[1])
        elif k == 'tab':
            hdr, rows, widths, red_all = b[1], b[2], b[3], b[4]
            add_three_line_table(doc, hdr, rows, widths=widths,
                                 font_size=10.5, red_all=red_all)
            sp = doc.add_paragraph()
            para_fmt(sp, indent_chars=0, line=1.0, before=0, after=0)
            sp.add_run('')
        elif k == 'note':
            table_note(b[1])
        elif k == 'eq':
            add_equation(doc, b[1], b[2], usable_cm=USABLE, red=True)
        elif k == 'toc':
            p = doc.add_paragraph()
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, indent_chars=0, line=1.5)
            add_field(p, ' TOC \\o "1-3" \\h \\z \\u ')
        elif k == 'pagebreak':
            p = doc.add_paragraph()
            para_fmt(p, indent_chars=0, line=1.0)
            p.add_run().add_break(WD_BREAK.PAGE)
        else:
            raise ValueError(k)


# ---------------- 封面 ----------------
for _ in range(4):
    p = doc.add_paragraph(); para_fmt(p, indent_chars=0, line=1.5)
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, line=1.6, before=0, after=10)
add_rich(p, '«' + TITLE + '»', cn=HEI, en=TIMES, size=18, base_bold=True)
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, line=1.6, before=0, after=0)
add_rich(p, '«' + SUBTITLE + '»', cn=HEI, en=TIMES, size=14, base_bold=True)
for _ in range(8):
    p = doc.add_paragraph(); para_fmt(p, indent_chars=0, line=1.5)
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, line=1.5)
add_rich(p, '«管理案例分析»', cn=HEI, en=TIMES, size=14)
p = doc.add_paragraph()
para_fmt(p, indent_chars=0, line=1.0)
p.add_run().add_break(WD_BREAK.PAGE)

# ---------------- 正文 ----------------
render(resolve_citations(FRONT))
render(resolve_citations(CH1))
render(resolve_citations(CH2))
render(resolve_citations(CH3))
render(resolve_citations(CH4))
render(resolve_citations(CH5))

# ---------------- 参考文献 ----------------
heading('参考文献', 1)
REFS = [REFDB[k] for k in _order]
_uncited = [k for k in REFDB if k not in _order]
if _uncited:
    raise SystemExit('存在未被正文引用的文献：%s' % _uncited)
print('引用文献数：', len(REFS))
for i, r in enumerate(REFS, 1):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_chars=0, line=1.4, before=0, after=2)
    pPr = p._element.get_or_add_pPr()
    ind = pPr.get_or_add_ind()
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:left'), '480')
    ind.set(qn('w:hangingChars'), '200')
    ind.set(qn('w:hanging'), '480')
    add_rich(p, '«[%d] %s»' % (i, r), cn=SONG, en=TIMES, size=10.5)

p = doc.add_paragraph()
para_fmt(p, indent_chars=0, line=1.0)
p.add_run().add_break(WD_BREAK.PAGE)
render(resolve_citations(APPENDIX))

# 修复默认模板 settings.xml 中缺少 percent 属性的 zoom 元素
_st = doc.settings.element
_zoom = _st.find(qn('w:zoom'))
if _zoom is not None and _zoom.get(qn('w:percent')) is None:
    _zoom.set(qn('w:percent'), '100')

out = os.path.join(BASE, '夜间经济场景理论案例分析_修订稿.docx')
doc.save(out)
print('saved:', out)
