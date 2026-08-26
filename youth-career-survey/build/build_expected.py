# -*- coding: utf-8 -*-
"""The expected-results report.

Every number is read from out/summary.json, produced by analysis/run.py over
the simulated dataset. The document says what it is on its first page and in
a banner on every page thereafter.
"""
import json
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
sys.path.insert(0, os.path.join(ROOT, "design"))
import items as I                                             # noqa: E402

OUT = os.path.join(ROOT, "预期结果_基于模拟数据.docx")
TAB = os.path.join(ROOT, "out")
with open(os.path.join(TAB, "summary.json"), encoding="utf-8") as fh:
    S = json.load(fh)

CN, CNB, EN = "宋体", "黑体", "Times New Roman"
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0x9C, 0x00, 0x00)


def _font(run, name=CN, size=10.5, bold=False, color=None, italic=False):
    run.font.name = EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), name)
    rf.set(qn("w:ascii"), EN)
    rf.set(qn("w:hAnsi"), EN)


def para(doc, text="", size=10.5, bold=False, name=CN, align=None,
         before=0, after=4, indent=0, first=0, color=None, spacing=1.3):
    p = doc.add_paragraph()
    f = p.paragraph_format
    f.space_before, f.space_after = Pt(before), Pt(after)
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    f.line_spacing = spacing
    if align is not None:
        p.alignment = align
    if indent:
        f.left_indent = Cm(indent)
    if first:
        f.first_line_indent = Cm(first)
    if text:
        _font(p.add_run(_MI.sub("\u2212", text) if _MI else text),
              name, size, bold, color)
    return p


def h1(doc, t):
    para(doc, t, 13.5, True, CNB, before=13, after=5)


def h2(doc, t):
    para(doc, t, 11.5, True, CNB, before=9, after=4)


def _shade(cell, hexv):
    pr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexv)
    pr.append(sh)


def _cell(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    f = p.paragraph_format
    f.space_before = f.space_after = Pt(1.5)
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    f.line_spacing = 1.02
    _font(p.add_run(text), CN, size, bold)


def _fix(tbl, widths):
    pr = tbl._tbl.tblPr
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    pr.append(lay)
    g = tbl._tbl.find(qn("w:tblGrid"))
    for gc, w in zip(g.findall(qn("w:gridCol")), widths):
        gc.set(qn("w:w"), str(int(round(w * 567))))
    for row in tbl.rows:
        for c, w in zip(row.cells, widths):
            c.width = Cm(w)


def _rule(tbl, row, sz="8"):
    for c in tbl.rows[row].cells:
        pr = c._tc.get_or_add_tcPr()
        b = OxmlElement("w:tcBorders")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), sz)
        bot.set(qn("w:color"), "000000")
        b.append(bot)
        pr.append(b)


def _top(tbl):
    for c in tbl.rows[0].cells:
        pr = c._tc.get_or_add_tcPr()
        b = OxmlElement("w:tcBorders")
        t = OxmlElement("w:top")
        t.set(qn("w:val"), "single")
        t.set(qn("w:sz"), "12")
        t.set(qn("w:color"), "000000")
        b.append(t)
        pr.append(b)


_MI = __import__("re").compile(r"(?<![\w–—-])-(?=[\d.])")


def mi(t):
    return _MI.sub("\u2212", str(t))


def read(name):
    rows = [l.rstrip("\n").split("\t")
            for l in open(os.path.join(TAB, name), encoding="utf-8")]
    return rows[0], rows[1:]


def table(doc, num, caption, name=None, header=None, rows=None, widths=None,
          note=None, size=9, first_left=True):
    if name:
        header, rows = read(name)
    para(doc, "表 %d  %s" % (num, caption), 10, True, CNB, before=8, after=3)
    k = len(header)
    widths = widths or [16.6 / k] * k
    tbl = doc.add_table(rows=len(rows) + 1, cols=k)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for j, h in enumerate(header):
        _cell(tbl.rows[0].cells[j], h, size, True)
        _shade(tbl.rows[0].cells[j], "F2F2F2")
    for i, r in enumerate(rows):
        for j, v in enumerate(r):
            al = (WD_ALIGN_PARAGRAPH.LEFT if (j == 0 and first_left)
                  else WD_ALIGN_PARAGRAPH.CENTER)
            _cell(tbl.rows[i + 1].cells[j], mi(v), size, False, al)
    _fix(tbl, widths)
    _top(tbl)
    _rule(tbl, 0)
    _rule(tbl, len(rows))
    if note:
        para(doc, mi("注：" + note), 8.5, False, CN, after=6, color=GREY,
             spacing=1.15)
    return tbl


def banner(doc):
    """A red strip repeated at the top of every page after the first."""
    sec = doc.sections[-1]
    hdr = sec.header
    hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run("本文件全部数值来自模拟数据，不含任何真实被试作答，"
                    "不可作为研究结果引用"), CN, 8.5, True, RED)


F = S["cfa"]
R = S["reliability"]
G = S["regression"]
M = S["mediation"]
P = S["power_curve"]
CM = S["cmb"]
N = S["meta"]["n"]
NAME = {"CA": "职业决策焦虑", "CE": "职业探索行为", "SE": "职业决策自我效能",
        "PA": "父母自主支持型生涯支持", "PD": "父母指导代办型生涯介入",
        "PF": "父母生涯参与频率", "CD": "职业决策困难"}


def f2(x):
    return "%.2f" % x


def f3(x):
    return "%.3f" % x


def pc(x):
    return "%.1f" % (100 * x)


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = EN
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CN)
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = s.bottom_margin = Cm(2.0)

    # ==================================================== warning page
    para(doc, "", after=40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), "FDE9E9")
    pr.append(sh)
    _font(p.add_run("本文件是预期结果，不是研究结果"), CNB, 20, True, RED)

    para(doc, "", after=14)
    for t in [
        "以下全部表格与数值，均由 sim/simulate.py 按本研究所假设的结构方程"
        "生成的 400 份模拟作答计算得到。没有任何一位真实被试参与。",
        "它的用途是：在正式发放问卷之前，把分析脚本跑通、把论文表格的结构定"
        "下来、并判断计划的样本量能不能检出所假设的效应。",
        "它不能用来支持任何实证结论，不能写进论文的结果部分，也不能提交给期"
        "刊或数据仓储。数据既然是按假设生成的，它当然“支持”假设——这是它的"
        "用途，也是它的界限。",
        "真实数据回收之后，请用完全相同的脚本重跑一遍，并以真实结果为准，"
        "无论方向是否与本文件一致。",
    ]:
        para(doc, t, 11, False, CN, after=8, first=0.78, spacing=1.5)

    para(doc, "", after=10)
    rows = [
        ["随机种子", str(S["meta"]["seed"])],
        ["模拟发放 / 有效", "%d / %d 份"
         % (S["sample"]["n_collected"], S["sample"]["n_valid"])],
        ["Bootstrap 次数", "%d" % S["meta"]["boots"]],
        ["检验力重复次数", "每个样本量 %d 次" % S["meta"]["power_reps"]],
        ["生成脚本", "sim/simulate.py"],
        ["分析脚本", "analysis/run.py"],
        ["运行耗时", "%.1f 秒" % S["meta"]["runtime_s"]],
    ]
    table(doc, 0, "", header=["项目", "取值"], rows=rows, widths=[5.4, 11.2],
          size=9.5)
    doc.paragraphs[-1]._p.getparent().remove(doc.paragraphs[-1]._p)

    doc.add_section(WD_SECTION.NEW_PAGE)
    banner(doc)

    # ==================================================== 1
    h1(doc, "一、研究模型与假设")
    para(doc, "本研究检验一个中心命题：父母在子女职业问题上的“介入方式”与"
              "“介入多少”，对职业决策困难的作用方向不同，而目前常用的量表"
              "只测后者。据此提出如下假设。", 10.5, first=0.74)
    rows = [
        ["H1", "职业决策焦虑正向预测职业决策困难。", "预期成立"],
        ["H2a", "职业探索行为在焦虑与困难之间起中介作用。", "预期成立"],
        ["H2b", "职业决策自我效能在焦虑与困难之间起中介作用。", "预期成立"],
        ["H2c", "焦虑 → 探索 → 自我效能 → 困难构成链式中介。", "预期成立"],
        ["H3a", "父母自主支持型生涯支持放大焦虑与困难的关系"
                "（交互项为正）。", "预期成立"],
        ["H3b", "父母指导代办型生涯介入缓冲焦虑与困难的关系"
                "（交互项为负）。", "预期成立"],
        ["H3c", "父母生涯参与频率不调节焦虑与困难的关系"
                "（交互项不显著）。", "预期成立（零假设）"],
    ]
    table(doc, 1, "研究假设", header=["编号", "内容", "模拟数据下的表现"],
          rows=rows, widths=[1.6, 11.0, 4.0])
    para(doc, "H3c 是一个方向明确的零假设，本身构成本研究的方法学贡献："
              "如果“介入多少”确实不调节，而“介入方式”的两种成分调节方向相反，"
              "那么以往用频率类量表得到的调节结果就是两种相反效应的混合，"
              "其符号取决于样本中两类家长的比例。", 10.5, first=0.74)

    # ==================================================== 2
    h1(doc, "二、样本与数据质量")
    para(doc, "按计划发放 %d 份，依据施测前确定的三条规则剔除后保留 %d 份"
              "有效问卷，有效率 %.1f%%。剔除情况为：%s。"
         % (S["sample"]["n_collected"], N,
            100.0 * N / S["sample"]["n_collected"],
            "；".join("%s %d 份" % (k, v)
                      for k, v in S["sample"]["excluded"].items())),
         10.5, first=0.74)
    hdr, rows = read("t1_sample.tsv")
    keep = [r for r in rows if r[0] in ("性别", "年级", "就读院校层次",
                                        "生源地", "目前的求职状态")]
    table(doc, 2, "有效样本的构成（节选）", header=hdr, rows=keep,
          widths=[3.6, 7.0, 3.0, 3.0],
          note="完整分布见 out/t1_sample.tsv。")

    # ==================================================== 3
    h1(doc, "三、测量模型")
    h2(doc, "3.1　信度与收敛效度")
    table(doc, 3, "各构念的信度与收敛效度", "t2_measurement.tsv",
          widths=[5.0, 1.9, 3.1, 2.4, 2.2, 2.0],
          note="七个构念的 Cronbach α 介于 %.2f 与 %.2f 之间，CR 均高于 .70。"
               "职业探索、职业决策自我效能与父母指导代办三个构念的 AVE 略低于"
               ".50，但 CR 达标，按 Fornell 与 Larcker 的判断标准仍可接受；"
               "论文中需就此说明。"
               % (min(v["alpha"] for v in R.values()),
                  max(v["alpha"] for v in R.values())))

    h2(doc, "3.2　验证性因子分析与竞争模型")
    para(doc, "对七个构念的 %d 个条目做验证性因子分析。假设的七因子模型拟合"
              "良好：χ²(%d) = %.1f，χ²/df = %.2f，CFI = %.3f，TLI = %.3f，"
              "RMSEA = %.3f，SRMR = %.3f。"
         % (sum(len(I.BY_KEY[k].items) for k in I.SUBSTANTIVE),
            F["df"], F["chi2"], F["ratio"], F["cfi"], F["tli"],
            F["rmsea"], F["srmr"]), 10.5, first=0.74)
    table(doc, 4, "竞争测量模型的比较", "t3_model_comparison.tsv",
          widths=[4.6, 2.4, 1.7, 1.7, 1.6, 1.6, 1.6, 1.4],
          note="关键的比较是第一行与第二行：把父母自主支持、指导代办与参与"
               "频率三个量表合并为一个“父母生涯支持”因子，拟合明显变差"
               "（CFI 由 %.3f 降至 %.3f）。这是“介入方式”与“介入多少”在"
               "测量层面就可以分开的直接证据，也是全文立论的前提。"
               % (F["cfi"], S["model_comparison"]["five"]["cfi"]))
    para(doc, "区分效度方面，HTMT 比值最大为 %.2f，低于 .85 的判断标准"
              "（完整矩阵见 out/t4_htmt.tsv）。" % S["htmt_max"],
         10.5, first=0.74)

    h2(doc, "3.3　共同方法偏差")
    para(doc, "采用三种程序性与统计性方法。程序上，问卷匿名作答，量表呈现"
              "顺序随机化，并混用五点与九点两种作答等级。统计上，Harman "
              "单因素检验中首个未旋转因子解释 %.1f%% 的变异，低于 40%% 的"
              "常用界限；标记变量法以理论无关的颜色偏好量表（α = %.2f，"
              "与各研究变量相关的绝对值最大为 %.2f）作偏相关校正，校正前后"
              "各相关系数变动的最大值为 %.3f，没有任何一个相关的显著性判断"
              "因此改变。共同方法偏差不构成严重威胁。"
         % (100 * CM["harman_first"], CM["marker_alpha"],
            CM["marker_max_r"], CM["max_shift"]), 10.5, first=0.74)

    # ==================================================== 4
    h1(doc, "四、描述统计与相关")
    table(doc, 5, "各变量的描述统计与相关矩阵", "t5_correlations.tsv",
          widths=[5.0, 1.1, 1.0, 1.0, 1.15, 1.15, 1.15, 1.15, 1.15, 1.15, 1.4],
          size=8.5,
          note="N = %d；列号与行号对应；对角线括号内为 Cronbach α；"
               "* p < .05，** p < .01，"
               "*** p < .001。请留意最后一行：父母生涯参与频率与职业决策困难"
               "的相关为 %.2f，接近于零；而自主支持与指导代办分别为 %.2f 与 "
               "%.2f，方向相反。频率之所以“无关”，是因为它把两种方向相反的"
               "介入方式加在了一起。"
               % (N, S["correlations"]["PF-CD"],
                  S["correlations"]["PA-CD"], S["correlations"]["PD-CD"]))

    # ==================================================== 5
    h1(doc, "五、假设检验")
    h2(doc, "5.1　链式中介")
    table(doc, 6, "焦虑经由探索与自我效能影响困难的中介效应",
          "t8_mediation.tsv", widths=[8.6, 3.0, 5.0],
          note="标准化系数；控制全部背景变量；偏差校正百分位 Bootstrap，"
               "重复 %d 次。三条间接路径的区间均不含零，H2a、H2b、H2c 成立。"
               "间接效应合计为 %.3f，占总效应 %.3f 的 %.1f%%。"
               % (S["meta"]["boots"], M["total_ind"], M["total"],
                  100 * M["total_ind"] / M["total"]))

    h2(doc, "5.2　三个调节效应")
    table(doc, 7, "以职业决策困难为因变量的层级回归", "t6_hierarchical.tsv",
          widths=[4.6, 4.0, 4.0, 4.0],
          note="标准化回归系数；模型 1 仅含 10 个背景变量；* p < .05，"
               "** p < .01，*** p < .001。方差膨胀因子最大为 %.2f，"
               "远低于 10，不存在多重共线性问题。" % G["max_vif"])
    table(doc, 8, "简单斜率检验", "t7_simple_slopes.tsv",
          widths=[5.0, 4.0, 4.0, 1.8, 1.8],
          note="表中数值为职业决策焦虑对职业决策困难的标准化回归斜率，"
               "分别取调节变量的均值上下一个标准差。")

    ca_pa = G["coef"]["CA×PA"]
    ca_pd = G["coef"]["CA×PD"]
    ca_pf = G["coef"]["CA×PF"]
    para(doc, "在这一次模拟抽样中，三个交互项的估计分别为：焦虑 × 自主支持 "
              "β = %.3f（p = %.3f），焦虑 × 指导代办 β = %.3f（p = %.3f），"
              "焦虑 × 参与频率 β = %.3f（p = %.3f）。方向全部符合假设，"
              "但第一个交互项没有达到显著。这不是偶然，下一节说明原因。"
         % (ca_pa["b"], ca_pa["p"], ca_pd["b"], ca_pd["p"],
            ca_pf["b"], ca_pf["p"]), 10.5, first=0.74)

    # ==================================================== 6
    h1(doc, "六、检验力：400 份不够")
    para(doc, "在每个样本量下重复抽取 %d 项独立研究，每项都走完整的分析流程，"
              "记录三个交互项达到 p < .05 的比例。结果如下。"
         % S["meta"]["power_reps"], 10.5, first=0.74)
    table(doc, 9, "不同有效样本量下三个交互项的检出率", "t9_power.tsv",
          widths=[2.6, 2.6, 2.0, 2.6, 2.0, 2.6, 2.2], size=8.5,
          note="β 为 %d 次重复的均值。焦虑 × 参与频率一栏的检出率始终在 5%% "
               "上下，正是名义显著性水平，说明它确实没有效应——这是 H3c 的"
               "证据，而不是检验力不足。"
          % S["meta"]["power_reps"])

    p400 = P["400"]
    p900 = P["900"]
    para(doc, "结论很直接。在 400 份有效样本下，两个关键交互项的检出率只有 "
              "%.1f%% 与 %.1f%%——大约一半的概率会做出一项什么也没测到的"
              "研究。要让两者都达到 80%% 的常规标准，需要约 %s 份有效样本；"
              "按本次模拟 %.1f%% 的有效率计算，需要发放约 %d 份。"
         % (100 * p400["CA×PA"]["power"], 100 * p400["CA×PD"]["power"],
            S["required_n"] or "1000 以上",
            100.0 * N / S["sample"]["n_collected"],
            int(round((S["required_n"] or 1000)
                      * S["sample"]["n_collected"] / N))),
         10.5, first=0.74)
    para(doc, "交互项的真实标准化系数只有 %.2f 上下，这在本领域是常见量级；"
              "问题不在效应太弱，而在于用量表均分做的调节分析会因测量误差"
              "而衰减，且交互项本身的标准误大。这也提示，以往在 400 人左右"
              "样本上报告出显著调节的研究，其本身的检出概率可能也在五成"
              "上下。"
         % abs(p900["CA×PA"]["mean_b"]), 10.5, first=0.74)

    para(doc, "可选的应对办法，按推荐程度排序：", 10.5, first=0.74,
         after=3)
    for t in [
        "1. 扩大样本至约 %s 份有效问卷。这是唯一确定有效的办法。可以跨校"
        "联合发放，或把年级放宽到大三下学期。" % (S["required_n"] or "900"),
        "2. 若样本量确实受限，则改变论证策略：把 H3c（频率不调节）与测量模型"
        "比较（表 4）作为主结果，把 H3a、H3b 作为探索性结果报告，并明确说明"
        "检验力限制。这样论文依然成立，但结论的力度下降。",
        "3. 不要在拿到数据后才决定报告哪一个。上述两条路线请在发放问卷之前"
        "选定，并写入分析计划。",
    ]:
        para(doc, t, 10.5, False, CN, after=3, indent=0.5)

    # ==================================================== 7
    h1(doc, "七、正式施测前的清单")
    for t in [
        "□ 确定并落实样本量方案（见第六节）。",
        "□ 完成伦理审查报批，取得批件编号，填入问卷首页与论文方法部分。",
        "□ 核对量表授权：职业决策自我效能对应的 CDMSE-SF 由 Mind Garden "
        "发行，若使用原条目须先购买许可；其余量表引用原文即可。",
        "□ 在问卷星中按《问卷》附录的“施测与质量控制要点”配置：随机化量表"
        "顺序、全部必答、记录时长、限答一次。",
        "□ 先做 30 人左右的预试，检查条目理解、作答时长与题项区分度，"
        "必要时修订措辞（修订后需重跑本流程）。",
        "□ 将剔除规则、假设与分析计划在数据回收前写定并存档，"
        "有条件可在 OSF 预注册。",
        "□ 回收后用 analysis/run.py 原样重跑，不修改脚本；"
        "以真实结果为准撰写论文。",
    ]:
        para(doc, t, 10.5, False, CN, after=4, indent=0.4)

    para(doc, "", after=16)
    para(doc, "本文件全部数值来自模拟数据。真实数据回收后作废。", 10.5,
         True, CNB, WD_ALIGN_PARAGRAPH.CENTER, color=RED)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
