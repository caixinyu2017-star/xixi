# -*- coding: utf-8 -*-
"""The questionnaire, as it will be printed and as it will be keyed into
Wenjuanxing. Item text and codes come from design/items.py, so the printed
form and the data file cannot drift apart.
"""
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
sys.path.insert(0, os.path.join(ROOT, "design"))
import items as I                                             # noqa: E402

OUT = os.path.join(ROOT, "高校毕业生职业决策调查问卷.docx")
CN = "宋体"
CNB = "黑体"
EN = "Times New Roman"


def _font(run, name=CN, size=10.5, bold=False, italic=False, color=None):
    run.font.name = EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), name)
    rf.set(qn("w:ascii"), EN)
    rf.set(qn("w:hAnsi"), EN)


def para(doc, text="", size=10.5, bold=False, name=CN, align=None,
         before=0, after=3, indent=0, first=0, spacing=1.25, color=None):
    p = doc.add_paragraph()
    f = p.paragraph_format
    f.space_before, f.space_after = Pt(before), Pt(after)
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    f.line_spacing = spacing
    if align is not None:
        p.alignment = align
    if indent:
        f.left_indent = Cm(indent)
    if first:
        f.first_line_indent = Cm(first)
    if text:
        _font(p.add_run(text), name, size, bold, color=color)
    return p


def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), "808080")
    b.append(bot)
    pr.append(b)


def _shade(cell, hexv):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexv)
    tcpr.append(sh)


def _cell(cell, text, size=9.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    f = p.paragraph_format
    f.space_before, f.space_after = Pt(1.5), Pt(1.5)
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    f.line_spacing = 1.05
    _font(p.add_run(text), CN, size, bold)


def _borders(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + side)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "9A9A9A")
        b.append(e)
    tblPr.append(b)


def _fix_layout(tbl, widths_cm):
    """python-docx needs an explicit grid and a fixed layout to honour widths."""
    tblPr = tbl._tbl.tblPr
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    grid_el = tbl._tbl.find(qn("w:tblGrid"))
    for gc, w in zip(grid_el.findall(qn("w:gridCol")), widths_cm):
        gc.set(qn("w:w"), str(int(round(w * 567))))
    for row in tbl.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def grid(doc, rows, k, first_w=None):
    """A Likert grid: statement rows against k numbered response columns."""
    if first_w is None:
        first_w = 10.6 if k == 5 else 8.5
    opt_w = (16.6 - first_w) / k
    tbl = doc.add_table(rows=len(rows) + 1, cols=k + 1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _borders(tbl)
    hdr = tbl.rows[0]
    _cell(hdr.cells[0], "题项", 9.5, True, WD_ALIGN_PARAGRAPH.CENTER)
    _shade(hdr.cells[0], "EFEFEF")
    for j in range(k):
        _cell(hdr.cells[j + 1], str(j + 1), 9.5, True)
        _shade(hdr.cells[j + 1], "EFEFEF")
    for i, (code, text) in enumerate(rows):
        r = tbl.rows[i + 1]
        _cell(r.cells[0], text, 9.5, False, WD_ALIGN_PARAGRAPH.LEFT)
        for j in range(k):
            _cell(r.cells[j + 1], "○", 9.5)
    _fix_layout(tbl, [first_w] + [opt_w] * k)
    return tbl


def section(doc, no, con, lead):
    para(doc, "%s、%s" % (no, con.cn), 12, True, CNB, before=10, after=2)
    para(doc, lead, 10, False, CN, after=2,
         color=RGBColor(0x44, 0x44, 0x44))
    txt, k = I.SCALES[con.scale]
    para(doc, txt, 9.5, False, CN, after=4,
         color=RGBColor(0x44, 0x44, 0x44))
    return k


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = EN
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CN)

    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)

    # ------------------------------------------------------------ cover
    para(doc, "高校毕业生职业决策状况调查问卷", 17, True, CNB,
         WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, "A Survey of Career Decision Making among Chinese "
              "University Graduating Students", 10.5, False, CN,
         WD_ALIGN_PARAGRAPH.CENTER, after=12)
    rule(doc)

    para(doc, "亲爱的同学：", 11, True, CNB, after=4)
    for t in [
        "您好！这是一项关于高校毕业生职业决策的学术研究。我们希望了解，在临近"
        "毕业、需要为自己的职业方向做出选择的时候，同学们的心理状态、探索行为，"
        "以及来自家庭的影响之间是怎样的关系。您的回答将帮助学校和就业指导工作"
        "更好地理解毕业生的真实处境。",
        "本问卷不记名，全部题目没有对错之分，也没有“标准答案”，请按照您的真实"
        "感受作答即可。全部作答约需 12 分钟。",
        "根据研究伦理要求，我们向您说明以下几点：",
    ]:
        para(doc, t, 10.5, False, CN, after=3, first=0.74)

    for t in [
        "（1）参与完全自愿。您可以在任何时候退出，无需说明理由，不会有任何影响。",
        "（2）问卷不收集姓名、学号、手机号、身份证号等可识别个人身份的信息。",
        "（3）数据仅用于学术研究，以汇总统计的形式呈现，不会出现任何个人的作答。",
        "（4）原始数据由研究者加密保存，保存期满后销毁。",
        "（5）本研究已通过嘉兴学院相关伦理审查程序。",
    ]:
        para(doc, t, 10.5, False, CN, after=2, indent=0.74)

    para(doc, "如对本研究有任何疑问，欢迎联系：嘉兴学院商学院  "
              "（邮箱：00008227@zjxu.edu.cn）。", 10.5, False, CN,
         after=6, first=0.74)
    para(doc, "衷心感谢您的参与！", 10.5, True, CNB,
         WD_ALIGN_PARAGRAPH.RIGHT, after=2)
    para(doc, "课题组  敬上", 10.5, False, CN,
         WD_ALIGN_PARAGRAPH.RIGHT, after=8)
    rule(doc)

    tbl = doc.add_table(rows=1, cols=2)
    _borders(tbl)
    _cell(tbl.rows[0].cells[0],
          "□ 我已阅读并理解上述说明，自愿参与本次调查。", 10.5, True,
          WD_ALIGN_PARAGRAPH.LEFT)
    _cell(tbl.rows[0].cells[1],
          "□ 我不愿参与（选择此项将结束作答）", 10.5, False,
          WD_ALIGN_PARAGRAPH.LEFT)
    _fix_layout(tbl, [9.6, 7.0])

    doc.add_section(WD_SECTION.NEW_PAGE)

    # ------------------------------------------------- Part 1: background
    para(doc, "第一部分  基本情况", 13, True, CNB, after=3)
    para(doc, "请在符合您情况的选项前打勾。", 10, False, CN, after=6,
         color=RGBColor(0x44, 0x44, 0x44))
    for i, (code, name, opts) in enumerate(I.DEMO, 1):
        p = para(doc, "", 10.5, after=3)
        _font(p.add_run("%d. %s：" % (i, name)), CN, 10.5, True)
        _font(p.add_run("   " + "    ".join("□ " + o.split(" ", 1)[1]
                                            for o in opts)), CN, 10.5)

    # ------------------------------------------------- Part 2: the scales
    para(doc, "第二部分  正式题项", 13, True, CNB, before=14, after=3)
    para(doc, "以下题目请根据每部分开头说明的等级作答，在相应数字下的圆圈中"
              "选择。", 10, False, CN, after=6,
         color=RGBColor(0x44, 0x44, 0x44))

    order = [
        (I.CA, "下面是一些同学在面对职业选择时可能有的感受。请判断每句话与您"
               "近来的情况有多符合。"),
        (I.CE, "请回想过去三个月，您做下面这些事情的频繁程度。"),
        (I.SE, "如果现在就要做下面这些事，您对自己能做好有多大信心？"),
        (I.PA, "下面是一些关于父母在您职业问题上的做法。请判断每句话与您家里"
               "的情况有多符合。"),
        (I.PD, "同样是关于父母的做法，请继续判断符合程度。"),
        (I.PF, "请回想过去半年，下列情况发生的频繁程度。"),
        (I.CD, "下面是同学们在做职业决定时常遇到的困难。请判断每句话在多大"
               "程度上符合您现在的情况。本部分为九点量表。"),
        (I.MK, "最后是几道与前面内容无关的题目，用于问卷质量检验。"),
    ]
    cn_num = "一二三四五六七八九十"
    for si, (con, lead) in enumerate(order):
        k = section(doc, cn_num[si], con, lead)
        rows = [(c, "%s. %s" % (c, t)) for c, t, _ in con.items]
        rev = [c for c, _, r in con.items if r]

        # attention checks sit inside a block, not at the end
        if con.key == "CE":
            pass
        if con.key == "CD":
            ac = I.ATTENTION[1]
            rows = rows[:6] + [(ac[0], "%s. %s" % (ac[0], ac[1]))] + rows[6:]
        grid(doc, rows, k)
        if rev:
            para(doc, "注：第 %s 题为反向计分题。" % "、".join(rev), 9,
                 False, CN, after=2, color=RGBColor(0x66, 0x66, 0x66))
        if con.key == "CE":
            ac = I.ATTENTION[0]
            para(doc, "", after=2)
            grid(doc, [(ac[0], "%s. %s" % (ac[0], ac[1]))], 5)

    para(doc, "问卷到此结束，再次感谢您的参与！", 11, True, CNB,
         WD_ALIGN_PARAGRAPH.CENTER, before=14, after=4)

    # ------------------------------------------------- annex: provenance
    doc.add_section(WD_SECTION.NEW_PAGE)
    para(doc, "附录  量表来源与使用授权说明", 13, True, CNB, after=3)
    para(doc, "本问卷题目为面向中国高校毕业生编写的中文测量条目，对应下列已"
              "发表工具所界定的构念，不是其原始条目的逐字翻译。正式施测前请"
              "逐项确认授权状态；标注 ★ 的量表需要事先取得许可。",
         10, False, CN, after=6, color=RGBColor(0x44, 0x44, 0x44))
    tbl = doc.add_table(rows=len(I.SOURCES) + 1, cols=3)
    _borders(tbl)
    for j, h in enumerate(["构念", "对应的已发表工具", "授权与使用说明"]):
        _cell(tbl.rows[0].cells[j], h, 9.5, True)
        _shade(tbl.rows[0].cells[j], "EFEFEF")
    for i, key in enumerate(["CA", "CE", "SE", "PA", "PD", "PF", "CD", "MK"]):
        src, lic = I.SOURCES[key]
        r = tbl.rows[i + 1]
        _cell(r.cells[0], I.BY_KEY[key].cn, 9, False,
              WD_ALIGN_PARAGRAPH.LEFT)
        _cell(r.cells[1], src, 9, False, WD_ALIGN_PARAGRAPH.LEFT)
        _cell(r.cells[2], lic, 9, False, WD_ALIGN_PARAGRAPH.LEFT)
    _fix_layout(tbl, [3.0, 7.4, 6.2])

    para(doc, "施测与质量控制要点", 12, True, CNB, before=12, after=3)
    for t in [
        "1. 发放方式：通过问卷星生成链接，由各学院辅导员在毕业年级班级群"
        "内投放，同一 IP 与同一微信账号限答一次。",
        "2. 题目呈现：第二部分各量表的呈现顺序随机化，量表内部题目顺序固定，"
        "以免破坏反向题与注意力检测题的位置设计。",
        "3. 强制作答：全部题目设为必答，避免缺失值；同时记录作答时长。",
        "4. 剔除规则（施测前即确定，不得事后调整）：注意力检测题 AC1、AC2 "
        "任一错答；作答时长短于 180 秒；全部条目作答标准差不足 0.30（连续"
        "选同一选项）。",
        "5. 计划有效样本：见《预期结果》一稿中的检验力分析结论。",
    ]:
        para(doc, t, 10.5, False, CN, after=3, indent=0.4)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
